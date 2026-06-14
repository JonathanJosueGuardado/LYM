
from __future__ import annotations

import base64
import hashlib
import html
import json
import os
import platform
import re
import secrets
import shutil
import socket
import sys
import tempfile
import time
import unicodedata
import uuid
from dataclasses import asdict, dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any, Optional

try:
    from cryptography.fernet import Fernet, InvalidToken
except Exception as exc:  # pragma: no cover
    print("ERROR: falta cryptography. Ejecuta: pip install cryptography")
    raise

try:  # La app puede compilar/importar sin PySide6; la GUI lo exige al ejecutar normal.
    from PySide6.QtCore import QDate, QPointF, QRectF, Qt, Signal, QTimer, QUrl
    from PySide6.QtGui import QAction, QColor, QDesktopServices, QFont, QIcon, QPainter, QPen, QPixmap
    from PySide6.QtWidgets import (
        QAbstractItemView, QAbstractSpinBox, QApplication, QCheckBox, QComboBox, QDateEdit,
        QDialog, QDoubleSpinBox, QFileDialog, QFormLayout, QFrame, QGridLayout,
        QGroupBox, QHBoxLayout, QHeaderView, QInputDialog, QLabel, QLineEdit,
        QListWidget, QMainWindow, QMessageBox, QPushButton, QScrollArea,
        QSpinBox, QSplitter, QStackedWidget, QStatusBar, QTableWidget,
        QTableWidgetItem, QTabWidget, QTextEdit, QVBoxLayout, QWidget,
    )
    PYSIDE_OK = True
except Exception:
    PYSIDE_OK = False

# =============================================================================
# CONSTANTES DEL SISTEMA
# =============================================================================

APP_NAME = "LYM AUTO CONTROL"
APP_SUBTITLE = "L & M Inversiones"
APP_VERSION = "v5.0_LEASING"
ADMIN_PASSWORD = "ADMIN123"  # Clave para abrir Configuración desde el login
APP_DIR_NAME = "LYM_AUTO_CONTROL"
SUBFOLDER_NAME = "CONTROL VEHICULOS LYM"

SUB_DATOS = "DATOS"
SUB_VEHICULOS = "VEHICULOS"
SUB_VEHICULOS_REGISTROS = "VEHICULOS_REGISTROS"
SUB_DOCUMENTOS = "DOCUMENTOS"
SUB_FOTOS = "FOTOS"
SUB_REPORTES = "REPORTES"
SUB_RESPALDOS = "RESPALDOS"
SUB_TEMP = "TEMP"
SUB_PLANTILLAS = "PLANTILLAS"
ALL_SUBFOLDERS = [
    SUB_DATOS, SUB_VEHICULOS, SUB_DOCUMENTOS, SUB_FOTOS,
    SUB_REPORTES, SUB_RESPALDOS, SUB_TEMP, SUB_PLANTILLAS,
]

F_KEY = "lym.key"
F_USUARIOS = "usuarios.json"
F_DISPOSITIVOS = "dispositivos.json"
F_COUNTERS = "correlativos.json"
F_LOCAL_CONFIG = "local_config.json"
F_AUDITORIA = "auditoria.enc"
F_CATALOG_MARCAS = "marcas.json"
F_CATALOG_SUBASTAS = "subastas.json"
F_CATALOG_TALLERES = "talleres.json"
F_CATALOG_NAVIERAS = "navieras.json"
F_CATALOG_ADUANAS = "aduanas.json"
F_CATALOG_PAISES = "paises_destino.json"
F_CATALOG_PROVEEDORES = "proveedores.json"
F_CATALOG_GESTORES = "gestores.json"
F_CATALOG_TRANSPORTISTAS_USA = "transportistas_usa.json"
F_CATALOG_TARIFAS = "tarifas.json"
F_CATALOG_ADUANAS_PAIS = "aduanas_por_pais.json"
F_CATALOG_ESTADOS_USA = "estados_usa.json"
F_HTML_TEMPLATE = "LYM_Inversiones_Presentacion_Actualizada_83_Vehiculos.html"

ROLE_ADMIN = "ADMIN"
ROLE_SUPERVISOR = "SUPERVISOR"
ROLE_OPERACIONES = "OPERACIONES"
ROLE_VENTAS = "VENTAS"
ROLE_CONTABILIDAD = "CONTABILIDAD"
ROLE_USUARIO = "USUARIO"
ROLES = [ROLE_ADMIN, ROLE_SUPERVISOR, ROLE_OPERACIONES, ROLE_VENTAS, ROLE_CONTABILIDAD, ROLE_USUARIO]

PERM_VIEW_ALL = "view_all"
PERM_VIEW_COSTS = "view_costs"
PERM_CREATE_PURCHASE = "create_purchase"
PERM_UPDATE_STAGE = "update_stage"
PERM_ADD_COSTS = "add_costs"
PERM_MARK_AVAILABLE = "mark_available"
PERM_MARK_SALE = "mark_sale"
PERM_GENERATE_REPORTS = "generate_reports"
PERM_CONFIG = "config"
PERM_EDIT_CRITICAL = "edit_critical"
PERM_DELETE_CANCEL = "delete_cancel"
ALL_PERMISSION_KEYS = [
    PERM_VIEW_ALL, PERM_VIEW_COSTS, PERM_CREATE_PURCHASE, PERM_UPDATE_STAGE,
    PERM_ADD_COSTS, PERM_MARK_AVAILABLE, PERM_MARK_SALE, PERM_GENERATE_REPORTS,
    PERM_CONFIG, PERM_EDIT_CRITICAL, PERM_DELETE_CANCEL,
]

PERM_LABELS = {
    PERM_VIEW_ALL: "Ver todo el inventario",
    PERM_VIEW_COSTS: "Ver costos y rentabilidad",
    PERM_CREATE_PURCHASE: "Crear compras",
    PERM_UPDATE_STAGE: "Actualizar etapas",
    PERM_ADD_COSTS: "Agregar costos",
    PERM_MARK_AVAILABLE: "Marcar disponible para venta",
    PERM_MARK_SALE: "Reservar / vender",
    PERM_GENERATE_REPORTS: "Generar reportes",
    PERM_CONFIG: "Acceso a configuración",
    PERM_EDIT_CRITICAL: "Editar datos críticos",
    PERM_DELETE_CANCEL: "Anular / eliminar",
}

def perm_label(key: str) -> str:
    return PERM_LABELS.get(key, key)

DEFAULT_MARCAS = [
    "TOYOTA", "HONDA", "HYUNDAI", "KIA", "NISSAN", "FORD", "CHEVROLET",
    "MAZDA", "MITSUBISHI", "JEEP", "BMW", "MERCEDES-BENZ", "LEXUS",
    "VOLKSWAGEN", "SUBARU", "DODGE", "RAM", "GMC", "AUDI", "OTRO",
]
DEFAULT_SUBASTAS = ["COPART", "IAAI", "MANHEIM", "ADESA", "DEALER", "COMPRA DIRECTA", "OTRO"]
DEFAULT_TALLERES = ["SIN ASIGNAR"]
DEFAULT_NAVIERAS = ["SIN ASIGNAR", "MIAMI", "SEABOARD", "KING OCEAN", "CROWLEY", "OTRA"]
DEFAULT_PROVEEDORES = ["SIN ASIGNAR", "OTRO"]
DEFAULT_GESTORES = ["SIN ASIGNAR", "GESTOR", "OTRO"]
DEFAULT_TRANSPORTISTAS_USA = ["SIN ASIGNAR", "GRÚA USA", "OTRO"]
DEFAULT_TARIFAS = ["ESTANDAR 1.0", "PREFERENCIAL 0.8", "PREMIUM 1.2"]
DEFAULT_ADUANAS = ["SAN BARTOLO", "LA HACHADURA", "ACAJUTLA", "OTRA"]
DEFAULT_PAISES_DESTINO = ["EL SALVADOR", "HONDURAS", "GUATEMALA", "NICARAGUA", "COSTA RICA", "PANAMA"]
ADUANAS_BY_COUNTRY = {
    "EL SALVADOR": ["ACAJUTLA", "LA HACHADURA", "SAN BARTOLO", "EL AMATILLO", "OTRA"],
    "HONDURAS": ["LA MESA", "PUERTO CORTES", "EL AMATILLO", "AGUA CALIENTE", "LAS MANOS", "OTRA"],
    "GUATEMALA": ["SANTO TOMAS DE CASTILLA", "PUERTO QUETZAL", "TECUN UMAN", "PEDRO DE ALVARADO", "OTRA"],
    "NICARAGUA": ["CORINTO", "EL GUASAULE", "LAS MANOS", "PEÑAS BLANCAS", "OTRA"],
    "COSTA RICA": ["LIMON", "CALDERA", "PEÑAS BLANCAS", "PASO CANOAS", "OTRA"],
    "PANAMA": ["BALBOA", "CRISTOBAL", "COLON", "PASO CANOAS", "OTRA"],
}
US_STATES = [
    "ALABAMA", "ALASKA", "ARIZONA", "ARKANSAS", "CALIFORNIA", "COLORADO",
    "CONNECTICUT", "DELAWARE", "FLORIDA", "GEORGIA", "HAWAII", "IDAHO",
    "ILLINOIS", "INDIANA", "IOWA", "KANSAS", "KENTUCKY", "LOUISIANA",
    "MAINE", "MARYLAND", "MASSACHUSETTS", "MICHIGAN", "MINNESOTA",
    "MISSISSIPPI", "MISSOURI", "MONTANA", "NEBRASKA", "NEVADA",
    "NEW HAMPSHIRE", "NEW JERSEY", "NEW MEXICO", "NEW YORK", "NORTH CAROLINA",
    "NORTH DAKOTA", "OHIO", "OKLAHOMA", "OREGON", "PENNSYLVANIA",
    "RHODE ISLAND", "SOUTH CAROLINA", "SOUTH DAKOTA", "TENNESSEE", "TEXAS",
    "UTAH", "VERMONT", "VIRGINIA", "WASHINGTON", "WEST VIRGINIA",
    "WISCONSIN", "WYOMING",
]

STAGE_COMPRADO = "COMPRADO"
STAGE_TRASLADO_USA = "TRASLADO_USA"
STAGE_MIAMI_NAVIERA = "MIAMI_NAVIERA"
STAGE_TRANSITO = "TRANSITO"
STAGE_ADUANA = "ADUANA"
STAGE_LEGALIZACION = "LEGALIZACION"
STAGE_TALLER = "TALLER"
STAGE_DISPONIBLE = "DISPONIBLE_VENTA"
STAGE_RESERVADO = "RESERVADO"
STAGE_VENDIDO = "VENDIDO"
STAGE_ANULADO = "ANULADO"

STAGES = [
    {"key": STAGE_COMPRADO, "label": "Comprado", "icon": "🧾", "color": "#F59A13"},
    {"key": STAGE_TRASLADO_USA, "label": "Traslado USA", "icon": "🚚", "color": "#8BC6B3"},
    {"key": STAGE_MIAMI_NAVIERA, "label": "Miami / Naviera", "icon": "🏢", "color": "#FFB547"},
    {"key": STAGE_TRANSITO, "label": "Tránsito", "icon": "🚢", "color": "#0E3A78"},
    {"key": STAGE_ADUANA, "label": "Aduana", "icon": "🛃", "color": "#374151"},
    {"key": STAGE_LEGALIZACION, "label": "Legalización", "icon": "📄", "color": "#7C3AED"},
    {"key": STAGE_TALLER, "label": "Taller", "icon": "🛠", "color": "#F97316"},
    {"key": STAGE_DISPONIBLE, "label": "Disponible", "icon": "🏷", "color": "#10B981"},
    {"key": STAGE_RESERVADO, "label": "Reservado", "icon": "🤝", "color": "#D97706"},
    {"key": STAGE_VENDIDO, "label": "Vendido", "icon": "✅", "color": "#08285A"},
]
STAGE_ORDER = [s["key"] for s in STAGES]
STAGE_META = {s["key"]: s for s in STAGES}

ALERT_LIMITS_DAYS = {
    STAGE_TRASLADO_USA: 10,
    STAGE_MIAMI_NAVIERA: 10,
    STAGE_TRANSITO: 35,
    STAGE_ADUANA: 7,
    STAGE_LEGALIZACION: 12,
    STAGE_TALLER: 20,
    STAGE_DISPONIBLE: 60,
}

# =============================================================================
# UTILIDADES
# =============================================================================

def _strip_accents(s: str) -> str:
    if not s:
        return ""
    nfkd = unicodedata.normalize("NFKD", str(s))
    return "".join(c for c in nfkd if not unicodedata.combining(c))


def _norm(s: str) -> str:
    return _strip_accents(str(s or "")).upper().strip()


def _safe_filename(s: str) -> str:
    s = _strip_accents(str(s or ""))
    s = re.sub(r"[^A-Za-z0-9_\-]+", "_", s).strip("_")
    return s or "SIN_NOMBRE"


def _parse_date(value: Any) -> Optional[date]:
    if not value:
        return None
    if isinstance(value, date) and not isinstance(value, datetime):
        return value
    if isinstance(value, datetime):
        return value.date()
    try:
        return datetime.fromisoformat(str(value)[:10]).date()
    except Exception:
        return None


def _fmt_date(value: Any) -> str:
    d = _parse_date(value)
    return d.strftime("%d/%m/%Y") if d else "—"


def _fmt_usd(value: Any) -> str:
    try:
        return f"$ {float(value or 0):,.2f}"
    except Exception:
        return "$ 0.00"


def _now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _today_iso() -> str:
    return date.today().isoformat()


def _app_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    try:
        return Path(__file__).resolve().parent
    except Exception:
        return Path.cwd()


def _local_config_dir() -> Path:
    if os.name == "nt":
        base = Path(os.environ.get("APPDATA", str(Path.home())))
    else:
        base = Path(os.environ.get("XDG_CONFIG_HOME", str(Path.home() / ".config")))
    d = base / APP_DIR_NAME
    d.mkdir(parents=True, exist_ok=True)
    return d


def local_config_path() -> Path:
    return _local_config_dir() / F_LOCAL_CONFIG


def _read_json_file(path: Path, default: Any) -> Any:
    try:
        if not path.exists():
            return default
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def _write_json_file(path: Path, data: Any) -> bool:
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        tmp = path.with_suffix(path.suffix + ".tmp")
        tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
        tmp.replace(path)
        return True
    except Exception as exc:
        print(f"[write_json] {path}: {exc}")
        return False


def _read_local_config() -> dict:
    return _read_json_file(local_config_path(), {})


def _write_local_config(data: dict) -> None:
    _write_json_file(local_config_path(), data)


# ---- Autologin local, ligado al device_key de la laptop ----
def _xor_encrypt(plain: str, device_key_hex: str) -> str:
    if not plain:
        return ""
    try:
        key = bytes.fromhex(device_key_hex)
    except Exception:
        key = device_key_hex.encode("utf-8")
    if not key:
        return ""
    pb = plain.encode("utf-8")
    out = bytes(b ^ key[i % len(key)] for i, b in enumerate(pb))
    return base64.b64encode(out).decode("ascii")


def _xor_decrypt(cipher_b64: str, device_key_hex: str) -> str:
    if not cipher_b64:
        return ""
    try:
        try:
            key = bytes.fromhex(device_key_hex)
        except Exception:
            key = device_key_hex.encode("utf-8")
        if not key:
            return ""
        cb = base64.b64decode(cipher_b64)
        out = bytes(b ^ key[i % len(key)] for i, b in enumerate(cb))
        return out.decode("utf-8")
    except Exception:
        return ""


def save_autologin(usuario: str, password: str, device_key: str, folder: str) -> None:
    data = _read_local_config()
    data["autologin"] = {
        "usuario": usuario,
        "password_enc": _xor_encrypt(password, device_key),
        "device_key": device_key,
        "folder": folder,
    }
    data["active_folder"] = folder
    _write_local_config(data)


def load_autologin(device_key: str) -> tuple[str, str]:
    data = _read_local_config()
    al = data.get("autologin") or {}
    if not al or al.get("device_key") != device_key:
        return "", ""
    folder = al.get("folder")
    if folder and Path(folder).exists():
        data["active_folder"] = folder
        _write_local_config(data)
    return al.get("usuario", ""), _xor_decrypt(al.get("password_enc", ""), device_key)


def clear_autologin() -> None:
    data = _read_local_config()
    if "autologin" in data:
        del data["autologin"]
        _write_local_config(data)


DASHBOARD_CARD_KEYS = [
    ("total", "Vehículos"),
    ("activos", "Activos"),
    ("disponibles", "Disponibles"),
    ("vendidos", "Vendidos"),
    ("capital", "Capital activo"),
    ("ganancia", "Ganancia esperada"),
    ("criticos", "Críticos"),
    ("compra_disp", "Compra → Disponible"),
]

def default_dashboard_settings() -> dict:
    return {"recent_table": True, "cards": {k: True for k, _ in DASHBOARD_CARD_KEYS}}

def get_dashboard_settings() -> dict:
    data = _read_local_config()
    saved = data.get("dashboard_settings") or {}
    base = default_dashboard_settings()
    # Compatibilidad con versiones anteriores
    if "kpis" in saved:
        visible = bool(saved.get("kpis", True))
        for k, _ in DASHBOARD_CARD_KEYS:
            base["cards"][k] = visible
        base["recent_table"] = bool(saved.get("recent_table", True))
        if not bool(saved.get("show_financial", True)):
            base["cards"]["capital"] = False; base["cards"]["ganancia"] = False
        if not bool(saved.get("show_critical", True)):
            base["cards"]["criticos"] = False
        return base
    if isinstance(saved, dict):
        base["recent_table"] = bool(saved.get("recent_table", base["recent_table"]))
        cards = saved.get("cards") or {}
        if isinstance(cards, dict):
            for k in base["cards"]:
                if k in cards:
                    base["cards"][k] = bool(cards[k])
    return base

def set_dashboard_settings(settings: dict) -> None:
    data = _read_local_config()
    base = default_dashboard_settings()
    base["recent_table"] = bool(settings.get("recent_table", True))
    cards = settings.get("cards") or {}
    for k in base["cards"]:
        base["cards"][k] = bool(cards.get(k, True))
    data["dashboard_settings"] = base
    _write_local_config(data)


def get_active_folder() -> Optional[Path]:
    data = _read_local_config()
    p = data.get("active_folder")
    if p and Path(p).exists():
        return Path(p)
    return None


def set_active_folder(path: str | Path) -> None:
    data = _read_local_config()
    data["active_folder"] = str(path)
    data["last_update"] = _now_iso()
    _write_local_config(data)


def get_data_folder() -> Optional[Path]:
    root = get_active_folder()
    if root is None:
        return None
    return root / SUBFOLDER_NAME


def system_file(name: str) -> Optional[Path]:
    df = get_data_folder()
    return None if df is None else df / name


def datos_file(name: str) -> Optional[Path]:
    df = get_data_folder()
    return None if df is None else df / SUB_DATOS / name


def rel_to_abs(rel_path: str) -> Optional[Path]:
    df = get_data_folder()
    if df is None or not rel_path:
        return None
    return df / rel_path.replace("/", os.sep)


def abs_to_rel(path: Path) -> str:
    df = get_data_folder()
    if df is None:
        return str(path)
    try:
        return str(path.relative_to(df)).replace("\\", "/")
    except Exception:
        return str(path)

# =============================================================================
# RECURSOS / LOGO / PLANTILLA HTML
# =============================================================================

class ResourceManager:
    LOGO_CANDIDATES = [
        "LOGO LYM.png", "LOGO_LYM.png", "LOGO LYM.jpg", "LOGO_LYM.jpg",
        "LOGO LYM.jpeg", "LOGO_LYM.jpeg", "LOGO LYM.webp", "LOGO_LYM.webp",
        "LOGO LYM", "LOGO_LYM",
    ]
    TEMPLATE_CANDIDATES = [F_HTML_TEMPLATE, "LYM_INVERSIONES_TEMPLATE.html", "PLANTILLA_LYM.html"]

    @classmethod
    def search_paths(cls) -> list[Path]:
        paths = []
        def add(p: Optional[Path]):
            if p and p not in paths:
                paths.append(p)
        app = _app_dir()
        add(app)
        add(app / SUB_PLANTILLAS)
        df = get_data_folder()
        if df:
            add(df)
            add(df / SUB_PLANTILLAS)
            add(df / SUB_REPORTES / SUB_PLANTILLAS)
        return paths

    @classmethod
    def find_resource(cls, candidates: list[str]) -> Optional[Path]:
        for folder in cls.search_paths():
            try:
                if not folder.exists():
                    continue
                for name in candidates:
                    p = folder / name
                    if p.exists() and p.is_file():
                        return p
            except Exception:
                continue
        return None

    @classmethod
    def find_logo(cls) -> Optional[Path]:
        return cls.find_resource(cls.LOGO_CANDIDATES)

    @classmethod
    def find_html_template(cls) -> Optional[Path]:
        return cls.find_resource(cls.TEMPLATE_CANDIDATES)

    @classmethod
    def path_to_data_uri(cls, p: Optional[Path]) -> str:
        if p is None or not p.exists():
            return ""
        try:
            data = p.read_bytes()
            ext = p.suffix.lower()
            if ext == ".webp" or data[:4] == b"RIFF":
                mime = "image/webp"
            elif ext in (".jpg", ".jpeg") or data[:3] == b"\xff\xd8\xff":
                mime = "image/jpeg"
            else:
                mime = "image/png"
            return f"data:{mime};base64,{base64.b64encode(data).decode('ascii')}"
        except Exception:
            return ""

    @classmethod
    def logo_data_uri(cls) -> str:
        return cls.path_to_data_uri(cls.find_logo())

    @classmethod
    def template_css(cls) -> str:
        p = cls.find_html_template()
        if not p:
            return ""
        try:
            text = p.read_text(encoding="utf-8", errors="ignore")
            m = re.search(r"<style[^>]*>(.*?)</style>", text, flags=re.S | re.I)
            return m.group(1) if m else ""
        except Exception:
            return ""

# =============================================================================
# BOOTSTRAP / CIFRADO
# =============================================================================

class CryptoManager:
    _fernet = None
    _key_cache = None

    @classmethod
    def key_path(cls) -> Optional[Path]:
        df = get_data_folder()
        return None if df is None else df / F_KEY

    @classmethod
    def get_or_create_key(cls) -> Optional[bytes]:
        kp = cls.key_path()
        if kp is None:
            return None
        kp.parent.mkdir(parents=True, exist_ok=True)
        if kp.exists():
            try:
                return kp.read_bytes()
            except Exception:
                return None
        key = Fernet.generate_key()
        kp.write_bytes(key)
        return key

    @classmethod
    def fernet(cls) -> Optional[Fernet]:
        key = cls.get_or_create_key()
        if key is None:
            return None
        if key != cls._key_cache:
            cls._key_cache = key
            cls._fernet = Fernet(key)
        return cls._fernet

    @classmethod
    def encrypt_bytes(cls, data: bytes) -> Optional[bytes]:
        f = cls.fernet()
        return f.encrypt(data) if f else None

    @classmethod
    def decrypt_bytes(cls, data: bytes) -> Optional[bytes]:
        f = cls.fernet()
        if not f:
            return None
        try:
            return f.decrypt(data)
        except (InvalidToken, Exception):
            return None


def save_encrypted_json_path(path: Path, obj: Any) -> bool:
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        raw = json.dumps(obj, ensure_ascii=False, indent=2, default=str).encode("utf-8")
        enc = CryptoManager.encrypt_bytes(raw)
        if enc is None:
            return False
        tmp = path.parent / f"~{uuid.uuid4().hex[:10]}{path.suffix or '.tmp'}"
        tmp.write_bytes(enc)
        tmp.replace(path)
        return True
    except Exception as exc:
        print(f"[save_enc] {path}: {exc}")
        return False


def load_encrypted_json_path(path: Path, default: Any = None) -> Any:
    try:
        if not path.exists():
            return default
        dec = CryptoManager.decrypt_bytes(path.read_bytes())
        if dec is None:
            return default
        return json.loads(dec.decode("utf-8"))
    except Exception as exc:
        print(f"[load_enc] {path}: {exc}")
        return default


def encrypt_file_to(src: Path, dst: Path) -> bool:
    try:
        data = src.read_bytes()
        enc = CryptoManager.encrypt_bytes(data)
        if enc is None:
            return False
        dst.parent.mkdir(parents=True, exist_ok=True)
        tmp = dst.parent / f"~{uuid.uuid4().hex[:10]}{dst.suffix or '.enc'}"
        tmp.write_bytes(enc)
        tmp.replace(dst)
        return True
    except Exception as exc:
        print(f"[encrypt_file] {src} -> {dst}: {exc}")
        return False


def decrypt_file_to_temp(rel_path: str, original_name: str = "") -> Optional[Path]:
    src = rel_to_abs(rel_path)
    df = get_data_folder()
    if not src or not src.exists() or not df:
        return None
    try:
        dec = CryptoManager.decrypt_bytes(src.read_bytes())
        if dec is None:
            return None
        tmp_dir = df / SUB_TEMP
        tmp_dir.mkdir(parents=True, exist_ok=True)
        name = original_name or src.stem
        if not Path(name).suffix:
            name += ".pdf"
        out = tmp_dir / f"{uuid.uuid4().hex[:8]}_{_safe_filename(Path(name).stem)}{Path(name).suffix}"
        out.write_bytes(dec)
        return out
    except Exception:
        return None


def ensure_subfolders() -> None:
    df = get_data_folder()
    if df is None:
        return
    df.mkdir(parents=True, exist_ok=True)
    for sub in ALL_SUBFOLDERS:
        (df / sub).mkdir(parents=True, exist_ok=True)
    (df / SUB_DATOS / SUB_VEHICULOS_REGISTROS).mkdir(parents=True, exist_ok=True)


def _write_json_system(name: str, data: Any) -> bool:
    p = system_file(name)
    if not p:
        return False
    return _write_json_file(p, data)


def _read_json_system(name: str, default: Any) -> Any:
    p = system_file(name)
    if not p:
        return default
    return _read_json_file(p, default)


def bootstrap_system() -> None:
    df = get_data_folder()
    if df is None:
        return
    ensure_subfolders()
    CryptoManager.get_or_create_key()
    catalogs = {
        F_CATALOG_MARCAS: DEFAULT_MARCAS,
        F_CATALOG_SUBASTAS: DEFAULT_SUBASTAS,
        F_CATALOG_TALLERES: DEFAULT_TALLERES,
        F_CATALOG_NAVIERAS: DEFAULT_NAVIERAS,
        F_CATALOG_ADUANAS: DEFAULT_ADUANAS,
        F_CATALOG_PAISES: DEFAULT_PAISES_DESTINO,
        F_CATALOG_ESTADOS_USA: US_STATES,
        F_USUARIOS: [],
        F_DISPOSITIVOS: [],
        F_COUNTERS: {},
    }
    for name, default in catalogs.items():
        p = system_file(name)
        if p and not p.exists():
            _write_json_file(p, default)
    p_adp = system_file(F_CATALOG_ADUANAS_PAIS)
    if p_adp and not p_adp.exists():
        _write_json_file(p_adp, ADUANAS_BY_COUNTRY)

# =============================================================================
# DISPOSITIVO / USUARIOS / PERMISOS
# =============================================================================

@dataclass
class DeviceInfo:
    uuid: str
    computer_name: str
    os_name: str
    os_version: str
    machine: str
    device_key: str

    def to_dict(self) -> dict:
        return asdict(self)


def collect_device_info() -> DeviceInfo:
    osn = platform.system()
    osv = platform.release()
    computer = socket.gethostname()
    machine = platform.machine() or "UNKNOWN"
    node = str(uuid.UUID(int=uuid.getnode()))
    raw = f"{node}|{computer}|{osn}|{osv}|{machine}"
    key = hashlib.sha256(raw.encode("utf-8")).hexdigest()[:32]
    return DeviceInfo(uuid=node, computer_name=computer, os_name=osn, os_version=osv, machine=machine, device_key=key)


def hash_password(password: str, salt: Optional[str] = None) -> tuple[str, str]:
    salt = salt or secrets.token_hex(16)
    h = hashlib.sha256(f"{salt}:{password}".encode("utf-8")).hexdigest()
    return h, salt


def verify_password(password: str, stored_hash: str, salt: str) -> bool:
    h, _ = hash_password(password, salt)
    return secrets.compare_digest(h, stored_hash or "")


def default_permissions_for_role(role: str) -> dict:
    role = _norm(role)
    if role in (ROLE_ADMIN, ROLE_SUPERVISOR):
        return {k: True for k in ALL_PERMISSION_KEYS}
    if role == ROLE_OPERACIONES:
        return {
            PERM_VIEW_ALL: True, PERM_VIEW_COSTS: False, PERM_CREATE_PURCHASE: True,
            PERM_UPDATE_STAGE: True, PERM_ADD_COSTS: False, PERM_MARK_AVAILABLE: True,
            PERM_MARK_SALE: False, PERM_GENERATE_REPORTS: True, PERM_CONFIG: False,
            PERM_EDIT_CRITICAL: False, PERM_DELETE_CANCEL: False,
        }
    if role == ROLE_CONTABILIDAD:
        return {
            PERM_VIEW_ALL: True, PERM_VIEW_COSTS: True, PERM_CREATE_PURCHASE: False,
            PERM_UPDATE_STAGE: False, PERM_ADD_COSTS: True, PERM_MARK_AVAILABLE: False,
            PERM_MARK_SALE: False, PERM_GENERATE_REPORTS: True, PERM_CONFIG: False,
            PERM_EDIT_CRITICAL: False, PERM_DELETE_CANCEL: False,
        }
    if role == ROLE_VENTAS:
        return {
            PERM_VIEW_ALL: True, PERM_VIEW_COSTS: False, PERM_CREATE_PURCHASE: False,
            PERM_UPDATE_STAGE: False, PERM_ADD_COSTS: False, PERM_MARK_AVAILABLE: False,
            PERM_MARK_SALE: True, PERM_GENERATE_REPORTS: True, PERM_CONFIG: False,
            PERM_EDIT_CRITICAL: False, PERM_DELETE_CANCEL: False,
        }
    return {
        PERM_VIEW_ALL: True, PERM_VIEW_COSTS: False, PERM_CREATE_PURCHASE: False,
        PERM_UPDATE_STAGE: False, PERM_ADD_COSTS: False, PERM_MARK_AVAILABLE: False,
        PERM_MARK_SALE: False, PERM_GENERATE_REPORTS: False, PERM_CONFIG: False,
        PERM_EDIT_CRITICAL: False, PERM_DELETE_CANCEL: False,
    }


def normalize_permissions(user: dict) -> dict:
    base = default_permissions_for_role(user.get("rol", ROLE_USUARIO))
    saved = user.get("permissions") or {}
    if isinstance(saved, dict):
        for k in ALL_PERMISSION_KEYS:
            if k in saved:
                base[k] = bool(saved[k])
    if _norm(user.get("rol")) == ROLE_ADMIN:
        return {k: True for k in ALL_PERMISSION_KEYS}
    return base


def user_has_permission(user: dict, perm: str) -> bool:
    return bool(normalize_permissions(user).get(perm, False))


def load_users() -> list[dict]:
    return _read_json_system(F_USUARIOS, [])


def save_users(users: list[dict]) -> bool:
    return _write_json_system(F_USUARIOS, users)


def find_user(username: str) -> Optional[dict]:
    u = (username or "").strip().lower()
    for user in load_users():
        if (user.get("usuario") or "").strip().lower() == u:
            return user
    return None


def upsert_user(user: dict) -> bool:
    users = load_users()
    u = (user.get("usuario") or "").strip().lower()
    for i, item in enumerate(users):
        if (item.get("usuario") or "").strip().lower() == u:
            users[i] = user
            return save_users(users)
    users.append(user)
    return save_users(users)


def create_user(usuario: str, password: str, rol: str, device: DeviceInfo, activo: bool = True, permissions: Optional[dict] = None) -> tuple[bool, str]:
    usuario = (usuario or "").strip()
    if not usuario or not password:
        return False, "Usuario y contraseña son obligatorios."
    rol = _norm(rol)
    if rol not in ROLES:
        return False, "Rol inválido."
    existing = find_user(usuario)
    h, salt = hash_password(password)
    user = {
        "usuario": usuario,
        "password_hash": h,
        "salt": salt,
        "rol": rol,
        "activo": bool(activo),
        "device_key": device.device_key if rol != ROLE_ADMIN else "",
        "device": device.to_dict(),
        "permissions": permissions if isinstance(permissions, dict) else default_permissions_for_role(rol),
        "fecha_creacion": existing.get("fecha_creacion") if existing else _now_iso(),
        "fecha_actualizacion": _now_iso(),
    }
    ok = upsert_user(user)
    if ok:
        register_device(usuario, device)
    return ok, "Usuario guardado." if ok else "No se pudo guardar usuario."


def authenticate(usuario: str, password: str, device: DeviceInfo) -> tuple[bool, str, Optional[dict]]:
    user = find_user(usuario)
    if not user:
        return False, "Usuario no encontrado.", None
    if not user.get("activo", True):
        return False, "Usuario desactivado.", None
    if not verify_password(password, user.get("password_hash", ""), user.get("salt", "")):
        return False, "Contraseña incorrecta.", None
    if _norm(user.get("rol")) != ROLE_ADMIN:
        dk = user.get("device_key")
        if dk and dk != device.device_key:
            return False, "Este usuario no está autorizado en esta laptop.", None
    user["permissions"] = normalize_permissions(user)
    return True, "OK", user


def register_device(usuario: str, device: DeviceInfo) -> None:
    devs = _read_json_system(F_DISPOSITIVOS, [])
    found = False
    for d in devs:
        if d.get("device_key") == device.device_key and d.get("usuario") == usuario:
            d.update({"estado": "ACTIVO", "ultima_actualizacion": _now_iso()})
            found = True
            break
    if not found:
        devs.append({"usuario": usuario, "estado": "ACTIVO", "fecha_creacion": _now_iso(), **device.to_dict()})
    _write_json_system(F_DISPOSITIVOS, devs)

# =============================================================================
# CATÁLOGOS
# =============================================================================

def load_catalog(filename: str, default: list[str]) -> list[str]:
    data = _read_json_system(filename, default)
    vals = sorted({_norm(x) for x in data if _norm(x)})
    if not vals:
        vals = default.copy()
    if data != vals:
        _write_json_system(filename, vals)
    return vals


def save_catalog(filename: str, values: list[str]) -> bool:
    clean = sorted({_norm(v) for v in values if _norm(v)})
    return _write_json_system(filename, clean)


def add_catalog_value(filename: str, value: str, default: list[str]) -> bool:
    v = _norm(value)
    if not v:
        return False
    vals = load_catalog(filename, default)
    if v not in vals:
        vals.append(v)
    return save_catalog(filename, vals)

def remove_catalog_value(filename: str, value: str, default: list[str]) -> bool:
    vals = load_catalog(filename, default)
    target = _norm(value)
    vals = [v for v in vals if _norm(v) != target]
    return save_catalog(filename, vals)

def load_aduanas_por_pais() -> dict[str, list[str]]:
    data = _read_json_system(F_CATALOG_ADUANAS_PAIS, ADUANAS_BY_COUNTRY)
    if not isinstance(data, dict):
        data = ADUANAS_BY_COUNTRY.copy()
    clean: dict[str, list[str]] = {}
    for pais, vals in data.items():
        p = _norm(pais)
        if not p:
            continue
        if isinstance(vals, list):
            clean[p] = sorted({_norm(v) for v in vals if _norm(v)}) or ["OTRA"]
    for pais, vals in ADUANAS_BY_COUNTRY.items():
        clean.setdefault(_norm(pais), vals.copy())
    _write_json_system(F_CATALOG_ADUANAS_PAIS, clean)
    return clean

def save_aduanas_por_pais(data: dict[str, list[str]]) -> bool:
    clean: dict[str, list[str]] = {}
    for pais, vals in (data or {}).items():
        p = _norm(pais)
        if p:
            clean[p] = sorted({_norm(v) for v in vals if _norm(v)}) or ["OTRA"]
    return _write_json_system(F_CATALOG_ADUANAS_PAIS, clean)

def aduanas_for_country(country: str) -> list[str]:
    data = load_aduanas_por_pais()
    vals = data.get(_norm(country), [])
    return vals.copy() if vals else ["OTRA"]

def add_aduana_for_country(country: str, aduana: str) -> bool:
    data = load_aduanas_por_pais()
    p = _norm(country); a = _norm(aduana)
    if not p or not a:
        return False
    vals = data.setdefault(p, [])
    if a not in vals:
        vals.append(a)
    return save_aduanas_por_pais(data)

def remove_aduana_for_country(country: str, aduana: str) -> bool:
    data = load_aduanas_por_pais()
    p = _norm(country); a = _norm(aduana)
    data[p] = [x for x in data.get(p, []) if _norm(x) != a]
    return save_aduanas_por_pais(data)

def load_tariffs() -> list[dict]:
    raw = _read_json_system(F_CATALOG_TARIFAS, DEFAULT_TARIFAS)
    result: list[dict] = []
    if isinstance(raw, list):
        for item in raw:
            if isinstance(item, dict):
                name = _norm(item.get("nombre") or item.get("name"))
                try: val = float(item.get("valor_usd") or item.get("valor") or item.get("monto") or 0)
                except Exception: val = 0.0
                tipo = _norm(item.get("tipo") or "USD") or "USD"
            else:
                txt = str(item or "").strip()
                m = re.search(r"(-?\d+(?:[.,]\d+)?)", txt)
                val = float(m.group(1).replace(',', '.')) if m else 0.0
                name = _norm(re.sub(r"-?\d+(?:[.,]\d+)?", "", txt).strip()) or _norm(txt)
                tipo = "USD"
            if name:
                result.append({"nombre": name, "valor_usd": max(0.0, round(val, 2)), "tipo": tipo})
    if not result:
        result = [{"nombre":"ESTANDAR", "valor_usd":1.00, "tipo":"USD"}, {"nombre":"PREFERENCIAL", "valor_usd":0.80, "tipo":"USD"}, {"nombre":"PREMIUM", "valor_usd":1.20, "tipo":"USD"}]
    save_tariffs(result)
    return result

def save_tariffs(items: list[dict]) -> bool:
    clean=[]
    for item in items or []:
        name=_norm(item.get("nombre") if isinstance(item, dict) else "")
        if not name: continue
        try: val=max(0.0, round(float(item.get("valor_usd") or 0), 2))
        except Exception: val=0.0
        clean.append({"nombre":name, "valor_usd":val, "tipo":_norm(item.get("tipo") or "USD") or "USD"})
    return _write_json_system(F_CATALOG_TARIFAS, clean)

def copy_report_to(path: Path, target: Path) -> tuple[bool, str]:
    try:
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(path, target)
        return True, "Copia guardada correctamente."
    except Exception as exc:
        return False, f"No se pudo guardar la copia: {exc}"

# =============================================================================
# VEHÍCULOS / REGISTROS / DOCUMENTOS
# =============================================================================

def vehicle_records_root() -> Optional[Path]:
    df = get_data_folder()
    return None if df is None else df / SUB_DATOS / SUB_VEHICULOS_REGISTROS


def vehicle_record_path(vehicle: dict) -> Optional[Path]:
    root = vehicle_records_root()
    if root is None:
        return None
    code = vehicle.get("codigo") or vehicle.get("id")
    year = str(vehicle.get("anio_compra") or (_parse_date(vehicle.get("fecha_compra")) or date.today()).year)
    return root / year / f"{_safe_filename(code)}.enc"


def load_vehicles() -> list[dict]:
    root = vehicle_records_root()
    records = []
    if not root or not root.exists():
        return []
    for p in root.rglob("*.enc"):
        obj = load_encrypted_json_path(p, default=None)
        if isinstance(obj, dict):
            veh = obj.get("vehicle") if isinstance(obj.get("vehicle"), dict) else obj
            if veh.get("id") and not veh.get("eliminado"):
                ensure_vehicle_runtime_fields(veh)
                records.append(veh)
    records.sort(key=lambda x: x.get("codigo", ""))
    return records


def save_vehicle(vehicle: dict) -> bool:
    ensure_vehicle_runtime_fields(vehicle)
    p = vehicle_record_path(vehicle)
    if p is None:
        return False
    payload = {
        "schema": "lym_vehicle_record_v1",
        "updated_at": _now_iso(),
        "vehicle_id": vehicle.get("id"),
        "codigo": vehicle.get("codigo"),
        "vehicle": vehicle,
    }
    return save_encrypted_json_path(p, payload)


def find_vehicle(vehicle_id_or_code: str) -> Optional[dict]:
    key = str(vehicle_id_or_code or "").strip().upper()
    for v in load_vehicles():
        if str(v.get("id", "")).upper() == key or str(v.get("codigo", "")).upper() == key:
            return v
    return None


def next_vehicle_code(fecha_compra: Optional[date] = None) -> str:
    y = (fecha_compra or date.today()).year
    p = system_file(F_COUNTERS)
    counters = _read_json_file(p, {}) if p else {}
    current = int(counters.get(str(y), 0)) + 1
    counters[str(y)] = current
    if p:
        _write_json_file(p, counters)
    return f"LYM-CV-{y}-{current:04d}"


def store_document(src_path: Path, vehicle_code: str, doc_type: str, original_hint: str = "") -> tuple[str, str]:
    df = get_data_folder()
    if df is None:
        return "", ""
    src = Path(src_path)
    if not src.exists():
        return "", ""
    ext = src.suffix or ".pdf"
    year = str(date.today().year)
    name_base = _safe_filename(f"{vehicle_code}_{doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    dst = df / SUB_DOCUMENTOS / year / _safe_filename(vehicle_code) / f"{name_base}{ext}.enc"
    if not encrypt_file_to(src, dst):
        return "", ""
    return abs_to_rel(dst), (original_hint or src.name)


def default_stage_record(stage_key: str) -> dict:
    meta = STAGE_META.get(stage_key, {"label": stage_key})
    return {
        "key": stage_key,
        "label": meta.get("label", stage_key),
        "status": "PENDIENTE",
        "fecha_inicio": None,
        "fecha_fin": None,
        "costo_usd": 0.0,
        "proveedor": "",
        "documento": "",
        "documento_nombre": "",
        "comentario": "",
        "usuario": "",
        "fecha_actualizacion": None,
    }


def ensure_vehicle_runtime_fields(vehicle: dict) -> dict:
    vehicle.setdefault("id", uuid.uuid4().hex)
    vehicle.setdefault("codigo", "")
    vehicle.setdefault("historial", [])
    vehicle.setdefault("gastos_extra", [])
    vehicle.setdefault("fotos", [])
    vehicle.setdefault("precio_venta_usd", 0.0)
    vehicle.setdefault("precio_minimo_usd", 0.0)
    vehicle.setdefault("cliente", "")
    vehicle.setdefault("estado_actual", STAGE_COMPRADO)
    etapas = vehicle.get("etapas")
    if not isinstance(etapas, dict):
        etapas = {}
    for s in STAGES:
        etapas.setdefault(s["key"], default_stage_record(s["key"]))
    vehicle["etapas"] = etapas
    return vehicle


def stage_index(stage_key: str) -> int:
    try:
        return STAGE_ORDER.index(stage_key)
    except ValueError:
        return 0


def next_stage_key(stage_key: str) -> Optional[str]:
    idx = stage_index(stage_key)
    if idx < len(STAGE_ORDER) - 1:
        return STAGE_ORDER[idx + 1]
    return None


def user_can_override_flow(user: dict) -> bool:
    return _norm(user.get("rol")) in (ROLE_ADMIN, ROLE_SUPERVISOR) or user_has_permission(user, PERM_EDIT_CRITICAL)


def min_sale_price(vehicle: dict) -> float:
    return max(vehicle_total_cost(vehicle), float(vehicle.get("precio_minimo_usd") or 0))


def vehicle_stage(vehicle: dict, stage_key: str) -> dict:
    ensure_vehicle_runtime_fields(vehicle)
    return vehicle["etapas"].setdefault(stage_key, default_stage_record(stage_key))


def vehicle_total_cost(vehicle: dict) -> float:
    ensure_vehicle_runtime_fields(vehicle)
    total = 0.0
    for st in vehicle.get("etapas", {}).values():
        try:
            total += float(st.get("costo_usd") or 0)
        except Exception:
            pass
    for g in vehicle.get("gastos_extra", []) or []:
        try:
            total += float(g.get("monto_usd") or 0)
        except Exception:
            pass
    return round(total, 2)


def vehicle_expected_profit(vehicle: dict) -> float:
    try:
        return round(float(vehicle.get("precio_venta_usd") or 0) - vehicle_total_cost(vehicle), 2)
    except Exception:
        return 0.0


def stage_duration_days(stage: dict) -> int:
    ini = _parse_date(stage.get("fecha_inicio"))
    if not ini:
        return 0
    fin = _parse_date(stage.get("fecha_fin")) or date.today()
    return max(0, (fin - ini).days)


def vehicle_days_from_purchase(vehicle: dict) -> int:
    fc = _parse_date(vehicle.get("fecha_compra"))
    if not fc:
        return 0
    fv = _parse_date(vehicle_stage(vehicle, STAGE_VENDIDO).get("fecha_fin"))
    return max(0, ((fv or date.today()) - fc).days)


def current_stage_days(vehicle: dict) -> int:
    st = vehicle_stage(vehicle, vehicle.get("estado_actual", STAGE_COMPRADO))
    return stage_duration_days(st)


def stage_alert_level(vehicle: dict, stage_key: Optional[str] = None) -> str:
    key = stage_key or vehicle.get("estado_actual", STAGE_COMPRADO)
    st = vehicle_stage(vehicle, key)
    if st.get("status") == "COMPLETADO" or key in (STAGE_COMPRADO, STAGE_VENDIDO, STAGE_RESERVADO):
        return "OK"
    days = stage_duration_days(st)
    limit = ALERT_LIMITS_DAYS.get(key)
    if not limit:
        return "OK"
    if days >= limit * 1.5:
        return "ROJO"
    if days >= limit:
        return "AMARILLO"
    return "OK"


def validate_purchase_data(data: dict) -> tuple[bool, str]:
    required = ["marca", "modelo", "anio", "millaje", "estado_usa", "subasta", "lote", "precio_ganado_usd", "fecha_compra"]
    for k in required:
        if data.get(k) in (None, ""):
            return False, f"Campo obligatorio faltante: {k}."
    try:
        anio = int(data.get("anio"))
        if anio < 1980 or anio > date.today().year + 1:
            return False, "Año del vehículo inválido."
    except Exception:
        return False, "Año del vehículo inválido."
    try:
        if int(data.get("millaje") or 0) < 0:
            return False, "El millaje no puede ser negativo."
    except Exception:
        return False, "Millaje inválido."
    try:
        if float(data.get("precio_ganado_usd") or 0) <= 0:
            return False, "El precio ganado debe ser mayor a cero."
    except Exception:
        return False, "Precio ganado inválido."
    fc = _parse_date(data.get("fecha_compra"))
    if not fc:
        return False, "Fecha de compra inválida."
    if fc > date.today():
        return False, "La fecha de compra no puede ser futura."
    return True, "OK"


def duplicate_lot_exists(subasta: str, lote: str, exclude_id: str = "") -> Optional[dict]:
    sub = _norm(subasta)
    lot = _norm(lote)
    for v in load_vehicles():
        if v.get("id") == exclude_id or v.get("estado_actual") == STAGE_ANULADO:
            continue
        if _norm(v.get("subasta")) == sub and _norm(v.get("lote")) == lot and lot:
            return v
    return None


def create_vehicle_purchase(data: dict, comprobante_src: Path, user: dict, device: DeviceInfo) -> tuple[bool, str, str]:
    ok, msg = validate_purchase_data(data)
    if not ok:
        return False, msg, ""
    if not user_has_permission(user, PERM_CREATE_PURCHASE):
        return False, "No tienes permiso para registrar compras.", ""
    if not comprobante_src or not Path(comprobante_src).exists():
        return False, "Debes subir el comprobante de compra.", ""
    dup = duplicate_lot_exists(data.get("subasta", ""), data.get("lote", ""))
    if dup:
        return False, f"Ya existe una compra con subasta/lote: {dup.get('codigo')}.", ""

    fc = _parse_date(data.get("fecha_compra")) or date.today()
    codigo = next_vehicle_code(fc)
    doc_rel, doc_name = store_document(Path(comprobante_src), codigo, "COMPROBANTE_COMPRA")
    if not doc_rel:
        return False, "No se pudo cifrar/guardar el comprobante de compra.", ""

    now = _now_iso()
    vehicle = {
        "id": uuid.uuid4().hex,
        "codigo": codigo,
        "anio_compra": fc.year,
        "marca": _norm(data.get("marca")),
        "modelo": _norm(data.get("modelo")),
        "anio": int(data.get("anio")),
        "millaje": int(data.get("millaje") or 0),
        "color": _norm(data.get("color", "")),
        "tipo": _norm(data.get("tipo", "")),
        "estado_usa": _norm(data.get("estado_usa")),
        "subasta": _norm(data.get("subasta")),
        "lote": str(data.get("lote") or "").strip().upper(),
        "precio_ganado_usd": round(float(data.get("precio_ganado_usd") or 0), 2),
        "fecha_compra": fc.isoformat(),
        "usuario_registro": user.get("usuario", ""),
        "computadora_registro": device.computer_name,
        "fecha_registro": now,
        "fecha_actualizacion": now,
        "estado_actual": STAGE_COMPRADO,
        "precio_venta_usd": 0.0,
        "precio_minimo_usd": 0.0,
        "observaciones": str(data.get("observaciones") or "").strip(),
        "etapas": {},
        "historial": [],
        "gastos_extra": [],
        "fotos": [],
    }
    ensure_vehicle_runtime_fields(vehicle)
    comp = vehicle_stage(vehicle, STAGE_COMPRADO)
    comp.update({
        "status": "COMPLETADO",
        "fecha_inicio": fc.isoformat(),
        "fecha_fin": fc.isoformat(),
        "costo_usd": vehicle["precio_ganado_usd"],
        "proveedor": vehicle["subasta"],
        "documento": doc_rel,
        "documento_nombre": doc_name,
        "comentario": "Compra registrada con comprobante.",
        "usuario": user.get("usuario", ""),
        "fecha_actualizacion": now,
    })
    vehicle["historial"].append({
        "fecha": now,
        "usuario": user.get("usuario", ""),
        "computadora": device.computer_name,
        "accion": "CREAR_COMPRA",
        "detalle": f"Compra {codigo} · {vehicle['marca']} {vehicle['modelo']} {vehicle['anio']} · lote {vehicle['lote']} · {_fmt_usd(vehicle['precio_ganado_usd'])}",
    })
    if save_vehicle(vehicle):
        log_audit("CREAR_COMPRA", user.get("usuario", ""), codigo, f"{vehicle['marca']} {vehicle['modelo']} {vehicle['anio']}")
        return True, f"Compra creada correctamente: {codigo}", vehicle["id"]
    return False, "No se pudo guardar el registro de compra.", ""


def validate_stage_dates(vehicle: dict, stage_key: str, fecha_inicio: date, fecha_fin: Optional[date], user: dict) -> tuple[bool, str]:
    fc = _parse_date(vehicle.get("fecha_compra")) or date.today()
    if fecha_inicio > date.today():
        return False, "No se permiten fechas futuras."
    if fecha_fin and fecha_fin > date.today():
        return False, "No se permiten fechas futuras."
    if fecha_inicio < fc:
        return False, "La fecha de la etapa no puede ser anterior a la fecha de compra."
    if fecha_fin and fecha_fin < fecha_inicio:
        return False, "La fecha final no puede ser anterior a la fecha de inicio."
    idx = stage_index(stage_key)
    if idx > 0:
        prev_key = STAGE_ORDER[idx - 1]
        prev = vehicle_stage(vehicle, prev_key)
        prev_start = _parse_date(prev.get("fecha_inicio"))
        if prev_start and fecha_inicio < prev_start:
            return False, f"La etapa {STAGE_META[stage_key]['label']} no puede iniciar antes de {STAGE_META[prev_key]['label']} ({_fmt_date(prev_start)})."
    return True, "OK"


def update_vehicle_stage(vehicle_id: str, stage_key: str, data: dict, document_src: Optional[Path], user: dict, device: DeviceInfo) -> tuple[bool, str]:
    if not user_has_permission(user, PERM_UPDATE_STAGE) and stage_key not in (STAGE_RESERVADO, STAGE_VENDIDO):
        return False, "No tienes permiso para actualizar etapas."
    if stage_key == STAGE_DISPONIBLE and not user_has_permission(user, PERM_MARK_AVAILABLE):
        return False, "No tienes permiso para marcar disponible para venta."
    if stage_key in (STAGE_RESERVADO, STAGE_VENDIDO) and not user_has_permission(user, PERM_MARK_SALE):
        return False, "No tienes permiso para reservar o vender."

    vehicle = find_vehicle(vehicle_id)
    if not vehicle:
        return False, "Vehículo no encontrado."
    if stage_key not in STAGE_ORDER:
        return False, "Etapa inválida."

    ini = _parse_date(data.get("fecha_inicio"))
    if not ini:
        return False, "Fecha de inicio inválida."
    fin = _parse_date(data.get("fecha_fin")) if data.get("fecha_fin") else None
    ok, msg = validate_stage_dates(vehicle, stage_key, ini, fin, user)
    if not ok:
        return False, msg

    try:
        costo = round(float(data.get("costo_usd") or 0), 2)
        if costo < 0:
            return False, "El costo no puede ser negativo."
    except Exception:
        return False, "Costo inválido."

    doc_rel = ""
    doc_name = ""
    if document_src and Path(document_src).exists():
        doc_rel, doc_name = store_document(Path(document_src), vehicle.get("codigo", "VEHICULO"), stage_key)
        if not doc_rel:
            return False, "No se pudo guardar el documento/evidencia."

    current_key = vehicle.get("estado_actual", STAGE_COMPRADO)
    current_idx = stage_index(current_key)
    target_idx = stage_index(stage_key)

    # Si se avanza de etapa, cerrar la etapa anterior si aún seguía abierta.
    if target_idx > current_idx:
        prev = vehicle_stage(vehicle, current_key)
        if prev.get("fecha_inicio") and not prev.get("fecha_fin") and current_key != stage_key:
            prev["fecha_fin"] = ini.isoformat()
            prev["status"] = "COMPLETADO"
            prev["fecha_actualizacion"] = _now_iso()

    st = vehicle_stage(vehicle, stage_key)
    prev_status = st.get("status")
    prospective_total = round(vehicle_total_cost(vehicle) - float(st.get("costo_usd") or 0) + costo, 2)
    if stage_key == STAGE_DISPONIBLE:
        try:
            pv = round(float(data.get("precio_venta_usd") or vehicle.get("precio_venta_usd") or 0), 2)
            pm = round(float(data.get("precio_minimo_usd") or vehicle.get("precio_minimo_usd") or 0), 2)
        except Exception:
            pv, pm = 0.0, 0.0
        if not user_can_override_flow(user):
            if pv and pv < prospective_total:
                return False, f"El precio de venta no puede ser menor al gasto total ({_fmt_usd(prospective_total)})."
            if pm and pm < prospective_total:
                return False, f"El precio mínimo no puede ser menor al gasto total ({_fmt_usd(prospective_total)})."
    st.update({
        "status": "COMPLETADO" if fin else "EN PROCESO",
        "fecha_inicio": ini.isoformat(),
        "fecha_fin": fin.isoformat() if fin else None,
        "costo_usd": costo,
        "proveedor": str(data.get("proveedor") or "").strip(),
        "comentario": str(data.get("comentario") or "").strip(),
        "usuario": user.get("usuario", ""),
        "fecha_actualizacion": _now_iso(),
    })
    if doc_rel:
        st["documento"] = doc_rel
        st["documento_nombre"] = doc_name

    if stage_key == STAGE_DISPONIBLE:
        try:
            vehicle["precio_venta_usd"] = round(float(data.get("precio_venta_usd") or vehicle.get("precio_venta_usd") or 0), 2)
            vehicle["precio_minimo_usd"] = max(round(float(data.get("precio_minimo_usd") or vehicle.get("precio_minimo_usd") or 0), 2), prospective_total if not user_can_override_flow(user) else 0)
        except Exception:
            pass
    if stage_key == STAGE_RESERVADO:
        try:
            precio_reserva = round(float(data.get("precio_reserva_usd") or data.get("precio_venta_usd") or vehicle.get("precio_venta_usd") or 0), 2)
        except Exception:
            precio_reserva = 0.0
        if precio_reserva and precio_reserva < min_sale_price(vehicle) and not user_can_override_flow(user):
            return False, f"No puedes reservar por debajo del costo/precio mínimo ({_fmt_usd(min_sale_price(vehicle))})."
        vehicle["precio_reserva_usd"] = precio_reserva
        vehicle["cliente"] = str(data.get("cliente") or vehicle.get("cliente") or "").strip()
        vehicle["monto_reserva_usd"] = round(float(data.get("monto_reserva_usd") or vehicle.get("monto_reserva_usd") or 0), 2)
    if stage_key == STAGE_VENDIDO:
        precio_real = round(float(data.get("precio_venta_real_usd") or data.get("precio_venta_usd") or vehicle.get("precio_venta_usd") or 0), 2)
        if precio_real < min_sale_price(vehicle) and not user_can_override_flow(user):
            return False, f"No puedes vender por debajo del costo/precio mínimo ({_fmt_usd(min_sale_price(vehicle))})."
        vehicle["precio_venta_real_usd"] = precio_real
        vehicle["cliente"] = str(data.get("cliente") or vehicle.get("cliente") or "").strip()

    # Si se está editando una etapa anterior, no retroceder el estado actual salvo permiso crítico.
    if target_idx >= current_idx or user_has_permission(user, PERM_EDIT_CRITICAL):
        vehicle["estado_actual"] = stage_key

    vehicle["fecha_actualizacion"] = _now_iso()
    vehicle.setdefault("historial", []).append({
        "fecha": _now_iso(),
        "usuario": user.get("usuario", ""),
        "computadora": device.computer_name,
        "accion": "ACTUALIZAR_ETAPA",
        "detalle": f"{STAGE_META[stage_key]['label']} · inicio={_fmt_date(ini)} · fin={_fmt_date(fin)} · costo={_fmt_usd(costo)} · estado_ant={prev_status}",
    })
    if save_vehicle(vehicle):
        log_audit("ACTUALIZAR_ETAPA", user.get("usuario", ""), vehicle.get("codigo", ""), STAGE_META[stage_key]["label"])
        return True, "Etapa actualizada correctamente."
    return False, "No se pudo guardar la etapa."


def add_extra_cost(vehicle_id: str, data: dict, document_src: Optional[Path], user: dict, device: DeviceInfo) -> tuple[bool, str]:
    if not user_has_permission(user, PERM_ADD_COSTS):
        return False, "No tienes permiso para agregar costos."
    vehicle = find_vehicle(vehicle_id)
    if not vehicle:
        return False, "Vehículo no encontrado."
    try:
        monto = round(float(data.get("monto_usd") or 0), 2)
    except Exception:
        return False, "Monto inválido."
    if monto <= 0:
        return False, "El monto debe ser mayor a cero."
    fecha = _parse_date(data.get("fecha")) or date.today()
    doc_rel, doc_name = "", ""
    if document_src and Path(document_src).exists():
        doc_rel, doc_name = store_document(Path(document_src), vehicle.get("codigo", "VEHICULO"), "GASTO_EXTRA")
    item = {
        "id": uuid.uuid4().hex,
        "fecha": fecha.isoformat(),
        "categoria": _norm(data.get("categoria") or "OTROS"),
        "descripcion": str(data.get("descripcion") or "").strip(),
        "monto_usd": monto,
        "proveedor": str(data.get("proveedor") or "").strip(),
        "documento": doc_rel,
        "documento_nombre": doc_name,
        "usuario": user.get("usuario", ""),
        "fecha_registro": _now_iso(),
    }
    vehicle.setdefault("gastos_extra", []).append(item)
    vehicle["historial"].append({"fecha": _now_iso(), "usuario": user.get("usuario", ""), "computadora": device.computer_name, "accion": "AGREGAR_GASTO", "detalle": f"{item['categoria']} · {_fmt_usd(monto)}"})
    vehicle["fecha_actualizacion"] = _now_iso()
    if save_vehicle(vehicle):
        return True, "Gasto agregado."
    return False, "No se pudo guardar el gasto."


def log_audit(accion: str, usuario: str, codigo: str = "", detalle: str = "") -> None:
    p = datos_file(F_AUDITORIA)
    if not p:
        return
    eventos = load_encrypted_json_path(p, default=[])
    if not isinstance(eventos, list):
        eventos = []
    eventos.append({
        "fecha": _now_iso(),
        "usuario": usuario,
        "computadora": socket.gethostname(),
        "accion": accion,
        "codigo": codigo,
        "detalle": detalle,
    })
    eventos = eventos[-5000:]
    save_encrypted_json_path(p, eventos)


def load_audit_events(limit: int = 1000) -> list[dict]:
    p = datos_file(F_AUDITORIA)
    if not p or not p.exists():
        return []
    eventos = load_encrypted_json_path(p, default=[])
    if not isinstance(eventos, list):
        return []
    eventos = list(reversed(eventos))
    return eventos[:max(1, limit)]


def load_registered_devices() -> list[dict]:
    data = _read_json_system(F_DISPOSITIVOS, [])
    return data if isinstance(data, list) else []


def update_user_permissions(username: str, permissions: dict) -> tuple[bool, str]:
    user = find_user(username)
    if not user:
        return False, "Usuario no encontrado."
    user["permissions"] = {k: bool(permissions.get(k, False)) for k in ALL_PERMISSION_KEYS}
    user["fecha_actualizacion"] = _now_iso()
    ok = upsert_user(user)
    if ok:
        log_audit("PERMISOS_USUARIO", username, "", json.dumps(user["permissions"], ensure_ascii=False))
    return ok, "Permisos actualizados." if ok else "No se pudieron actualizar permisos."


def set_user_active(username: str, active: bool) -> tuple[bool, str]:
    user = find_user(username)
    if not user:
        return False, "Usuario no encontrado."
    user["activo"] = bool(active)
    user["fecha_actualizacion"] = _now_iso()
    ok = upsert_user(user)
    if ok:
        log_audit("USUARIO_ACTIVO", username, "", "ACTIVO" if active else "INACTIVO")
    return ok, "Usuario actualizado." if ok else "No se pudo actualizar usuario."


def create_backup_snapshot(usuario: str = "", detalle: str = "Manual") -> tuple[bool, str, Optional[Path]]:
    df = get_data_folder()
    if df is None or not df.exists():
        return False, "Carpeta del sistema no disponible.", None
    backup_dir = df / SUB_RESPALDOS
    backup_dir.mkdir(parents=True, exist_ok=True)
    out = backup_dir / f"LYM_BACKUP_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    try:
        import zipfile
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zf:
            for p in df.rglob('*'):
                if not p.is_file():
                    continue
                try:
                    rel = p.relative_to(df)
                except Exception:
                    continue
                if rel.parts and rel.parts[0] == SUB_RESPALDOS:
                    continue
                zf.write(p, rel.as_posix())
        log_audit("BACKUP", usuario, "", f"{detalle} · {out.name}")
        return True, "Backup creado correctamente.", out
    except Exception as exc:
        return False, f"No se pudo crear el backup: {exc}", None


def list_backups() -> list[Path]:
    df = get_data_folder()
    if df is None:
        return []
    backup_dir = df / SUB_RESPALDOS
    if not backup_dir.exists():
        return []
    return sorted(backup_dir.glob('*.zip'), key=lambda p: p.stat().st_mtime, reverse=True)

# =============================================================================
# KPIS / REPORTES
# =============================================================================

def compute_kpis(vehicles: list[dict]) -> dict:
    total = len(vehicles)
    vendidos = [v for v in vehicles if v.get("estado_actual") == STAGE_VENDIDO]
    activos = [v for v in vehicles if v.get("estado_actual") not in (STAGE_VENDIDO, STAGE_ANULADO)]
    disponibles = [v for v in vehicles if v.get("estado_actual") == STAGE_DISPONIBLE]
    reservados = [v for v in vehicles if v.get("estado_actual") == STAGE_RESERVADO]
    criticos = [v for v in activos if stage_alert_level(v) == "ROJO" or (v.get("estado_actual") == STAGE_DISPONIBLE and current_stage_days(v) >= 90)]
    capital = sum(vehicle_total_cost(v) for v in activos)
    precio_pub = sum(float(v.get("precio_venta_usd") or 0) for v in disponibles + reservados)
    ganancia_esp = sum(vehicle_expected_profit(v) for v in disponibles + reservados)
    prom_compra_disp = _avg([
        ( _parse_date(vehicle_stage(v, STAGE_DISPONIBLE).get("fecha_inicio")) - _parse_date(v.get("fecha_compra")) ).days
        for v in vehicles
        if _parse_date(vehicle_stage(v, STAGE_DISPONIBLE).get("fecha_inicio")) and _parse_date(v.get("fecha_compra"))
    ])
    prom_disp_venta = _avg([
        ( _parse_date(vehicle_stage(v, STAGE_VENDIDO).get("fecha_inicio")) - _parse_date(vehicle_stage(v, STAGE_DISPONIBLE).get("fecha_inicio")) ).days
        for v in vendidos
        if _parse_date(vehicle_stage(v, STAGE_VENDIDO).get("fecha_inicio")) and _parse_date(vehicle_stage(v, STAGE_DISPONIBLE).get("fecha_inicio"))
    ])
    by_stage = {s["key"]: 0 for s in STAGES}
    for v in vehicles:
        if v.get("estado_actual") in by_stage:
            by_stage[v.get("estado_actual")] += 1
    by_brand: dict[str, int] = {}
    for v in vehicles:
        by_brand[_norm(v.get("marca")) or "SIN MARCA"] = by_brand.get(_norm(v.get("marca")) or "SIN MARCA", 0) + 1
    return {
        "total": total, "activos": len(activos), "vendidos": len(vendidos),
        "disponibles": len(disponibles), "reservados": len(reservados), "criticos": len(criticos),
        "capital": round(capital, 2), "precio_publicado": round(precio_pub, 2), "ganancia_esperada": round(ganancia_esp, 2),
        "prom_compra_disp": round(prom_compra_disp, 1), "prom_disp_venta": round(prom_disp_venta, 1),
        "by_stage": by_stage, "by_brand": by_brand,
    }


def _avg(values: list[int | float]) -> float:
    vals = [float(v) for v in values if v is not None]
    return sum(vals) / len(vals) if vals else 0.0


def _html_attr(s: Any) -> str:
    return html.escape(str(s or ""), quote=True)


def _stage_badge(stage_key: str) -> str:
    meta = STAGE_META.get(stage_key, {"label": stage_key, "color": "#64748B"})
    return f'<span class="badge" style="background:{meta["color"]}22;color:{meta["color"]};border:1px solid {meta["color"]}55">{html.escape(meta["label"])}</span>'


def generate_html_report(vehicles: list[dict], user: Optional[dict] = None) -> Optional[Path]:
    df = get_data_folder()
    if df is None:
        return None
    out_dir = df / SUB_REPORTES
    out_dir.mkdir(parents=True, exist_ok=True)
    kpis = compute_kpis(vehicles)
    logo_uri = ResourceManager.logo_data_uri()
    template_css = ResourceManager.template_css()
    generated = datetime.now().strftime("%d/%m/%Y %H:%M")

    stage_rows = []
    max_stage = max(kpis["by_stage"].values() or [1]) or 1
    for s in STAGES:
        n = kpis["by_stage"].get(s["key"], 0)
        pct = (n / max_stage) * 100 if max_stage else 0
        stage_rows.append(f"""
        <div class="stage-row">
          <div class="stage-left"><span class="dot" style="background:{s['color']}"></span><b>{html.escape(s['label'])}</b></div>
          <div class="barline"><span style="width:{pct:.1f}%;background:{s['color']}"></span></div>
          <strong>{n}</strong>
        </div>""")

    brand_rows = []
    sorted_brands = sorted(kpis["by_brand"].items(), key=lambda x: x[1], reverse=True)[:12]
    max_brand = max([x[1] for x in sorted_brands] or [1])
    for marca, n in sorted_brands:
        pct = n / max_brand * 100 if max_brand else 0
        brand_rows.append(f"""
        <div class="stage-row">
          <div class="stage-left"><span class="dot orange-dot"></span><b>{html.escape(marca)}</b></div>
          <div class="barline"><span style="width:{pct:.1f}%"></span></div>
          <strong>{n}</strong>
        </div>""")

    available_cards = []
    priority_map = {"ROJO": 0, "AMARILLO": 1, "OK": 2}
    disponibles = [v for v in vehicles if v.get("estado_actual") in (STAGE_DISPONIBLE, STAGE_RESERVADO)]
    disponibles.sort(key=lambda v: (priority_map.get(stage_alert_level(v), 2), -current_stage_days(v)))
    for v in disponibles:
        days = current_stage_days(v)
        alert = stage_alert_level(v)
        css_alert = "critico" if alert == "ROJO" or days >= 90 else "alerta" if alert == "AMARILLO" or days >= 60 else "ok"
        available_cards.append(f"""
        <article class="vehicle-card {css_alert}" data-days="{days}" data-profit="{vehicle_expected_profit(v):.2f}">
          <div class="vehicle-top">
            <span class="code">{html.escape(v.get('codigo',''))}</span>
            {_stage_badge(v.get('estado_actual',''))}
          </div>
          <h3>{html.escape(v.get('marca',''))} {html.escape(v.get('modelo',''))} {html.escape(str(v.get('anio','')))}</h3>
          <p>Millaje: <b>{int(v.get('millaje') or 0):,}</b> · Lote: <b>{html.escape(v.get('lote',''))}</b></p>
          <div class="money-grid"><span>Costo total</span><b>{_fmt_usd(vehicle_total_cost(v))}</b><span>Precio venta</span><b>{_fmt_usd(v.get('precio_venta_usd'))}</b><span>Ganancia esperada</span><b>{_fmt_usd(vehicle_expected_profit(v))}</b><span>Días disponible</span><b>{days}</b></div>
        </article>""")

    table_rows = []
    for v in sorted(vehicles, key=lambda x: (stage_index(x.get("estado_actual", "")), x.get("codigo", ""))):
        table_rows.append(f"""
        <tr data-stage="{_html_attr(v.get('estado_actual'))}" data-days="{vehicle_days_from_purchase(v)}">
          <td><b>{html.escape(v.get('codigo',''))}</b></td>
          <td>{html.escape(v.get('marca',''))}</td>
          <td>{html.escape(v.get('modelo',''))}</td>
          <td>{html.escape(str(v.get('anio','')))}</td>
          <td>{int(v.get('millaje') or 0):,}</td>
          <td>{html.escape(v.get('subasta',''))}</td>
          <td>{html.escape(v.get('lote',''))}</td>
          <td>{_fmt_usd(vehicle_total_cost(v))}</td>
          <td>{_stage_badge(v.get('estado_actual',''))}</td>
          <td>{vehicle_days_from_purchase(v)}</td>
        </tr>""")

    css = f"""
    {template_css}
    :root{{--navy:#08285a;--navy2:#0e3a78;--orange:#f59a13;--orange2:#ffb547;--bg:#f5f8fc;--card:#fff;--ink:#0b172a;--muted:#637083;--line:rgba(8,40,90,.12);--green:#10B981;--red:#DC2626;--amber:#F59E0B;}}
    body{{margin:0;font-family:Inter,Segoe UI,Arial,sans-serif;background:linear-gradient(135deg,#f6f9ff,#eef4fb 55%,#fff8ef);color:var(--ink)}}
    .report{{max-width:1240px;margin:auto;padding:42px 24px 70px}}
    .hero{{display:grid;grid-template-columns:1.25fr .55fr;gap:24px;align-items:center;min-height:460px}}
    .logo-box{{background:rgba(255,255,255,.78);border:1px solid var(--line);border-radius:34px;display:grid;place-items:center;min-height:360px;box-shadow:0 28px 80px rgba(8,40,90,.16)}}
    .logo-box img{{max-width:300px;max-height:300px;filter:drop-shadow(0 18px 35px rgba(8,40,90,.20))}}
    .eyebrow{{display:inline-flex;gap:8px;align-items:center;color:var(--orange);font-weight:900;text-transform:uppercase;letter-spacing:.16em;background:rgba(245,154,19,.10);border:1px solid rgba(245,154,19,.20);border-radius:99px;padding:8px 13px;font-size:12px}}
    h1{{font-size:clamp(42px,7vw,84px);line-height:.95;color:var(--navy);letter-spacing:-.06em;margin:18px 0}}
    h2{{font-size:38px;color:var(--navy);letter-spacing:-.04em;margin:10px 0 18px}}
    .lead{{font-size:20px;line-height:1.5;color:#334155;max-width:850px}}
    .kpis{{display:grid;grid-template-columns:repeat(4,1fr);gap:18px;margin:24px 0}}
    .card{{background:rgba(255,255,255,.86);border:1px solid var(--line);border-radius:26px;padding:20px;box-shadow:0 20px 55px rgba(8,40,90,.10)}}
    .kpi .label{{color:var(--muted);font-size:12px;text-transform:uppercase;letter-spacing:.09em;font-weight:900}}
    .kpi .value{{font-size:38px;color:var(--navy);font-weight:950;line-height:1.05;margin-top:8px}}
    .grid2{{display:grid;grid-template-columns:1fr 1fr;gap:20px;margin:24px 0}}
    .stage-row{{display:grid;grid-template-columns:190px 1fr 50px;gap:12px;align-items:center;padding:11px 0;border-bottom:1px solid rgba(8,40,90,.08)}}
    .stage-left{{display:flex;align-items:center;gap:10px;color:#1f2a3d}}
    .dot{{width:13px;height:13px;border-radius:99px;display:inline-block}}.orange-dot{{background:var(--orange)}}
    .barline{{height:12px;background:#e9eef7;border-radius:99px;overflow:hidden}}.barline span{{height:100%;display:block;background:linear-gradient(90deg,var(--orange),var(--navy));border-radius:inherit}}
    .section-head{{display:flex;justify-content:space-between;align-items:center;gap:20px;margin-top:36px}}
    .action-btn{{border:none;background:linear-gradient(135deg,var(--orange),var(--orange2));color:#08285a;font-weight:950;padding:13px 18px;border-radius:999px;cursor:pointer;box-shadow:0 16px 32px rgba(245,154,19,.25)}}
    .vehicles{{display:grid;grid-template-columns:repeat(3,minmax(260px,1fr));gap:18px;margin-top:18px}}
    .vehicle-card{{background:#fff;border:1px solid var(--line);border-radius:24px;padding:18px;box-shadow:0 18px 46px rgba(8,40,90,.10);position:relative;overflow:hidden}}
    .vehicle-card:before{{content:"";position:absolute;inset:0 0 auto;height:6px;background:var(--green)}}.vehicle-card.alerta:before{{background:var(--amber)}}.vehicle-card.critico:before{{background:var(--red)}}
    .vehicle-top{{display:flex;justify-content:space-between;align-items:center;gap:12px}}.code{{font-weight:950;color:var(--navy)}}
    .badge{{border-radius:999px;padding:5px 9px;font-size:11px;font-weight:950;white-space:nowrap}}
    .money-grid{{display:grid;grid-template-columns:1fr auto;gap:8px;margin-top:14px;color:#526071;font-size:13px}}.money-grid b{{color:var(--navy)}}
    table{{width:100%;border-collapse:separate;border-spacing:0;background:#fff;border-radius:22px;overflow:hidden;box-shadow:0 18px 46px rgba(8,40,90,.10)}}
    th{{background:var(--navy);color:#fff;text-align:left;padding:12px;font-size:12px;text-transform:uppercase;letter-spacing:.07em}}td{{padding:12px;border-bottom:1px solid rgba(8,40,90,.08);font-weight:700;color:#263346}}
    .note{{background:linear-gradient(135deg,var(--navy),var(--navy2));color:white;border-radius:28px;padding:24px;margin:28px 0;box-shadow:0 24px 70px rgba(8,40,90,.22)}}.note h2{{color:white}}.note p{{color:rgba(255,255,255,.84)}}
    @media(max-width:980px){{.hero,.grid2{{grid-template-columns:1fr}}.kpis{{grid-template-columns:repeat(2,1fr)}}.vehicles{{grid-template-columns:1fr}}}}
    """
    html_text = f"""<!DOCTYPE html>
<html lang="es"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0"><title>LYM Auto Control · Reporte Gerencial</title><style>{css}</style></head>
<body><main class="report">
<section class="hero">
  <div><span class="eyebrow">Reporte gerencial · Corte {generated}</span><h1>Inventario de vehículos <span style="color:var(--orange)">L & M</span></h1><p class="lead">Control operativo y financiero de compras vehiculares: trazabilidad por etapa, capital invertido, disponibles para venta, unidades críticas y rotación.</p></div>
  <div class="logo-box">{'<img src="'+logo_uri+'" alt="Logo LYM">' if logo_uri else '<h2>L&M</h2>'}</div>
</section>
<section class="kpis">
  <div class="card kpi"><div class="label">Vehículos totales</div><div class="value">{kpis['total']}</div></div>
  <div class="card kpi"><div class="label">Activos</div><div class="value">{kpis['activos']}</div></div>
  <div class="card kpi"><div class="label">Disponibles</div><div class="value">{kpis['disponibles']}</div></div>
  <div class="card kpi"><div class="label">Vendidos</div><div class="value">{kpis['vendidos']}</div></div>
  <div class="card kpi"><div class="label">Capital activo</div><div class="value">{_fmt_usd(kpis['capital'])}</div></div>
  <div class="card kpi"><div class="label">Ganancia esperada</div><div class="value">{_fmt_usd(kpis['ganancia_esperada'])}</div></div>
  <div class="card kpi"><div class="label">Críticos</div><div class="value">{kpis['criticos']}</div></div>
  <div class="card kpi"><div class="label">Compra → disponible</div><div class="value">{kpis['prom_compra_disp']} días</div></div>
</section>
<section class="grid2"><div class="card"><h2>Estatus operativo</h2>{''.join(stage_rows)}</div><div class="card"><h2>Marcas en inventario</h2>{''.join(brand_rows)}</div></section>
<section class="section-head"><div><h2>Vehículos listos para venta</h2><p>Prioridad automática: disponibles antiguos y críticos aparecen primero.</p></div><button class="action-btn" onclick="reorganizar()">Reorganizar inventario antiguo</button></section>
<div class="vehicles" id="vehicles">{''.join(available_cards) if available_cards else '<div class="card">No hay vehículos disponibles/reservados.</div>'}</div>
<section class="note"><h2>Lectura ejecutiva</h2><p>El reporte prioriza capital detenido, atrasos por etapa y unidades disponibles para venta. Los vehículos con más días disponibles deben revisarse para promoción, ajuste de precio o seguimiento comercial.</p></section>
<section><h2>Detalle general</h2><table><thead><tr><th>Código</th><th>Marca</th><th>Modelo</th><th>Año</th><th>Millaje</th><th>Subasta</th><th>Lote</th><th>Costo total</th><th>Estado</th><th>Días</th></tr></thead><tbody>{''.join(table_rows)}</tbody></table></section>
</main><script>
function reorganizar(){{
 const box=document.getElementById('vehicles');
 const cards=[...box.querySelectorAll('.vehicle-card')];
 cards.sort((a,b)=>(parseInt(b.dataset.days||0)-parseInt(a.dataset.days||0)) || ((parseFloat(b.dataset.profit||0)-parseFloat(a.dataset.profit||0))));
 cards.forEach(c=>box.appendChild(c));
}}
</script></body></html>"""
    out = out_dir / f"LYM_REPORTE_GERENCIAL_{datetime.now().strftime('%Y-%m-%d_%H%M')}.html"
    out.write_text(html_text, encoding="utf-8")
    log_audit("GENERAR_REPORTE_HTML", (user or {}).get("usuario", ""), "", out.name)
    return out



def generate_inventory_excel(vehicles: list[dict], user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    """Genera un reporte .xlsx de inventario. Usa openpyxl solo dentro de la app del usuario."""
    df = get_data_folder()
    if df is None:
        return False, "Carpeta del sistema no disponible.", None
    out_dir = df / SUB_REPORTES
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"LYM_INVENTARIO_{datetime.now().strftime('%Y-%m-%d_%H%M')}.xlsx"
    try:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except Exception:
            return False, "Para generar Excel instala: pip install openpyxl", None
        wb = Workbook()
        ws = wb.active
        ws.title = "Inventario LYM"
        headers = ["Código", "Lote", "Fecha compra", "Marca", "Modelo", "Año", "Millaje", "Estado USA", "Subasta", "Estado actual", "Costo total", "Precio venta", "Ganancia esperada", "Días compra", "Alerta"]
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
        ws.cell(1,1).value = "INVENTARIO VEHICULAR L&M INVERSIONES"
        ws.cell(1,1).font = Font(size=16, bold=True, color="FFFFFF")
        ws.cell(1,1).fill = PatternFill("solid", fgColor="08285A")
        ws.cell(1,1).alignment = Alignment(horizontal="center")
        ws.append(headers)
        for cell in ws[2]:
            cell.fill = PatternFill("solid", fgColor="F59A13")
            cell.font = Font(bold=True, color="08285A")
            cell.alignment = Alignment(horizontal="center")
        for v in vehicles:
            ws.append([
                v.get("codigo"), v.get("lote"), _fmt_date(v.get("fecha_compra")), v.get("marca"), v.get("modelo"), v.get("anio"),
                int(v.get("millaje") or 0), v.get("estado_usa"), v.get("subasta"), STAGE_META.get(v.get("estado_actual"),{}).get("label",v.get("estado_actual")),
                vehicle_total_cost(v), float(v.get("precio_venta_usd") or 0), vehicle_expected_profit(v), vehicle_days_from_purchase(v), stage_alert_level(v)
            ])
        for row in ws.iter_rows(min_row=3, max_row=ws.max_row, min_col=11, max_col=13):
            for cell in row:
                cell.number_format = '$#,##0.00'
        thin = Side(style="thin", color="D9E2EF")
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=len(headers)):
            for cell in row:
                cell.border = Border(bottom=thin)
                cell.alignment = Alignment(vertical="center")
        for col in range(1, len(headers)+1):
            ws.column_dimensions[get_column_letter(col)].width = min(24, max(12, len(str(headers[col-1]))+4))
        ws.freeze_panes = "A3"
        wb.save(out)
        log_audit("GENERAR_EXCEL_INVENTARIO", (user or {}).get("usuario", ""), "", out.name)
        return True, "Reporte Excel de inventario generado.", out
    except Exception as exc:
        return False, f"Error generando Excel: {exc}", None


def generate_kpi_html_report(vehicles: list[dict], user: Optional[dict] = None) -> Optional[Path]:
    df = get_data_folder()
    if df is None:
        return None
    out_dir = df / SUB_REPORTES
    out_dir.mkdir(parents=True, exist_ok=True)
    k = compute_kpis(vehicles)
    logo_uri = ResourceManager.logo_data_uri()
    rows = []
    for s in STAGES:
        vals = [stage_duration_days(vehicle_stage(v, s["key"])) for v in vehicles if vehicle_stage(v, s["key"]).get("fecha_inicio")]
        avg = _avg(vals)
        count = sum(1 for v in vehicles if v.get("estado_actual") == s["key"])
        rows.append(f"<tr><td>{html.escape(s['label'])}</td><td>{count}</td><td>{avg:.1f} días</td><td>{ALERT_LIMITS_DAYS.get(s['key'], '—')}</td></tr>")
    html_text = f"""<!doctype html><html lang='es'><head><meta charset='utf-8'><title>LYM KPI tiempos</title>
<style>:root{{--navy:#08285a;--orange:#f59a13;--bg:#f4f7fb;--text:#0b172a}}body{{font-family:Segoe UI,Arial;background:var(--bg);color:var(--text);margin:0}}main{{max-width:1180px;margin:0 auto;padding:34px}}.hero{{display:flex;justify-content:space-between;align-items:center;background:linear-gradient(135deg,#08285a,#0e3a78);color:white;border-radius:28px;padding:32px}}.logo{{background:white;border-radius:22px;padding:18px}}.logo img{{max-width:160px}}.cards{{display:grid;grid-template-columns:repeat(4,1fr);gap:16px;margin:24px 0}}.card{{background:white;border-radius:18px;padding:20px;box-shadow:0 10px 28px rgba(8,40,90,.08)}}.value{{font-size:30px;font-weight:900;color:var(--navy)}}h1,h2{{margin:0 0 8px}}table{{width:100%;border-collapse:collapse;background:white;border-radius:18px;overflow:hidden}}th{{background:var(--navy);color:white;text-align:left;padding:12px}}td{{padding:12px;border-bottom:1px solid #e2e8f0}}.orange{{color:var(--orange)}}</style></head><body><main>
<section class='hero'><div><h1>Reporte KPI de tiempos <span class='orange'>L&M</span></h1><p>Corte {datetime.now().strftime('%d/%m/%Y %H:%M')} · Control de duración por etapas y rotación.</p></div><div class='logo'>{'<img src="'+logo_uri+'">' if logo_uri else '<h2>L&M</h2>'}</div></section>
<section class='cards'><div class='card'><b>Total unidades</b><div class='value'>{k['total']}</div></div><div class='card'><b>Activas</b><div class='value'>{k['activos']}</div></div><div class='card'><b>Compra → Disponible</b><div class='value'>{k['prom_compra_disp']} días</div></div><div class='card'><b>Disponible → Venta</b><div class='value'>{k['prom_disp_venta']} días</div></div></section>
<section class='card'><h2>Promedio y control por etapa</h2><table><thead><tr><th>Etapa</th><th>Unidades actuales</th><th>Promedio histórico</th><th>Límite alerta</th></tr></thead><tbody>{''.join(rows)}</tbody></table></section>
</main></body></html>"""
    out = out_dir / f"LYM_KPI_TIEMPOS_{datetime.now().strftime('%Y-%m-%d_%H%M')}.html"
    out.write_text(html_text, encoding="utf-8")
    log_audit("GENERAR_KPI_HTML", (user or {}).get("usuario", ""), "", out.name)
    return out



def generate_kpi_excel_report(vehicles: list[dict], user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    df = get_data_folder()
    if df is None:
        return False, "Carpeta del sistema no disponible.", None
    out_dir = df / SUB_REPORTES
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"LYM_KPI_TIEMPOS_{datetime.now().strftime('%Y-%m-%d_%H%M')}.xlsx"
    try:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except Exception:
            return False, "Para generar Excel instala: pip install openpyxl", None
        wb = Workbook()
        ws = wb.active
        ws.title = "KPI Tiempos"
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)
        ws.cell(1,1).value = "LYM AUTO CONTROL - KPI DE TIEMPOS"
        ws.cell(1,1).font = Font(size=16, bold=True, color="FFFFFF")
        ws.cell(1,1).fill = PatternFill("solid", fgColor="08285A")
        ws.cell(1,1).alignment = Alignment(horizontal="center")
        headers = ["Etapa", "Unidades actuales", "Registros con fecha", "Promedio días", "Máximo días", "Límite alerta", "Observación"]
        ws.append(headers)
        for cell in ws[2]:
            cell.fill = PatternFill("solid", fgColor="F59A13")
            cell.font = Font(bold=True, color="08285A")
            cell.alignment = Alignment(horizontal="center")
        for s in STAGES:
            vals = [stage_duration_days(vehicle_stage(v, s["key"])) for v in vehicles if vehicle_stage(v, s["key"]).get("fecha_inicio")]
            avg = _avg(vals)
            maxv = max(vals or [0])
            count_current = sum(1 for v in vehicles if v.get("estado_actual") == s["key"])
            limit = ALERT_LIMITS_DAYS.get(s["key"], "—")
            obs = "Revisar" if isinstance(limit, int) and avg >= limit else "OK"
            ws.append([s["label"], count_current, len(vals), round(avg, 2), maxv, limit, obs])
        ws2 = wb.create_sheet("Detalle vehículos")
        headers2 = ["Código", "Vehículo", "Estado actual", "Compra", "Días compra", "Compra→Disponible", "Disponible→Venta", "Costo total", "Precio venta", "Ganancia esperada"]
        ws2.append(headers2)
        for c in ws2[1]:
            c.fill = PatternFill("solid", fgColor="08285A"); c.font = Font(bold=True, color="FFFFFF")
        for v in vehicles:
            disp = _parse_date(vehicle_stage(v, STAGE_DISPONIBLE).get("fecha_inicio"))
            compra = _parse_date(v.get("fecha_compra"))
            vendido = _parse_date(vehicle_stage(v, STAGE_VENDIDO).get("fecha_inicio"))
            ws2.append([
                v.get("codigo"), f"{v.get('marca')} {v.get('modelo')} {v.get('anio')}", STAGE_META.get(v.get("estado_actual"),{}).get("label",v.get("estado_actual")),
                _fmt_date(compra), vehicle_days_from_purchase(v), max(0,(disp-compra).days) if disp and compra else "", max(0,(vendido-disp).days) if vendido and disp else "",
                vehicle_total_cost(v), float(v.get("precio_venta_usd") or 0), vehicle_expected_profit(v)
            ])
        thin = Side(style="thin", color="D9E2EF")
        for wsx in [ws, ws2]:
            for row in wsx.iter_rows():
                for cell in row:
                    cell.border = Border(bottom=thin)
                    cell.alignment = Alignment(vertical="center")
            for col in range(1, wsx.max_column + 1):
                wsx.column_dimensions[get_column_letter(col)].width = 22
        for row in ws2.iter_rows(min_row=2, min_col=8, max_col=10):
            for cell in row:
                cell.number_format = '$#,##0.00'
        wb.save(out)
        log_audit("GENERAR_EXCEL_KPI", (user or {}).get("usuario", ""), "", out.name)
        return True, "Reporte Excel KPI de tiempos generado.", out
    except Exception as exc:
        return False, f"Error generando Excel KPI: {exc}", None

# =============================================================================
# GUI PySide6
# =============================================================================

if PYSIDE_OK:
    LYM_STYLESHEET = """
    QWidget { font-family: 'Segoe UI', Arial, sans-serif; font-size: 10pt; color:#0b172a; }
    QMainWindow, QDialog { background:#eff4f9; }
    QLineEdit, QComboBox, QDateEdit, QSpinBox, QDoubleSpinBox, QTextEdit {
        background:#ffffff; border:1px solid rgba(8,40,90,.16); border-radius:12px; padding:8px 10px; min-height:26px;
    }
    QLineEdit:focus, QComboBox:focus, QDateEdit:focus, QSpinBox:focus, QDoubleSpinBox:focus, QTextEdit:focus { border:2px solid #f59a13; }
    QPushButton { background:#08285a; color:#fff; border:none; border-radius:12px; padding:10px 14px; font-weight:800; }
    QPushButton:hover { background:#0e3a78; }
    QPushButton#orange { background:#f59a13; color:#08285a; }
    QPushButton#orange:hover { background:#ffb547; }
    QPushButton#ghost { background:#fff; color:#08285a; border:1px solid rgba(8,40,90,.16); }
    QPushButton#danger { background:#dc2626; color:#fff; }
    QTableWidget { background:#fff; border:1px solid rgba(8,40,90,.12); border-radius:14px; gridline-color:rgba(8,40,90,.08); selection-background-color:#fff3df; selection-color:#08285a; }
    QHeaderView::section { background:#08285a; color:#fff; padding:8px; border:none; font-weight:900; }
    QTabWidget::pane { border:1px solid rgba(8,40,90,.12); background:#fff; border-radius:12px; }
    QTabBar::tab { background:#eaf1f8; color:#08285a; padding:9px 14px; border-top-left-radius:10px; border-top-right-radius:10px; font-weight:800; }
    QTabBar::tab:selected { background:#f59a13; color:#08285a; }
    QGroupBox { font-weight:900; color:#08285a; border:1px solid rgba(8,40,90,.14); border-radius:14px; margin-top:12px; padding:14px; background:rgba(255,255,255,.82); }
    QGroupBox::title { subcontrol-origin: margin; left:12px; padding:0 6px; }
    """

    def make_title(text: str, sub: str = "") -> QWidget:
        w = QWidget()
        lay = QVBoxLayout(w); lay.setContentsMargins(0, 0, 0, 12)
        title = QLabel(text); title.setStyleSheet("font-size:28px;font-weight:950;color:#08285a;letter-spacing:-1px;")
        lay.addWidget(title)
        if sub:
            lab = QLabel(sub); lab.setStyleSheet("color:#637083;font-weight:650;font-size:12pt;")
            lay.addWidget(lab)
        return w

    class AnimatedLogoWidget(QWidget):
        def __init__(self, max_size: int = 300):
            super().__init__()
            self.max_size = max_size
            self._tick = 0
            lay = QVBoxLayout(self)
            lay.setContentsMargins(0, 0, 0, 0)
            self.label = QLabel()
            self.label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lay.addWidget(self.label)
            self.pix = QPixmap(str(ResourceManager.find_logo() or ""))
            if self.pix.isNull():
                self.label.setText("L&M")
                self.label.setStyleSheet("font-size:56px;color:#08285a;font-weight:950;background:white;border-radius:30px;padding:28px;")
            else:
                self._apply_size(max_size)
                self.timer = QTimer(self)
                self.timer.timeout.connect(self.animate)
                self.timer.start(180)

        def _apply_size(self, size: int):
            if not self.pix.isNull():
                self.label.setPixmap(self.pix.scaled(size, size, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))

        def animate(self):
            self._tick = (self._tick + 1) % 24
            extra = 14 if self._tick < 12 else 0
            self._apply_size(self.max_size + extra)

    class KpiCard(QFrame):
        def __init__(self, label: str, value: str, accent: str = "#f59a13"):
            super().__init__()
            self.setObjectName("KpiCard")
            self.setStyleSheet("QFrame#KpiCard{background:#fff;border:1px solid #d9e2ef;border-radius:16px;}")
            lay = QVBoxLayout(self); lay.setContentsMargins(16,14,16,14)
            top = QHBoxLayout()
            marker = QLabel(); marker.setFixedSize(10,10); marker.setStyleSheet(f"background:{accent};border-radius:5px;")
            l = QLabel(label.upper()); l.setStyleSheet("color:#64748b;font-weight:900;font-size:9pt;letter-spacing:1px;")
            top.addWidget(marker); top.addWidget(l); top.addStretch(1)
            v = QLabel(value); v.setStyleSheet("color:#08285a;font-weight:950;font-size:24pt;")
            lay.addLayout(top); lay.addStretch(1); lay.addWidget(v)

    class MoneyEdit(QLineEdit):
        def __init__(self):
            super().__init__()
            self.setAlignment(Qt.AlignmentFlag.AlignRight)
            self.setPlaceholderText("$ 0.00")
            self._minimum = 0.0
            self._maximum = 999999999.0
            self._sanitizing = False
            self.textChanged.connect(self._sanitize_text)
        def setPrefix(self, prefix: str): pass
        def setDecimals(self, decimals: int): pass
        def setRange(self, minimum: float, maximum: float):
            self._minimum = float(minimum); self._maximum = float(maximum)
        def setValue(self, value: float):
            try:
                self.setText(f"$ {float(value or 0):,.2f}")
            except Exception:
                self.setText("$ 0.00")
        def _sanitize_text(self):
            if self._sanitizing:
                return
            raw = self.text()
            allowed = "0123456789.,$ "
            clean = "".join(ch for ch in raw if ch in allowed)
            # Solo permitir un separador decimal final; las comas se usan para miles.
            if clean != raw:
                pos = self.cursorPosition()
                self._sanitizing = True
                self.setText(clean)
                self.setCursorPosition(max(0, min(pos-1, len(clean))))
                self._sanitizing = False
        def value(self) -> float:
            raw = self.text().replace("$", "").replace(",", "").strip()
            if raw in ("", ".", "-"):
                return 0.0
            try:
                val = float(raw)
            except Exception:
                return 0.0
            return max(self._minimum, min(self._maximum, val))

    def configure_date_edit(widget: QDateEdit, initial: Optional[str] = None, minimum: Optional[str] = None):
        widget.setCalendarPopup(True)
        widget.setDisplayFormat("dd/MM/yyyy")
        widget.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
        widget.setMaximumDate(QDate.currentDate())
        if minimum:
            qmin = QDate.fromString(minimum, "yyyy-MM-dd")
            if qmin.isValid():
                widget.setMinimumDate(qmin)
        if initial:
            qd = QDate.fromString(initial, "yyyy-MM-dd")
            if not qd.isValid() or qd > QDate.currentDate():
                qd = QDate.currentDate()
            if qd < widget.minimumDate():
                qd = widget.minimumDate()
            widget.setDate(qd)
        return widget

    class TimelineWidget(QWidget):
        stageDoubleClicked = Signal(str)
        def __init__(self, vehicle: Optional[dict] = None):
            super().__init__()
            self.vehicle = vehicle or {}
            self.setMinimumHeight(240)
            self.node_positions: dict[str, QPointF] = {}

        def set_vehicle(self, vehicle: dict):
            self.vehicle = vehicle or {}
            self.update()

        def paintEvent(self, event):
            painter = QPainter(self)
            painter.setRenderHint(QPainter.RenderHint.Antialiasing)
            w, h = self.width(), self.height()
            left, right = 55, w - 55
            y = h // 2 - 8
            n = len(STAGES)
            if n < 2:
                return
            self.node_positions = {}
            painter.setPen(QPen(QColor("#dbe4ef"), 12, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap))
            painter.drawLine(left, y, right, y)
            current_key = self.vehicle.get("estado_actual", STAGE_COMPRADO)
            current_idx = stage_index(current_key)
            for i in range(n - 1):
                x1 = left + (right - left) * i / (n - 1)
                x2 = left + (right - left) * (i + 1) / (n - 1)
                key = STAGES[i]["key"]
                if i < current_idx or vehicle_stage(self.vehicle, key).get("status") == "COMPLETADO":
                    painter.setPen(QPen(QColor(STAGES[i]["color"]), 12, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap))
                    painter.drawLine(int(x1), y, int(x2), y)
            for i, s in enumerate(STAGES):
                x = left + (right - left) * i / (n - 1)
                self.node_positions[s["key"]] = QPointF(x, y)
                st = vehicle_stage(self.vehicle, s["key"]) if self.vehicle else default_stage_record(s["key"])
                completed = st.get("status") == "COMPLETADO"
                active = s["key"] == current_key
                pending = not st.get("fecha_inicio")
                color = QColor(s["color"] if (completed or active) else "#94a3b8")
                if active and stage_alert_level(self.vehicle, s["key"]) == "ROJO":
                    color = QColor("#dc2626")
                radius = 19 if active else 15
                painter.setPen(QPen(QColor("#ffffff"), 4))
                painter.setBrush(color)
                painter.drawEllipse(QPointF(x, y), radius, radius)
                painter.setPen(QColor("white" if (completed or active) else "#e2e8f0"))
                painter.setFont(QFont("Segoe UI Emoji", 10))
                painter.drawText(QRectF(x-16, y-14, 32, 28), Qt.AlignmentFlag.AlignCenter, s["icon"])
                painter.setPen(QColor("#08285a" if (completed or active) else "#64748b"))
                painter.setFont(QFont("Segoe UI", 8, QFont.Weight.Bold))
                painter.drawText(QRectF(x - 60, y + 26, 120, 40), Qt.AlignmentFlag.AlignCenter | Qt.TextFlag.TextWordWrap, s["label"])
                date_text = _fmt_date(st.get("fecha_inicio")) if st.get("fecha_inicio") else "Pendiente"
                painter.setPen(QColor("#637083"))
                painter.setFont(QFont("Segoe UI", 7))
                painter.drawText(QRectF(x - 55, y - 60, 110, 30), Qt.AlignmentFlag.AlignCenter | Qt.TextFlag.TextWordWrap, date_text)

        def mouseDoubleClickEvent(self, event):
            pos = event.position()
            nearest_key, nearest_dist = "", 10**9
            for key, p in self.node_positions.items():
                d = (p.x() - pos.x()) ** 2 + (p.y() - pos.y()) ** 2
                if d < nearest_dist:
                    nearest_key, nearest_dist = key, d
            if nearest_key:
                self.stageDoubleClicked.emit(nearest_key)

    class AdminGate(QDialog):
        def __init__(self, parent=None):
            super().__init__(parent)
            self.setWindowTitle("Acceso restringido")
            self.setMinimumWidth(420)
            lay=QVBoxLayout(self); lay.setContentsMargins(18,18,18,18); lay.setSpacing(12)
            logo=QLabel(); logo.setAlignment(Qt.AlignmentFlag.AlignCenter); pix=QPixmap(str(ResourceManager.find_logo() or ""))
            if not pix.isNull(): logo.setPixmap(pix.scaled(190,190,Qt.AspectRatioMode.KeepAspectRatio,Qt.TransformationMode.SmoothTransformation)); lay.addWidget(logo)
            lay.addWidget(QLabel("Ingresa la clave de administrador:"))
            self.ed=QLineEdit(); self.ed.setEchoMode(QLineEdit.EchoMode.Password); lay.addWidget(self.ed)
            btns=QHBoxLayout(); cancel=QPushButton("Cancelar"); cancel.setObjectName("ghost"); cancel.clicked.connect(self.reject); ok=QPushButton("Entrar"); ok.setObjectName("orange"); ok.clicked.connect(self.try_key)
            btns.addWidget(cancel); btns.addWidget(ok); lay.addLayout(btns); self.ed.returnPressed.connect(self.try_key)
        def try_key(self):
            if self.ed.text()==ADMIN_PASSWORD: self.accept()
            else:
                QMessageBox.warning(self,"Acceso denegado","Clave incorrecta."); self.ed.clear()

    class LoginDialog(QDialog):
        def __init__(self):
            super().__init__()
            self.setWindowTitle(f"{APP_NAME} · Login")
            self.setMinimumSize(1100, 660)
            self.device = collect_device_info()
            self.user: Optional[dict] = None
            self._build()
            QTimer.singleShot(80, self._try_autofill)

        def _build(self):
            outer = QHBoxLayout(self); outer.setContentsMargins(0,0,0,0); outer.setSpacing(0)
            left = QFrame(); left.setStyleSheet("QFrame{background:#101827;}")
            ll = QVBoxLayout(left); ll.setContentsMargins(58,50,58,50); ll.setSpacing(18)
            logo_frame=QFrame(); logo_frame.setStyleSheet("background:#ffffff;border-radius:18px;")
            lf=QVBoxLayout(logo_frame); lf.setContentsMargins(28,28,28,28)
            logo=QLabel(); logo.setAlignment(Qt.AlignmentFlag.AlignCenter); pix=QPixmap(str(ResourceManager.find_logo() or ""))
            if not pix.isNull(): logo.setPixmap(pix.scaled(310,310,Qt.AspectRatioMode.KeepAspectRatio,Qt.TransformationMode.SmoothTransformation))
            else: logo.setText("L&M"); logo.setStyleSheet("font-size:56px;color:#08285a;font-weight:950;")
            lf.addWidget(logo)
            ll.addWidget(logo_frame)
            brand=QLabel("LYM AUTO CONTROL"); brand.setAlignment(Qt.AlignmentFlag.AlignCenter); brand.setStyleSheet("color:white;font-size:32px;font-weight:950;letter-spacing:1px;")
            tagline=QLabel("Control de inventario vehicular"); tagline.setAlignment(Qt.AlignmentFlag.AlignCenter); tagline.setStyleSheet("color:#ffcf7a;font-size:15pt;")
            version=QLabel(f"v{APP_VERSION}"); version.setAlignment(Qt.AlignmentFlag.AlignCenter); version.setStyleSheet("color:#94a3b8;")
            ll.addWidget(brand); ll.addWidget(tagline); ll.addStretch(1); ll.addWidget(version)
            outer.addWidget(left,1)

            right=QFrame(); right.setStyleSheet("QFrame{background:#f3f5f8;}")
            rl=QVBoxLayout(right); rl.setContentsMargins(42,42,42,42); rl.addStretch(1)
            card=QFrame(); card.setMaximumWidth(500); card.setStyleSheet("QFrame{background:white;border:1px solid #d9e2ef;border-radius:20px;}")
            cl=QVBoxLayout(card); cl.setContentsMargins(36,36,36,36); cl.setSpacing(14)
            title=QLabel("Iniciar sesión"); title.setStyleSheet("font-size:30px;font-weight:950;color:#08285a;"); cl.addWidget(title)
            sub=QLabel(f"Bienvenido a {APP_NAME} · {APP_SUBTITLE}"); sub.setStyleSheet("color:#475569;"); cl.addWidget(sub)
            cl.addSpacing(8)
            cl.addWidget(QLabel("Usuario")); self.user_edit=QLineEdit(); self.user_edit.setPlaceholderText("Tu usuario"); cl.addWidget(self.user_edit)
            cl.addWidget(QLabel("Contraseña")); self.pass_edit=QLineEdit(); self.pass_edit.setEchoMode(QLineEdit.EchoMode.Password); self.pass_edit.setPlaceholderText("Tu contraseña"); cl.addWidget(self.pass_edit)
            self.remember=QCheckBox("Recordar usuario y contraseña en esta laptop"); self.remember.setChecked(True); cl.addWidget(self.remember)
            folder=get_active_folder(); folder_txt=str(folder) if folder else "No configurada"
            self.lbl_folder=QLabel(f"<small style='color:#64748b;'>Carpeta del sistema:<br>{folder_txt}</small>"); self.lbl_folder.setTextFormat(Qt.TextFormat.RichText); self.lbl_folder.setWordWrap(True); cl.addWidget(self.lbl_folder)
            self.login_btn=QPushButton("🔓  Entrar"); self.login_btn.setObjectName("orange"); self.login_btn.clicked.connect(self.login); cl.addWidget(self.login_btn)
            bts=QHBoxLayout(); self.admin_btn=QPushButton("⚙️  Configuración"); self.admin_btn.setObjectName("ghost"); self.admin_btn.clicked.connect(self.open_config_gate); self.folder_btn=QPushButton("📂  Carpeta…"); self.folder_btn.setObjectName("ghost"); self.folder_btn.clicked.connect(self.select_folder); bts.addWidget(self.admin_btn); bts.addWidget(self.folder_btn); cl.addLayout(bts)
            self.status=QLabel(""); self.status.setWordWrap(True); self.status.setStyleSheet("color:#dc2626;font-weight:700;"); cl.addWidget(self.status)
            info=QLabel(f"<small style='color:#94a3b8;'>Equipo: {self.device.computer_name} · {self.device.os_name} {self.device.os_version}<br>Device key: {self.device.device_key[:12]}…</small>"); info.setTextFormat(Qt.TextFormat.RichText); info.setAlignment(Qt.AlignmentFlag.AlignCenter); cl.addWidget(info)
            rl.addWidget(card, alignment=Qt.AlignmentFlag.AlignCenter); rl.addStretch(2)
            outer.addWidget(right,1)
            self.pass_edit.returnPressed.connect(self.login); self.user_edit.returnPressed.connect(lambda:self.pass_edit.setFocus())

        def _try_autofill(self):
            u,p=load_autologin(self.device.device_key)
            if u: self.user_edit.setText(u)
            if p: self.pass_edit.setText(p)

        def select_folder(self):
            path = QFileDialog.getExistingDirectory(self, "Selecciona carpeta raíz compartida", str(Path.home()))
            if path:
                set_active_folder(path); bootstrap_system()
                self.lbl_folder.setText(f"<small style='color:#64748b;'>Carpeta del sistema:<br>{get_active_folder()}</small>")
                QMessageBox.information(self,"Carpeta configurada",f"Carpeta del sistema configurada en:\n{get_active_folder()}")

        def ensure_folder(self) -> bool:
            if not get_active_folder():
                QMessageBox.warning(self,"Carpeta requerida","Selecciona primero la carpeta del sistema.")
                return False
            bootstrap_system(); return True

        def open_config_gate(self):
            gate=AdminGate(self)
            if gate.exec()!=QDialog.DialogCode.Accepted: return
            if not self.ensure_folder():
                self.select_folder()
                if not get_active_folder(): return
            admin_stub={"usuario":"_ADMIN_SETUP_","rol":ROLE_ADMIN,"permissions":{k:True for k in ALL_PERMISSION_KEYS}}
            dlg=ConfigDialog(self, admin_stub, self.device); dlg.exec()

        def login(self):
            if not self.ensure_folder(): return
            usuario=self.user_edit.text().strip(); password=self.pass_edit.text()
            if not usuario or not password:
                self.status.setText("Ingresa usuario y contraseña."); return
            self.status.setStyleSheet("color:#1d4ed8;font-weight:700;"); self.status.setText("Validando usuario…")
            QApplication.processEvents()
            ok,msg,user=authenticate(usuario,password,self.device)
            if not ok:
                self.status.setStyleSheet("color:#dc2626;font-weight:700;"); self.status.setText(msg); return
            if self.remember.isChecked(): save_autologin(user.get("usuario",""), password, self.device.device_key, str(get_active_folder()))
            else: clear_autologin()
            self.user=user; log_audit("LOGIN", user.get("usuario", ""), "", "Inicio de sesión"); self.accept()

    class CreateUserDialog(QDialog):
        def __init__(self, parent, device: DeviceInfo, first_admin: bool = False):
            super().__init__(parent)
            self.device = device
            self.first_admin = first_admin
            self.setWindowTitle("Crear usuario")
            self.setMinimumWidth(460)
            lay = QVBoxLayout(self)
            form = QFormLayout()
            self.usuario = QLineEdit()
            self.password = QLineEdit(); self.password.setEchoMode(QLineEdit.EchoMode.Password)
            self.rol = QComboBox(); self.rol.addItems(ROLES)
            if first_admin:
                self.rol.setCurrentText(ROLE_ADMIN); self.rol.setEnabled(False)
            self.activo = QCheckBox("Usuario activo")
            self.activo.setChecked(True)
            form.addRow("Usuario:", self.usuario); form.addRow("Contraseña:", self.password); form.addRow("Rol:", self.rol); form.addRow("Estado:", self.activo)
            lay.addLayout(form)
            info = QLabel("Los permisos iniciales se basan en el rol. Luego pueden personalizarse en Configuración > Permisos.")
            info.setWordWrap(True); info.setStyleSheet("color:#637083;")
            lay.addWidget(info)
            btn = QPushButton("Guardar usuario"); btn.setObjectName("orange"); btn.clicked.connect(self.save)
            lay.addWidget(btn)

        def save(self):
            ok, msg = create_user(self.usuario.text(), self.password.text(), self.rol.currentText(), self.device, activo=self.activo.isChecked())
            if not ok:
                QMessageBox.warning(self, "Validación", msg); return
            self.accept()

    class DashboardCustomizeDialog(QDialog):
        def __init__(self, parent=None):
            super().__init__(parent); self.setWindowTitle("Personalizar dashboard"); self.setMinimumWidth(520)
            s=get_dashboard_settings(); lay=QVBoxLayout(self)
            lay.addWidget(QLabel("Activa o desactiva cada tarjeta del dashboard:"))
            self.card_checks = {}
            for key, label in DASHBOARD_CARD_KEYS:
                chk = QCheckBox(label)
                chk.setChecked(bool(s.get("cards", {}).get(key, True)))
                self.card_checks[key] = chk
                lay.addWidget(chk)
            self.chk_table=QCheckBox("Mostrar tabla de vehículos recientes/críticos"); self.chk_table.setChecked(s.get("recent_table", True)); lay.addWidget(self.chk_table)
            btn=QPushButton("Guardar personalización"); btn.setObjectName("orange"); btn.clicked.connect(self.save); lay.addWidget(btn)
        def save(self):
            set_dashboard_settings({"recent_table":self.chk_table.isChecked(),"cards":{k:c.isChecked() for k,c in self.card_checks.items()}}); self.accept()

    class DashboardPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main = main; self._build()
        def _build(self):
            lay = QVBoxLayout(self)
            head=QHBoxLayout(); head.addWidget(make_title("Dashboard gerencial", "Resumen limpio de inventario, disponibilidad, capital y alertas.")); head.addStretch(1)
            bcustom=QPushButton("Personalizar dashboard"); bcustom.setObjectName("ghost"); bcustom.clicked.connect(self.customize); head.addWidget(bcustom, alignment=Qt.AlignmentFlag.AlignTop)
            lay.addLayout(head)
            self.kpi_wrap=QWidget(); self.kpi_grid = QGridLayout(self.kpi_wrap); lay.addWidget(self.kpi_wrap)
            self.table_label=QLabel("Vehículos recientes / críticos")
            self.table = QTableWidget(0, 6); self.table.setHorizontalHeaderLabels(["Código", "Vehículo", "Estado", "Costo", "Días", "Alerta"])
            self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            self.table.cellDoubleClicked.connect(self.open_row)
            lay.addWidget(self.table_label); lay.addWidget(self.table)
        def customize(self):
            dlg=DashboardCustomizeDialog(self)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def refresh(self):
            s=get_dashboard_settings()
            self.kpi_wrap.setVisible(s.get("kpis", True))
            self.table.setVisible(s.get("recent_table", True)); self.table_label.setVisible(s.get("recent_table", True))
            while self.kpi_grid.count():
                item = self.kpi_grid.takeAt(0)
                if item.widget(): item.widget().deleteLater()
            vehicles = load_vehicles(); k = compute_kpis(vehicles)
            all_cards = {
                "total": ("Vehículos", str(k["total"]), "#08285a"),
                "activos": ("Activos", str(k["activos"]), "#0e3a78"),
                "disponibles": ("Disponibles", str(k["disponibles"]), "#10B981"),
                "vendidos": ("Vendidos", str(k["vendidos"]), "#08285a"),
                "capital": ("Capital activo", _fmt_usd(k["capital"]), "#f59a13"),
                "ganancia": ("Ganancia esperada", _fmt_usd(k["ganancia_esperada"]), "#10B981"),
                "criticos": ("Críticos", str(k["criticos"]), "#dc2626"),
                "compra_disp": ("Compra→Disponible", f"{max(0,k['prom_compra_disp'])} días", "#7C3AED"),
            }
            visible_cards = s.get("cards", {})
            cards = [all_cards[key] for key, _ in DASHBOARD_CARD_KEYS if visible_cards.get(key, True)]
            for idx, (label, value, color) in enumerate(cards):
                self.kpi_grid.addWidget(KpiCard(label, value, color), idx//4, idx%4)
            candidates = sorted(vehicles, key=lambda v: (stage_alert_level(v) != "ROJO", -vehicle_days_from_purchase(v)))[:25]
            self.table.setRowCount(len(candidates)); self._ids = []
            for r, v in enumerate(candidates):
                self._ids.append(v.get("id"))
                vals = [v.get("codigo"), f"{v.get('marca')} {v.get('modelo')} {v.get('anio')}", STAGE_META.get(v.get("estado_actual"),{}).get("label", v.get("estado_actual")), _fmt_usd(vehicle_total_cost(v)), str(vehicle_days_from_purchase(v)), stage_alert_level(v)]
                for c, val in enumerate(vals): self.table.setItem(r, c, QTableWidgetItem(str(val)))
        def open_row(self, row, col):
            if row < len(getattr(self, "_ids", [])):
                self.main.open_vehicle_detail(self._ids[row])

    class ReporteriaPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main=main; self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Reportería", "Genera reportes separados: inventario, KPI de tiempos y Excel operativo."))
            grid=QGridLayout(); lay.addLayout(grid)
            items=[
                ("Reporte Excel de inventario", "Exporta la tabla general de vehículos a Excel.", self.report_excel),
                ("Reporte HTML de inventario", "Reporte visual de inventario y disponibles para venta.", self.report_html_inventory),
                ("Reporte HTML KPI de tiempos", "Mide días por etapa, ciclo completo y atrasos en HTML.", self.report_html_kpi),
                ("Reporte Excel KPI de tiempos", "Mide tiempos por etapa y detalle por vehículo en Excel.", self.report_excel_kpi),
            ]
            for i,(title,desc,fn) in enumerate(items):
                card=QFrame(); card.setStyleSheet("QFrame{background:white;border:1px solid #d9e2ef;border-radius:18px;}"); cl=QVBoxLayout(card); lab=QLabel(title); lab.setStyleSheet("font-size:16pt;font-weight:900;color:#08285a;"); d=QLabel(desc); d.setWordWrap(True); d.setStyleSheet("color:#64748b;"); b=QPushButton("Generar"); b.setObjectName("orange"); b.clicked.connect(fn); cl.addWidget(lab); cl.addWidget(d); cl.addStretch(1); cl.addWidget(b); grid.addWidget(card, i//2, i%2)
            lay.addStretch(1)
        def _after_report(self, path: Path, title: str):
            msg = QMessageBox(self)
            msg.setWindowTitle(title)
            msg.setText(f"Reporte generado correctamente:\n{path.name}")
            msg.setInformativeText("¿Qué deseas hacer?")
            open_btn = msg.addButton("Abrir", QMessageBox.ButtonRole.AcceptRole)
            copy_btn = msg.addButton("Guardar una copia…", QMessageBox.ButtonRole.ActionRole)
            close_btn = msg.addButton("Cerrar", QMessageBox.ButtonRole.RejectRole)
            msg.exec()
            clicked = msg.clickedButton()
            if clicked == open_btn:
                QDesktopServices.openUrl(QUrl.fromLocalFile(str(path)))
            elif clicked == copy_btn:
                target, _ = QFileDialog.getSaveFileName(self, "Guardar copia", str(Path.home() / path.name), "Archivos (*" + path.suffix + ")")
                if target:
                    ok, m = copy_report_to(path, Path(target))
                    if ok: QMessageBox.information(self, "Copia guardada", m)
                    else: QMessageBox.warning(self, "Copia", m)
        def report_excel(self):
            ok,msg,out=generate_inventory_excel(load_vehicles(), self.main.user)
            if not ok or not out: QMessageBox.warning(self,"Excel",msg); return
            self._after_report(out, "Excel de inventario")
        def report_html_inventory(self):
            out=generate_html_report(load_vehicles(), self.main.user)
            if out: self._after_report(out, "HTML de inventario")
        def report_html_kpi(self):
            out=generate_kpi_html_report(load_vehicles(), self.main.user)
            if out: self._after_report(out, "HTML KPI de tiempos")
        def report_excel_kpi(self):
            ok,msg,out=generate_kpi_excel_report(load_vehicles(), self.main.user)
            if not ok or not out: QMessageBox.warning(self,"Excel KPI",msg); return
            self._after_report(out, "Excel KPI de tiempos")

    class TarifaDialog(QDialog):
        def __init__(self, parent=None, item: Optional[dict] = None):
            super().__init__(parent)
            self.setWindowTitle("Agregar / editar tarifa")
            self.setMinimumWidth(420)
            lay=QVBoxLayout(self); form=QFormLayout()
            self.nombre=QLineEdit((item or {}).get("nombre", ""))
            self.valor=MoneyEdit(); self.valor.setValue(float((item or {}).get("valor_usd") or 0))
            self.tipo=QComboBox(); self.tipo.addItems(["USD", "POR_VALOR", "POR_CBM", "POR_UNIDAD"]); self.tipo.setCurrentText((item or {}).get("tipo", "USD"))
            form.addRow("Nombre:", self.nombre); form.addRow("Valor:", self.valor); form.addRow("Tipo:", self.tipo)
            lay.addLayout(form)
            b=QPushButton("Guardar tarifa"); b.setObjectName("orange"); b.clicked.connect(self.accept); lay.addWidget(b)
        def data(self) -> dict:
            return {"nombre": _norm(self.nombre.text()), "valor_usd": self.valor.value(), "tipo": _norm(self.tipo.currentText())}

    class CatalogosPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main=main; self._build(); self.refresh_all()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Catálogos", "Navieras, talleres, proveedores, gestores, transportistas, tarifas y aduanas por país."))
            self.tabs=QTabWidget(); lay.addWidget(self.tabs)
            self.simple_tabs = {}
            simple_defs = [
                ("Navieras", F_CATALOG_NAVIERAS, DEFAULT_NAVIERAS),
                ("Talleres", F_CATALOG_TALLERES, DEFAULT_TALLERES),
                ("Proveedores", F_CATALOG_PROVEEDORES, DEFAULT_PROVEEDORES),
                ("Gestores", F_CATALOG_GESTORES, DEFAULT_GESTORES),
                ("Transportistas USA", F_CATALOG_TRANSPORTISTAS_USA, DEFAULT_TRANSPORTISTAS_USA),
            ]
            for title, fn, default in simple_defs:
                self.tabs.addTab(self._build_simple_tab(title, fn, default), title)
            self.tabs.addTab(self._build_aduanas_tab(), "Aduanas")
            self.tabs.addTab(self._build_tarifas_tab(), "Tarifas")

        def _build_simple_tab(self, title: str, filename: str, default: list[str]) -> QWidget:
            w=QWidget(); lay=QVBoxLayout(w)
            row=QHBoxLayout(); badd=QPushButton(f"Agregar {title[:-1] if title.endswith('s') else title}"); badd.setObjectName("orange"); bedit=QPushButton("Editar seleccionado"); bedit.setObjectName("ghost"); bdel=QPushButton("Borrar seleccionado"); bdel.setObjectName("danger")
            row.addWidget(badd); row.addWidget(bedit); row.addWidget(bdel); row.addStretch(1); lay.addLayout(row)
            table=QTableWidget(0,1); table.setHorizontalHeaderLabels(["Nombre"]); table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); lay.addWidget(table)
            self.simple_tabs[title]=(table, filename, default)
            badd.clicked.connect(lambda _,t=title:self.simple_add(t)); bedit.clicked.connect(lambda _,t=title:self.simple_edit(t)); bdel.clicked.connect(lambda _,t=title:self.simple_delete(t))
            return w

        def _selected_simple_value(self, title: str) -> str:
            table,_,_=self.simple_tabs[title]; row=table.currentRow(); item=table.item(row,0) if row>=0 else None
            return item.text() if item else ""
        def simple_add(self, title: str):
            table,fn,default=self.simple_tabs[title]
            text,ok=QInputDialog.getText(self,"Agregar",f"Nuevo valor para {title}:")
            if ok and text.strip(): add_catalog_value(fn,text,default); self.refresh_simple(title)
        def simple_edit(self, title: str):
            old=self._selected_simple_value(title)
            if not old: QMessageBox.information(self,"Catálogos","Selecciona una línea."); return
            table,fn,default=self.simple_tabs[title]
            text,ok=QInputDialog.getText(self,"Editar",f"Editar valor de {title}:", text=old)
            if ok and text.strip(): remove_catalog_value(fn,old,default); add_catalog_value(fn,text,default); self.refresh_simple(title)
        def simple_delete(self, title: str):
            val=self._selected_simple_value(title)
            if not val: QMessageBox.information(self,"Catálogos","Selecciona una línea para borrar."); return
            if QMessageBox.question(self,"Borrar",f"¿Borrar '{val}' de {title}?") != QMessageBox.StandardButton.Yes: return
            table,fn,default=self.simple_tabs[title]; remove_catalog_value(fn,val,default); self.refresh_simple(title)
        def refresh_simple(self, title: str):
            table,fn,default=self.simple_tabs[title]; vals=load_catalog(fn, default); table.setRowCount(len(vals))
            for r,val in enumerate(vals): table.setItem(r,0,QTableWidgetItem(val))

        def _build_aduanas_tab(self) -> QWidget:
            w=QWidget(); lay=QVBoxLayout(w)
            top=QHBoxLayout(); self.aduana_pais=QComboBox(); self.aduana_pais.setEditable(True); self.aduana_pais.addItems(load_catalog(F_CATALOG_PAISES, DEFAULT_PAISES_DESTINO)); self.aduana_pais.currentTextChanged.connect(self.refresh_aduanas_tab)
            top.addWidget(QLabel("País:")); top.addWidget(self.aduana_pais); top.addStretch(1); lay.addLayout(top)
            row=QHBoxLayout(); badd=QPushButton("Agregar aduana"); badd.setObjectName("orange"); bdel=QPushButton("Borrar seleccionado"); bdel.setObjectName("danger"); row.addWidget(badd); row.addWidget(bdel); row.addStretch(1); lay.addLayout(row)
            self.aduanas_table=QTableWidget(0,1); self.aduanas_table.setHorizontalHeaderLabels(["Aduana"]); self.aduanas_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.aduanas_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.aduanas_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); lay.addWidget(self.aduanas_table)
            badd.clicked.connect(self.add_aduana); bdel.clicked.connect(self.delete_aduana)
            return w
        def refresh_aduanas_tab(self):
            if not hasattr(self,'aduanas_table'): return
            vals=aduanas_for_country(self.aduana_pais.currentText()); self.aduanas_table.setRowCount(len(vals))
            for r,val in enumerate(vals): self.aduanas_table.setItem(r,0,QTableWidgetItem(val))
        def add_aduana(self):
            pais=self.aduana_pais.currentText()
            text,ok=QInputDialog.getText(self,"Agregar aduana",f"Nueva aduana para {pais}:")
            if ok and text.strip(): add_aduana_for_country(pais,text); self.refresh_aduanas_tab()
        def delete_aduana(self):
            row=self.aduanas_table.currentRow(); item=self.aduanas_table.item(row,0) if row>=0 else None
            if not item: QMessageBox.information(self,"Aduanas","Selecciona una aduana para borrar."); return
            if QMessageBox.question(self,"Borrar",f"¿Borrar '{item.text()}' de {self.aduana_pais.currentText()}?") != QMessageBox.StandardButton.Yes: return
            remove_aduana_for_country(self.aduana_pais.currentText(), item.text()); self.refresh_aduanas_tab()

        def _build_tarifas_tab(self) -> QWidget:
            w=QWidget(); lay=QVBoxLayout(w)
            row=QHBoxLayout(); badd=QPushButton("Agregar tarifa"); badd.setObjectName("orange"); bedit=QPushButton("Editar seleccionado"); bedit.setObjectName("ghost"); bdel=QPushButton("Borrar seleccionado"); bdel.setObjectName("danger")
            row.addWidget(badd); row.addWidget(bedit); row.addWidget(bdel); row.addStretch(1); lay.addLayout(row)
            self.tarifas_table=QTableWidget(0,3); self.tarifas_table.setHorizontalHeaderLabels(["Nombre","Valor","Tipo"]); self.tarifas_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.tarifas_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.tarifas_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); lay.addWidget(self.tarifas_table)
            badd.clicked.connect(self.add_tarifa); bedit.clicked.connect(self.edit_tarifa); bdel.clicked.connect(self.delete_tarifa)
            return w
        def refresh_tarifas(self):
            vals=load_tariffs(); self.tarifas_table.setRowCount(len(vals))
            for r,item in enumerate(vals):
                self.tarifas_table.setItem(r,0,QTableWidgetItem(item.get('nombre','')))
                self.tarifas_table.setItem(r,1,QTableWidgetItem(_fmt_usd(item.get('valor_usd'))))
                self.tarifas_table.setItem(r,2,QTableWidgetItem(item.get('tipo','USD')))
        def add_tarifa(self):
            dlg=TarifaDialog(self)
            if dlg.exec()==QDialog.DialogCode.Accepted:
                vals=load_tariffs(); vals.append(dlg.data()); save_tariffs(vals); self.refresh_tarifas()
        def edit_tarifa(self):
            row=self.tarifas_table.currentRow(); vals=load_tariffs()
            if row<0 or row>=len(vals): QMessageBox.information(self,"Tarifas","Selecciona una tarifa."); return
            dlg=TarifaDialog(self, vals[row])
            if dlg.exec()==QDialog.DialogCode.Accepted:
                vals[row]=dlg.data(); save_tariffs(vals); self.refresh_tarifas()
        def delete_tarifa(self):
            row=self.tarifas_table.currentRow(); vals=load_tariffs()
            if row<0 or row>=len(vals): QMessageBox.information(self,"Tarifas","Selecciona una tarifa."); return
            if QMessageBox.question(self,"Borrar",f"¿Borrar tarifa '{vals[row].get('nombre')}'?") != QMessageBox.StandardButton.Yes: return
            vals.pop(row); save_tariffs(vals); self.refresh_tarifas()
        def refresh_all(self):
            for title in list(self.simple_tabs.keys()): self.refresh_simple(title)
            self.refresh_aduanas_tab(); self.refresh_tarifas()

    class PurchasePage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main = main; self.comprobante_path = ""; self._build()
        def _build(self):
            lay = QVBoxLayout(self); lay.addWidget(make_title("Agregar nueva compra", "Código automático: LYM-CV-2026-0001, con comprobante y usuario registrador."))
            scroll = QScrollArea(); scroll.setWidgetResizable(True); content = QWidget(); scroll.setWidget(content); form_lay = QVBoxLayout(content)
            g1 = QGroupBox("Datos del vehículo"); f1 = QFormLayout(g1)
            self.marca = QComboBox(); self.marca.setEditable(True)
            self.modelo = QLineEdit(); self.anio = QSpinBox(); self.anio.setRange(1980, date.today().year + 1); self.anio.setValue(date.today().year)
            self.millaje = QSpinBox(); self.millaje.setRange(0, 999999); self.color = QLineEdit(); self.tipo = QComboBox(); self.tipo.setEditable(True); self.tipo.addItems(["SEDAN", "SUV", "PICKUP", "VAN", "CAMIONETA", "DEPORTIVO", "OTRO"])
            for w in (self.marca, self.tipo): w.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
            f1.addRow("Marca:", self.marca); f1.addRow("Modelo:", self.modelo); f1.addRow("Año:", self.anio); f1.addRow("Millaje:", self.millaje); f1.addRow("Color:", self.color); f1.addRow("Tipo:", self.tipo)
            g2 = QGroupBox("Compra / Subasta"); f2 = QFormLayout(g2)
            self.estado_usa = QComboBox(); self.subasta = QComboBox(); self.subasta.setEditable(True)
            self.btn_add_subasta = QPushButton("+ Agregar subasta"); self.btn_add_subasta.setObjectName("ghost"); self.btn_add_subasta.clicked.connect(self.add_subasta)
            subrow = QWidget(); sublay = QHBoxLayout(subrow); sublay.setContentsMargins(0,0,0,0); sublay.addWidget(self.subasta); sublay.addWidget(self.btn_add_subasta)
            self.lote = QLineEdit(); self.precio = MoneyEdit(); self.precio.setRange(0, 9999999); self.precio.setDecimals(2); self.precio.setPrefix("$ ")
            self.fecha = configure_date_edit(QDateEdit()); self.fecha.setDate(QDate.currentDate())
            self.usuario_reg = QLineEdit(self.main.user.get("usuario", "")); self.usuario_reg.setReadOnly(True)
            f2.addRow("Estado USA:", self.estado_usa); f2.addRow("Subasta:", subrow); f2.addRow("Número de lote:", self.lote); f2.addRow("Precio ganado USD:", self.precio); f2.addRow("Fecha compra:", self.fecha); f2.addRow("Usuario registra:", self.usuario_reg)
            g3 = QGroupBox("Comprobante y observaciones"); f3 = QFormLayout(g3)
            self.comp_label = QLineEdit(); self.comp_label.setReadOnly(True)
            bcomp = QPushButton("Subir comprobante de compra"); bcomp.setObjectName("ghost"); bcomp.clicked.connect(self.pick_comprobante)
            crow = QWidget(); clay = QHBoxLayout(crow); clay.setContentsMargins(0,0,0,0); clay.addWidget(self.comp_label); clay.addWidget(bcomp)
            self.obs = QTextEdit(); self.obs.setMinimumHeight(80)
            f3.addRow("Comprobante:", crow); f3.addRow("Observaciones:", self.obs)
            form_lay.addWidget(g1); form_lay.addWidget(g2); form_lay.addWidget(g3)
            btn = QPushButton("Crear compra"); btn.setObjectName("orange"); btn.clicked.connect(self.save_purchase); form_lay.addWidget(btn)
            lay.addWidget(scroll); self.refresh_catalogs()
        def refresh_catalogs(self):
            self.marca.clear(); self.marca.addItems(load_catalog(F_CATALOG_MARCAS, DEFAULT_MARCAS))
            self.subasta.clear(); self.subasta.addItems(load_catalog(F_CATALOG_SUBASTAS, DEFAULT_SUBASTAS))
            self.estado_usa.clear(); self.estado_usa.addItems(load_catalog(F_CATALOG_ESTADOS_USA, US_STATES))
        def add_subasta(self):
            text, ok = QInputDialog.getText(self, "Agregar subasta", "Nombre de subasta:")
            if ok and text.strip():
                add_catalog_value(F_CATALOG_SUBASTAS, text, DEFAULT_SUBASTAS); self.refresh_catalogs(); self.subasta.setCurrentText(_norm(text))
        def pick_comprobante(self):
            path, _ = QFileDialog.getOpenFileName(self, "Selecciona comprobante", str(Path.home()), "Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if path: self.comprobante_path = path; self.comp_label.setText(path)
        def save_purchase(self):
            data = {
                "marca": self.marca.currentText(), "modelo": self.modelo.text(), "anio": self.anio.value(),
                "millaje": self.millaje.value(), "color": self.color.text(), "tipo": self.tipo.currentText(),
                "estado_usa": self.estado_usa.currentText(), "subasta": self.subasta.currentText(), "lote": self.lote.text(),
                "precio_ganado_usd": self.precio.value(), "fecha_compra": self.fecha.date().toPython().isoformat(),
                "observaciones": self.obs.toPlainText(),
            }
            ok, msg, vid = create_vehicle_purchase(data, Path(self.comprobante_path) if self.comprobante_path else Path(), self.main.user, self.main.device)
            if not ok: QMessageBox.warning(self, "Validación", msg); return
            QMessageBox.information(self, "Compra creada", msg)
            self.main.refresh_all(); self.main.open_vehicle_detail(vid)

    class InventoryPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main = main; self._ids=[]; self._build()
        def _build(self):
            lay = QVBoxLayout(self); lay.addWidget(make_title("Inventario vehicular", "Doble click sobre un código CV para abrir expediente, línea de tiempo y datos completos."))
            filt = QHBoxLayout(); self.search = QLineEdit(); self.search.setPlaceholderText("Buscar código, marca, modelo, lote..."); self.search.textChanged.connect(self.refresh)
            self.status = QComboBox(); self.status.addItem("TODOS")
            for s in STAGES: self.status.addItem(s["label"], s["key"])
            self.status.currentIndexChanged.connect(self.refresh)
            filt.addWidget(self.search); filt.addWidget(self.status); lay.addLayout(filt)
            self.table = QTableWidget(0, 13)
            self.table.setHorizontalHeaderLabels(["CV", "Marca", "Modelo", "Año", "Millaje", "Estado USA", "Subasta", "Lote", "Precio ganado", "Costo total", "Precio venta", "Estado", "Días"])
            self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
            self.table.cellDoubleClicked.connect(self.open_row)
            lay.addWidget(self.table)
        def refresh(self):
            text = _norm(self.search.text()); st = self.status.currentData()
            vehicles = load_vehicles()
            if st: vehicles = [v for v in vehicles if v.get("estado_actual") == st]
            if text:
                vehicles = [v for v in vehicles if text in _norm(" ".join(str(v.get(k,"")) for k in ["codigo","marca","modelo","lote","subasta","estado_usa"] ))]
            self._ids=[]; self.table.setRowCount(len(vehicles))
            for r,v in enumerate(vehicles):
                self._ids.append(v.get("id"))
                vals=[v.get("codigo"),v.get("marca"),v.get("modelo"),v.get("anio"),f"{int(v.get('millaje') or 0):,}",v.get("estado_usa"),v.get("subasta"),v.get("lote"),_fmt_usd(v.get("precio_ganado_usd")),_fmt_usd(vehicle_total_cost(v)),_fmt_usd(v.get("precio_venta_usd")),STAGE_META.get(v.get("estado_actual"),{}).get("label",v.get("estado_actual")),str(vehicle_days_from_purchase(v))]
                for c,val in enumerate(vals): self.table.setItem(r,c,QTableWidgetItem(str(val)))
        def open_row(self,row,col):
            if row < len(self._ids): self.main.open_vehicle_detail(self._ids[row])
        def generate_report(self):
            if not user_has_permission(self.main.user, PERM_GENERATE_REPORTS):
                QMessageBox.warning(self,"Permiso","No tienes permiso para generar reportes."); return
            out=generate_html_report(load_vehicles(), self.main.user)
            if out:
                QMessageBox.information(self,"Reporte generado",f"Reporte HTML generado:\n{out}")
                QDesktopServices.openUrl(QUrl.fromLocalFile(str(out)))

    class StageUpdateDialog(QDialog):
        def __init__(self, parent, vehicle: dict, stage_key: str, user: dict, device: DeviceInfo, mode: str = "advance"):
            super().__init__(parent)
            self.vehicle = vehicle
            self.stage_key = stage_key
            self.user = user
            self.device = device
            self.mode = mode
            self.doc_path = ""
            self.setWindowTitle(f"{('Pasar a' if mode=='advance' else 'Editar')} etapa · {STAGE_META[stage_key]['label']}")
            self.setMinimumWidth(660)
            self._build()

        def _combo(self, values: list[str]) -> QComboBox:
            cb = QComboBox(); cb.setEditable(True); cb.addItems(values); return cb

        def _build(self):
            lay = QVBoxLayout(self)
            meta = STAGE_META[self.stage_key]
            current_label = STAGE_META.get(self.vehicle.get("estado_actual", STAGE_COMPRADO), {}).get("label", self.vehicle.get("estado_actual"))
            title = QLabel(f"<b>{self.vehicle.get('codigo')}</b> · {self.vehicle.get('marca')} {self.vehicle.get('modelo')} {self.vehicle.get('anio')}")
            lay.addWidget(title)
            sub = QLabel(f"Etapa actual: <b>{current_label}</b> → Nueva etapa: <b>{meta['label']}</b>")
            sub.setTextFormat(Qt.TextFormat.RichText); sub.setStyleSheet("color:#475569;")
            lay.addWidget(sub)
            form = QFormLayout(); form.setVerticalSpacing(12)
            st = vehicle_stage(self.vehicle, self.stage_key)
            min_stage_date = (_parse_date(vehicle_stage(self.vehicle, STAGE_ORDER[stage_index(self.stage_key)-1]).get("fecha_inicio")) if stage_index(self.stage_key) > 0 else _parse_date(self.vehicle.get("fecha_compra"))) or _parse_date(self.vehicle.get("fecha_compra")) or date.today()
            self.fecha_evento = configure_date_edit(QDateEdit(), st.get("fecha_inicio") or _today_iso(), min_stage_date.isoformat())
            date_labels = {
                STAGE_TRASLADO_USA: "Fecha salida subasta:",
                STAGE_MIAMI_NAVIERA: "Fecha llegada Miami:",
                STAGE_TRANSITO: "Fecha salida naviera:",
                STAGE_ADUANA: "Fecha llegada aduana:",
                STAGE_LEGALIZACION: "Fecha inicio legalización:",
                STAGE_TALLER: "Fecha ingreso taller:",
                STAGE_DISPONIBLE: "Fecha disponible venta:",
                STAGE_RESERVADO: "Fecha reserva:",
                STAGE_VENDIDO: "Fecha venta:",
            }
            form.addRow(date_labels.get(self.stage_key, "Fecha:"), self.fecha_evento)
            self.costo = MoneyEdit(); self.costo.setPrefix("$ "); self.costo.setDecimals(2); self.costo.setRange(0,9999999); self.costo.setValue(float(st.get("costo_usd") or 0))
            self.proveedor_widget = None
            self.extra_widgets = {}
            if self.stage_key == STAGE_TRASLADO_USA:
                self.proveedor_widget = self._combo(load_catalog(F_CATALOG_TRANSPORTISTAS_USA, DEFAULT_TRANSPORTISTAS_USA))
                form.addRow("Transportista USA:", self.proveedor_widget); form.addRow("Costo traslado:", self.costo)
            elif self.stage_key in (STAGE_MIAMI_NAVIERA, STAGE_TRANSITO):
                self.proveedor_widget = self._combo(load_catalog(F_CATALOG_NAVIERAS, DEFAULT_NAVIERAS))
                form.addRow("Naviera:", self.proveedor_widget); form.addRow("Costo naviera/bodega:", self.costo)
            elif self.stage_key == STAGE_ADUANA:
                self.pais = self._combo(load_catalog(F_CATALOG_PAISES, DEFAULT_PAISES_DESTINO)); self.aduana = self._combo(aduanas_for_country(self.pais.currentText()))
                self.pais.currentTextChanged.connect(self.refresh_aduanas_by_country)
                form.addRow("País destino:", self.pais); form.addRow("Aduana:", self.aduana); form.addRow("Gasto aduana:", self.costo)
            elif self.stage_key == STAGE_LEGALIZACION:
                self.proveedor_widget = self._combo(load_catalog(F_CATALOG_GESTORES, DEFAULT_GESTORES))
                form.addRow("Gestor legal:", self.proveedor_widget); form.addRow("Costo legalización:", self.costo)
            elif self.stage_key == STAGE_TALLER:
                self.proveedor_widget = self._combo(load_catalog(F_CATALOG_TALLERES, DEFAULT_TALLERES))
                form.addRow("Taller:", self.proveedor_widget); form.addRow("Gasto taller/repuestos:", self.costo)
            elif self.stage_key == STAGE_DISPONIBLE:
                form.addRow("Gasto final/preparación:", self.costo)
                self.precio_venta = MoneyEdit(); self.precio_venta.setPrefix("$ "); self.precio_venta.setRange(0,9999999); self.precio_venta.setDecimals(2); self.precio_venta.setValue(float(self.vehicle.get("precio_venta_usd") or 0))
                self.precio_min = MoneyEdit(); self.precio_min.setPrefix("$ "); self.precio_min.setRange(0,9999999); self.precio_min.setDecimals(2); self.precio_min.setValue(max(float(self.vehicle.get("precio_minimo_usd") or 0), vehicle_total_cost(self.vehicle)))
                self.costo_total_lbl = QLabel()
                self.costo_total_lbl.setTextFormat(Qt.TextFormat.RichText); self.costo_total_lbl.setStyleSheet("color:#08285a;background:#fff7e6;padding:8px;border-radius:8px;")
                try: self.costo.textChanged.connect(self._update_total_disponible)
                except Exception: pass
                self._update_total_disponible()
                form.addRow("Costo total:", self.costo_total_lbl); form.addRow("Precio venta:", self.precio_venta); form.addRow("Precio mínimo:", self.precio_min)
            elif self.stage_key == STAGE_RESERVADO:
                self.cliente = QLineEdit(self.vehicle.get("cliente") or "")
                self.monto_reserva = MoneyEdit(); self.monto_reserva.setPrefix("$ "); self.monto_reserva.setRange(0,9999999); self.monto_reserva.setDecimals(2); self.monto_reserva.setValue(float(self.vehicle.get("monto_reserva_usd") or 0))
                self.precio_reserva = MoneyEdit(); self.precio_reserva.setPrefix("$ "); self.precio_reserva.setRange(0,9999999); self.precio_reserva.setDecimals(2); self.precio_reserva.setValue(float(self.vehicle.get("precio_venta_usd") or min_sale_price(self.vehicle)))
                form.addRow("Cliente:", self.cliente); form.addRow("Monto reserva:", self.monto_reserva); form.addRow("Precio acordado:", self.precio_reserva)
            elif self.stage_key == STAGE_VENDIDO:
                self.cliente = QLineEdit(self.vehicle.get("cliente") or "")
                self.precio_real = MoneyEdit(); self.precio_real.setPrefix("$ "); self.precio_real.setRange(0,9999999); self.precio_real.setDecimals(2); self.precio_real.setValue(float(self.vehicle.get("precio_venta_real_usd") or self.vehicle.get("precio_reserva_usd") or self.vehicle.get("precio_venta_usd") or min_sale_price(self.vehicle)))
                form.addRow("Cliente:", self.cliente); form.addRow("Precio venta real:", self.precio_real)
            else:
                self.proveedor_widget = self._combo(load_catalog(F_CATALOG_PROVEEDORES, DEFAULT_PROVEEDORES))
                form.addRow("Responsable/proveedor:", self.proveedor_widget); form.addRow("Costo:", self.costo)
            self.doc_label = QLineEdit(st.get("documento_nombre") or ""); self.doc_label.setReadOnly(True)
            bdoc = QPushButton("Subir PDF/evidencia"); bdoc.setObjectName("ghost"); bdoc.clicked.connect(self.pick_doc)
            drow = QWidget(); dlay = QHBoxLayout(drow); dlay.setContentsMargins(0,0,0,0); dlay.addWidget(self.doc_label); dlay.addWidget(bdoc)
            form.addRow("Documento:", drow)
            self.comentario = QTextEdit(st.get("comentario") or ""); self.comentario.setMinimumHeight(90)
            form.addRow("Comentario:", self.comentario)
            lay.addLayout(form)
            note = QLabel("La fecha registrada abre esta etapa y cierra automáticamente la etapa anterior para calcular días reales.")
            note.setWordWrap(True); note.setStyleSheet("color:#64748b;")
            lay.addWidget(note)
            btn = QPushButton("Guardar avance"); btn.setObjectName("orange"); btn.clicked.connect(self.save); lay.addWidget(btn)

        def _update_total_disponible(self):
            if hasattr(self, "costo_total_lbl"):
                base = vehicle_total_cost(self.vehicle)
                final = base + (self.costo.value() if hasattr(self, "costo") else 0)
                self.costo_total_lbl.setText(f"Costo acumulado actual: <b>{_fmt_usd(base)}</b><br>Con gasto final/preparación: <b>{_fmt_usd(final)}</b>")

        def refresh_aduanas_by_country(self):
            if hasattr(self, "aduana") and hasattr(self, "pais"):
                current = self.aduana.currentText()
                self.aduana.clear(); self.aduana.addItems(aduanas_for_country(self.pais.currentText()))
                idx = self.aduana.findText(current)
                if idx >= 0: self.aduana.setCurrentIndex(idx)

        def pick_doc(self):
            path,_=QFileDialog.getOpenFileName(self,"Selecciona evidencia",str(Path.home()),"Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if path: self.doc_path=path; self.doc_label.setText(path)

        def _proveedor_text(self) -> str:
            if self.stage_key == STAGE_ADUANA:
                return f"{self.pais.currentText()} · {self.aduana.currentText()}"
            if self.proveedor_widget:
                return self.proveedor_widget.currentText()
            return ""

        def save(self):
            fecha = self.fecha_evento.date().toPython().isoformat()
            # VENDIDO se cierra el mismo día; las demás etapas quedan como etapa actual abierta.
            fecha_fin = fecha if self.stage_key == STAGE_VENDIDO else None
            data={"fecha_inicio": fecha, "fecha_fin": fecha_fin, "costo_usd":self.costo.value() if hasattr(self,'costo') else 0, "proveedor":self._proveedor_text(), "comentario":self.comentario.toPlainText()}
            if hasattr(self,"precio_venta"):
                total_con_final = vehicle_total_cost(self.vehicle) - float(vehicle_stage(self.vehicle, self.stage_key).get("costo_usd") or 0) + (self.costo.value() if hasattr(self, "costo") else 0)
                if not user_can_override_flow(self.user):
                    if self.precio_min.value() < total_con_final:
                        QMessageBox.warning(self,"Validación",f"El precio mínimo no puede ser menor al gasto total {_fmt_usd(total_con_final)}."); return
                    if self.precio_venta.value() and self.precio_venta.value() < total_con_final:
                        QMessageBox.warning(self,"Validación",f"El precio de venta no puede ser menor al gasto total {_fmt_usd(total_con_final)}."); return
                data["precio_venta_usd"]=self.precio_venta.value(); data["precio_minimo_usd"]=self.precio_min.value()
            if hasattr(self,"precio_reserva"):
                data["precio_reserva_usd"]=self.precio_reserva.value(); data["monto_reserva_usd"]=self.monto_reserva.value(); data["cliente"]=self.cliente.text()
            if hasattr(self,"precio_real"):
                data["precio_venta_real_usd"]=self.precio_real.value(); data["cliente"]=self.cliente.text()
            ok,msg=update_vehicle_stage(self.vehicle.get("id"), self.stage_key, data, Path(self.doc_path) if self.doc_path else None, self.user, self.device)
            if not ok: QMessageBox.warning(self,"Validación",msg); return
            QMessageBox.information(self,"Guardado",msg); self.accept()

    class VehicleDetailDialog(QDialog):
        def __init__(self, parent, vehicle_id: str, user: dict, device: DeviceInfo):
            super().__init__(parent); self.vehicle_id=vehicle_id; self.user=user; self.device=device
            self.setWindowTitle("Expediente del vehículo"); self.setMinimumSize(1280,820); self._build(); self.refresh()
        def _build(self):
            lay=QVBoxLayout(self)
            self.header=QLabel(); self.header.setStyleSheet("font-size:24px;font-weight:950;color:#08285a;")
            lay.addWidget(self.header)
            self.timeline=TimelineWidget(); self.timeline.stageDoubleClicked.connect(self.open_stage_detail); lay.addWidget(self.timeline)
            btns=QHBoxLayout()
            self.next_btn=QPushButton("Pasar a siguiente etapa"); self.next_btn.setObjectName("orange"); self.next_btn.clicked.connect(self.advance_next)
            self.edit_btn=QPushButton("Editar etapa seleccionada"); self.edit_btn.setObjectName("ghost"); self.edit_btn.clicked.connect(self.edit_selected_stage)
            bdoc=QPushButton("Abrir comprobante compra"); bdoc.setObjectName("ghost"); bdoc.clicked.connect(self.open_purchase_doc)
            btns.addWidget(self.next_btn); btns.addWidget(self.edit_btn); btns.addWidget(bdoc); btns.addStretch(1); lay.addLayout(btns)
            self.tabs=QTabWidget(); lay.addWidget(self.tabs)
            self.summary=QTextEdit(); self.summary.setReadOnly(True); self.tabs.addTab(self.summary,"Resumen")
            self.stage_table=QTableWidget(0,7); self.stage_table.setHorizontalHeaderLabels(["Etapa","Estado","Inicio","Fin","Días","Costo","Dato etapa"]); self.stage_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.stage_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.stage_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.stage_table.cellDoubleClicked.connect(lambda r,c:self.open_stage_detail(STAGE_ORDER[r] if r<len(STAGE_ORDER) else STAGE_COMPRADO)); self.tabs.addTab(self.stage_table,"Línea de tiempo / etapas")
            self.hist=QTextEdit(); self.hist.setReadOnly(True); self.tabs.addTab(self.hist,"Historial")
        def refresh(self):
            self.vehicle=find_vehicle(self.vehicle_id)
            if not self.vehicle: return
            v=self.vehicle
            self.header.setText(f"{v.get('codigo')} · {v.get('marca')} {v.get('modelo')} {v.get('anio')} · {STAGE_META.get(v.get('estado_actual'),{}).get('label',v.get('estado_actual'))}")
            self.timeline.set_vehicle(v)
            nxt = next_stage_key(v.get("estado_actual", STAGE_COMPRADO))
            self.next_btn.setEnabled(bool(nxt))
            self.next_btn.setText(f"Pasar a siguiente etapa: {STAGE_META[nxt]['label']}" if nxt else "Proceso finalizado")
            resumen=f"""Código: {v.get('codigo')}
Vehículo: {v.get('marca')} {v.get('modelo')} {v.get('anio')}
Millaje: {int(v.get('millaje') or 0):,}
Estado USA: {v.get('estado_usa')}
Subasta: {v.get('subasta')}
Lote: {v.get('lote')}
Precio ganado: {_fmt_usd(v.get('precio_ganado_usd'))}
Fecha compra: {_fmt_date(v.get('fecha_compra'))}
Usuario registró: {v.get('usuario_registro')}

Estado actual: {STAGE_META.get(v.get('estado_actual'),{}).get('label',v.get('estado_actual'))}
Días desde compra: {vehicle_days_from_purchase(v)}
Días en etapa actual: {current_stage_days(v)}
Alerta etapa actual: {stage_alert_level(v)}

Costo total acumulado: {_fmt_usd(vehicle_total_cost(v))}
Precio venta publicado: {_fmt_usd(v.get('precio_venta_usd'))}
Precio mínimo permitido: {_fmt_usd(min_sale_price(v))}
Ganancia esperada: {_fmt_usd(vehicle_expected_profit(v))}
Precio venta real: {_fmt_usd(v.get('precio_venta_real_usd'))}
Cliente: {v.get('cliente','')}

Observaciones:
{v.get('observaciones','')}
"""
            self.summary.setPlainText(resumen)
            self.stage_table.setRowCount(len(STAGES))
            for r,s in enumerate(STAGES):
                st=vehicle_stage(v,s["key"]); vals=[s["label"],st.get("status"),_fmt_date(st.get("fecha_inicio")),_fmt_date(st.get("fecha_fin")),str(stage_duration_days(st)),_fmt_usd(st.get("costo_usd")),st.get("proveedor","")]
                for c,val in enumerate(vals): self.stage_table.setItem(r,c,QTableWidgetItem(str(val)))
            hist="\n".join([f"[{h.get('fecha')}] {h.get('usuario')} · {h.get('accion')} · {h.get('detalle')}" for h in v.get('historial',[])])
            self.hist.setPlainText(hist)
        def advance_next(self):
            v = find_vehicle(self.vehicle_id)
            if not v: return
            nxt = next_stage_key(v.get("estado_actual", STAGE_COMPRADO))
            if not nxt:
                QMessageBox.information(self,"Etapas","El vehículo ya no tiene una etapa siguiente."); return
            dlg=StageUpdateDialog(self,v,nxt,self.user,self.device,mode="advance")
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def selected_stage_key(self) -> str:
            row=self.stage_table.currentRow()
            if row < 0: return self.vehicle.get("estado_actual", STAGE_COMPRADO)
            return STAGE_ORDER[row]
        def edit_selected_stage(self):
            key=self.selected_stage_key()
            if not user_can_override_flow(self.user):
                QMessageBox.warning(self,"Permiso","Solo ADMIN/SUPERVISOR o usuarios con permiso especial pueden editar etapas anteriores o corregir el flujo."); return
            dlg=StageUpdateDialog(self,self.vehicle,key,self.user,self.device,mode="edit")
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def open_stage_detail(self, stage_key: str):
            if stage_key == next_stage_key(self.vehicle.get("estado_actual", STAGE_COMPRADO)):
                self.advance_next(); return
            st=vehicle_stage(self.vehicle, stage_key)
            msg = (f"Etapa: {STAGE_META[stage_key]['label']}\n"
                   f"Estado: {st.get('status')}\n"
                   f"Inicio: {_fmt_date(st.get('fecha_inicio'))}\n"
                   f"Fin: {_fmt_date(st.get('fecha_fin'))}\n"
                   f"Días: {stage_duration_days(st)}\n"
                   f"Costo: {_fmt_usd(st.get('costo_usd'))}\n"
                   f"Dato: {st.get('proveedor','')}\n\n"
                   "Para corregir esta etapa usa el botón 'Editar etapa seleccionada' con permiso especial.")
            QMessageBox.information(self,"Detalle de etapa", msg)
        def open_purchase_doc(self):
            st=vehicle_stage(self.vehicle,STAGE_COMPRADO); p=decrypt_file_to_temp(st.get("documento",""),st.get("documento_nombre","comprobante.pdf"))
            if p: QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
            else: QMessageBox.warning(self,"Documento","No se pudo abrir el comprobante.")

    class ConfigDialog(QDialog):
        def __init__(self, parent, user: dict, device: DeviceInfo):
            super().__init__(parent); self.user=user; self.device=device; self.setWindowTitle("Configuración LYM"); self.setMinimumSize(1080,720); self._build(); self.refresh_all()
        def _build(self):
            lay=QVBoxLayout(self)
            self.tabs=QTabWidget(); lay.addWidget(self.tabs)

            # Usuarios
            self.tab_users=QWidget(); ul=QVBoxLayout(self.tab_users)
            self.users_table=QTableWidget(0,5); self.users_table.setHorizontalHeaderLabels(["Usuario","Rol","Activo","Device","Actualizado"]); self.users_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.users_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.users_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            ul.addWidget(self.users_table)
            ub=QHBoxLayout()
            bcreate=QPushButton("Crear usuario"); bcreate.setObjectName("orange"); bcreate.clicked.connect(self.add_user)
            btoggle=QPushButton("Activar / Desactivar"); btoggle.setObjectName("ghost"); btoggle.clicked.connect(self.toggle_user)
            ub.addWidget(bcreate); ub.addWidget(btoggle); ub.addStretch(1); ul.addLayout(ub)
            self.tabs.addTab(self.tab_users, "Usuarios")

            # Permisos
            self.tab_perms=QWidget(); pl=QVBoxLayout(self.tab_perms)
            top=QHBoxLayout(); self.perm_user_combo=QComboBox(); self.perm_user_combo.currentIndexChanged.connect(self.load_perm_user); top.addWidget(QLabel("Usuario:")); top.addWidget(self.perm_user_combo); top.addStretch(1); pl.addLayout(top)
            self.perm_role=QLabel("Rol:")
            pl.addWidget(self.perm_role)
            pw=QWidget(); self.perm_grid=QGridLayout(pw); self.perm_checks={}
            for i,key in enumerate(ALL_PERMISSION_KEYS):
                chk=QCheckBox(perm_label(key)); self.perm_checks[key]=chk; self.perm_grid.addWidget(chk, i//2, i%2)
            pl.addWidget(pw)
            bsave_perm=QPushButton("Guardar permisos"); bsave_perm.setObjectName("orange"); bsave_perm.clicked.connect(self.save_permissions); pl.addWidget(bsave_perm)
            self.tabs.addTab(self.tab_perms, "Permisos")

            # Auditoría
            self.tab_audit=QWidget(); al=QVBoxLayout(self.tab_audit)
            self.audit_table=QTableWidget(0,5); self.audit_table.setHorizontalHeaderLabels(["Fecha","Usuario","Acción","Código","Detalle"]); self.audit_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.audit_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            al.addWidget(self.audit_table)
            brefa=QPushButton("Actualizar auditoría"); brefa.setObjectName("ghost"); brefa.clicked.connect(self.refresh_audit); al.addWidget(brefa)
            self.tabs.addTab(self.tab_audit, "Auditoría")

            # Backups
            self.tab_backup=QWidget(); bl=QVBoxLayout(self.tab_backup)
            self.backup_table=QTableWidget(0,3); self.backup_table.setHorizontalHeaderLabels(["Archivo","Fecha","Tamaño"]); self.backup_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.backup_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            bl.addWidget(self.backup_table)
            bb=QHBoxLayout(); bcreate_backup=QPushButton("Crear backup ahora"); bcreate_backup.setObjectName("orange"); bcreate_backup.clicked.connect(self.create_backup_now)
            bopen_backup=QPushButton("Abrir carpeta backups"); bopen_backup.setObjectName("ghost"); bopen_backup.clicked.connect(self.open_backup_folder)
            bb.addWidget(bcreate_backup); bb.addWidget(bopen_backup); bb.addStretch(1); bl.addLayout(bb)
            self.tabs.addTab(self.tab_backup, "Respaldos")

            # Catálogos
            self.tab_catalog=QWidget(); cl=QVBoxLayout(self.tab_catalog)
            row=QHBoxLayout(); self.catalog_type=QComboBox();
            self.catalog_map={"Subastas":(F_CATALOG_SUBASTAS, DEFAULT_SUBASTAS), "Marcas":(F_CATALOG_MARCAS, DEFAULT_MARCAS), "Estados USA":(F_CATALOG_ESTADOS_USA, US_STATES), "Talleres":(F_CATALOG_TALLERES, DEFAULT_TALLERES), "Navieras":(F_CATALOG_NAVIERAS, DEFAULT_NAVIERAS), "Aduanas":(F_CATALOG_ADUANAS, DEFAULT_ADUANAS), "Países destino":(F_CATALOG_PAISES, DEFAULT_PAISES_DESTINO), "Proveedores":(F_CATALOG_PROVEEDORES, DEFAULT_PROVEEDORES), "Tarifas":(F_CATALOG_TARIFAS, DEFAULT_TARIFAS)}
            self.catalog_type.addItems(self.catalog_map.keys()); self.catalog_type.currentIndexChanged.connect(self.refresh_catalog_list)
            row.addWidget(QLabel("Catálogo:")); row.addWidget(self.catalog_type); row.addStretch(1); cl.addLayout(row)
            self.catalog_list=QListWidget(); cl.addWidget(self.catalog_list)
            cab=QHBoxLayout(); badd=QPushButton("Agregar valor"); badd.setObjectName("orange"); badd.clicked.connect(self.add_catalog_item); bdel=QPushButton("Borrar seleccionado"); bdel.setObjectName("danger"); bdel.clicked.connect(self.delete_catalog_item); cab.addWidget(badd); cab.addWidget(bdel); cab.addStretch(1); cl.addLayout(cab)
            self.tabs.addTab(self.tab_catalog, "Catálogos")

            # Dispositivos
            self.tab_dev=QWidget(); dl=QVBoxLayout(self.tab_dev)
            self.dev_table=QTableWidget(0,5); self.dev_table.setHorizontalHeaderLabels(["Usuario","Equipo","OS","Device key","Estado"]); self.dev_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.dev_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            dl.addWidget(self.dev_table)
            self.tabs.addTab(self.tab_dev, "Dispositivos")

            self.tabs.addTab(self._build_folder_tab(), "Carpeta del sistema")
            self.tabs.addTab(self._build_visual_tab(), "Visualización y reportes")
            self.tabs.addTab(self._build_security_tab(), "Seguridad")

        def _build_folder_tab(self):
            w=QWidget(); lay=QVBoxLayout(w); lay.setContentsMargins(12,12,12,12)
            self.folder_config_label=QLabel(f"Carpeta actual:\n{get_data_folder()}"); self.folder_config_label.setWordWrap(True); self.folder_config_label.setStyleSheet("background:#f1f5f9;color:#475569;padding:12px;border-radius:8px;")
            lay.addWidget(self.folder_config_label)
            row=QHBoxLayout(); bchange=QPushButton("Cambiar carpeta"); bchange.setObjectName("orange"); bchange.clicked.connect(self.change_folder); bensure=QPushButton("Verificar / crear estructura"); bensure.setObjectName("ghost"); bensure.clicked.connect(self.ensure_structure); row.addWidget(bchange); row.addWidget(bensure); row.addStretch(1); lay.addLayout(row); lay.addStretch(1); return w

        def _build_visual_tab(self):
            w=QWidget(); lay=QVBoxLayout(w); lay.setContentsMargins(12,12,12,12)
            lay.addWidget(QLabel("Preferencias visuales del dashboard y reportería."))
            b=QPushButton("Personalizar dashboard"); b.setObjectName("orange"); b.clicked.connect(lambda: DashboardCustomizeDialog(self).exec())
            lay.addWidget(b); lay.addStretch(1); return w

        def _build_security_tab(self):
            w=QWidget(); lay=QVBoxLayout(w); lay.setContentsMargins(12,12,12,12)
            kp = CryptoManager.key_path()
            info=QLabel(f"Estado del cifrado:\nAlgoritmo: Fernet\nLlave: {kp}\nEstado: {'Presente' if kp and kp.exists() else 'Faltante'}\n\nNo borres ni edites manualmente esta llave.")
            info.setWordWrap(True); info.setStyleSheet("background:#f1f5f9;color:#475569;padding:14px;border-radius:8px;")
            lay.addWidget(info)
            b=QPushButton("Verificar cifrado"); b.setObjectName("orange"); b.clicked.connect(self.verify_crypto); lay.addWidget(b); lay.addStretch(1); return w

        def change_folder(self):
            d=QFileDialog.getExistingDirectory(self,"Selecciona carpeta raíz",str(Path.home()))
            if not d: return
            set_active_folder(d); bootstrap_system(); self.folder_config_label.setText(f"Carpeta actual:\n{get_data_folder()}"); log_audit("CAMBIAR_CARPETA", self.user.get("usuario",""), "", str(get_data_folder())); QMessageBox.information(self,"Carpeta", "Carpeta actualizada. Cierra y vuelve a abrir el sistema para refrescar todo.")

        def ensure_structure(self):
            bootstrap_system(); QMessageBox.information(self,"Estructura", "Estructura verificada/creada correctamente.")

        def verify_crypto(self):
            try:
                f=CryptoManager.fernet(); test=b"LYM_TEST_"+secrets.token_bytes(8); enc=f.encrypt(test); dec=f.decrypt(enc)
                QMessageBox.information(self,"Cifrado", "Cifrado verificado correctamente." if dec==test else "La verificación falló.")
            except Exception as exc:
                QMessageBox.critical(self,"Cifrado", f"Error verificando cifrado: {exc}")

        def _selected_username(self) -> str:
            row = self.users_table.currentRow()
            if row < 0: return ""
            item = self.users_table.item(row, 0)
            return item.text().strip() if item else ""

        def refresh_all(self):
            self.refresh_users(); self.refresh_perm_users(); self.refresh_audit(); self.refresh_backups(); self.refresh_catalog_list(); self.refresh_devices()

        def refresh_users(self):
            users=load_users(); self.users_table.setRowCount(len(users))
            for r,u in enumerate(users):
                vals=[u.get("usuario"),u.get("rol"),"SI" if u.get("activo",True) else "NO",u.get("device_key","")[:10],u.get("fecha_actualizacion","")]
                for c,val in enumerate(vals): self.users_table.setItem(r,c,QTableWidgetItem(str(val)))

        def refresh_perm_users(self):
            current=self.perm_user_combo.currentText()
            self.perm_user_combo.blockSignals(True); self.perm_user_combo.clear(); self.perm_user_combo.addItems([u.get("usuario","") for u in load_users()]); self.perm_user_combo.blockSignals(False)
            idx=self.perm_user_combo.findText(current)
            if idx>=0: self.perm_user_combo.setCurrentIndex(idx)
            self.load_perm_user()

        def load_perm_user(self):
            username=self.perm_user_combo.currentText().strip()
            user=find_user(username)
            if not user:
                self.perm_role.setText("Rol:")
                for chk in self.perm_checks.values(): chk.setChecked(False)
                return
            perms=normalize_permissions(user)
            self.perm_role.setText(f"Rol: {user.get('rol')} · Puedes personalizar permisos por usuario")
            locked = _norm(user.get("rol")) == ROLE_ADMIN
            for key,chk in self.perm_checks.items():
                chk.setChecked(bool(perms.get(key, False))); chk.setEnabled(not locked)

        def save_permissions(self):
            username=self.perm_user_combo.currentText().strip()
            if not username: return
            user=find_user(username)
            if user and _norm(user.get("rol")) == ROLE_ADMIN:
                QMessageBox.information(self,"Permisos","El rol ADMIN siempre mantiene todos los permisos.")
                return
            perms={k: chk.isChecked() for k,chk in self.perm_checks.items()}
            ok,msg=update_user_permissions(username, perms)
            if not ok: QMessageBox.warning(self,"Permisos",msg); return
            QMessageBox.information(self,"Permisos",msg)
            self.refresh_users()

        def add_user(self):
            dlg=CreateUserDialog(self,self.device)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh_all()

        def toggle_user(self):
            username=self._selected_username()
            if not username:
                QMessageBox.information(self,"Usuarios","Selecciona un usuario."); return
            user=find_user(username)
            if not user: return
            ok,msg=set_user_active(username, not bool(user.get("activo", True)))
            if not ok: QMessageBox.warning(self,"Usuarios",msg); return
            self.refresh_users()

        def refresh_audit(self):
            eventos=load_audit_events(500)
            self.audit_table.setRowCount(len(eventos))
            for r,e in enumerate(eventos):
                vals=[e.get("fecha"),e.get("usuario"),e.get("accion"),e.get("codigo"),e.get("detalle")]
                for c,val in enumerate(vals): self.audit_table.setItem(r,c,QTableWidgetItem(str(val)))

        def create_backup_now(self):
            ok,msg,out=create_backup_snapshot(self.user.get("usuario",""), "Manual")
            if not ok: QMessageBox.warning(self,"Backups",msg); return
            QMessageBox.information(self,"Backups",f"{msg}\n{out}")
            self.refresh_backups()

        def refresh_backups(self):
            files=list_backups(); self.backup_table.setRowCount(len(files))
            for r,p in enumerate(files):
                vals=[p.name, datetime.fromtimestamp(p.stat().st_mtime).strftime('%d/%m/%Y %H:%M'), f"{p.stat().st_size/1024:.1f} KB"]
                for c,val in enumerate(vals): self.backup_table.setItem(r,c,QTableWidgetItem(str(val)))

        def open_backup_folder(self):
            df=get_data_folder()
            if not df: return
            folder=df / SUB_RESPALDOS
            folder.mkdir(parents=True, exist_ok=True)
            QDesktopServices.openUrl(QUrl.fromLocalFile(str(folder)))

        def refresh_catalog_list(self):
            label=self.catalog_type.currentText(); filename, default = self.catalog_map[label]
            self.catalog_list.clear(); self.catalog_list.addItems(load_catalog(filename, default))

        def add_catalog_item(self):
            label=self.catalog_type.currentText(); filename, default = self.catalog_map[label]
            text,ok=QInputDialog.getText(self,"Agregar valor",f"Nuevo valor para {label}:")
            if ok and text.strip():
                add_catalog_value(filename, text, default); self.refresh_catalog_list()

        def delete_catalog_item(self):
            item=self.catalog_list.currentItem()
            if not item:
                QMessageBox.information(self,"Catálogos","Selecciona una línea para borrar."); return
            label=self.catalog_type.currentText(); filename, default = self.catalog_map[label]
            if QMessageBox.question(self,"Borrar",f"¿Borrar '{item.text()}' del catálogo {label}?") != QMessageBox.StandardButton.Yes:
                return
            remove_catalog_value(filename, item.text(), default); self.refresh_catalog_list()

        def refresh_devices(self):
            devs=load_registered_devices()
            self.dev_table.setRowCount(len(devs))
            for r,d in enumerate(devs):
                vals=[d.get("usuario"), d.get("computer_name"), f"{d.get('os_name')} {d.get('os_version')}", d.get("device_key"), d.get("estado","")]
                for c,val in enumerate(vals): self.dev_table.setItem(r,c,QTableWidgetItem(str(val)))

    class MainWindow(QMainWindow):
        def __init__(self, user: dict, device: DeviceInfo):
            super().__init__(); self.user=user; self.device=device; self.setWindowTitle(f"{APP_NAME} · {user.get('usuario')}"); self.resize(1440,880); self.setWindowIcon(QIcon(str(ResourceManager.find_logo() or ""))); self._build(); self.refresh_all()
        def _build(self):
            central=QWidget(); self.setCentralWidget(central); main=QHBoxLayout(central); main.setContentsMargins(0,0,0,0)
            side=QFrame(); side.setFixedWidth(250); side.setStyleSheet("QFrame{background:#08285a;} QLabel{color:white;} QPushButton{background:transparent;color:white;text-align:left;padding:13px 16px;border-radius:0;font-weight:850;} QPushButton:hover{background:#0e3a78;border-left:5px solid #f59a13;}")
            sl=QVBoxLayout(side); sl.setContentsMargins(14,22,14,18)
            logo=QLabel(); logo.setAlignment(Qt.AlignmentFlag.AlignCenter); pix=QPixmap(str(ResourceManager.find_logo() or ""))
            if not pix.isNull(): logo.setPixmap(pix.scaled(150,150,Qt.AspectRatioMode.KeepAspectRatio,Qt.TransformationMode.SmoothTransformation))
            else: logo.setText("L&M"); logo.setStyleSheet("font-size:30px;font-weight:950;color:#f59a13;")
            sl.addWidget(logo)
            title=QLabel("LYM AUTO CONTROL"); title.setAlignment(Qt.AlignmentFlag.AlignCenter); title.setStyleSheet("font-weight:950;color:#ffcf7a;font-size:14pt;")
            sub=QLabel(f"Rol: {self.user.get('rol')}"); sub.setAlignment(Qt.AlignmentFlag.AlignCenter); sub.setStyleSheet("font-weight:700;color:#d7e6ff;font-size:10pt;")
            sl.addWidget(title); sl.addWidget(sub)
            self.stack=QStackedWidget(); self.pages=[]
            def add_page(name, widget):
                idx=self.stack.addWidget(widget); self.pages.append(widget); b=QPushButton(name); b.clicked.connect(lambda _,i=idx:self.stack.setCurrentIndex(i)); sl.addWidget(b)
            self.dashboard=DashboardPage(self); self.purchase=PurchasePage(self); self.inventory=InventoryPage(self); self.reporteria=ReporteriaPage(self); self.catalogos=CatalogosPage(self)
            add_page("🏠  Inicio", self.dashboard); add_page("➕  Nueva compra", self.purchase); add_page("🚗  Inventario / CV", self.inventory); add_page("📊  Reportería", self.reporteria); add_page("📚  Catálogos", self.catalogos)
            if user_has_permission(self.user, PERM_CONFIG):
                bconf=QPushButton("⚙️  Configuración"); bconf.clicked.connect(self.open_config); sl.addWidget(bconf)
            sl.addStretch(1); bclose=QPushButton("🚪  Cerrar"); bclose.clicked.connect(self.close); sl.addWidget(bclose)
            main.addWidget(side); main.addWidget(self.stack,1); self.setStatusBar(QStatusBar()); self.statusBar().showMessage(f"Usuario: {self.user.get('usuario')} · Carpeta: {get_data_folder()}")
        def refresh_all(self):
            bootstrap_system(); self.dashboard.refresh(); self.inventory.refresh(); self.purchase.refresh_catalogs(); self.catalogos.refresh_all()
        def open_vehicle_detail(self, vehicle_id: str):
            dlg=VehicleDetailDialog(self,vehicle_id,self.user,self.device); dlg.exec(); self.refresh_all()
        def open_config(self):
            dlg=ConfigDialog(self,self.user,self.device); dlg.exec(); self.refresh_all()


# =============================================================================
# LYM AUTO CONTROL V4 - COMPRA VEHICULAR + COSTOS + COTIZACIONES LEASING
# Bloque agregado sin borrar la base V3.3.1: redefine constantes, funciones y pantallas.
# =============================================================================

APP_VERSION = "2.0.0_LEASING"
F_COTIZACIONES = "cotizaciones.enc"
F_CATALOG_TASAS_LEASING = "tasas_leasing.json"
DEFAULT_TASAS_LEASING = ["1.00", "1.80", "2.50", "3.00"]

# Permisos comerciales nuevos
PERM_VIEW_QUOTES = "view_quotes"
PERM_CREATE_QUOTES = "create_quotes"
PERM_CLOSE_QUOTES = "close_quotes"
PERM_GENERATE_PROPOSALS = "generate_proposals"
for _p in [PERM_VIEW_QUOTES, PERM_CREATE_QUOTES, PERM_CLOSE_QUOTES, PERM_GENERATE_PROPOSALS]:
    if _p not in ALL_PERMISSION_KEYS:
        ALL_PERMISSION_KEYS.append(_p)
PERM_LABELS.update({
    PERM_VIEW_QUOTES: "Ver cotizaciones y clientes",
    PERM_CREATE_QUOTES: "Crear / editar cotizaciones",
    PERM_CLOSE_QUOTES: "Cerrar venta / marcar vendido",
    PERM_GENERATE_PROPOSALS: "Generar propuestas leasing",
})

# Flujo operativo nuevo. Se conservan las constantes legacy para migrar datos viejos.
STAGE_COMPRADO = "COMPRADO"
STAGE_TRASLADO_USA = "TRASLADO_USA"
STAGE_TRANSITO = "TRANSITO"
STAGE_ADUANA = "ADUANA"
STAGE_PREPARACION = "PREPARACION_FINAL"
STAGE_PRECIO_FINAL = "PRECIO_FINAL"
STAGE_DISPONIBLE = "DISPONIBLE_VENTA"
STAGE_ANULADO = "ANULADO"

LEGACY_STAGE_MAP = {
    "MIAMI_NAVIERA": STAGE_TRANSITO,
    "LEGALIZACION": STAGE_PREPARACION,
    "TALLER": STAGE_PREPARACION,
    "RESERVADO": STAGE_DISPONIBLE,
    "VENDIDO": STAGE_DISPONIBLE,
}

STAGES = [
    {"key": STAGE_COMPRADO, "label": "Comprado", "icon": "🧾", "color": "#F59A13"},
    {"key": STAGE_TRASLADO_USA, "label": "Traslado USA", "icon": "🚚", "color": "#8BC6B3"},
    {"key": STAGE_TRANSITO, "label": "Tránsito", "icon": "🚢", "color": "#0E3A78"},
    {"key": STAGE_ADUANA, "label": "Aduana", "icon": "🛃", "color": "#374151"},
    {"key": STAGE_PREPARACION, "label": "Legalización / Taller", "icon": "🛠", "color": "#7C3AED"},
    {"key": STAGE_PRECIO_FINAL, "label": "Precio final", "icon": "💰", "color": "#F97316"},
    {"key": STAGE_DISPONIBLE, "label": "Disponible", "icon": "🏷", "color": "#10B981"},
]
STAGE_ORDER = [s["key"] for s in STAGES]
STAGE_META = {s["key"]: s for s in STAGES}
ALERT_LIMITS_DAYS = {
    STAGE_TRASLADO_USA: 10,
    STAGE_TRANSITO: 35,
    STAGE_ADUANA: 7,
    STAGE_PREPARACION: 20,
    STAGE_PRECIO_FINAL: 5,
    STAGE_DISPONIBLE: 60,
}

COMM_NO_DISPONIBLE = "NO_DISPONIBLE"
COMM_DISPONIBLE = "DISPONIBLE"
COMM_APARTADO = "APARTADO"
COMM_VENDIDO = "VENDIDO"

QUOTE_NUEVA = "NUEVA"
QUOTE_ENVIADA = "COTIZACION_ENVIADA"
QUOTE_SEGUIMIENTO = "EN_SEGUIMIENTO"
QUOTE_NEGOCIACION = "NEGOCIACION"
QUOTE_GANADA = "COMPRA_CERRADA"
QUOTE_PERDIDA = "PERDIDA"
QUOTE_REOFERTAR = "VEHICULO_VENDIDO_REOFERTAR"
QUOTE_STATUS_LABELS = {
    QUOTE_NUEVA: "Nuevo interesado",
    QUOTE_ENVIADA: "Cotización enviada",
    QUOTE_SEGUIMIENTO: "En seguimiento",
    QUOTE_NEGOCIACION: "Negociando",
    QUOTE_GANADA: "Compra cerrada",
    QUOTE_PERDIDA: "Perdida",
    QUOTE_REOFERTAR: "Vehículo vendido / Reofertar",
}

# Guardamos referencias legacy para compatibilidad de datos antiguos.
_LEGACY_default_permissions_for_role = default_permissions_for_role
_LEGACY_ensure_vehicle_runtime_fields = ensure_vehicle_runtime_fields
_LEGACY_vehicle_total_cost = vehicle_total_cost
_LEGACY_stage_index = stage_index
_LEGACY_next_stage_key = next_stage_key


def default_permissions_for_role(role: str) -> dict:
    role_norm = _norm(role)
    base = {k: False for k in ALL_PERMISSION_KEYS}
    if role_norm in (ROLE_ADMIN, ROLE_SUPERVISOR):
        return {k: True for k in ALL_PERMISSION_KEYS}
    if role_norm == ROLE_OPERACIONES:
        for k in [PERM_VIEW_ALL, PERM_CREATE_PURCHASE, PERM_UPDATE_STAGE, PERM_MARK_AVAILABLE, PERM_GENERATE_REPORTS, PERM_VIEW_QUOTES]:
            base[k] = True
    elif role_norm == ROLE_CONTABILIDAD:
        for k in [PERM_VIEW_ALL, PERM_VIEW_COSTS, PERM_ADD_COSTS, PERM_GENERATE_REPORTS, PERM_VIEW_QUOTES]:
            base[k] = True
    elif role_norm == ROLE_VENTAS:
        for k in [PERM_VIEW_ALL, PERM_VIEW_QUOTES, PERM_CREATE_QUOTES, PERM_CLOSE_QUOTES, PERM_GENERATE_PROPOSALS, PERM_GENERATE_REPORTS]:
            base[k] = True
    else:
        base[PERM_VIEW_ALL] = True
        base[PERM_VIEW_QUOTES] = True
    return base


def normalize_permissions(user: dict) -> dict:
    base = default_permissions_for_role(user.get("rol", ROLE_USUARIO))
    saved = user.get("permissions") or {}
    if isinstance(saved, dict):
        for k in ALL_PERMISSION_KEYS:
            if k in saved:
                base[k] = bool(saved[k])
    if _norm(user.get("rol")) == ROLE_ADMIN:
        return {k: True for k in ALL_PERMISSION_KEYS}
    return base


def user_has_permission(user: dict, perm: str) -> bool:
    return bool(normalize_permissions(user).get(perm, False))


def stage_index(stage_key: str) -> int:
    stage_key = LEGACY_STAGE_MAP.get(stage_key, stage_key)
    try:
        return STAGE_ORDER.index(stage_key)
    except ValueError:
        return 0


def next_stage_key(stage_key: str) -> Optional[str]:
    key = LEGACY_STAGE_MAP.get(stage_key, stage_key)
    idx = stage_index(key)
    if idx < len(STAGE_ORDER) - 1:
        return STAGE_ORDER[idx + 1]
    return None


def _quote_file() -> Optional[Path]:
    return datos_file(F_COTIZACIONES)


def load_quotes() -> list[dict]:
    p = _quote_file()
    if not p or not p.exists():
        return []
    data = load_encrypted_json_path(p, default=[])
    if not isinstance(data, list):
        return []
    for q in data:
        ensure_quote_runtime_fields(q)
    return data


def save_quotes(quotes: list[dict]) -> bool:
    p = _quote_file()
    if not p:
        return False
    return save_encrypted_json_path(p, quotes)


def ensure_quote_runtime_fields(q: dict) -> dict:
    q.setdefault("id", uuid.uuid4().hex)
    q.setdefault("cliente", {})
    q.setdefault("vehicle_snapshot", {})
    q.setdefault("leasing", {})
    q.setdefault("legal", {})
    q.setdefault("seguimientos", [])
    q.setdefault("estado", QUOTE_NUEVA)
    q.setdefault("fecha_cotizacion", _today_iso())
    q.setdefault("ultima_gestion", q.get("fecha_cotizacion") or _today_iso())
    q.setdefault("propuestas", [])
    return q


def find_quote(quote_id: str) -> Optional[dict]:
    for q in load_quotes():
        if q.get("id") == quote_id:
            return q
    return None


def quotes_for_vehicle(vehicle_id: str) -> list[dict]:
    return [q for q in load_quotes() if q.get("vehicle_id") == vehicle_id]


def upsert_quote(quote: dict) -> bool:
    quotes = load_quotes()
    found = False
    for i, q in enumerate(quotes):
        if q.get("id") == quote.get("id"):
            quotes[i] = quote
            found = True
            break
    if not found:
        quotes.append(quote)
    return save_quotes(quotes)


def quote_days_without_purchase(q: dict) -> int:
    if q.get("estado") in (QUOTE_GANADA, QUOTE_PERDIDA):
        return 0
    base = _parse_date(q.get("ultima_gestion")) or _parse_date(q.get("fecha_cotizacion")) or date.today()
    return max(0, (date.today() - base).days)


def quote_alert_level(q: dict) -> str:
    if q.get("estado") == QUOTE_GANADA:
        return "CERRADA"
    if q.get("estado") == QUOTE_PERDIDA:
        return "PERDIDA"
    days = quote_days_without_purchase(q)
    if days >= 14:
        return "ROJO"
    if days >= 7:
        return "AMARILLO"
    return "VERDE"


def _quote_alert_color(level: str) -> str:
    return {"VERDE": "#DCFCE7", "AMARILLO": "#FEF9C3", "ROJO": "#FEE2E2", "CERRADA": "#DBEAFE", "PERDIDA": "#F1F5F9"}.get(level, "#FFFFFF")


def store_document_named(src_path: Path, vehicle_code: str, doc_type: str, final_base: str, original_hint: str = "") -> tuple[str, str]:
    df = get_data_folder()
    if df is None:
        return "", ""
    src = Path(src_path)
    if not src.exists():
        return "", ""
    ext = src.suffix or ".pdf"
    year = str(date.today().year)
    base = _safe_filename(final_base or f"{vehicle_code}_{doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    dst = df / SUB_DOCUMENTOS / year / _safe_filename(vehicle_code) / f"{base}{ext}.enc"
    # Evitar sobrescritura si repiten OC
    if dst.exists():
        dst = dst.with_name(f"{base}_{datetime.now().strftime('%H%M%S')}{ext}.enc")
    if not encrypt_file_to(src, dst):
        return "", ""
    return abs_to_rel(dst), (original_hint or f"{base}{ext}")


def _store_cost_doc(src_path: str, vehicle_code: str, doc_type: str, label: str) -> tuple[str, str]:
    if not src_path:
        return "", ""
    p = Path(src_path)
    if not p.exists():
        return "", ""
    return store_document_named(p, vehicle_code, doc_type, f"{vehicle_code}_{doc_type}_{label}_{datetime.now().strftime('%Y%m%d_%H%M%S')}", p.name)


def ensure_vehicle_runtime_fields(vehicle: dict) -> dict:
    vehicle.setdefault("id", uuid.uuid4().hex)
    vehicle.setdefault("codigo", "")
    vehicle.setdefault("historial", [])
    vehicle.setdefault("gastos_extra", [])
    vehicle.setdefault("fotos", [])
    vehicle.setdefault("precio_venta_usd", 0.0)
    vehicle.setdefault("precio_minimo_usd", 0.0)
    vehicle.setdefault("cliente", "")
    current = vehicle.get("estado_actual", STAGE_COMPRADO)
    if current in ("VENDIDO", "RESERVADO"):
        vehicle.setdefault("estado_comercial", COMM_VENDIDO if current == "VENDIDO" else COMM_APARTADO)
        vehicle["estado_actual"] = STAGE_DISPONIBLE
    else:
        vehicle["estado_actual"] = LEGACY_STAGE_MAP.get(current, current if current in STAGE_ORDER else STAGE_COMPRADO)
    if vehicle.get("estado_actual") == STAGE_DISPONIBLE:
        vehicle.setdefault("estado_comercial", COMM_DISPONIBLE)
    else:
        vehicle.setdefault("estado_comercial", COMM_NO_DISPONIBLE)
    etapas = vehicle.get("etapas") if isinstance(vehicle.get("etapas"), dict) else {}
    for s in STAGES:
        etapas.setdefault(s["key"], default_stage_record(s["key"]))
    vehicle["etapas"] = etapas
    vehicle.setdefault("caracteristicas", "")
    vehicle.setdefault("oc_compra_numero", "")
    vehicle.setdefault("oc_compra_documento", "")
    vehicle.setdefault("oc_compra_documento_nombre", "")
    vehicle.setdefault("foto_principal", "")
    vehicle.setdefault("foto_principal_nombre", "")
    vehicle.setdefault("gastos_detallados", [])
    vehicle.setdefault("precio_final", {})
    return vehicle


def vehicle_stage(vehicle: dict, stage_key: str) -> dict:
    ensure_vehicle_runtime_fields(vehicle)
    key = LEGACY_STAGE_MAP.get(stage_key, stage_key)
    return vehicle["etapas"].setdefault(key, default_stage_record(key))


def vehicle_total_cost(vehicle: dict) -> float:
    ensure_vehicle_runtime_fields(vehicle)
    costs = vehicle.get("gastos_detallados")
    if isinstance(costs, list) and costs:
        total = 0.0
        for g in costs:
            try:
                total += float(g.get("monto_usd") or 0)
            except Exception:
                pass
        return round(total, 2)
    # Fallback para registros viejos que todavía dependen de costo por etapa.
    total = 0.0
    for st in vehicle.get("etapas", {}).values():
        try:
            total += float(st.get("costo_usd") or 0)
        except Exception:
            pass
    for g in vehicle.get("gastos_extra", []) or []:
        try:
            total += float(g.get("monto_usd") or 0)
        except Exception:
            pass
    return round(total, 2)


def vehicle_expected_profit(vehicle: dict) -> float:
    try:
        pf = vehicle.get("precio_final") or {}
        venta_neta = float(pf.get("venta_neta_usd") or 0)
        if venta_neta:
            return round(venta_neta - vehicle_total_cost(vehicle), 2)
        return round(float(vehicle.get("precio_venta_usd") or 0) - vehicle_total_cost(vehicle), 2)
    except Exception:
        return 0.0


def min_sale_price(vehicle: dict) -> float:
    return max(vehicle_total_cost(vehicle), float(vehicle.get("precio_minimo_usd") or 0))


def validate_purchase_data(data: dict) -> tuple[bool, str]:
    required = ["marca", "modelo", "anio", "millaje", "estado_usa", "subasta", "lote", "precio_ganado_usd", "fecha_compra"]
    for k in required:
        if data.get(k) in (None, ""):
            return False, f"Campo obligatorio faltante: {k}."
    try:
        anio = int(data.get("anio"))
        if anio < 1980 or anio > date.today().year + 1:
            return False, "Año del vehículo inválido."
    except Exception:
        return False, "Año del vehículo inválido."
    try:
        if int(data.get("millaje") or 0) < 0:
            return False, "El millaje no puede ser negativo."
    except Exception:
        return False, "Millaje inválido."
    try:
        if float(data.get("precio_ganado_usd") or 0) <= 0:
            return False, "El precio ganado debe ser mayor a cero."
    except Exception:
        return False, "Precio ganado inválido."
    fc = _parse_date(data.get("fecha_compra"))
    if not fc:
        return False, "Fecha de compra inválida."
    if fc > date.today():
        return False, "La fecha de compra no puede ser futura."
    return True, "OK"


def create_vehicle_purchase(data: dict, comprobante_src: Path, user: dict, device: DeviceInfo) -> tuple[bool, str, str]:
    ok, msg = validate_purchase_data(data)
    if not ok:
        return False, msg, ""
    if not user_has_permission(user, PERM_CREATE_PURCHASE):
        return False, "No tienes permiso para registrar compras.", ""
    if not comprobante_src or not Path(comprobante_src).exists():
        return False, "Debes subir el comprobante de compra.", ""
    dup = duplicate_lot_exists(data.get("subasta", ""), data.get("lote", ""))
    if dup:
        return False, f"Ya existe una compra con subasta/lote: {dup.get('codigo')}.", ""
    fc = _parse_date(data.get("fecha_compra")) or date.today()
    codigo = next_vehicle_code(fc)
    doc_rel, doc_name = store_document(Path(comprobante_src), codigo, "COMPROBANTE_COMPRA")
    if not doc_rel:
        return False, "No se pudo cifrar/guardar el comprobante de compra.", ""
    oc_rel = oc_name = ""
    oc_num = str(data.get("oc_compra_numero") or "").strip().upper()
    oc_src = data.get("oc_compra_src") or ""
    if oc_num and oc_src and Path(oc_src).exists():
        oc_rel, oc_name = store_document_named(Path(oc_src), codigo, "OC_COMPRA", f"OC_{oc_num}", Path(oc_src).name)
    now = _now_iso()
    vehicle = {
        "id": uuid.uuid4().hex,
        "codigo": codigo,
        "anio_compra": fc.year,
        "marca": _norm(data.get("marca")),
        "modelo": _norm(data.get("modelo")),
        "anio": int(data.get("anio")),
        "millaje": int(data.get("millaje") or 0),
        "color": _norm(data.get("color", "")),
        "tipo": _norm(data.get("tipo", "")),
        "estado_usa": _norm(data.get("estado_usa")),
        "subasta": _norm(data.get("subasta")),
        "lote": str(data.get("lote") or "").strip().upper(),
        "precio_ganado_usd": round(float(data.get("precio_ganado_usd") or 0), 2),
        "fecha_compra": fc.isoformat(),
        "usuario_registro": user.get("usuario", ""),
        "computadora_registro": device.computer_name,
        "fecha_registro": now,
        "fecha_actualizacion": now,
        "estado_actual": STAGE_COMPRADO,
        "estado_comercial": COMM_NO_DISPONIBLE,
        "precio_venta_usd": 0.0,
        "precio_minimo_usd": 0.0,
        "observaciones": str(data.get("observaciones") or "").strip(),
        "caracteristicas": str(data.get("caracteristicas") or "").strip(),
        "oc_compra_numero": oc_num,
        "oc_compra_documento": oc_rel,
        "oc_compra_documento_nombre": oc_name,
        "etapas": {},
        "historial": [],
        "gastos_extra": [],
        "gastos_detallados": [],
        "fotos": [],
        "precio_final": {},
    }
    ensure_vehicle_runtime_fields(vehicle)
    vehicle["gastos_detallados"].append({
        "id": uuid.uuid4().hex,
        "source": "purchase",
        "stage_key": STAGE_COMPRADO,
        "categoria": "COMPRA",
        "subcategoria": "PRECIO_GANADO",
        "fecha": fc.isoformat(),
        "descripcion": f"Precio ganado en {vehicle['subasta']} lote {vehicle['lote']}",
        "monto_usd": vehicle["precio_ganado_usd"],
        "proveedor": vehicle["subasta"],
        "oc_numero": oc_num,
        "oc_documento": oc_rel,
        "oc_documento_nombre": oc_name,
        "comprobante": doc_rel,
        "comprobante_nombre": doc_name,
        "usuario": user.get("usuario", ""),
        "fecha_registro": now,
    })
    comp = vehicle_stage(vehicle, STAGE_COMPRADO)
    comp.update({
        "status": "COMPLETADO",
        "fecha_inicio": fc.isoformat(),
        "fecha_fin": fc.isoformat(),
        "costo_usd": vehicle["precio_ganado_usd"],
        "proveedor": vehicle["subasta"],
        "documento": doc_rel,
        "documento_nombre": doc_name,
        "comentario": "Compra registrada con comprobante y características base.",
        "usuario": user.get("usuario", ""),
        "fecha_actualizacion": now,
        "extra": {"oc_compra_numero": oc_num, "oc_compra_documento": oc_rel},
    })
    vehicle["historial"].append({
        "fecha": now,
        "usuario": user.get("usuario", ""),
        "computadora": device.computer_name,
        "accion": "CREAR_COMPRA",
        "detalle": f"Compra {codigo} · {vehicle['marca']} {vehicle['modelo']} {vehicle['anio']} · lote {vehicle['lote']} · {_fmt_usd(vehicle['precio_ganado_usd'])}",
    })
    if save_vehicle(vehicle):
        log_audit("CREAR_COMPRA", user.get("usuario", ""), codigo, f"{vehicle['marca']} {vehicle['modelo']} {vehicle['anio']}")
        return True, f"Compra creada correctamente: {codigo}", vehicle["id"]
    return False, "No se pudo guardar el registro de compra.", ""


def _sync_stage_cost_items(vehicle: dict, stage_key: str, cost_items: list[dict], user: dict, device: DeviceInfo) -> float:
    ensure_vehicle_runtime_fields(vehicle)
    source = f"stage:{stage_key}"
    keep = [g for g in vehicle.get("gastos_detallados", []) if g.get("source") != source]
    total = 0.0
    for raw in cost_items or []:
        try:
            monto = round(float(raw.get("monto_usd") or 0), 2)
        except Exception:
            monto = 0.0
        if monto <= 0:
            continue
        label = _safe_filename(f"{raw.get('categoria','GASTO')}_{raw.get('subcategoria','DETALLE')}")
        comp_rel = comp_name = oc_rel = oc_name = ""
        if raw.get("comprobante_src"):
            comp_rel, comp_name = _store_cost_doc(raw.get("comprobante_src"), vehicle.get("codigo", "VEHICULO"), "COMPROBANTE", label)
        if raw.get("oc_src"):
            oc_rel, oc_name = _store_cost_doc(raw.get("oc_src"), vehicle.get("codigo", "VEHICULO"), "OC", label)
        item = {
            "id": raw.get("id") or uuid.uuid4().hex,
            "source": source,
            "stage_key": stage_key,
            "categoria": _norm(raw.get("categoria") or stage_key),
            "subcategoria": _norm(raw.get("subcategoria") or "GENERAL"),
            "fecha": (_parse_date(raw.get("fecha")) or date.today()).isoformat(),
            "descripcion": str(raw.get("descripcion") or "").strip(),
            "monto_usd": monto,
            "proveedor": str(raw.get("proveedor") or "").strip(),
            "oc_numero": str(raw.get("oc_numero") or "").strip().upper(),
            "oc_documento": oc_rel or raw.get("oc_documento", ""),
            "oc_documento_nombre": oc_name or raw.get("oc_documento_nombre", ""),
            "comprobante": comp_rel or raw.get("comprobante", ""),
            "comprobante_nombre": comp_name or raw.get("comprobante_nombre", ""),
            "usuario": user.get("usuario", ""),
            "fecha_registro": _now_iso(),
        }
        keep.append(item)
        total += monto
    vehicle["gastos_detallados"] = keep
    return round(total, 2)


def update_vehicle_stage(vehicle_id: str, stage_key: str, data: dict, document_src: Optional[Path], user: dict, device: DeviceInfo) -> tuple[bool, str]:
    stage_key = LEGACY_STAGE_MAP.get(stage_key, stage_key)
    if not user_has_permission(user, PERM_UPDATE_STAGE):
        return False, "No tienes permiso para actualizar etapas."
    if stage_key == STAGE_DISPONIBLE and not user_has_permission(user, PERM_MARK_AVAILABLE):
        return False, "No tienes permiso para marcar disponible para venta."
    vehicle = find_vehicle(vehicle_id)
    if not vehicle:
        return False, "Vehículo no encontrado."
    ensure_vehicle_runtime_fields(vehicle)
    current_saved_key = LEGACY_STAGE_MAP.get(vehicle.get("estado_actual", STAGE_COMPRADO), vehicle.get("estado_actual", STAGE_COMPRADO))
    if vehicle.get("estado_comercial") == COMM_VENDIDO:
        return False, "Este vehículo ya fue vendido. El flujo operativo está cerrado y no se pueden editar etapas."
    if current_saved_key == STAGE_DISPONIBLE and stage_key != STAGE_PRECIO_FINAL:
        return False, "Este vehículo ya está disponible para la venta. La etapa operativa queda cerrada; solo se permite recalcular precio desde el módulo autorizado."
    if stage_key not in STAGE_ORDER:
        return False, "Etapa inválida."
    current_key = LEGACY_STAGE_MAP.get(vehicle.get("estado_actual", STAGE_COMPRADO), vehicle.get("estado_actual", STAGE_COMPRADO))
    current_idx = stage_index(current_key)
    target_idx = stage_index(stage_key)
    if target_idx > current_idx + 1 and not user_can_override_flow(user):
        return False, "No puedes saltar etapas del flujo operativo."
    ini = _parse_date(data.get("fecha_inicio"))
    if not ini:
        return False, "Fecha de inicio inválida."
    fin = _parse_date(data.get("fecha_fin")) if data.get("fecha_fin") else None
    ok, msg = validate_stage_dates(vehicle, stage_key, ini, fin, user)
    if not ok:
        return False, msg
    cost_items = data.get("cost_items") if isinstance(data.get("cost_items"), list) else None
    if cost_items is not None:
        costo = _sync_stage_cost_items(vehicle, stage_key, cost_items, user, device)
    else:
        try:
            costo = round(float(data.get("costo_usd") or 0), 2)
            if costo < 0:
                return False, "El costo no puede ser negativo."
        except Exception:
            return False, "Costo inválido."
    doc_rel = doc_name = ""
    if document_src and Path(document_src).exists():
        doc_rel, doc_name = store_document(Path(document_src), vehicle.get("codigo", "VEHICULO"), stage_key)
        if not doc_rel:
            return False, "No se pudo guardar el documento/evidencia."
    if target_idx > current_idx:
        prev = vehicle_stage(vehicle, current_key)
        if prev.get("fecha_inicio") and not prev.get("fecha_fin") and current_key != stage_key:
            prev["fecha_fin"] = ini.isoformat()
            prev["status"] = "COMPLETADO"
            prev["fecha_actualizacion"] = _now_iso()
    st = vehicle_stage(vehicle, stage_key)
    prev_status = st.get("status")
    st.update({
        "status": "COMPLETADO" if fin else "EN PROCESO",
        "fecha_inicio": ini.isoformat(),
        "fecha_fin": fin.isoformat() if fin else None,
        "costo_usd": costo,
        "proveedor": str(data.get("proveedor") or "").strip(),
        "comentario": str(data.get("comentario") or "").strip(),
        "usuario": user.get("usuario", ""),
        "fecha_actualizacion": _now_iso(),
        "extra": data.get("extra") if isinstance(data.get("extra"), dict) else st.get("extra", {}),
    })
    if doc_rel:
        st["documento"] = doc_rel
        st["documento_nombre"] = doc_name
    if stage_key == STAGE_PRECIO_FINAL:
        pf = data.get("precio_final") if isinstance(data.get("precio_final"), dict) else {}
        if pf:
            vehicle["precio_final"] = pf
            vehicle["precio_venta_usd"] = round(float(pf.get("precio_venta_cliente_usd") or 0), 2)
            vehicle["precio_minimo_usd"] = round(float(pf.get("precio_minimo_usd") or vehicle_total_cost(vehicle)), 2)
        if data.get("foto_principal_src") and Path(data.get("foto_principal_src")).exists():
            frel, fname = store_document_named(Path(data.get("foto_principal_src")), vehicle.get("codigo", "VEHICULO"), "FOTO_PROPUESTA", f"{vehicle.get('codigo')}_FOTO_PRINCIPAL", Path(data.get("foto_principal_src")).name)
            if frel:
                vehicle["foto_principal"] = frel
                vehicle["foto_principal_nombre"] = fname
    if stage_key == STAGE_DISPONIBLE:
        if float(vehicle.get("precio_venta_usd") or 0) <= 0:
            return False, "Antes de marcar disponible debes definir el precio final del cliente."
        vehicle["estado_comercial"] = COMM_DISPONIBLE
    if target_idx >= current_idx or user_has_permission(user, PERM_EDIT_CRITICAL):
        vehicle["estado_actual"] = stage_key
    vehicle["fecha_actualizacion"] = _now_iso()
    vehicle.setdefault("historial", []).append({
        "fecha": _now_iso(),
        "usuario": user.get("usuario", ""),
        "computadora": device.computer_name,
        "accion": "ACTUALIZAR_ETAPA",
        "detalle": f"{STAGE_META[stage_key]['label']} · inicio={_fmt_date(ini)} · fin={_fmt_date(fin)} · costo={_fmt_usd(costo)} · estado_ant={prev_status}",
    })
    if save_vehicle(vehicle):
        log_audit("ACTUALIZAR_ETAPA", user.get("usuario", ""), vehicle.get("codigo", ""), STAGE_META[stage_key]["label"])
        return True, "Etapa actualizada correctamente."
    return False, "No se pudo guardar la etapa."


def can_advance_from_stage(vehicle: dict, stage_key: str) -> tuple[bool, str]:
    ensure_vehicle_runtime_fields(vehicle)
    st = vehicle_stage(vehicle, stage_key)
    extra = st.get("extra") if isinstance(st.get("extra"), dict) else {}
    if stage_key == STAGE_COMPRADO:
        return True, "OK"
    if not st.get("fecha_inicio"):
        return False, "Primero debes guardar datos de la etapa actual."
    if stage_key == STAGE_TRASLADO_USA:
        if not extra.get("fecha_llegada_yarda") or not extra.get("naviera_entregada"):
            return False, "Antes de avanzar debes registrar la naviera y la fecha de llegada a yarda de consolidación."
    if stage_key == STAGE_TRANSITO:
        if not extra.get("fecha_salida_naviera"):
            return False, "Antes de avanzar debes registrar la fecha de salida de la naviera/barco."
    if stage_key == STAGE_ADUANA:
        if not extra.get("fecha_liberacion_aduana"):
            return False, "Antes de avanzar debes registrar la fecha de liberación de aduana."
    if stage_key == STAGE_PREPARACION:
        if not extra.get("legalizacion_fin") or not extra.get("taller_salida"):
            return False, "Antes de avanzar deben estar completas Legalización y Taller."
    if stage_key == STAGE_PRECIO_FINAL:
        if float(vehicle.get("precio_venta_usd") or 0) <= 0:
            return False, "Antes de avanzar debes guardar el precio final del cliente."
        if not vehicle.get("foto_principal"):
            return False, "Antes de avanzar debes subir la foto principal que usará la propuesta."
    return True, "OK"


def calculate_leasing(precio_vehiculo: float, ingreso_cliente: float, pago_inicial: float, plazo_meses: int, tasa_mensual_pct: float, seguro: float, gps: float, iva_pct: float = 13.0) -> dict:
    precio_vehiculo = round(float(precio_vehiculo or 0), 2)
    ingreso_cliente = round(float(ingreso_cliente or 0), 2)
    pago_inicial = round(float(pago_inicial or 0), 2)
    plazo_meses = int(plazo_meses or 0)
    tasa = float(tasa_mensual_pct or 0) / 100.0
    seguro = round(float(seguro or 0), 2)
    gps = round(float(gps or 0), 2)
    iva_pct = float(iva_pct or 13.0)
    monto_leasing = max(0.0, round(precio_vehiculo - pago_inicial, 2))
    if plazo_meses <= 0:
        cuota_base = 0.0
    elif tasa == 0:
        cuota_base = monto_leasing / plazo_meses
    else:
        cuota_base = monto_leasing * tasa * (1 + tasa) ** plazo_meses / (((1 + tasa) ** plazo_meses) - 1)
    cuota_base = round(cuota_base, 2)
    cuota_total = round(cuota_base + seguro + gps, 2)
    iva_cuota = round(cuota_base * (iva_pct / 100.0), 2)
    cuota_total_iva = round(cuota_total + iva_cuota, 2)
    pct_ingreso = round((cuota_total_iva / ingreso_cliente * 100.0), 2) if ingreso_cliente else 0.0
    if pct_ingreso <= 30:
        riesgo = "VERDE"
        riesgo_texto = "Riesgo bajo"
    elif pct_ingreso <= 40:
        riesgo = "AMARILLO"
        riesgo_texto = "Riesgo medio"
    else:
        riesgo = "ROJO"
        riesgo_texto = "Riesgo alto"
    return {
        "precio_vehiculo": precio_vehiculo,
        "ingreso_cliente": ingreso_cliente,
        "pago_inicial": pago_inicial,
        "plazo_meses": plazo_meses,
        "tasa_mensual_pct": round(float(tasa_mensual_pct or 0), 4),
        "seguro_mensual": seguro,
        "gps_mensual": gps,
        "iva_pct": iva_pct,
        "monto_leasing": monto_leasing,
        "cuota_base": cuota_base,
        "cuota_total_sin_iva": cuota_total,
        "iva_sobre_cuota": iva_cuota,
        "cuota_total_con_iva": cuota_total_iva,
        "pct_ingreso": pct_ingreso,
        "riesgo": riesgo,
        "riesgo_texto": riesgo_texto,
    }


def leasing_table(precio_vehiculo: float, ingreso_cliente: float, pago_inicial: float, tasa_mensual_pct: float, seguro: float, gps: float, iva_pct: float = 13.0) -> list[dict]:
    rows = []
    for i, plazo in enumerate([12, 24, 36, 48, 60, 72, 84], start=1):
        calc = calculate_leasing(precio_vehiculo, ingreso_cliente, pago_inicial, plazo, tasa_mensual_pct, seguro, gps, iva_pct)
        rows.append({"anio": i, "plazo": plazo, **calc})
    return rows


def calculate_legal_fees(valor_vehiculo: float, pct: float = 1.5, base_fija: float = 140.0, tope: float = 365.0) -> dict:
    valor = round(float(valor_vehiculo or 0), 2)
    valor_pct = round(valor * (float(pct or 0) / 100.0), 2)
    subtotal = round(valor_pct + float(base_fija or 0), 2)
    total = round(min(subtotal, float(tope or subtotal)), 2)
    return {
        "valor_leasing_vehiculo": valor,
        "pct": float(pct or 0),
        "valor_pct": valor_pct,
        "base_fija": round(float(base_fija or 0), 2),
        "subtotal": subtotal,
        "tope": round(float(tope or 0), 2),
        "valor_legales_iva_incluido": total,
        "aplica_tope": "SI" if subtotal > float(tope or subtotal) else "NO",
    }


def create_quote(data: dict, user: dict, device: DeviceInfo, quote_id: str = "") -> tuple[bool, str, str]:
    if not user_has_permission(user, PERM_CREATE_QUOTES):
        return False, "No tienes permiso para crear cotizaciones.", ""
    vehicle = find_vehicle(data.get("vehicle_id", ""))
    if not vehicle:
        return False, "Vehículo no encontrado.", ""
    ensure_vehicle_runtime_fields(vehicle)
    if vehicle.get("estado_actual") != STAGE_DISPONIBLE or vehicle.get("estado_comercial") == COMM_VENDIDO:
        return False, "Solo puedes cotizar vehículos disponibles y no vendidos.", ""
    cliente_nombre = str(data.get("cliente_nombre") or "").strip()
    telefono = str(data.get("telefono") or "").strip()
    if not cliente_nombre or not telefono:
        return False, "Nombre y teléfono del cliente son obligatorios.", ""
    def _pick_num(key: str, default: float) -> float:
        value = data.get(key)
        if value is None or value == "":
            return default
        return float(value)
    precio = round(_pick_num("precio_vehiculo", float(vehicle.get("precio_venta_usd") or 0)), 2)
    ingreso = round(_pick_num("ingreso_cliente", 0), 2)
    prima_pct = round(_pick_num("prima_pct", 20), 4)
    comision = round(_pick_num("comision_usd", 100), 2)
    pago_inicial = round(_pick_num("pago_inicial", ((precio * prima_pct / 100.0) + comision)), 2)
    plazo = int(_pick_num("plazo_meses", 60))
    tasa = _pick_num("tasa_mensual_pct", 2.5)
    seguro = round(_pick_num("seguro_mensual", 80), 2)
    gps = round(_pick_num("gps_mensual", 20), 2)
    iva_pct = _pick_num("iva_pct", 13)
    leasing = calculate_leasing(precio, ingreso, pago_inicial, plazo, tasa, seguro, gps, iva_pct)
    legal = calculate_legal_fees(precio)
    quote = find_quote(quote_id) if quote_id else None
    now = _now_iso()
    if not quote:
        quote = {"id": uuid.uuid4().hex, "fecha_creacion": now, "creado_por": user.get("usuario", "")}
    ensure_quote_runtime_fields(quote)
    quote.update({
        "vehicle_id": vehicle.get("id"),
        "vehicle_code": vehicle.get("codigo"),
        "fecha_cotizacion": data.get("fecha_cotizacion") or quote.get("fecha_cotizacion") or _today_iso(),
        "ultima_gestion": data.get("ultima_gestion") or _today_iso(),
        "estado": data.get("estado") or QUOTE_ENVIADA,
        "cliente": {
            "nombre": cliente_nombre,
            "telefono": telefono,
            "correo": str(data.get("correo") or "").strip(),
            "medio_contacto": _norm(data.get("medio_contacto") or "WHATSAPP"),
        },
        "vehicle_snapshot": {
            "codigo": vehicle.get("codigo"),
            "marca": vehicle.get("marca"),
            "modelo": vehicle.get("modelo"),
            "anio": vehicle.get("anio"),
            "millaje": vehicle.get("millaje"),
            "color": vehicle.get("color"),
            "precio_venta_usd": precio,
            "foto_principal": vehicle.get("foto_principal"),
            "foto_principal_nombre": vehicle.get("foto_principal_nombre"),
            "caracteristicas": vehicle.get("caracteristicas", ""),
        },
        "leasing": {**leasing, "prima_pct": prima_pct, "comision_usd": comision},
        "legal": legal,
        "fecha_actualizacion": now,
        "actualizado_por": user.get("usuario", ""),
        "comentario": str(data.get("comentario") or quote.get("comentario") or "").strip(),
    })
    quote.setdefault("seguimientos", []).append({
        "fecha": now,
        "usuario": user.get("usuario", ""),
        "accion": "COTIZACION_GUARDADA",
        "comentario": f"Cotización para {vehicle.get('codigo')} · cuota {_fmt_usd(leasing['cuota_total_con_iva'])}",
    })
    if upsert_quote(quote):
        log_audit("GUARDAR_COTIZACION", user.get("usuario", ""), vehicle.get("codigo", ""), f"Cliente {cliente_nombre}")
        return True, "Cotización guardada correctamente.", quote.get("id")
    return False, "No se pudo guardar la cotización.", ""


def add_quote_followup(quote_id: str, comentario: str, estado: str, user: dict) -> tuple[bool, str]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada."
    ensure_quote_runtime_fields(quote)
    now = _now_iso()
    quote["estado"] = estado or quote.get("estado", QUOTE_SEGUIMIENTO)
    quote["ultima_gestion"] = date.today().isoformat()
    quote["fecha_actualizacion"] = now
    quote.setdefault("seguimientos", []).append({"fecha": now, "usuario": user.get("usuario", ""), "accion": "SEGUIMIENTO", "comentario": comentario})
    if upsert_quote(quote):
        return True, "Seguimiento guardado."
    return False, "No se pudo guardar el seguimiento."


def mark_quote_won_and_vehicle_sold(quote_id: str, user: dict, device: DeviceInfo) -> tuple[bool, str]:
    if not user_has_permission(user, PERM_CLOSE_QUOTES):
        return False, "No tienes permiso para cerrar ventas."
    quotes = load_quotes()
    idx = next((i for i, q in enumerate(quotes) if q.get("id") == quote_id), -1)
    if idx < 0:
        return False, "Cotización no encontrada."
    quote = quotes[idx]
    vehicle = find_vehicle(quote.get("vehicle_id", ""))
    if not vehicle:
        return False, "Vehículo no encontrado."
    ensure_vehicle_runtime_fields(vehicle)
    now = _now_iso()
    quote["estado"] = QUOTE_GANADA
    quote["fecha_cierre"] = date.today().isoformat()
    quote["ultima_gestion"] = date.today().isoformat()
    quote.setdefault("seguimientos", []).append({"fecha": now, "usuario": user.get("usuario", ""), "accion": "VENTA_CERRADA", "comentario": "Cliente compró el vehículo."})
    quotes[idx] = quote
    for i, q in enumerate(quotes):
        if q.get("vehicle_id") == vehicle.get("id") and q.get("id") != quote_id and q.get("estado") not in (QUOTE_GANADA, QUOTE_PERDIDA):
            q["estado"] = QUOTE_REOFERTAR
            q["ultima_gestion"] = date.today().isoformat()
            q.setdefault("seguimientos", []).append({"fecha": now, "usuario": user.get("usuario", ""), "accion": "VEHICULO_VENDIDO", "comentario": "El carro cotizado se vendió. Cliente queda para reofertar otro disponible."})
            quotes[i] = q
    vehicle["estado_comercial"] = COMM_VENDIDO
    vehicle["fecha_venta"] = date.today().isoformat()
    vehicle["venta_quote_id"] = quote_id
    vehicle["cliente"] = quote.get("cliente", {}).get("nombre", "")
    vehicle["precio_venta_real_usd"] = quote.get("leasing", {}).get("precio_vehiculo") or vehicle.get("precio_venta_usd")
    vehicle.setdefault("historial", []).append({"fecha": now, "usuario": user.get("usuario", ""), "computadora": device.computer_name, "accion": "VENTA_CERRADA", "detalle": f"Venta cerrada por cotización {quote_id} · {vehicle.get('cliente')}"})
    ok1 = save_vehicle(vehicle)
    ok2 = save_quotes(quotes)
    if ok1 and ok2:
        log_audit("VENTA_CERRADA", user.get("usuario", ""), vehicle.get("codigo", ""), vehicle.get("cliente", ""))
        return True, "Venta cerrada. El vehículo quedó vendido y las demás cotizaciones quedaron para reofertar."
    return False, "No se pudo cerrar la venta correctamente."


def generate_quote_proposal_html(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada.", None
    df = get_data_folder()
    if df is None:
        return False, "Carpeta del sistema no disponible.", None
    out_dir = df / SUB_REPORTES / "PROPUESTAS"
    out_dir.mkdir(parents=True, exist_ok=True)
    ensure_quote_runtime_fields(quote)
    cliente = quote.get("cliente", {})
    snap = quote.get("vehicle_snapshot", {})
    leasing = quote.get("leasing", {})
    legal = quote.get("legal", {})
    logo_uri = ResourceManager.logo_data_uri()
    caracteristicas = [x.strip() for x in str(snap.get("caracteristicas") or "").replace(";", "\n").splitlines() if x.strip()]
    if not caracteristicas:
        caracteristicas = ["Transmisión automática", "Aire acondicionado", "Vehículo revisado por L&M Inversiones"]
    bullets = "".join(f"<li>{html.escape(x)}</li>" for x in caracteristicas[:18])
    rows = "".join(f"<tr><td>{r['plazo']}</td><td>{_fmt_usd(r['cuota_base'])}</td><td>{_fmt_usd(r['cuota_total_sin_iva'])}</td><td>{_fmt_usd(r['cuota_total_con_iva'])}</td></tr>" for r in leasing_table(leasing.get("precio_vehiculo", 0), leasing.get("ingreso_cliente", 0), leasing.get("pago_inicial", 0), leasing.get("tasa_mensual_pct", 0), leasing.get("seguro_mensual", 0), leasing.get("gps_mensual", 0), leasing.get("iva_pct", 13)))
    html_text = f"""<!doctype html><html lang='es'><head><meta charset='utf-8'><title>Propuesta Leasing {html.escape(snap.get('codigo',''))}</title>
<style>
@page {{ size: letter; margin: 18mm; }}
body{{font-family:Segoe UI,Arial,sans-serif;color:#111827;margin:0;background:#f7f9fc}}.page{{background:white;max-width:900px;margin:18px auto;padding:34px 42px;box-shadow:0 10px 30px #0001}}.top{{display:flex;justify-content:space-between;align-items:flex-start;border-bottom:5px solid #08285a;padding-bottom:12px}}.logo img{{max-width:110px}}h1{{text-align:center;color:#08285a;text-decoration:underline;font-size:22px}}h2{{color:#08285a;margin-bottom:8px}}.vehicle{{text-align:center;font-weight:900;color:#08285a;font-size:20px}}.diamond{{color:#0e7490}}.cols{{display:grid;grid-template-columns:1fr 1fr;gap:18px}}li{{margin:7px 0}}.money{{display:grid;grid-template-columns:1fr 1fr 1fr;gap:10px;margin:18px 0}}.box{{border:1px solid #cbd5e1;padding:10px;border-radius:8px;background:#f8fafc}}.box b{{display:block;color:#08285a;font-size:18px}}table{{border-collapse:collapse;width:100%;margin:12px 0}}th{{background:#08285a;color:white;padding:9px;text-align:left}}td{{border:1px solid #111827;padding:8px}}.note{{color:#dc2626;font-weight:800}}.sign{{margin-top:40px;display:flex;justify-content:space-between}}.conditions li{{font-size:12px;line-height:1.35}}.footer{{text-align:center;color:#64748b;margin-top:28px}}</style></head><body><section class='page'>
<div class='top'><div class='logo'>{'<img src="'+logo_uri+'">' if logo_uri else '<b>L&M INVERSIONES</b>'}</div><div><b>L&amp;M Inversiones, S.A. de C.V.</b><br>Ayudando a lograr tus sueños</div></div>
<p><b>{html.escape(cliente.get('nombre','Cliente'))}</b><br>Presente.</p><h1>Propuesta de Arrendamiento Vehicular</h1>
<p>Por medio de la presente, <b>L&amp;M Inversiones, S.A. de C.V.</b> tiene el agrado de presentarle la propuesta de arrendamiento vehicular modalidad leasing.</p>
<p class='vehicle'><span class='diamond'>◆</span> {html.escape(str(snap.get('anio','')))} {html.escape(snap.get('marca',''))} {html.escape(snap.get('modelo',''))} <span class='diamond'>◆</span></p>
<h2>Características destacadas</h2><div class='cols'><ul>{bullets}</ul><ul><li>Millaje: <b>{int(snap.get('millaje') or 0):,}</b></li><li>Color: <b>{html.escape(str(snap.get('color') or ''))}</b></li><li>Código interno: <b>{html.escape(str(snap.get('codigo') or ''))}</b></li></ul></div>
<p>Compartimos ante usted el detalle de plazo y cuota mensual. El cálculo incluye seguro estimado, servicio GPS e IVA aplicable sobre la cuota base del arrendamiento.</p>
<div class='money'><div class='box'>MONTO LEASING<b>{_fmt_usd(leasing.get('monto_leasing'))}</b></div><div class='box'>Costo legal<b>{_fmt_usd(legal.get('valor_legales_iva_incluido'))}</b></div><div class='box'>Opción de compra<b>{_fmt_usd(500)}</b></div></div>
<table><thead><tr><th>Plazo</th><th>Tasa rentabilidad</th><th>Cuota base</th><th>Cuota total mensual incluye seguro, GPS e IVA</th></tr></thead><tbody><tr><td>{int(leasing.get('plazo_meses') or 0)} meses</td><td>{leasing.get('tasa_mensual_pct',0)}%</td><td>{_fmt_usd(leasing.get('cuota_base'))}</td><td><b>{_fmt_usd(leasing.get('cuota_total_con_iva'))}</b></td></tr></tbody></table>
<p class='note'>La cuota mensual incluye IVA, seguro y GPS. El IVA mensual se aplica únicamente sobre la cuota base del arrendamiento.</p>
<h2>Opciones comparativas</h2><table><thead><tr><th>Plazo</th><th>Cuota financiamiento</th><th>Cuota total</th><th>Cuota total con IVA</th></tr></thead><tbody>{rows}</tbody></table>
<p>Quedamos atentos a cualquier consulta o duda.</p><p>San Salvador, {date.today().strftime('%d/%m/%Y')}</p>
<div class='sign'><div>Atentamente,<br><br><b>L&amp;M Inversiones, S.A. de C.V.</b></div><div>X____________________________<br>Aceptado por cliente</div></div>
<h2>Condiciones y vigencia de la oferta</h2><ul class='conditions'><li>Esta propuesta tiene vigencia de 15 días calendario a partir de la fecha de emisión.</li><li>La propuesta se formalizará mediante contrato de arrendamiento donde se establecerán condiciones de uso, pagos mensuales y opción de adquisición.</li><li>El valor del seguro es estimado y puede variar según aseguradora y uso del vehículo.</li><li>Los gastos administrativos y legales se cancelan al momento de la firma.</li><li>Esta propuesta es informativa y no constituye compromiso contractual hasta su formalización.</li></ul>
<div class='footer'>Ayudando a lograr tus sueños</div></section></body></html>"""
    out = out_dir / f"PROPUESTA_LEASING_{_safe_filename(snap.get('codigo','CV'))}_{_safe_filename(cliente.get('nombre','CLIENTE'))}_{datetime.now().strftime('%Y%m%d_%H%M')}.html"
    out.write_text(html_text, encoding="utf-8")
    quote.setdefault("propuestas", []).append({"fecha": _now_iso(), "usuario": (user or {}).get("usuario", ""), "path": abs_to_rel(out), "nombre": out.name, "tipo": "HTML_IMPRIMIBLE"})
    upsert_quote(quote)
    log_audit("GENERAR_PROPUESTA", (user or {}).get("usuario", ""), quote.get("vehicle_code", ""), out.name)
    return True, "Propuesta generada correctamente.", out


def vehicle_days_from_purchase(vehicle: dict) -> int:
    fc = _parse_date(vehicle.get("fecha_compra"))
    if not fc:
        return 0
    end = _parse_date(vehicle.get("fecha_venta")) if vehicle.get("estado_comercial") == COMM_VENDIDO else date.today()
    return max(0, ((end or date.today()) - fc).days)


def current_stage_days(vehicle: dict) -> int:
    st = vehicle_stage(vehicle, vehicle.get("estado_actual", STAGE_COMPRADO))
    return stage_duration_days(st)


def stage_alert_level(vehicle: dict, stage_key: Optional[str] = None) -> str:
    key = LEGACY_STAGE_MAP.get(stage_key or vehicle.get("estado_actual", STAGE_COMPRADO), stage_key or vehicle.get("estado_actual", STAGE_COMPRADO))
    st = vehicle_stage(vehicle, key)
    if st.get("status") == "COMPLETADO" or key == STAGE_COMPRADO:
        return "OK"
    days = stage_duration_days(st)
    limit = ALERT_LIMITS_DAYS.get(key)
    if not limit:
        return "OK"
    if days >= limit * 1.5:
        return "ROJO"
    if days >= limit:
        return "AMARILLO"
    return "OK"


def compute_kpis(vehicles: list[dict]) -> dict:
    for v in vehicles:
        ensure_vehicle_runtime_fields(v)
    total = len(vehicles)
    vendidos = [v for v in vehicles if v.get("estado_comercial") == COMM_VENDIDO]
    activos = [v for v in vehicles if v.get("estado_actual") != STAGE_ANULADO and v.get("estado_comercial") != COMM_VENDIDO]
    disponibles = [v for v in vehicles if v.get("estado_actual") == STAGE_DISPONIBLE and v.get("estado_comercial") == COMM_DISPONIBLE]
    criticos = [v for v in activos if stage_alert_level(v) == "ROJO" or (v.get("estado_actual") == STAGE_DISPONIBLE and current_stage_days(v) >= 90)]
    capital = sum(vehicle_total_cost(v) for v in activos)
    precio_pub = sum(float(v.get("precio_venta_usd") or 0) for v in disponibles)
    ganancia_esp = sum(vehicle_expected_profit(v) for v in disponibles)
    prom_compra_disp = _avg([
        (_parse_date(vehicle_stage(v, STAGE_DISPONIBLE).get("fecha_inicio")) - _parse_date(v.get("fecha_compra"))).days
        for v in vehicles
        if _parse_date(vehicle_stage(v, STAGE_DISPONIBLE).get("fecha_inicio")) and _parse_date(v.get("fecha_compra"))
    ])
    by_stage = {s["key"]: 0 for s in STAGES}
    for v in vehicles:
        if v.get("estado_comercial") == COMM_VENDIDO:
            continue
        if v.get("estado_actual") in by_stage:
            by_stage[v.get("estado_actual")] += 1
    by_brand: dict[str, int] = {}
    for v in vehicles:
        by_brand[_norm(v.get("marca")) or "SIN MARCA"] = by_brand.get(_norm(v.get("marca")) or "SIN MARCA", 0) + 1
    quotes = load_quotes()
    q_levels = {"VERDE": 0, "AMARILLO": 0, "ROJO": 0}
    for q in quotes:
        lvl = quote_alert_level(q)
        if lvl in q_levels:
            q_levels[lvl] += 1
    return {
        "total": total, "activos": len(activos), "vendidos": len(vendidos),
        "disponibles": len(disponibles), "reservados": sum(1 for v in vehicles if v.get("estado_comercial") == COMM_APARTADO),
        "criticos": len(criticos), "capital": round(capital, 2), "precio_publicado": round(precio_pub, 2),
        "ganancia_esperada": round(ganancia_esp, 2), "prom_compra_disp": round(prom_compra_disp, 1),
        "prom_disp_venta": 0.0, "by_stage": by_stage, "by_brand": by_brand,
        "quotes_total": len(quotes), "quotes_verde": q_levels["VERDE"], "quotes_amarillo": q_levels["AMARILLO"],
        "quotes_rojo": q_levels["ROJO"], "quotes_ganadas": sum(1 for q in quotes if q.get("estado") == QUOTE_GANADA),
        "quotes_reofertar": sum(1 for q in quotes if q.get("estado") == QUOTE_REOFERTAR),
    }


def _add_logo_to_worksheet(ws, start_cell: str = "A1"):
    try:
        from openpyxl.drawing.image import Image as XLImage
        logo = ResourceManager.find_logo()
        if logo and logo.exists():
            img = XLImage(str(logo))
            img.height = 80
            img.width = 80
            ws.add_image(img, start_cell)
            return True
    except Exception:
        return False
    return False


def generate_inventory_excel(vehicles: list[dict], user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    df = get_data_folder()
    if df is None:
        return False, "Carpeta del sistema no disponible.", None
    out_dir = df / SUB_REPORTES
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"LYM_INVENTARIO_COMPRA_VEHICULAR_{datetime.now().strftime('%Y-%m-%d_%H%M')}.xlsx"
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        wb = Workbook()
        ws = wb.active
        ws.title = "Inventario"
        headers = ["Código", "Lote", "Fecha compra", "Marca", "Modelo", "Año", "Millaje", "Subasta", "Estado operativo", "Estado comercial", "Costo total", "Venta cliente", "Ganancia neta base", "Días compra", "Alerta", "OC compra"]
        ws.merge_cells(start_row=1, start_column=1, end_row=3, end_column=2)
        _add_logo_to_worksheet(ws, "A1")
        ws.merge_cells(start_row=1, start_column=3, end_row=1, end_column=len(headers))
        ws.cell(1,3).value = "INVENTARIO VEHICULAR L&M INVERSIONES"
        ws.cell(1,3).font = Font(size=16, bold=True, color="FFFFFF")
        ws.cell(1,3).fill = PatternFill("solid", fgColor="08285A")
        ws.cell(1,3).alignment = Alignment(horizontal="center")
        ws.append([]); ws.append([]); ws.append(headers)
        for cell in ws[4]:
            cell.fill = PatternFill("solid", fgColor="F59A13")
            cell.font = Font(bold=True, color="08285A")
            cell.alignment = Alignment(horizontal="center")
        for v in vehicles:
            ensure_vehicle_runtime_fields(v)
            ws.append([v.get("codigo"), v.get("lote"), _fmt_date(v.get("fecha_compra")), v.get("marca"), v.get("modelo"), v.get("anio"), int(v.get("millaje") or 0), v.get("subasta"), STAGE_META.get(v.get("estado_actual"),{}).get("label",v.get("estado_actual")), v.get("estado_comercial"), vehicle_total_cost(v), float(v.get("precio_venta_usd") or 0), vehicle_expected_profit(v), vehicle_days_from_purchase(v), stage_alert_level(v), v.get("oc_compra_numero", "")])
        for row in ws.iter_rows(min_row=5, min_col=11, max_col=13):
            for cell in row:
                cell.number_format = '$#,##0.00'
        thin = Side(style="thin", color="D9E2EF")
        for row in ws.iter_rows(min_row=4, max_row=ws.max_row, min_col=1, max_col=len(headers)):
            for cell in row:
                cell.border = Border(bottom=thin)
                cell.alignment = Alignment(vertical="center")
        for col in range(1, len(headers)+1):
            ws.column_dimensions[get_column_letter(col)].width = min(28, max(12, len(str(headers[col-1]))+5))
        ws.freeze_panes = "A5"
        wb.save(out)
        log_audit("GENERAR_EXCEL_INVENTARIO", (user or {}).get("usuario", ""), "", out.name)
        return True, "Reporte Excel de inventario generado.", out
    except Exception as exc:
        return False, f"Error generando Excel: {exc}", None


def generate_quotes_excel_report(user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    df = get_data_folder()
    if df is None:
        return False, "Carpeta del sistema no disponible.", None
    out_dir = df / SUB_REPORTES
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"LYM_COTIZACIONES_GENERAL_{datetime.now().strftime('%Y-%m-%d_%H%M')}.xlsx"
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        wb = Workbook()
        ws = wb.active
        ws.title = "Cotizaciones"
        headers = ["Cliente", "Teléfono", "Correo", "Medio", "Vehículo cotizado", "Código", "Fecha cotización", "Última gestión", "Días sin compra", "Color", "Estado", "Precio", "Prima", "Monto leasing", "Tasa", "Plazo", "Cuota final", "Legales"]
        ws.merge_cells(start_row=1, start_column=1, end_row=3, end_column=2)
        _add_logo_to_worksheet(ws, "A1")
        ws.merge_cells(start_row=1, start_column=3, end_row=1, end_column=len(headers))
        ws.cell(1,3).value = "REPORTE GENERAL DE COTIZACIONES L&M INVERSIONES"
        ws.cell(1,3).font = Font(size=16, bold=True, color="FFFFFF")
        ws.cell(1,3).fill = PatternFill("solid", fgColor="08285A")
        ws.cell(1,3).alignment = Alignment(horizontal="center")
        ws.append([]); ws.append([]); ws.append(headers)
        for c in ws[4]:
            c.fill = PatternFill("solid", fgColor="F59A13"); c.font = Font(bold=True, color="08285A"); c.alignment = Alignment(horizontal="center")
        fill_map = {"VERDE": "DCFCE7", "AMARILLO": "FEF9C3", "ROJO": "FEE2E2", "CERRADA": "DBEAFE", "PERDIDA": "F1F5F9"}
        for q in load_quotes():
            ensure_quote_runtime_fields(q); cl=q.get("cliente",{}); snap=q.get("vehicle_snapshot",{}); le=q.get("leasing",{}); lg=q.get("legal",{}); lvl=quote_alert_level(q)
            ws.append([cl.get("nombre"), cl.get("telefono"), cl.get("correo"), cl.get("medio_contacto"), f"{snap.get('anio','')} {snap.get('marca','')} {snap.get('modelo','')}", q.get("vehicle_code"), _fmt_date(q.get("fecha_cotizacion")), _fmt_date(q.get("ultima_gestion")), quote_days_without_purchase(q), lvl, QUOTE_STATUS_LABELS.get(q.get("estado"), q.get("estado")), le.get("precio_vehiculo"), le.get("pago_inicial"), le.get("monto_leasing"), le.get("tasa_mensual_pct"), le.get("plazo_meses"), le.get("cuota_total_con_iva"), lg.get("valor_legales_iva_incluido")])
            for cell in ws[ws.max_row]:
                cell.fill = PatternFill("solid", fgColor=fill_map.get(lvl, "FFFFFF"))
        for row in ws.iter_rows(min_row=5, min_col=12, max_col=18):
            for cell in row:
                if cell.column not in (15, 16):
                    cell.number_format = '$#,##0.00'
        thin = Side(style="thin", color="D9E2EF")
        for row in ws.iter_rows(min_row=4, max_row=ws.max_row, min_col=1, max_col=len(headers)):
            for cell in row:
                cell.border = Border(bottom=thin); cell.alignment = Alignment(vertical="center")
        for col in range(1, len(headers)+1):
            ws.column_dimensions[get_column_letter(col)].width = 18 if col not in (5, 11) else 28
        ws.freeze_panes = "A5"
        wb.save(out)
        log_audit("GENERAR_EXCEL_COTIZACIONES", (user or {}).get("usuario", ""), "", out.name)
        return True, "Reporte Excel de cotizaciones generado.", out
    except Exception as exc:
        return False, f"Error generando Excel de cotizaciones: {exc}", None


def generate_quotes_html_report(user: Optional[dict] = None) -> Optional[Path]:
    df = get_data_folder()
    if df is None:
        return None
    out_dir = df / SUB_REPORTES
    out_dir.mkdir(parents=True, exist_ok=True)
    quotes = load_quotes()
    logo_uri = ResourceManager.logo_data_uri()
    k = compute_kpis(load_vehicles())
    rows = []
    for q in quotes:
        cl=q.get("cliente",{}); snap=q.get("vehicle_snapshot",{}); le=q.get("leasing",{}); lvl=quote_alert_level(q); color=_quote_alert_color(lvl)
        rows.append(f"<tr style='background:{color}'><td><b>{html.escape(cl.get('nombre',''))}</b><br>{html.escape(cl.get('telefono',''))}</td><td>{html.escape(cl.get('medio_contacto',''))}</td><td>{html.escape(str(snap.get('anio','')))} {html.escape(snap.get('marca',''))} {html.escape(snap.get('modelo',''))}<br><b>{html.escape(q.get('vehicle_code',''))}</b></td><td>{_fmt_date(q.get('fecha_cotizacion'))}</td><td>{_fmt_date(q.get('ultima_gestion'))}</td><td>{quote_days_without_purchase(q)}</td><td><b>{lvl}</b></td><td>{html.escape(QUOTE_STATUS_LABELS.get(q.get('estado'), q.get('estado','')))}</td><td>{_fmt_usd(le.get('cuota_total_con_iva'))}</td></tr>")
    html_text = f"""<!doctype html><html lang='es'><head><meta charset='utf-8'><title>LYM Cotizaciones</title><style>
body{{font-family:Segoe UI,Arial;background:#f5f8fc;margin:0;color:#0b172a}}main{{max-width:1240px;margin:auto;padding:34px}}.hero{{display:flex;justify-content:space-between;align-items:center;background:linear-gradient(135deg,#08285a,#0e3a78);color:white;border-radius:28px;padding:28px}}.logo{{background:white;border-radius:20px;padding:12px}}.logo img{{max-width:120px}}.cards{{display:grid;grid-template-columns:repeat(6,1fr);gap:14px;margin:22px 0}}.card{{background:white;border-radius:18px;padding:16px;box-shadow:0 14px 35px #08285a14}}.value{{font-size:28px;font-weight:950;color:#08285a}}table{{width:100%;border-collapse:collapse;background:white;border-radius:18px;overflow:hidden;box-shadow:0 14px 35px #08285a14}}th{{background:#08285a;color:white;padding:10px;text-align:left}}td{{padding:10px;border-bottom:1px solid #dbe4ef}}.legend{{display:flex;gap:10px;flex-wrap:wrap;margin:14px 0}}.pill{{border-radius:99px;padding:7px 12px;font-weight:800}}</style></head><body><main>
<section class='hero'><div><h1>Reporte general de cotizaciones L&M</h1><p>Corte {datetime.now().strftime('%d/%m/%Y %H:%M')} · Seguimiento por cliente aunque el carro cotizado ya se haya vendido.</p></div><div class='logo'>{'<img src="'+logo_uri+'">' if logo_uri else '<b>L&M</b>'}</div></section>
<section class='cards'><div class='card'>Total<div class='value'>{k['quotes_total']}</div></div><div class='card'>Verdes<div class='value'>{k['quotes_verde']}</div></div><div class='card'>Amarillas<div class='value'>{k['quotes_amarillo']}</div></div><div class='card'>Rojas<div class='value'>{k['quotes_rojo']}</div></div><div class='card'>Ganadas<div class='value'>{k['quotes_ganadas']}</div></div><div class='card'>Reofertar<div class='value'>{k['quotes_reofertar']}</div></div></section>
<div class='legend'><span class='pill' style='background:#DCFCE7'>Verde: 0-6 días desde última gestión</span><span class='pill' style='background:#FEF9C3'>Amarillo: 7-13 días</span><span class='pill' style='background:#FEE2E2'>Rojo: 14+ días</span></div>
<table><thead><tr><th>Cliente</th><th>Medio</th><th>Vehículo</th><th>Fecha cotización</th><th>Última gestión</th><th>Días</th><th>Color</th><th>Estado</th><th>Cuota final</th></tr></thead><tbody>{''.join(rows) if rows else '<tr><td colspan="9">No hay cotizaciones registradas.</td></tr>'}</tbody></table>
</main></body></html>"""
    out = out_dir / f"LYM_COTIZACIONES_GENERAL_{datetime.now().strftime('%Y-%m-%d_%H%M')}.html"
    out.write_text(html_text, encoding="utf-8")
    log_audit("GENERAR_HTML_COTIZACIONES", (user or {}).get("usuario", ""), "", out.name)
    return out



# =============================================================================
# V4 UI - pantallas específicas para compra vehicular, precio final y cotizaciones
# =============================================================================
if PYSIDE_OK:
    class PurchasePage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main = main; self.comprobante_path = ""; self.oc_path = ""; self._build()
        def _build(self):
            lay = QVBoxLayout(self); lay.addWidget(make_title("Compra vehicular · Nueva compra", "Etapa 1: compra, características, comprobante y OC inicial del vehículo."))
            scroll = QScrollArea(); scroll.setWidgetResizable(True); content = QWidget(); scroll.setWidget(content); form_lay = QVBoxLayout(content)
            g1 = QGroupBox("Datos del vehículo"); f1 = QFormLayout(g1)
            self.marca = QComboBox(); self.marca.setEditable(True)
            self.modelo = QLineEdit(); self.anio = QSpinBox(); self.anio.setRange(1980, date.today().year + 1); self.anio.setValue(date.today().year)
            self.millaje = QSpinBox(); self.millaje.setRange(0, 999999); self.color = QLineEdit(); self.tipo = QComboBox(); self.tipo.setEditable(True); self.tipo.addItems(["SEDAN", "SUV", "PICKUP", "VAN", "CAMIONETA", "DEPORTIVO", "OTRO"])
            self.caracteristicas = QTextEdit(); self.caracteristicas.setMinimumHeight(125); self.caracteristicas.setPlaceholderText("Pega aquí motor, transmisión, aire acondicionado, sensores, cámara, pantalla, blind spot, rines, etc. No repitas año ni millaje porque ya están arriba.")
            f1.addRow("Marca:", self.marca); f1.addRow("Modelo:", self.modelo); f1.addRow("Año:", self.anio); f1.addRow("Millaje:", self.millaje); f1.addRow("Color:", self.color); f1.addRow("Tipo:", self.tipo); f1.addRow("Características para propuesta:", self.caracteristicas)
            g2 = QGroupBox("Compra / Subasta"); f2 = QFormLayout(g2)
            self.estado_usa = QComboBox(); self.subasta = QComboBox(); self.subasta.setEditable(True)
            self.btn_add_subasta = QPushButton("+ Agregar subasta"); self.btn_add_subasta.setObjectName("ghost"); self.btn_add_subasta.clicked.connect(self.add_subasta)
            subrow = QWidget(); sublay = QHBoxLayout(subrow); sublay.setContentsMargins(0,0,0,0); sublay.addWidget(self.subasta); sublay.addWidget(self.btn_add_subasta)
            self.lote = QLineEdit(); self.precio = MoneyEdit(); self.precio.setRange(0, 9999999)
            self.fecha = configure_date_edit(QDateEdit()); self.fecha.setDate(QDate.currentDate())
            self.usuario_reg = QLineEdit(self.main.user.get("usuario", "")); self.usuario_reg.setReadOnly(True)
            f2.addRow("Estado USA:", self.estado_usa); f2.addRow("Subasta:", subrow); f2.addRow("Número de lote:", self.lote); f2.addRow("Precio ganado USD:", self.precio); f2.addRow("Fecha compra:", self.fecha); f2.addRow("Usuario registra:", self.usuario_reg)
            g3 = QGroupBox("Comprobantes y OC"); f3 = QFormLayout(g3)
            self.comp_label = QLineEdit(); self.comp_label.setReadOnly(True)
            bcomp = QPushButton("Subir comprobante de compra"); bcomp.setObjectName("ghost"); bcomp.clicked.connect(self.pick_comprobante)
            crow = QWidget(); clay = QHBoxLayout(crow); clay.setContentsMargins(0,0,0,0); clay.addWidget(self.comp_label); clay.addWidget(bcomp)
            self.oc_num = QLineEdit(); self.oc_num.setPlaceholderText("Ejemplo: OC-000123")
            self.oc_label = QLineEdit(); self.oc_label.setReadOnly(True)
            boc = QPushButton("Subir PDF de OC"); boc.setObjectName("ghost"); boc.clicked.connect(self.pick_oc)
            ocrow = QWidget(); oclay = QHBoxLayout(ocrow); oclay.setContentsMargins(0,0,0,0); oclay.addWidget(self.oc_label); oclay.addWidget(boc)
            self.obs = QTextEdit(); self.obs.setMinimumHeight(80)
            f3.addRow("Comprobante:", crow); f3.addRow("Número OC:", self.oc_num); f3.addRow("PDF OC:", ocrow); f3.addRow("Observaciones:", self.obs)
            form_lay.addWidget(g1); form_lay.addWidget(g2); form_lay.addWidget(g3)
            btn = QPushButton("Crear compra"); btn.setObjectName("orange"); btn.clicked.connect(self.save_purchase); form_lay.addWidget(btn)
            lay.addWidget(scroll); self.refresh_catalogs()
        def refresh_catalogs(self):
            self.marca.clear(); self.marca.addItems(load_catalog(F_CATALOG_MARCAS, DEFAULT_MARCAS))
            self.subasta.clear(); self.subasta.addItems(load_catalog(F_CATALOG_SUBASTAS, DEFAULT_SUBASTAS))
            self.estado_usa.clear(); self.estado_usa.addItems(load_catalog(F_CATALOG_ESTADOS_USA, US_STATES))
        def add_subasta(self):
            text, ok = QInputDialog.getText(self, "Agregar subasta", "Nombre de subasta:")
            if ok and text.strip(): add_catalog_value(F_CATALOG_SUBASTAS, text, DEFAULT_SUBASTAS); self.refresh_catalogs(); self.subasta.setCurrentText(_norm(text))
        def pick_comprobante(self):
            path, _ = QFileDialog.getOpenFileName(self, "Selecciona comprobante", str(Path.home()), "Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if path: self.comprobante_path = path; self.comp_label.setText(path)
        def pick_oc(self):
            path, _ = QFileDialog.getOpenFileName(self, "Selecciona PDF de OC", str(Path.home()), "PDF (*.pdf);;Todos (*.*)")
            if path: self.oc_path = path; self.oc_label.setText(path)
        def save_purchase(self):
            if not self.comprobante_path:
                QMessageBox.warning(self, "Validación", "Debes subir el comprobante de compra."); return
            if not self.oc_num.text().strip() or not self.oc_path:
                resp = QMessageBox.question(self, "OC faltante", "No agregaste número/PDF de OC. ¿Deseas continuar de todos modos?")
                if resp != QMessageBox.StandardButton.Yes: return
            data = {"marca": self.marca.currentText(), "modelo": self.modelo.text(), "anio": self.anio.value(), "millaje": self.millaje.value(), "color": self.color.text(), "tipo": self.tipo.currentText(), "estado_usa": self.estado_usa.currentText(), "subasta": self.subasta.currentText(), "lote": self.lote.text(), "precio_ganado_usd": self.precio.value(), "fecha_compra": self.fecha.date().toPython().isoformat(), "observaciones": self.obs.toPlainText(), "caracteristicas": self.caracteristicas.toPlainText(), "oc_compra_numero": self.oc_num.text(), "oc_compra_src": self.oc_path}
            ok, msg, vid = create_vehicle_purchase(data, Path(self.comprobante_path) if self.comprobante_path else Path(), self.main.user, self.main.device)
            if not ok: QMessageBox.warning(self, "Validación", msg); return
            QMessageBox.information(self, "Compra creada", msg); self.main.refresh_all(); self.main.open_vehicle_detail(vid)

    class StageUpdateDialog(QDialog):
        def __init__(self, parent, vehicle: dict, stage_key: str, user: dict, device: DeviceInfo, mode: str = "advance"):
            super().__init__(parent); self.vehicle=vehicle; self.stage_key=LEGACY_STAGE_MAP.get(stage_key, stage_key); self.user=user; self.device=device; self.mode=mode; self.doc_path=""; self.cost_lines=[]; self.foto_path=""; self.setWindowTitle(f"{STAGE_META[self.stage_key]['label']} · {vehicle.get('codigo')}"); self.setMinimumSize(880, 720); self._build()
        def _combo(self, values: list[str]) -> QComboBox:
            cb=QComboBox(); cb.setEditable(True); cb.addItems(values); return cb
        def _stage_cost(self, sub: str) -> dict:
            source=f"stage:{self.stage_key}"; subn=_norm(sub)
            for g in self.vehicle.get("gastos_detallados", []) or []:
                if g.get("source")==source and _norm(g.get("subcategoria"))==subn: return g
            return {}
        def _pick_line_doc(self, line: dict, kind: str):
            title = "Selecciona PDF de OC" if kind == "oc" else "Selecciona comprobante"
            path,_=QFileDialog.getOpenFileName(self,title,str(Path.home()),"Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if path:
                line[kind+"_src"] = path
                line[kind+"_label"].setText(path)
        def _add_cost_line(self, form: QFormLayout, label: str, categoria: str, subcategoria: str, required: bool=False, proveedor: str=""):
            old=self._stage_cost(subcategoria); row=QWidget(); lay=QHBoxLayout(row); lay.setContentsMargins(0,0,0,0)
            money=MoneyEdit(); money.setRange(0,9999999); money.setValue(float(old.get("monto_usd") or 0)); oc=QLineEdit(old.get("oc_numero", "")); oc.setPlaceholderText("OC")
            doc_label=QLineEdit(old.get("comprobante_nombre", "")); doc_label.setReadOnly(True); bdoc=QPushButton("Comprobante"); bdoc.setObjectName("ghost")
            oc_label=QLineEdit(old.get("oc_documento_nombre", "")); oc_label.setReadOnly(True); boc=QPushButton("PDF OC"); boc.setObjectName("ghost")
            line={"categoria":categoria,"subcategoria":subcategoria,"required":required,"monto":money,"oc":oc,"doc_src":"","oc_src":"","doc_label":doc_label,"oc_label":oc_label,"old":old,"proveedor":proveedor}
            bdoc.clicked.connect(lambda _,ln=line:self._pick_line_doc(ln,"doc")); boc.clicked.connect(lambda _,ln=line:self._pick_line_doc(ln,"oc"))
            lay.addWidget(money,1); lay.addWidget(oc,1); lay.addWidget(doc_label,2); lay.addWidget(bdoc); lay.addWidget(oc_label,2); lay.addWidget(boc)
            form.addRow(("* " if required else "") + label + ":", row); self.cost_lines.append(line)
        def _collect_costs(self, default_fecha: str, proveedor: str="") -> tuple[list[dict], list[str]]:
            items=[]; missing=[]
            for line in self.cost_lines:
                val=line["monto"].value()
                if line.get("required") and val <= 0: missing.append(line["subcategoria"])
                if val > 0:
                    old=line.get("old") or {}
                    items.append({"id": old.get("id"), "categoria": line["categoria"], "subcategoria": line["subcategoria"], "descripcion": line["subcategoria"].replace("_", " ").title(), "monto_usd": val, "proveedor": proveedor or line.get("proveedor") or "", "oc_numero": line["oc"].text(), "comprobante_src": line.get("doc_src", ""), "oc_src": line.get("oc_src", ""), "comprobante": old.get("comprobante", ""), "comprobante_nombre": old.get("comprobante_nombre", ""), "oc_documento": old.get("oc_documento", ""), "oc_documento_nombre": old.get("oc_documento_nombre", ""), "fecha": default_fecha})
            return items, missing
        def _build(self):
            lay=QVBoxLayout(self); v=self.vehicle; st=vehicle_stage(v,self.stage_key); extra=st.get("extra") if isinstance(st.get("extra"), dict) else {}
            lay.addWidget(QLabel(f"<b>{v.get('codigo')}</b> · {v.get('marca')} {v.get('modelo')} {v.get('anio')} · Costo actual: <b>{_fmt_usd(vehicle_total_cost(v))}</b>"))
            scroll=QScrollArea(); scroll.setWidgetResizable(True); content=QWidget(); scroll.setWidget(content); main=QVBoxLayout(content); lay.addWidget(scroll)
            form=QFormLayout(); main.addLayout(form)
            min_date = _stage_min_start_date(v, self.stage_key)
            initial_date = st.get("fecha_inicio") or _stage_default_start_date(v, self.stage_key) or _today_iso()
            self.fecha_evento = configure_date_edit(QDateEdit(), initial_date, min_date)
            locked = self.stage_key==STAGE_TRASLADO_USA and st.get("fecha_inicio") and not user_can_override_flow(self.user)
            if locked: self.fecha_evento.setEnabled(False)
            label_map={STAGE_TRASLADO_USA:"Fecha salida de subasta", STAGE_TRANSITO:"Fecha llegada a naviera/yarda", STAGE_ADUANA:"Fecha llegada a aduana", STAGE_PREPARACION:"Fecha inicio preparación", STAGE_PRECIO_FINAL:"Fecha definición precio", STAGE_DISPONIBLE:"Fecha disponible"}
            form.addRow(label_map.get(self.stage_key,"Fecha"), self.fecha_evento)
            self.comentario=QTextEdit(st.get("comentario") or ""); self.comentario.setMinimumHeight(70)
            self.proveedor_widget=None; self.extra_widgets={}
            if self.stage_key == STAGE_TRASLADO_USA:
                self.transportista=self._combo(load_catalog(F_CATALOG_TRANSPORTISTAS_USA, DEFAULT_TRANSPORTISTAS_USA)); self.transportista.setCurrentText(st.get("proveedor") or "")
                self.naviera_entregada=self._combo(load_catalog(F_CATALOG_NAVIERAS, DEFAULT_NAVIERAS)); self.naviera_entregada.setCurrentText(extra.get("naviera_entregada") or "")
                self.fecha_llegada_yarda=OptionalDateEdit(extra.get("fecha_llegada_yarda") or "", self._date(self.fecha_evento), "Ya llegó")
                form.addRow("Transportista USA:", self.transportista); form.addRow("Naviera entregada:", self.naviera_entregada); form.addRow("Fecha llegada a yarda de consolidación:", self.fecha_llegada_yarda)
                self._add_cost_line(form,"Grúa / traslado USA","TRASLADO_USA","GRUA_TRASLADO_USA",False)
            elif self.stage_key == STAGE_TRANSITO:
                prev=vehicle_stage(v, STAGE_TRASLADO_USA); pextra=prev.get("extra") if isinstance(prev.get("extra"),dict) else {}
                self.naviera=self._combo(load_catalog(F_CATALOG_NAVIERAS, DEFAULT_NAVIERAS)); self.naviera.setCurrentText(extra.get("naviera") or pextra.get("naviera_entregada") or "")
                self.fecha_salida_naviera=OptionalDateEdit(extra.get("fecha_salida_naviera") or "", self._date(self.fecha_evento), "Ya salió")
                self.motivo_extra=QTextEdit(extra.get("motivo_extra") or ""); self.motivo_extra.setMinimumHeight(60)
                form.addRow("Naviera:", self.naviera); form.addRow("Fecha salida de naviera / barco:", self.fecha_salida_naviera); form.addRow("Motivo costo extra:", self.motivo_extra)
                self._add_cost_line(form,"Costo extra tránsito","TRANSITO","COSTO_EXTRA_TRANSITO",False)
            elif self.stage_key == STAGE_ADUANA:
                self.pais=self._combo(load_catalog(F_CATALOG_PAISES, DEFAULT_PAISES_DESTINO)); self.pais.setCurrentText(extra.get("pais") or "EL SALVADOR")
                self.aduana=self._combo(aduanas_for_country(self.pais.currentText())); self.aduana.setCurrentText(extra.get("aduana") or "")
                self.pais.currentTextChanged.connect(self._refresh_aduanas)
                self.fecha_liberacion=OptionalDateEdit(extra.get("fecha_liberacion_aduana") or "", self._date(self.fecha_evento), "Liberada")
                form.addRow("País destino:", self.pais); form.addRow("Aduana:", self.aduana); form.addRow("Fecha liberación aduana:", self.fecha_liberacion)
                self._add_cost_line(form,"Naviera · Grúa interna","ADUANA","NAVIERA_GRUA_INTERNA",True)
                self._add_cost_line(form,"Naviera · Flete","ADUANA","NAVIERA_FLETE",True)
                self._add_cost_line(form,"Naviera · BL","ADUANA","NAVIERA_BL",True)
                self._add_cost_line(form,"Pago de impuestos","ADUANA","IMPUESTOS_ADUANA",True)
                self._add_cost_line(form,"Almacenamiento","ADUANA","ALMACENAMIENTO_ADUANA",False)
                self._add_cost_line(form,"Servicio trámite aduanal","ADUANA","TRAMITE_ADUANAL",True)
            elif self.stage_key == STAGE_PREPARACION:
                gleg=QGroupBox("Legalización"); fl=QFormLayout(gleg); main.addWidget(gleg)
                self.emision_pedido=OptionalDateEdit(extra.get("emision_pedido") or "", min_date, "Solicitada")
                self.emision_obtenida=OptionalDateEdit(extra.get("emision_obtenida") or "", min_date, "Obtenida")
                self.cita_pedido=OptionalDateEdit(extra.get("cita_pedido") or "", min_date, "Solicitada")
                self.cita_asignada=OptionalDateEdit(extra.get("cita_asignada") or "", min_date, "Asignada")
                self.placas_ingreso=OptionalDateEdit(extra.get("placas_ingreso") or "", min_date, "Ingresado")
                self.placas_entrega=OptionalDateEdit(extra.get("placas_entrega") or "", min_date, "Entregadas")
                self.legal_fin=configure_date_edit(QDateEdit(), extra.get("legalizacion_fin") or _today_iso(), min_date)
                for lab,w in [("Fecha pedimos emisión",self.emision_pedido),("Fecha obtuvimos emisión",self.emision_obtenida),("Fecha pedimos cita",self.cita_pedido),("Fecha asignada cita",self.cita_asignada),("Fecha ingresamos placas",self.placas_ingreso),("Fecha entregaron placas",self.placas_entrega),("Fin legalización",self.legal_fin)]: fl.addRow(lab+":", w)
                self._add_cost_line(fl,"Emisiones","LEGALIZACION","EMISIONES",False); self._add_cost_line(fl,"Cita","LEGALIZACION","CITA",False); self._add_cost_line(fl,"Placas","LEGALIZACION","PLACAS",False)
                gt=QGroupBox("Taller"); ft=QFormLayout(gt); main.addWidget(gt)
                self.taller=self._combo(load_catalog(F_CATALOG_TALLERES, DEFAULT_TALLERES)); self.taller.setCurrentText(extra.get("taller") or "")
                self.taller_ingreso=configure_date_edit(QDateEdit(), extra.get("taller_ingreso") or _today_iso(), min_date); self.taller_salida=configure_date_edit(QDateEdit(), extra.get("taller_salida") or _today_iso(), min_date)
                self.motivo_taller=QTextEdit(extra.get("motivo_taller") or ""); self.motivo_taller.setMinimumHeight(70)
                ft.addRow("Taller:", self.taller); ft.addRow("Fecha ingreso taller:", self.taller_ingreso); ft.addRow("Fecha salida taller:", self.taller_salida); ft.addRow("Motivo / comentarios:", self.motivo_taller)
                self._add_cost_line(ft,"Enderezado y pintura","TALLER","ENDEREZADO_PINTURA",False); self._add_cost_line(ft,"Servicios mecánicos","TALLER","SERVICIOS_MECANICOS",False); self._add_cost_line(ft,"Grúa local","TALLER","GRUA_LOCAL",False)
                self.repuestos_table=QTableWidget(0,3); self.repuestos_table.setHorizontalHeaderLabels(["Descripción repuesto/servicio", "Valor USD", "OC"]); self.repuestos_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
                for g in self.vehicle.get("gastos_detallados",[]) or []:
                    if g.get("source")==f"stage:{self.stage_key}" and _norm(g.get("subcategoria"))=="REPUESTO_DETALLE":
                        r=self.repuestos_table.rowCount(); self.repuestos_table.insertRow(r); self.repuestos_table.setItem(r,0,QTableWidgetItem(g.get("descripcion",""))); self.repuestos_table.setItem(r,1,QTableWidgetItem(str(g.get("monto_usd",0)))); self.repuestos_table.setItem(r,2,QTableWidgetItem(g.get("oc_numero","")))
                brow=QHBoxLayout(); badd=QPushButton("+ Agregar repuesto"); badd.setObjectName("orange"); bdel=QPushButton("Borrar línea"); bdel.setObjectName("danger"); badd.clicked.connect(self._add_repuesto_row); bdel.clicked.connect(lambda:self.repuestos_table.removeRow(self.repuestos_table.currentRow()) if self.repuestos_table.currentRow()>=0 else None); brow.addWidget(badd); brow.addWidget(bdel); brow.addStretch(1)
                main.addWidget(QLabel("Detalle de repuestos / gastos del taller:")); main.addWidget(self.repuestos_table); main.addLayout(brow)
            elif self.stage_key == STAGE_PRECIO_FINAL:
                self._build_price_section(main, form)
            elif self.stage_key == STAGE_DISPONIBLE:
                info=QLabel(f"Precio cliente final: <b>{_fmt_usd(v.get('precio_venta_usd'))}</b><br>Al confirmar esta etapa el carro caerá automáticamente en Cotizaciones como disponible."); info.setTextFormat(Qt.TextFormat.RichText); info.setStyleSheet("background:#dcfce7;padding:12px;border-radius:10px;color:#14532d;"); main.addWidget(info)
            form.addRow("Comentario:", self.comentario)
            btn=QPushButton("Guardar avance / datos de esta etapa"); btn.setObjectName("orange"); btn.clicked.connect(self.save); lay.addWidget(btn)
        def _refresh_aduanas(self):
            cur=self.aduana.currentText(); self.aduana.clear(); self.aduana.addItems(aduanas_for_country(self.pais.currentText())); idx=self.aduana.findText(cur); self.aduana.setCurrentIndex(idx if idx>=0 else 0)
        def _add_repuesto_row(self):
            r=self.repuestos_table.rowCount(); self.repuestos_table.insertRow(r); self.repuestos_table.setItem(r,0,QTableWidgetItem("")); self.repuestos_table.setItem(r,1,QTableWidgetItem("0")); self.repuestos_table.setItem(r,2,QTableWidgetItem(""))
        def _build_price_section(self, main, form):
            self.cost_preview=QTableWidget(0,4); self.cost_preview.setHorizontalHeaderLabels(["Etapa","Categoría","Descripción","Monto"]); self.cost_preview.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.cost_preview.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            costs=self.vehicle.get("gastos_detallados",[]) or []; self.cost_preview.setRowCount(len(costs))
            for r,g in enumerate(costs):
                vals=[STAGE_META.get(g.get("stage_key"),{}).get("label",g.get("stage_key")), g.get("subcategoria"), g.get("descripcion"), _fmt_usd(g.get("monto_usd"))]
                for c,val in enumerate(vals): self.cost_preview.setItem(r,c,QTableWidgetItem(str(val)))
            main.addWidget(QLabel("Preview de gastos desglosados. Si editas gastos en etapas anteriores, este resumen se actualiza automáticamente.")); main.addWidget(self.cost_preview)
            self.margen=QDoubleSpinBox(); self.margen.setRange(0,100); self.margen.setDecimals(2); self.margen.setSuffix(" %"); self.margen.setValue(float((self.vehicle.get("precio_final") or {}).get("margen_pct") or 15))
            self.iva_duca=MoneyEdit(); self.iva_duca.setRange(0,999999); self.iva_duca.setValue(float((self.vehicle.get("precio_final") or {}).get("iva_duca_usd") or 0))
            self.iva_pct=QDoubleSpinBox(); self.iva_pct.setRange(0,30); self.iva_pct.setDecimals(2); self.iva_pct.setSuffix(" %"); self.iva_pct.setValue(float((self.vehicle.get("precio_final") or {}).get("iva_pct") or 13))
            self.pago_cuenta_pct=QDoubleSpinBox(); self.pago_cuenta_pct.setRange(0,20); self.pago_cuenta_pct.setDecimals(2); self.pago_cuenta_pct.setSuffix(" %"); self.pago_cuenta_pct.setValue(float((self.vehicle.get("precio_final") or {}).get("pago_cuenta_pct") or 1.75))
            self.vts=MoneyEdit(); self.vts.setRange(0,999); self.vts.setValue(float((self.vehicle.get("precio_final") or {}).get("vts_usd") or 2.07))
            self.foto_label=QLineEdit(self.vehicle.get("foto_principal_nombre", "")); self.foto_label.setReadOnly(True); bfoto=QPushButton("Subir foto principal propuesta"); bfoto.setObjectName("ghost"); bfoto.clicked.connect(self._pick_foto)
            frow=QWidget(); fl=QHBoxLayout(frow); fl.setContentsMargins(0,0,0,0); fl.addWidget(self.foto_label); fl.addWidget(bfoto)
            self.calc_lbl=QLabel(); self.calc_lbl.setTextFormat(Qt.TextFormat.RichText); self.calc_lbl.setStyleSheet("background:#fff7ed;border:1px solid #fed7aa;border-radius:10px;padding:12px;color:#08285a;")
            for w in [self.margen,self.iva_pct,self.pago_cuenta_pct]: w.valueChanged.connect(self._refresh_price_calc)
            try: self.iva_duca.textChanged.connect(self._refresh_price_calc); self.vts.textChanged.connect(self._refresh_price_calc)
            except Exception: pass
            form.addRow("Margen deseado:", self.margen); form.addRow("IVA DUCA manual:", self.iva_duca); form.addRow("IVA venta:", self.iva_pct); form.addRow("Pago a cuenta:", self.pago_cuenta_pct); form.addRow("VTS fijo:", self.vts); form.addRow("Foto principal:", frow); main.addWidget(self.calc_lbl); self._refresh_price_calc()
        def _pick_foto(self):
            path,_=QFileDialog.getOpenFileName(self,"Selecciona foto principal",str(Path.home()),"Imágenes (*.png *.jpg *.jpeg *.webp);;Todos (*.*)")
            if path: self.foto_path=path; self.foto_label.setText(path)
        def _price_calc(self) -> dict:
            costo=vehicle_total_cost(self.vehicle)
            iva_pct=self.iva_pct.value()/100.0
            pago_pct=self.pago_cuenta_pct.value()/100.0
            iva_duca=self.iva_duca.value()
            vts=self.vts.value()
            manual=float(self.precio_manual.value() or 0) if hasattr(self, "precio_manual") else 0.0
            if manual > 0:
                precio_cliente=round(manual,2)
                venta_neta=round(max(0.0, (precio_cliente - vts) / (1 + iva_pct)),2) if iva_pct > -1 else max(0.0, precio_cliente-vts)
                iva_venta=round(max(0.0, precio_cliente - vts - venta_neta),2)
                margen_pct=round(((venta_neta / costo) - 1) * 100, 2) if costo else 0.0
            else:
                margen_pct=self.margen.value()
                venta_neta=round(costo*(1+(margen_pct/100.0)),2)
                iva_venta=round(venta_neta*iva_pct,2)
                precio_cliente=round(venta_neta+iva_venta+vts,2)
            iva_pagar=max(0.0, round(iva_venta-iva_duca,2))
            pago_cuenta=round(venta_neta*pago_pct,2)
            utilidad=round(venta_neta-costo,2)
            utilidad_neta_est=round(utilidad-iva_pagar-pago_cuenta-vts,2)
            return {"costo_total_usd":costo,"margen_pct":margen_pct,"venta_neta_usd":venta_neta,"iva_pct":self.iva_pct.value(),"iva_venta_usd":iva_venta,"iva_duca_usd":iva_duca,"iva_pagar_usd":iva_pagar,"pago_cuenta_pct":self.pago_cuenta_pct.value(),"pago_cuenta_usd":pago_cuenta,"vts_usd":vts,"precio_venta_cliente_usd":precio_cliente,"precio_minimo_usd":precio_cliente,"precio_redondeado_manual_usd":manual if manual>0 else 0.0,"utilidad_bruta_usd":utilidad,"utilidad_neta_estimada_usd":utilidad_neta_est}
        def _refresh_price_calc(self):
            pf=self._price_calc(); self.calc_lbl.setText(f"Costo total: <b>{_fmt_usd(pf['costo_total_usd'])}</b><br>Venta neta: <b>{_fmt_usd(pf['venta_neta_usd'])}</b> · IVA venta: <b>{_fmt_usd(pf['iva_venta_usd'])}</b> · IVA a pagar: <b>{_fmt_usd(pf['iva_pagar_usd'])}</b><br>Pago a cuenta: <b>{_fmt_usd(pf['pago_cuenta_usd'])}</b> · VTS: <b>{_fmt_usd(pf['vts_usd'])}</b><br>Precio cliente final: <b style='font-size:18px'>{_fmt_usd(pf['precio_venta_cliente_usd'])}</b><br>Utilidad bruta contra costo: <b>{_fmt_usd(pf['utilidad_bruta_usd'])}</b> · Utilidad neta estimada interna: <b>{_fmt_usd(pf['utilidad_neta_estimada_usd'])}</b>")
        def _date(self, w): return w.date().toPython().isoformat()
        def save(self):
            fecha=self._date(self.fecha_evento); fin=None; proveedor=""; extra={}; cost_items=[]
            if self.stage_key==STAGE_TRASLADO_USA:
                proveedor=self.transportista.currentText(); llegada=self._date(self.fecha_llegada_yarda)
                if _parse_date(llegada) < _parse_date(fecha): QMessageBox.warning(self,"Fechas","La llegada a yarda no puede ser menor que la salida de subasta."); return
                fin=llegada; extra={"transportista":proveedor,"naviera_entregada":self.naviera_entregada.currentText(),"fecha_llegada_yarda":llegada}; cost_items,missing=self._collect_costs(fecha, proveedor)
            elif self.stage_key==STAGE_TRANSITO:
                proveedor=self.naviera.currentText(); salida=self._date(self.fecha_salida_naviera)
                if _parse_date(salida) < _parse_date(fecha): QMessageBox.warning(self,"Fechas","La salida de naviera no puede ser menor que la llegada a naviera/yarda."); return
                fin=salida; extra={"naviera":proveedor,"fecha_salida_naviera":salida,"motivo_extra":self.motivo_extra.toPlainText()}; cost_items,missing=self._collect_costs(fecha, proveedor)
            elif self.stage_key==STAGE_ADUANA:
                proveedor=f"{self.pais.currentText()} · {self.aduana.currentText()}"; lib=self._date(self.fecha_liberacion)
                if _parse_date(lib) < _parse_date(fecha): QMessageBox.warning(self,"Fechas","La liberación no puede ser menor que la llegada a aduana."); return
                fin=lib; extra={"pais":self.pais.currentText(),"aduana":self.aduana.currentText(),"fecha_liberacion_aduana":lib}; cost_items,missing=self._collect_costs(fecha, proveedor)
            elif self.stage_key==STAGE_PREPARACION:
                proveedor=self.taller.currentText(); fin=max([self._date(self.legal_fin), self._date(self.taller_salida)])
                extra={"emision_pedido":self._date(self.emision_pedido),"emision_obtenida":self._date(self.emision_obtenida),"cita_pedido":self._date(self.cita_pedido),"cita_asignada":self._date(self.cita_asignada),"placas_ingreso":self._date(self.placas_ingreso),"placas_entrega":self._date(self.placas_entrega),"legalizacion_fin":self._date(self.legal_fin),"taller":self.taller.currentText(),"taller_ingreso":self._date(self.taller_ingreso),"taller_salida":self._date(self.taller_salida),"motivo_taller":self.motivo_taller.toPlainText()}; cost_items,missing=self._collect_costs(fecha, proveedor)
                for r in range(self.repuestos_table.rowCount()):
                    desc=(self.repuestos_table.item(r,0).text() if self.repuestos_table.item(r,0) else "").strip(); valtxt=(self.repuestos_table.item(r,1).text() if self.repuestos_table.item(r,1) else "0").replace(",",""); oc=(self.repuestos_table.item(r,2).text() if self.repuestos_table.item(r,2) else "")
                    try: val=float(valtxt or 0)
                    except Exception: val=0
                    if desc and val>0: cost_items.append({"categoria":"TALLER","subcategoria":"REPUESTO_DETALLE","descripcion":desc,"monto_usd":val,"proveedor":proveedor,"oc_numero":oc,"fecha":fecha})
            elif self.stage_key==STAGE_PRECIO_FINAL:
                pf=self._price_calc(); proveedor="INTERNO"; extra={"precio_final_guardado":True}; data={"fecha_inicio":fecha,"fecha_fin":fecha,"proveedor":proveedor,"comentario":self.comentario.toPlainText(),"extra":extra,"precio_final":pf,"foto_principal_src":self.foto_path}
                ok,msg=update_vehicle_stage(self.vehicle.get("id"), self.stage_key, data, None, self.user, self.device)
                if not ok: QMessageBox.warning(self,"Validación",msg); return
                QMessageBox.information(self,"Guardado",msg); self.accept(); return
            elif self.stage_key==STAGE_DISPONIBLE:
                fin=fecha; proveedor="VENTAS"; extra={"disponible_para_cotizar":True}; missing=[]
            else:
                missing=[]
            if missing:
                resp=QMessageBox.question(self,"Gastos en cero", "Estos gastos obligatorios están en cero:\n"+"\n".join(missing)+"\n\n¿Seguro que no aplican o no los tuviste?")
                if resp != QMessageBox.StandardButton.Yes: return
            data={"fecha_inicio":fecha,"fecha_fin":fin,"proveedor":proveedor,"comentario":self.comentario.toPlainText(),"extra":extra,"cost_items":cost_items}
            ok,msg=update_vehicle_stage(self.vehicle.get("id"), self.stage_key, data, None, self.user, self.device)
            if not ok: QMessageBox.warning(self,"Validación",msg); return
            QMessageBox.information(self,"Guardado",msg); self.accept()

    class VehicleDetailDialog(QDialog):
        def __init__(self, parent, vehicle_id: str, user: dict, device: DeviceInfo):
            super().__init__(parent); self.vehicle_id=vehicle_id; self.user=user; self.device=device; self.setWindowTitle("Expediente del vehículo"); self.setMinimumSize(1320,850); self._build(); self.refresh()
        def _build(self):
            lay=QVBoxLayout(self); self.header=QLabel(); self.header.setStyleSheet("font-size:24px;font-weight:950;color:#08285a;"); lay.addWidget(self.header)
            self.timeline=TimelineWidget(); self.timeline.stageDoubleClicked.connect(self.open_stage_detail); lay.addWidget(self.timeline)
            btns=QHBoxLayout(); self.current_btn=QPushButton("Guardar / editar etapa actual"); self.current_btn.setObjectName("orange"); self.current_btn.clicked.connect(self.edit_current_stage); self.next_btn=QPushButton("Pasar a siguiente etapa"); self.next_btn.setObjectName("orange"); self.next_btn.clicked.connect(self.advance_next); self.edit_btn=QPushButton("Corregir etapa seleccionada"); self.edit_btn.setObjectName("ghost"); self.edit_btn.clicked.connect(self.edit_selected_stage); bdoc=QPushButton("Abrir comprobante compra"); bdoc.setObjectName("ghost"); bdoc.clicked.connect(self.open_purchase_doc); btns.addWidget(self.current_btn); btns.addWidget(self.next_btn); btns.addWidget(self.edit_btn); btns.addWidget(bdoc); btns.addStretch(1); lay.addLayout(btns)
            self.tabs=QTabWidget(); lay.addWidget(self.tabs)
            self.summary=QTextEdit(); self.summary.setReadOnly(True); self.tabs.addTab(self.summary,"Resumen")
            self.stage_table=QTableWidget(0,7); self.stage_table.setHorizontalHeaderLabels(["Etapa","Estado","Inicio","Fin","Días","Costo","Dato etapa"]); self.stage_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.stage_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.stage_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.stage_table.cellDoubleClicked.connect(lambda r,c:self.open_stage_detail(STAGE_ORDER[r] if r<len(STAGE_ORDER) else STAGE_COMPRADO)); self.tabs.addTab(self.stage_table,"Línea de tiempo / etapas")
            self.cost_table=QTableWidget(0,8); self.cost_table.setHorizontalHeaderLabels(["Etapa","Categoría","Subcategoría","Descripción","Proveedor","OC","Monto","Usuario"]); self.cost_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.cost_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.tabs.addTab(self.cost_table,"Costos detallados")
            self.quote_table=QTableWidget(0,6); self.quote_table.setHorizontalHeaderLabels(["Cliente","Teléfono","Fecha","Última gestión","Estado","Color"]); self.quote_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.quote_table.cellDoubleClicked.connect(self.open_quote_row); self.tabs.addTab(self.quote_table,"Cotizaciones del carro")
            self.hist=QTextEdit(); self.hist.setReadOnly(True); self.tabs.addTab(self.hist,"Historial")
        def refresh(self):
            self.vehicle=find_vehicle(self.vehicle_id)
            if not self.vehicle: return
            v=self.vehicle; ensure_vehicle_runtime_fields(v); self.timeline.set_vehicle(v)
            self.header.setText(f"{v.get('codigo')} · {v.get('marca')} {v.get('modelo')} {v.get('anio')} · {STAGE_META.get(v.get('estado_actual'),{}).get('label',v.get('estado_actual'))} · Comercial: {v.get('estado_comercial')}")
            nxt=next_stage_key(v.get("estado_actual",STAGE_COMPRADO)); self.next_btn.setEnabled(bool(nxt)); self.next_btn.setText(f"Pasar a siguiente etapa: {STAGE_META[nxt]['label']}" if nxt else "Proceso operativo finalizado")
            pf=v.get("precio_final") or {}
            resumen=f"""Código: {v.get('codigo')}\nVehículo: {v.get('marca')} {v.get('modelo')} {v.get('anio')}\nMillaje: {int(v.get('millaje') or 0):,}\nEstado USA: {v.get('estado_usa')}\nSubasta: {v.get('subasta')}\nLote: {v.get('lote')}\nOC compra: {v.get('oc_compra_numero','')}\nPrecio ganado: {_fmt_usd(v.get('precio_ganado_usd'))}\nFecha compra: {_fmt_date(v.get('fecha_compra'))}\nUsuario registró: {v.get('usuario_registro')}\n\nEstado operativo: {STAGE_META.get(v.get('estado_actual'),{}).get('label',v.get('estado_actual'))}\nEstado comercial: {v.get('estado_comercial')}\nDías desde compra: {vehicle_days_from_purchase(v)}\nDías en etapa actual: {current_stage_days(v)}\nAlerta etapa actual: {stage_alert_level(v)}\n\nCosto total acumulado: {_fmt_usd(vehicle_total_cost(v))}\nVenta neta interna: {_fmt_usd(pf.get('venta_neta_usd'))}\nPrecio cliente final: {_fmt_usd(v.get('precio_venta_usd'))}\nUtilidad bruta esperada: {_fmt_usd(vehicle_expected_profit(v))}\nCliente venta: {v.get('cliente','')}\n\nCaracterísticas para propuesta:\n{v.get('caracteristicas','')}\n\nObservaciones:\n{v.get('observaciones','')}"""
            self.summary.setPlainText(resumen)
            self.stage_table.setRowCount(len(STAGES))
            for r,s in enumerate(STAGES):
                st=vehicle_stage(v,s["key"]); vals=[s["label"],st.get("status"),_fmt_date(st.get("fecha_inicio")),_fmt_date(st.get("fecha_fin")),str(stage_duration_days(st)),_fmt_usd(st.get("costo_usd")),st.get("proveedor","")]
                for c,val in enumerate(vals): self.stage_table.setItem(r,c,QTableWidgetItem(str(val)))
            costs=v.get("gastos_detallados",[]) or []; self.cost_table.setRowCount(len(costs))
            for r,g in enumerate(costs):
                vals=[STAGE_META.get(g.get("stage_key"),{}).get("label",g.get("stage_key")),g.get("categoria"),g.get("subcategoria"),g.get("descripcion"),g.get("proveedor"),g.get("oc_numero"),_fmt_usd(g.get("monto_usd")),g.get("usuario")]
                for c,val in enumerate(vals): self.cost_table.setItem(r,c,QTableWidgetItem(str(val)))
            self._quote_ids=[]; qs=quotes_for_vehicle(v.get("id")); self.quote_table.setRowCount(len(qs))
            for r,q in enumerate(qs):
                self._quote_ids.append(q.get("id")); cl=q.get("cliente",{}); lvl=quote_alert_level(q); vals=[cl.get("nombre"),cl.get("telefono"),_fmt_date(q.get("fecha_cotizacion")),_fmt_date(q.get("ultima_gestion")),QUOTE_STATUS_LABELS.get(q.get("estado"),q.get("estado")),lvl]
                for c,val in enumerate(vals):
                    item=QTableWidgetItem(str(val)); item.setBackground(QColor(_quote_alert_color(lvl))); self.quote_table.setItem(r,c,item)
            hist="\n".join([f"[{h.get('fecha')}] {h.get('usuario')} · {h.get('accion')} · {h.get('detalle')}" for h in v.get('historial',[])])
            self.hist.setPlainText(hist)
        def edit_current_stage(self):
            key=self.vehicle.get("estado_actual",STAGE_COMPRADO); dlg=StageUpdateDialog(self,self.vehicle,key,self.user,self.device,mode="current");
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def advance_next(self):
            v=find_vehicle(self.vehicle_id)
            if not v: return
            current=v.get("estado_actual",STAGE_COMPRADO); ok,msg=can_advance_from_stage(v,current)
            if not ok: QMessageBox.warning(self,"Antes de avanzar",msg); return
            nxt=next_stage_key(current)
            if not nxt: QMessageBox.information(self,"Etapas","El vehículo ya no tiene etapa siguiente."); return
            dlg=StageUpdateDialog(self,v,nxt,self.user,self.device,mode="advance")
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def selected_stage_key(self) -> str:
            row=self.stage_table.currentRow(); return STAGE_ORDER[row] if 0 <= row < len(STAGE_ORDER) else self.vehicle.get("estado_actual",STAGE_COMPRADO)
        def edit_selected_stage(self):
            key=self.selected_stage_key()
            if key != self.vehicle.get("estado_actual") and not user_can_override_flow(self.user): QMessageBox.warning(self,"Permiso","Solo ADMIN/SUPERVISOR o permiso especial pueden corregir etapas anteriores."); return
            dlg=StageUpdateDialog(self,self.vehicle,key,self.user,self.device,mode="edit")
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def open_stage_detail(self, stage_key: str):
            row = STAGE_ORDER.index(stage_key) if stage_key in STAGE_ORDER else -1
            if row >= 0: self.stage_table.selectRow(row)
            if stage_key == self.vehicle.get("estado_actual"): self.edit_current_stage()
            elif stage_key == next_stage_key(self.vehicle.get("estado_actual",STAGE_COMPRADO)): self.advance_next()
            else: self.edit_selected_stage()
        def open_purchase_doc(self):
            st=vehicle_stage(self.vehicle,STAGE_COMPRADO); p=decrypt_file_to_temp(st.get("documento",""),st.get("documento_nombre","comprobante.pdf"))
            if p: QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
            else: QMessageBox.warning(self,"Documento","No se pudo abrir el comprobante.")
        def open_quote_row(self,row,col):
            if row < len(getattr(self,"_quote_ids",[])):
                dlg=QuoteDetailDialog(self,self._quote_ids[row],self.user,self.device); dlg.exec(); self.refresh()

    class QuoteEditorDialog(QDialog):
        def __init__(self, parent, user: dict, device: DeviceInfo, vehicle_id: str = "", quote_id: str = "", preset_client: Optional[dict] = None):
            super().__init__(parent); self.user=user; self.device=device; self.vehicle_id=vehicle_id; self.quote_id=quote_id; self.preset_client=preset_client or {}; self.setWindowTitle("Cotización leasing"); self.setMinimumSize(940,760); self._build(); self.refresh_calc()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Cotización leasing", "Simulador de cuota y gastos legales en base al valor del carro."))
            form=QFormLayout(); lay.addLayout(form)
            self.vehicle_combo=QComboBox(); self.vehicle_combo.setEditable(False); self._vehicle_ids=[]
            for v in load_vehicles():
                ensure_vehicle_runtime_fields(v)
                if v.get("estado_actual")==STAGE_DISPONIBLE and v.get("estado_comercial")==COMM_DISPONIBLE:
                    self._vehicle_ids.append(v.get("id")); self.vehicle_combo.addItem(f"{v.get('codigo')} · {v.get('marca')} {v.get('modelo')} {v.get('anio')} · {_fmt_usd(v.get('precio_venta_usd'))}")
            if self.vehicle_id and self.vehicle_id in self._vehicle_ids: self.vehicle_combo.setCurrentIndex(self._vehicle_ids.index(self.vehicle_id))
            self.nombre=QLineEdit(self.preset_client.get("nombre", "")); self.telefono=QLineEdit(self.preset_client.get("telefono", "")); self.correo=QLineEdit(self.preset_client.get("correo", "")); self.medio=QComboBox(); self.medio.setEditable(True); self.medio.addItems(["WHATSAPP","FACEBOOK","TIKTOK","PAGINA WEB","REFERIDO","LLAMADA","OTRO"]); self.medio.setCurrentText(self.preset_client.get("medio_contacto", "WHATSAPP"))
            self.ingreso=MoneyEdit(); self.ingreso.setRange(0,999999); self.ingreso.setValue(1900)
            self.precio=MoneyEdit(); self.precio.setRange(0,999999); self.prima_pct=QDoubleSpinBox(); self.prima_pct.setRange(0,90); self.prima_pct.setDecimals(2); self.prima_pct.setSuffix(" %"); self.prima_pct.setValue(20)
            self.comision=MoneyEdit(); self.comision.setRange(0,9999); self.comision.setValue(100); self.plazo=QSpinBox(); self.plazo.setRange(1,120); self.plazo.setValue(60)
            self.tasa=QComboBox(); self.tasa.setEditable(True); self.tasa.addItems(load_catalog(F_CATALOG_TASAS_LEASING, DEFAULT_TASAS_LEASING)); self.tasa.setCurrentText("2.50")
            self.seguro=MoneyEdit(); self.seguro.setRange(0,9999); self.seguro.setValue(80); self.gps=MoneyEdit(); self.gps.setRange(0,9999); self.gps.setValue(20); self.iva=QDoubleSpinBox(); self.iva.setRange(0,30); self.iva.setDecimals(2); self.iva.setSuffix(" %"); self.iva.setValue(13)
            self.comentario=QTextEdit(); self.comentario.setMinimumHeight(60)
            for label,w in [("Vehículo:",self.vehicle_combo),("Cliente:",self.nombre),("Teléfono:",self.telefono),("Correo:",self.correo),("Medio contacto:",self.medio),("Ingreso mensual cliente:",self.ingreso),("Precio vehículo:",self.precio),("Prima mínima:",self.prima_pct),("Comisión editable:",self.comision),("Plazo meses:",self.plazo),("Tasa rentabilidad mensual:",self.tasa),("Seguro mensual incluye IVA:",self.seguro),("GPS mensual incluye IVA:",self.gps),("IVA sobre cuota base:",self.iva),("Comentario:",self.comentario)]: form.addRow(label,w)
            self.vehicle_combo.currentIndexChanged.connect(self._vehicle_changed); self._vehicle_changed()
            for w in [self.ingreso,self.precio,self.comision,self.seguro,self.gps]:
                try: w.textChanged.connect(self.refresh_calc)
                except Exception: pass
            self.prima_pct.valueChanged.connect(self.refresh_calc); self.plazo.valueChanged.connect(self.refresh_calc); self.iva.valueChanged.connect(self.refresh_calc); self.tasa.currentTextChanged.connect(self.refresh_calc)
            self.result=QLabel(); self.result.setTextFormat(Qt.TextFormat.RichText); self.result.setStyleSheet("background:#eef6ff;border:1px solid #bfdbfe;border-radius:10px;padding:12px;color:#08285a;"); lay.addWidget(self.result)
            self.table=QTableWidget(0,4); self.table.setHorizontalHeaderLabels(["Plazo","Cuota financiamiento","Cuota total","Cuota total con IVA"]); self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); lay.addWidget(self.table)
            btns=QHBoxLayout(); bsave=QPushButton("Guardar cotización"); bsave.setObjectName("orange"); bsave.clicked.connect(self.save); btns.addWidget(bsave); btns.addStretch(1); lay.addLayout(btns)
        def _vehicle_changed(self):
            if 0 <= self.vehicle_combo.currentIndex() < len(self._vehicle_ids):
                v=find_vehicle(self._vehicle_ids[self.vehicle_combo.currentIndex()]); self.precio.setValue(float((v or {}).get("precio_venta_usd") or 0))
            self.refresh_calc()
        def _tasa_value(self):
            try: return float(str(self.tasa.currentText()).replace("%","").replace(",","."))
            except Exception: return 2.5
        def _calc_data(self):
            precio=self.precio.value(); pago_inicial=round(precio*self.prima_pct.value()/100.0 + self.comision.value(),2); calc=calculate_leasing(precio,self.ingreso.value(),pago_inicial,self.plazo.value(),self._tasa_value(),self.seguro.value(),self.gps.value(),self.iva.value()); legal=calculate_legal_fees(precio); return calc,legal,pago_inicial
        def refresh_calc(self):
            calc,legal,pago=self._calc_data(); self.result.setText(f"Monto leasing: <b>{_fmt_usd(calc['monto_leasing'])}</b> · Prima + comisión: <b>{_fmt_usd(pago)}</b><br>Cuota base: <b>{_fmt_usd(calc['cuota_base'])}</b> · Cuota total con IVA: <b style='font-size:18px'>{_fmt_usd(calc['cuota_total_con_iva'])}</b><br>% ingreso destinado a cuota: <b>{calc['pct_ingreso']}%</b> · Riesgo: <b>{calc['riesgo_texto']} / {calc['riesgo']}</b><br>Gastos legales: <b>{_fmt_usd(legal['valor_legales_iva_incluido'])}</b> · Aplica tope: <b>{legal['aplica_tope']}</b>")
            rows=leasing_table(self.precio.value(),self.ingreso.value(),pago,self._tasa_value(),self.seguro.value(),self.gps.value(),self.iva.value()); self.table.setRowCount(len(rows))
            for r,row in enumerate(rows):
                vals=[row['plazo'],_fmt_usd(row['cuota_base']),_fmt_usd(row['cuota_total_sin_iva']),_fmt_usd(row['cuota_total_con_iva'])]
                for c,val in enumerate(vals): self.table.setItem(r,c,QTableWidgetItem(str(val)))
        def save(self):
            if not self._vehicle_ids: QMessageBox.warning(self,"Cotización","No hay vehículos disponibles para cotizar."); return
            data={"vehicle_id":self._vehicle_ids[self.vehicle_combo.currentIndex()],"cliente_nombre":self.nombre.text(),"telefono":self.telefono.text(),"correo":self.correo.text(),"medio_contacto":self.medio.currentText(),"ingreso_cliente":self.ingreso.value(),"precio_vehiculo":self.precio.value(),"prima_pct":self.prima_pct.value(),"comision_usd":self.comision.value(),"plazo_meses":self.plazo.value(),"tasa_mensual_pct":self._tasa_value(),"seguro_mensual":self.seguro.value(),"gps_mensual":self.gps.value(),"iva_pct":self.iva.value(),"comentario":self.comentario.toPlainText(),"fecha_cotizacion":date.today().isoformat()}
            ok,msg,qid=create_quote(data,self.user,self.device,self.quote_id)
            if not ok: QMessageBox.warning(self,"Cotización",msg); return
            QMessageBox.information(self,"Cotización",msg); self.quote_id=qid; self.accept()

    class QuoteDetailDialog(QDialog):
        def __init__(self, parent, quote_id: str, user: dict, device: DeviceInfo):
            super().__init__(parent); self.quote_id=quote_id; self.user=user; self.device=device; self.setWindowTitle("Detalle de cotización"); self.setMinimumSize(900,680); self._build(); self.refresh()
        def _build(self):
            lay=QVBoxLayout(self); self.summary=QTextEdit(); self.summary.setReadOnly(True); lay.addWidget(self.summary)
            btns=QHBoxLayout(); bseg=QPushButton("Agregar seguimiento"); bseg.setObjectName("orange"); bseg.clicked.connect(self.add_follow); bprop=QPushButton("Generar propuesta"); bprop.setObjectName("orange"); bprop.clicked.connect(self.generate_prop); breof=QPushButton("Ofrecer otro carro"); breof.setObjectName("ghost"); breof.clicked.connect(self.reoffer); bwon=QPushButton("Marcar compra / vendido"); bwon.setObjectName("danger"); bwon.clicked.connect(self.mark_won); btns.addWidget(bseg); btns.addWidget(bprop); btns.addWidget(breof); btns.addWidget(bwon); btns.addStretch(1); lay.addLayout(btns)
            self.hist=QTextEdit(); self.hist.setReadOnly(True); lay.addWidget(self.hist)
        def refresh(self):
            self.quote=find_quote(self.quote_id)
            if not self.quote: return
            q=self.quote; cl=q.get("cliente",{}); s=q.get("vehicle_snapshot",{}); le=q.get("leasing",{}); lg=q.get("legal",{}); lvl=quote_alert_level(q)
            self.summary.setPlainText(f"Cliente: {cl.get('nombre')}\nTeléfono: {cl.get('telefono')}\nCorreo: {cl.get('correo')}\nMedio: {cl.get('medio_contacto')}\n\nVehículo cotizado: {s.get('anio')} {s.get('marca')} {s.get('modelo')} · {q.get('vehicle_code')}\nPrecio vehículo: {_fmt_usd(le.get('precio_vehiculo'))}\nPrima/comisión: {_fmt_usd(le.get('pago_inicial'))}\nMonto leasing: {_fmt_usd(le.get('monto_leasing'))}\nPlazo: {le.get('plazo_meses')} meses\nTasa: {le.get('tasa_mensual_pct')}%\nCuota final con IVA: {_fmt_usd(le.get('cuota_total_con_iva'))}\nGastos legales: {_fmt_usd(lg.get('valor_legales_iva_incluido'))}\n\nEstado: {QUOTE_STATUS_LABELS.get(q.get('estado'),q.get('estado'))}\nColor seguimiento: {lvl}\nDías sin compra/gestión: {quote_days_without_purchase(q)}")
            self.hist.setPlainText("\n".join([f"[{h.get('fecha')}] {h.get('usuario')} · {h.get('accion')} · {h.get('comentario')}" for h in q.get("seguimientos",[])]))
        def add_follow(self):
            comentario, ok = QInputDialog.getMultiLineText(self,"Seguimiento","Comentario del seguimiento:")
            if not ok or not comentario.strip(): return
            states=list(QUOTE_STATUS_LABELS.keys()); state, ok2=QInputDialog.getItem(self,"Estado","Nuevo estado:",[QUOTE_STATUS_LABELS[x] for x in states],0,False)
            estado=states[[QUOTE_STATUS_LABELS[x] for x in states].index(state)] if ok2 else QUOTE_SEGUIMIENTO
            ok,msg=add_quote_followup(self.quote_id,comentario,estado,self.user); QMessageBox.information(self,"Seguimiento",msg) if ok else QMessageBox.warning(self,"Seguimiento",msg); self.refresh()
        def generate_prop(self):
            ok,msg,path=generate_quote_proposal_html(self.quote_id,self.user)
            if not ok or not path: QMessageBox.warning(self,"Propuesta",msg); return
            QMessageBox.information(self,"Propuesta",f"{msg}\n{path.name}"); QDesktopServices.openUrl(QUrl.fromLocalFile(str(path))); self.refresh()
        def reoffer(self):
            cl=(self.quote or {}).get("cliente",{})
            dlg=QuoteEditorDialog(self,self.user,self.device,preset_client=cl)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def mark_won(self):
            if QMessageBox.question(self,"Cerrar venta","¿Confirmas que este cliente compró el vehículo? Las demás cotizaciones de este carro quedarán para reofertar.") != QMessageBox.StandardButton.Yes: return
            ok,msg=mark_quote_won_and_vehicle_sold(self.quote_id,self.user,self.device); QMessageBox.information(self,"Venta",msg) if ok else QMessageBox.warning(self,"Venta",msg); self.refresh()

    class CotizacionesPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main=main; self._vehicle_ids=[]; self._quote_ids=[]; self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Cotizaciones", "Vehículos disponibles, clientes interesados y cotizaciones generales sin borrar historial."))
            self.tabs=QTabWidget(); lay.addWidget(self.tabs)
            wv=QWidget(); vl=QVBoxLayout(wv); brow=QHBoxLayout(); bnew=QPushButton("Nueva cotización"); bnew.setObjectName("orange"); bnew.clicked.connect(self.new_quote); brow.addWidget(bnew); brow.addStretch(1); vl.addLayout(brow)
            self.vehicle_table=QTableWidget(0,7); self.vehicle_table.setHorizontalHeaderLabels(["Código","Vehículo","Precio","Costo","Ganancia","Cotizaciones","Días disponible"]); self.vehicle_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.vehicle_table.cellDoubleClicked.connect(self.open_vehicle_quotes); vl.addWidget(self.vehicle_table); self.tabs.addTab(wv,"Vehículos disponibles")
            wg=QWidget(); gl=QVBoxLayout(wg); self.quote_table=QTableWidget(0,9); self.quote_table.setHorizontalHeaderLabels(["Cliente","Teléfono","Medio","Vehículo cotizado","Fecha","Última gestión","Días","Estado","Color"]); self.quote_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.quote_table.cellDoubleClicked.connect(self.open_quote); gl.addWidget(self.quote_table); self.tabs.addTab(wg,"Cotizaciones general")
        def refresh(self):
            vehicles=[v for v in load_vehicles() if ensure_vehicle_runtime_fields(v) and v.get("estado_actual")==STAGE_DISPONIBLE and v.get("estado_comercial")==COMM_DISPONIBLE]
            self._vehicle_ids=[]; self.vehicle_table.setRowCount(len(vehicles))
            for r,v in enumerate(vehicles):
                self._vehicle_ids.append(v.get("id")); vals=[v.get("codigo"),f"{v.get('marca')} {v.get('modelo')} {v.get('anio')}",_fmt_usd(v.get("precio_venta_usd")),_fmt_usd(vehicle_total_cost(v)),_fmt_usd(vehicle_expected_profit(v)),str(len(quotes_for_vehicle(v.get("id")))),str(current_stage_days(v))]
                for c,val in enumerate(vals): self.vehicle_table.setItem(r,c,QTableWidgetItem(str(val)))
            qs=load_quotes(); self._quote_ids=[]; self.quote_table.setRowCount(len(qs))
            for r,q in enumerate(qs):
                self._quote_ids.append(q.get("id")); cl=q.get("cliente",{}); snap=q.get("vehicle_snapshot",{}); lvl=quote_alert_level(q); vals=[cl.get("nombre"),cl.get("telefono"),cl.get("medio_contacto"),f"{snap.get('anio','')} {snap.get('marca','')} {snap.get('modelo','')} · {q.get('vehicle_code')}",_fmt_date(q.get("fecha_cotizacion")),_fmt_date(q.get("ultima_gestion")),str(quote_days_without_purchase(q)),QUOTE_STATUS_LABELS.get(q.get("estado"),q.get("estado")),lvl]
                for c,val in enumerate(vals):
                    item=QTableWidgetItem(str(val)); item.setBackground(QColor(_quote_alert_color(lvl))); self.quote_table.setItem(r,c,item)
        def new_quote(self):
            dlg=QuoteEditorDialog(self,self.main.user,self.main.device)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh(); self.main.refresh_all()
        def open_vehicle_quotes(self,row,col):
            if row < len(self._vehicle_ids):
                dlg=QuoteEditorDialog(self,self.main.user,self.main.device,vehicle_id=self._vehicle_ids[row])
                if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh(); self.main.refresh_all()
        def open_quote(self,row,col):
            if row < len(self._quote_ids):
                dlg=QuoteDetailDialog(self,self._quote_ids[row],self.main.user,self.main.device); dlg.exec(); self.refresh(); self.main.refresh_all()

    class ReporteriaPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main=main; self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Reportería", "Inventario, tiempos, cotizaciones y KPIs comerciales con logo de la empresa."))
            grid=QGridLayout(); lay.addLayout(grid)
            items=[("Excel inventario compra vehicular","Detalle operativo con costos, estado comercial y OC.",self.report_excel_inventory),("HTML inventario gerencial","Reporte visual de capital, etapas y disponibles.",self.report_html_inventory),("Excel KPI tiempos","Tiempos por etapa y ciclo compra→disponible.",self.report_excel_kpi),("HTML KPI tiempos","Dashboard visual de tiempos operativos.",self.report_html_kpi),("Excel cotizaciones general","Clientes, carros cotizados, colores y cuotas.",self.report_excel_quotes),("HTML cotizaciones general","Dashboard comercial verde/amarillo/rojo.",self.report_html_quotes)]
            for i,(title,desc,fn) in enumerate(items):
                card=QFrame(); card.setStyleSheet("QFrame{background:white;border:1px solid #d9e2ef;border-radius:18px;}"); cl=QVBoxLayout(card); lab=QLabel(title); lab.setStyleSheet("font-size:16pt;font-weight:900;color:#08285a;"); d=QLabel(desc); d.setWordWrap(True); d.setStyleSheet("color:#64748b;"); b=QPushButton("Generar"); b.setObjectName("orange"); b.clicked.connect(fn); cl.addWidget(lab); cl.addWidget(d); cl.addStretch(1); cl.addWidget(b); grid.addWidget(card,i//2,i%2)
            lay.addStretch(1)
        def _after_report(self,path:Path,title:str):
            msg=QMessageBox(self); msg.setWindowTitle(title); msg.setText(f"Reporte generado correctamente:\n{path.name}"); open_btn=msg.addButton("Abrir",QMessageBox.ButtonRole.AcceptRole); copy_btn=msg.addButton("Guardar una copia…",QMessageBox.ButtonRole.ActionRole); msg.addButton("Cerrar",QMessageBox.ButtonRole.RejectRole); msg.exec(); clicked=msg.clickedButton()
            if clicked==open_btn: QDesktopServices.openUrl(QUrl.fromLocalFile(str(path)))
            elif clicked==copy_btn:
                target,_=QFileDialog.getSaveFileName(self,"Guardar copia",str(Path.home()/path.name),"Archivos (*"+path.suffix+")")
                if target:
                    ok,m=copy_report_to(path,Path(target)); QMessageBox.information(self,"Copia",m) if ok else QMessageBox.warning(self,"Copia",m)
        def report_excel_inventory(self):
            ok,msg,out=generate_inventory_excel(load_vehicles(),self.main.user); QMessageBox.warning(self,"Excel",msg) if not ok or not out else self._after_report(out,"Excel inventario")
        def report_html_inventory(self):
            out=generate_html_report(load_vehicles(),self.main.user); self._after_report(out,"HTML inventario") if out else None
        def report_excel_kpi(self):
            ok,msg,out=generate_kpi_excel_report(load_vehicles(),self.main.user); QMessageBox.warning(self,"Excel KPI",msg) if not ok or not out else self._after_report(out,"Excel KPI")
        def report_html_kpi(self):
            out=generate_kpi_html_report(load_vehicles(),self.main.user); self._after_report(out,"HTML KPI") if out else None
        def report_excel_quotes(self):
            ok,msg,out=generate_quotes_excel_report(self.main.user); QMessageBox.warning(self,"Excel cotizaciones",msg) if not ok or not out else self._after_report(out,"Excel cotizaciones")
        def report_html_quotes(self):
            out=generate_quotes_html_report(self.main.user); self._after_report(out,"HTML cotizaciones") if out else None

    class MainWindow(QMainWindow):
        def __init__(self, user: dict, device: DeviceInfo):
            super().__init__(); self.user=user; self.device=device; self.setWindowTitle(f"{APP_NAME} · {user.get('usuario')} · v{APP_VERSION}"); self.resize(1480,900); self.setWindowIcon(QIcon(str(ResourceManager.find_logo() or ""))); self._build(); self.refresh_all()
        def _build(self):
            central=QWidget(); self.setCentralWidget(central); main=QHBoxLayout(central); main.setContentsMargins(0,0,0,0)
            side=QFrame(); side.setFixedWidth(265); side.setStyleSheet("QFrame{background:#08285a;} QLabel{color:white;} QPushButton{background:transparent;color:white;text-align:left;padding:11px 14px;border-radius:0;font-weight:850;} QPushButton:hover{background:#0e3a78;border-left:5px solid #f59a13;}")
            sl=QVBoxLayout(side); sl.setContentsMargins(14,20,14,16)
            logo=QLabel(); logo.setAlignment(Qt.AlignmentFlag.AlignCenter); pix=QPixmap(str(ResourceManager.find_logo() or ""))
            if not pix.isNull(): logo.setPixmap(pix.scaled(145,145,Qt.AspectRatioMode.KeepAspectRatio,Qt.TransformationMode.SmoothTransformation))
            else: logo.setText("L&M"); logo.setStyleSheet("font-size:30px;font-weight:950;color:#f59a13;")
            sl.addWidget(logo); title=QLabel("LYM AUTO CONTROL"); title.setAlignment(Qt.AlignmentFlag.AlignCenter); title.setStyleSheet("font-weight:950;color:#ffcf7a;font-size:14pt;"); sub=QLabel(f"Rol: {self.user.get('rol')}"); sub.setAlignment(Qt.AlignmentFlag.AlignCenter); sub.setStyleSheet("font-weight:700;color:#d7e6ff;font-size:10pt;"); sl.addWidget(title); sl.addWidget(sub)
            self.stack=QStackedWidget(); self.pages=[]
            def add_page(name, widget):
                idx=self.stack.addWidget(widget); self.pages.append(widget); b=QPushButton(name); b.clicked.connect(lambda _,i=idx:self.stack.setCurrentIndex(i)); sl.addWidget(b)
            def section(text):
                lab=QLabel(text); lab.setStyleSheet("color:#ffcf7a;font-weight:950;margin-top:10px;padding:8px 6px;border-bottom:1px solid rgba(255,255,255,.12);"); sl.addWidget(lab)
            self.dashboard=DashboardPage(self); self.purchase=PurchasePage(self); self.inventory=InventoryPage(self); self.cotizaciones=CotizacionesPage(self); self.reporteria=ReporteriaPage(self); self.catalogos=CatalogosPage(self)
            add_page("🏠  Inicio", self.dashboard); section("COMPRA VEHICULAR"); add_page("➕  Nueva compra", self.purchase); add_page("🚗  Inventario / CV", self.inventory); section("COMERCIAL"); add_page("💬  Cotizaciones", self.cotizaciones); section("GERENCIA"); add_page("📊  Reportería", self.reporteria); add_page("📚  Catálogos", self.catalogos)
            if user_has_permission(self.user, PERM_CONFIG):
                bconf=QPushButton("⚙️  Configuración"); bconf.clicked.connect(self.open_config); sl.addWidget(bconf)
            sl.addStretch(1); bclose=QPushButton("🚪  Cerrar"); bclose.clicked.connect(self.close); sl.addWidget(bclose); main.addWidget(side); main.addWidget(self.stack,1); self.setStatusBar(QStatusBar()); self.statusBar().showMessage(f"Usuario: {self.user.get('usuario')} · Carpeta: {get_data_folder()}")
        def refresh_all(self):
            bootstrap_system(); self.dashboard.refresh(); self.inventory.refresh(); self.purchase.refresh_catalogs(); self.cotizaciones.refresh(); self.catalogos.refresh_all()
        def open_vehicle_detail(self, vehicle_id: str):
            dlg=VehicleDetailDialog(self,vehicle_id,self.user,self.device); dlg.exec(); self.refresh_all()
        def open_config(self):
            dlg=ConfigDialog(self,self.user,self.device); dlg.exec(); self.refresh_all()



# =============================================================================
# LYM AUTO CONTROL V4.1 - CORRECCIONES DE COTIZACIÓN, GASTOS Y PROPUESTA
# =============================================================================
APP_VERSION = "2.2.0_LEASING"


def _legal_doc_flag(rel: str = "", src: str = "") -> str:
    return "✅ Cargado" if rel or src else "❌ Vacío"


def _short_doc_status(rel: str = "", src: str = "") -> str:
    return "✅" if rel or src else "❌"


def _aduana_impuestos_amount(vehicle: dict) -> float:
    """IVA DUCA / impuestos pagados en aduana para prellenar precio final."""
    ensure_vehicle_runtime_fields(vehicle)
    total = 0.0
    for g in vehicle.get("gastos_detallados", []) or []:
        if _norm(g.get("stage_key")) == _norm(STAGE_ADUANA) and _norm(g.get("subcategoria")) == "IMPUESTOS_ADUANA":
            try:
                total += float(g.get("monto_usd") or 0)
            except Exception:
                pass
    return round(total, 2)


# Propuesta rediseñada para parecerse al PDF modelo: logo arriba, línea azul,
# vehículo centrado, características en dos columnas, resumen financiero claro y condiciones.
def generate_quote_proposal_html(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada.", None
    df = get_data_folder()
    if df is None:
        return False, "Carpeta del sistema no disponible.", None
    out_dir = df / SUB_REPORTES / "PROPUESTAS"
    out_dir.mkdir(parents=True, exist_ok=True)
    ensure_quote_runtime_fields(quote)
    cliente = quote.get("cliente", {})
    snap = quote.get("vehicle_snapshot", {})
    leasing = quote.get("leasing", {})
    legal = quote.get("legal", {})
    logo_uri = ResourceManager.logo_data_uri()
    vehicle_name = f"{snap.get('anio','')} {snap.get('marca','')} {snap.get('modelo','')}".strip()
    raw_chars = str(snap.get("caracteristicas") or "").replace(";", "\n")
    caracteristicas = [x.strip(" •-\t") for x in raw_chars.splitlines() if x.strip(" •-\t")]
    base_chars = [
        f"Año {snap.get('anio','')}",
        f"{int(snap.get('millaje') or 0):,} millas",
    ]
    if snap.get("color"):
        base_chars.append(f"Color {snap.get('color')}")
    if not caracteristicas:
        caracteristicas = ["Transmisión automática", "Aire acondicionado", "Vehículo revisado por L&M Inversiones"]
    full_chars = []
    for x in base_chars + caracteristicas:
        if x and x not in full_chars:
            full_chars.append(x)
    left = full_chars[: max(1, (len(full_chars) + 1) // 2)]
    right = full_chars[max(1, (len(full_chars) + 1) // 2):]
    bullets_left = "".join(f"<li>{html.escape(x)}</li>" for x in left[:12])
    bullets_right = "".join(f"<li>{html.escape(x)}</li>" for x in right[:12])
    tasa = leasing.get("tasa_mensual_pct", 0)
    fecha = date.today().strftime('%d de %B de %Y')
    css = """
@page { size: letter; margin: 13mm; }
*{box-sizing:border-box} body{font-family:'Arial Narrow',Arial,sans-serif;background:#eef2f7;margin:0;color:#111827}.sheet{background:#fff;width:8.5in;min-height:11in;margin:16px auto;padding:24px 34px 20px;box-shadow:0 10px 35px rgba(8,40,90,.16);position:relative}.brand{display:grid;grid-template-columns:118px 1fr;gap:16px;align-items:end}.logo{width:105px;height:105px;display:flex;align-items:center;justify-content:center}.logo img{max-width:100%;max-height:100%}.blue-line{height:8px;border-top:4px solid #08285a;border-bottom:1px solid #94a3b8;box-shadow:0 2px 4px rgba(8,40,90,.25);margin-bottom:8px}.tag{font-size:11px;color:#08285a;font-style:italic;text-align:center}.to{font-size:17px;font-weight:800;line-height:1.25;margin:4px 0 14px}.title{text-align:center;text-decoration:underline;font-weight:900;font-size:20px;margin:4px 0 16px}.intro{font-size:15px;line-height:1.38}.vehicle{text-align:center;font-weight:950;font-size:18px;color:#111827;margin:12px 0}.diamond{color:#0477a3;font-size:16px;margin:0 8px}.section-title{font-weight:950;font-size:16px;margin:8px 0}.features{display:grid;grid-template-columns:1fr 1fr;gap:22px;margin:2px 0 16px}.features ul{margin:0;padding-left:22px}.features li{font-size:15px;margin:7px 0}.pitch{font-size:15px;line-height:1.34;margin:14px 0}.finance-line{display:grid;grid-template-columns:1.2fr 1fr 1fr;gap:8px;align-items:end;margin:12px 0 8px}.finance-line .item{font-size:15px;font-weight:800}.finance-line b{text-decoration:underline;font-weight:950}.finance-line .center{text-align:center}.finance-line .right{text-align:right}.quote-table{width:100%;border-collapse:collapse;font-size:14px;margin-top:6px}.quote-table th,.quote-table td{border:1.5px solid #111;padding:7px 8px;text-align:left}.quote-table th{font-weight:950;background:#f8fafc}.quote-table td{height:28px}.red-note{color:#dc2626;font-weight:950;text-decoration:underline;font-size:13px;margin-top:4px}.closing{font-size:15px;margin:28px 0}.city{font-size:15px;margin-top:24px}.footer-word{position:absolute;bottom:22px;left:0;right:0;text-align:center;font-family:serif;color:#374151;font-size:16px}.page-num{position:absolute;right:20px;bottom:18px;font-size:13px}.signature{font-size:15px;margin-top:30px;line-height:1.4}.accept{display:flex;justify-content:flex-start;margin:36px 0 10px;font-size:14px}.conditions{margin-top:12px}.conditions h2{font-size:16px;text-align:center;margin:8px 0 12px}.conditions li{font-size:13px;line-height:1.34;margin:8px 0}.small{font-size:12px;color:#374151}.no-print-note{max-width:8.5in;margin:0 auto 10px;color:#475569;font-size:12px}@media print{body{background:white}.sheet{box-shadow:none;margin:0;width:auto;page-break-after:always}.no-print-note{display:none}}
"""
    html_text = f"""<!doctype html><html lang='es'><head><meta charset='utf-8'><title>Propuesta Leasing {html.escape(vehicle_name)}</title><style>{css}</style></head><body>
<div class='no-print-note'>Propuesta HTML imprimible. Usa Ctrl+P y guardar como PDF para enviarla al cliente.</div>
<section class='sheet'>
  <div class='brand'><div><div class='logo'>{'<img src="'+logo_uri+'" alt="Logo L&M">' if logo_uri else '<b>L&M</b>'}</div><div class='tag'>Ayudando a lograr tus sueños</div></div><div><div class='blue-line'></div></div></div>
  <p class='to'>{html.escape(cliente.get('nombre') or 'Cliente')}<br>Presente.</p>
  <div class='title'>Propuesta de Arrendamiento Vehicular</div>
  <p class='intro'>Por medio de la presente, <b>L&amp;M Inversiones, S.A. de C.V.</b> tiene el agrado de presentarle la propuesta de contrato de arrendamiento (Leasing) opción: <b>RENT A CAR.</b></p>
  <div class='vehicle'><span class='diamond'>♦</span>{html.escape(vehicle_name.upper())}<span class='diamond'>♦</span></div>
  <div class='section-title'>Características destacadas:</div>
  <div class='features'><ul>{bullets_left}</ul><ul>{bullets_right}</ul></div>
  <p class='pitch'>Un vehículo eficiente y equipado para manejar con estilo, seguridad y confianza todos los días.<br>Compartimos ante usted el detalle de plazo y cuota mensual para <b>{html.escape(vehicle_name.upper())}</b>. El cálculo incluye seguro estimado, servicio GPS y una tasa del <b>{tasa}%</b> mensual.</p>
  <div class='finance-line'><div class='item'>MONTO LEASING: <b>{_fmt_usd(leasing.get('monto_leasing'))}</b></div><div class='item center'>Costo Legal: <b>{_fmt_usd(legal.get('valor_legales_iva_incluido'))}</b> <span class='small'>Incluye IVA</span></div><div class='item right'>Opción de compra: <b>{_fmt_usd(500)}</b></div></div>
  <table class='quote-table'><thead><tr><th>Plazo<br>(meses)</th><th>Tasa rentabilidad<br>(%)</th><th>Cuota base<br>(US$)</th><th>Cuota total mensual incluye seguro<br>Cobertura total, GPS e IVA<br>(US$)</th></tr></thead><tbody><tr><td>{int(leasing.get('plazo_meses') or 0)}</td><td>{tasa}%</td><td>{_fmt_usd(leasing.get('cuota_base')).replace('$ ', '')}</td><td><b>{_fmt_usd(leasing.get('cuota_total_con_iva')).replace('$ ', '')}</b></td></tr></tbody></table>
  <div class='red-note'>La cuota mensual incluye IVA, seguro y GPS.</div>
  <p class='closing'>Quedamos atentos a cualquier consulta o duda.</p><p class='city'>San Salvador, {fecha}</p>
  <div class='footer-word'>Ayudando a lograr tus sueños</div><div class='page-num'>1</div>
</section>
<section class='sheet'>
  <div class='signature'>Atentamente,<br><br><b>L&amp;M Inversiones, S.A. de C.V.</b><br>guillermo.moreno@lyminversiones.com<br>Tel: (503) 7475 5821</div>
  <div class='accept'><div>X____________________________ &nbsp;&nbsp;&nbsp; ___________<br>&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;Aceptado por cliente</div></div>
  <div class='conditions'><h2>Condiciones y Vigencia de la Oferta</h2><ul>
    <li>Esta propuesta forma parte de una oferta especial de arrendamiento vehicular (leasing) válida por 15 días calendario a partir de la fecha de emisión.</li>
    <li>La tasa de interés mensual indicada corresponde a la cotización solicitada y brinda condiciones accesibles y transparentes para nuestros clientes.</li>
    <li>La propuesta se formalizará mediante contrato de arrendamiento donde se establecerán las condiciones de uso, pagos mensuales y opción de adquisición al finalizar el plazo.</li>
    <li>El valor del seguro mostrado es estimado y puede variar según las características y el uso del vehículo; el costo final se definirá con base en la cotización de la aseguradora.</li>
    <li>Los gastos administrativos y legales derivados de la formalización del contrato deberán ser cancelados al momento de la firma.</li>
    <li>Las cuotas mensuales incluyen IVA, seguro y GPS, reflejando el monto total a pagar. No se aplican cobros adicionales fuera de los valores detallados en esta propuesta.</li>
    <li>Durante el período de arrendamiento, el cliente deberá mantener el vehículo en buen estado y cumplir con los servicios preventivos o correctivos recomendados.</li>
    <li>L&M Inversiones, S.A. de C.V. se reserva el derecho de modificar los términos comerciales o administrativos de esta propuesta.</li>
    <li>Esta propuesta tiene carácter informativo y no constituye un compromiso contractual hasta la formalización del contrato de arrendamiento financiero.</li>
  </ul></div>
  <div class='footer-word'>Ayudando a lograr tus sueños</div><div class='page-num'>2</div>
</section>
</body></html>"""
    out = out_dir / f"PROPUESTA_LEASING_{_safe_filename(snap.get('codigo','CV'))}_{_safe_filename(cliente.get('nombre','CLIENTE'))}_{datetime.now().strftime('%Y%m%d_%H%M')}.html"
    out.write_text(html_text, encoding="utf-8")
    quote.setdefault("propuestas", []).append({"fecha": _now_iso(), "usuario": (user or {}).get("usuario", ""), "path": abs_to_rel(out), "nombre": out.name, "tipo": "HTML_IMPRIMIBLE"})
    upsert_quote(quote)
    log_audit("GENERAR_PROPUESTA", (user or {}).get("usuario", ""), quote.get("vehicle_code", ""), out.name)
    return True, "Propuesta generada correctamente.", out


if PYSIDE_OK:
    class OptionalDateEdit(QWidget):
        def __init__(self, initial: str = "", minimum: str = "", label: str = "Usar fecha"):
            super().__init__()
            self.chk = QCheckBox(label)
            self.date_edit = configure_date_edit(QDateEdit(), initial or _today_iso(), minimum)
            self.chk.setChecked(bool(initial))
            self.date_edit.setEnabled(bool(initial))
            self.chk.toggled.connect(self.date_edit.setEnabled)
            lay = QHBoxLayout(self); lay.setContentsMargins(0,0,0,0)
            lay.addWidget(self.chk); lay.addWidget(self.date_edit, 1)
        def date_iso(self) -> str:
            return self.date_edit.date().toPython().isoformat() if self.chk.isChecked() else ""
        def date_value(self) -> Optional[date]:
            return _parse_date(self.date_iso()) if self.chk.isChecked() else None
        def setMinimumDate(self, qdate):
            self.date_edit.setMinimumDate(qdate)

    class CostBreakdownDialog(QDialog):
        def __init__(self, parent, vehicle: dict):
            super().__init__(parent)
            self.vehicle = vehicle
            self.setWindowTitle(f"Visualizar gastos · {vehicle.get('codigo','')}")
            self.setMinimumSize(980, 620)
            self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint)
            lay = QVBoxLayout(self)
            total = vehicle_total_cost(vehicle)
            head = QLabel(f"<b>{vehicle.get('codigo')}</b> · {vehicle.get('marca')} {vehicle.get('modelo')} {vehicle.get('anio')} · Total gastos: <b>{_fmt_usd(total)}</b>")
            head.setTextFormat(Qt.TextFormat.RichText); head.setStyleSheet("font-size:18px;color:#08285a;padding:8px;")
            lay.addWidget(head)
            self.table = QTableWidget(0, 10)
            self.table.setHorizontalHeaderLabels(["Etapa","Categoría","Subcategoría","Descripción","Fecha","Proveedor","OC","Comprobante","OC PDF","Monto"])
            self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
            lay.addWidget(self.table)
            costs = vehicle.get("gastos_detallados", []) or []
            self.table.setRowCount(len(costs))
            for r,g in enumerate(costs):
                vals=[STAGE_META.get(g.get('stage_key'),{}).get('label',g.get('stage_key')),g.get('categoria'),g.get('subcategoria'),g.get('descripcion'),_fmt_date(g.get('fecha')),g.get('proveedor'),g.get('oc_numero'),_short_doc_status(g.get('comprobante')), _short_doc_status(g.get('oc_documento')), _fmt_usd(g.get('monto_usd'))]
                for c,val in enumerate(vals):
                    item=QTableWidgetItem(str(val)); item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable); self.table.setItem(r,c,item)
            btn=QPushButton("Cerrar"); btn.setObjectName("ghost"); btn.clicked.connect(self.accept); lay.addWidget(btn)

    class WorkshopItemDialog(QDialog):
        def __init__(self, parent=None, item: Optional[dict] = None):
            super().__init__(parent)
            self.item = dict(item or {})
            self.doc_src = self.item.get("comprobante_src", "")
            self.oc_src = self.item.get("oc_src", "")
            self.setWindowTitle("Repuesto / gasto de taller")
            self.setMinimumWidth(560)
            lay = QVBoxLayout(self); form = QFormLayout(); lay.addLayout(form)
            self.desc = QLineEdit(self.item.get("descripcion", "")); self.desc.setPlaceholderText("Ejemplo: Puerta, vidrio, pintura bumper...")
            self.monto = MoneyEdit(); self.monto.setRange(0,999999); self.monto.setValue(float(self.item.get("monto_usd") or 0))
            self.oc = QLineEdit(self.item.get("oc_numero", "")); self.oc.setPlaceholderText("Número OC si aplica")
            self.doc_status = QLabel(_legal_doc_flag(self.item.get("comprobante", ""), self.doc_src)); self.doc_status.setStyleSheet("font-weight:900;color:#08285a;")
            self.oc_status = QLabel(_legal_doc_flag(self.item.get("oc_documento", ""), self.oc_src)); self.oc_status.setStyleSheet("font-weight:900;color:#08285a;")
            bdoc = QPushButton("Subir comprobante"); bdoc.setObjectName("ghost"); bdoc.clicked.connect(self.pick_doc)
            boc = QPushButton("Subir PDF OC"); boc.setObjectName("ghost"); boc.clicked.connect(self.pick_oc)
            drow=QWidget(); dl=QHBoxLayout(drow); dl.setContentsMargins(0,0,0,0); dl.addWidget(self.doc_status); dl.addWidget(bdoc)
            orow=QWidget(); ol=QHBoxLayout(orow); ol.setContentsMargins(0,0,0,0); ol.addWidget(self.oc_status); ol.addWidget(boc)
            form.addRow("Descripción:", self.desc); form.addRow("Valor USD:", self.monto); form.addRow("OC:", self.oc); form.addRow("Comprobante:", drow); form.addRow("PDF OC:", orow)
            btns=QHBoxLayout(); save=QPushButton("Guardar línea"); save.setObjectName("orange"); cancel=QPushButton("Cancelar"); cancel.setObjectName("ghost"); save.clicked.connect(self.accept); cancel.clicked.connect(self.reject); btns.addWidget(cancel); btns.addWidget(save); lay.addLayout(btns)
        def pick_doc(self):
            path,_=QFileDialog.getOpenFileName(self,"Selecciona comprobante",str(Path.home()),"Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if path: self.doc_src=path; self.doc_status.setText("✅ Cargado")
        def pick_oc(self):
            path,_=QFileDialog.getOpenFileName(self,"Selecciona PDF de OC",str(Path.home()),"Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if path: self.oc_src=path; self.oc_status.setText("✅ Cargado")
        def data(self) -> dict:
            d = dict(self.item)
            d.update({"descripcion": self.desc.text().strip(), "monto_usd": self.monto.value(), "oc_numero": self.oc.text().strip().upper(), "comprobante_src": self.doc_src, "oc_src": self.oc_src})
            return d

    class StageUpdateDialog(QDialog):
        def __init__(self, parent, vehicle: dict, stage_key: str, user: dict, device: DeviceInfo, mode: str = "advance"):
            super().__init__(parent)
            self.vehicle=vehicle; self.stage_key=LEGACY_STAGE_MAP.get(stage_key, stage_key); self.user=user; self.device=device; self.mode=mode
            self.cost_lines=[]; self.foto_path=""; self.repuestos=[]
            self.setWindowTitle(f"{STAGE_META[self.stage_key]['label']} · {vehicle.get('codigo')}")
            self.setMinimumSize(960, 760)
            self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint)
            self._build()
        def _combo(self, values: list[str]) -> QComboBox:
            cb=QComboBox(); cb.setEditable(True); cb.addItems(values); return cb
        def _stage_cost(self, sub: str) -> dict:
            source=f"stage:{self.stage_key}"; subn=_norm(sub)
            for g in self.vehicle.get("gastos_detallados", []) or []:
                if g.get("source")==source and _norm(g.get("subcategoria"))==subn:
                    return g
            return {}
        def _status_label(self, loaded=False) -> QLabel:
            lab=QLabel("✅" if loaded else "❌")
            lab.setFixedWidth(28); lab.setAlignment(Qt.AlignmentFlag.AlignCenter); lab.setStyleSheet("font-size:18px;font-weight:950;")
            return lab
        def _pick_line_doc(self, line: dict, kind: str):
            title = "Selecciona PDF de OC" if kind == "oc" else "Selecciona comprobante"
            path,_=QFileDialog.getOpenFileName(self,title,str(Path.home()),"Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if path:
                line[kind+"_src"] = path
                line[kind+"_label"].setText("✅")
        def _add_cost_line(self, form: QFormLayout, label: str, categoria: str, subcategoria: str, required: bool=False, proveedor: str=""):
            old=self._stage_cost(subcategoria); row=QWidget(); lay=QHBoxLayout(row); lay.setContentsMargins(0,0,0,0)
            money=MoneyEdit(); money.setRange(0,9999999); money.setValue(float(old.get("monto_usd") or 0))
            oc=QLineEdit(old.get("oc_numero", "")); oc.setPlaceholderText("OC"); oc.setMaximumWidth(110)
            doc_label=self._status_label(bool(old.get("comprobante"))); bdoc=QPushButton("Comp."); bdoc.setObjectName("ghost"); bdoc.setMaximumWidth(80)
            oc_label=self._status_label(bool(old.get("oc_documento"))); boc=QPushButton("PDF OC"); boc.setObjectName("ghost"); boc.setMaximumWidth(80)
            line={"categoria":categoria,"subcategoria":subcategoria,"required":required,"monto":money,"oc":oc,"doc_src":"","oc_src":"","doc_label":doc_label,"oc_label":oc_label,"old":old,"proveedor":proveedor}
            bdoc.clicked.connect(lambda _,ln=line:self._pick_line_doc(ln,"doc")); boc.clicked.connect(lambda _,ln=line:self._pick_line_doc(ln,"oc"))
            lay.addWidget(money,2); lay.addWidget(oc,1); lay.addWidget(doc_label); lay.addWidget(bdoc); lay.addWidget(oc_label); lay.addWidget(boc)
            form.addRow(("* " if required else "") + label + ":", row); self.cost_lines.append(line)
        def _collect_costs(self, default_fecha: str, proveedor: str="") -> tuple[list[dict], list[str]]:
            items=[]; missing=[]
            for line in self.cost_lines:
                val=line["monto"].value()
                if line.get("required") and val <= 0:
                    missing.append(line["subcategoria"])
                if val > 0:
                    old=line.get("old") or {}
                    items.append({"id": old.get("id"), "categoria": line["categoria"], "subcategoria": line["subcategoria"], "descripcion": line["subcategoria"].replace("_", " ").title(), "monto_usd": val, "proveedor": proveedor or line.get("proveedor") or "", "oc_numero": line["oc"].text(), "comprobante_src": line.get("doc_src", ""), "oc_src": line.get("oc_src", ""), "comprobante": old.get("comprobante", ""), "comprobante_nombre": old.get("comprobante_nombre", ""), "oc_documento": old.get("oc_documento", ""), "oc_documento_nombre": old.get("oc_documento_nombre", ""), "fecha": default_fecha})
            return items, missing
        def _build(self):
            lay=QVBoxLayout(self); v=self.vehicle; st=vehicle_stage(v,self.stage_key); extra=st.get("extra") if isinstance(st.get("extra"), dict) else {}
            head=QLabel(f"<b>{v.get('codigo')}</b> · {v.get('marca')} {v.get('modelo')} {v.get('anio')} · Costo actual: <b>{_fmt_usd(vehicle_total_cost(v))}</b>")
            head.setTextFormat(Qt.TextFormat.RichText); lay.addWidget(head)
            top_btns=QHBoxLayout(); bg=QPushButton("Visualizar gastos"); bg.setObjectName("ghost"); bg.clicked.connect(self.open_costs); top_btns.addWidget(bg); top_btns.addStretch(1); lay.addLayout(top_btns)
            scroll=QScrollArea(); scroll.setWidgetResizable(True); content=QWidget(); scroll.setWidget(content); main=QVBoxLayout(content); lay.addWidget(scroll)
            form=QFormLayout(); main.addLayout(form)
            min_date = _stage_min_start_date(v, self.stage_key)
            initial_date = st.get("fecha_inicio") or _stage_default_start_date(v, self.stage_key) or _today_iso()
            self.fecha_evento = configure_date_edit(QDateEdit(), initial_date, min_date)
            locked = self.stage_key==STAGE_TRASLADO_USA and st.get("fecha_inicio") and not user_can_override_flow(self.user)
            if locked: self.fecha_evento.setEnabled(False)
            label_map={STAGE_TRASLADO_USA:"Fecha salida de subasta", STAGE_TRANSITO:"Fecha llegada a naviera/yarda", STAGE_ADUANA:"Fecha llegada a aduana", STAGE_PREPARACION:"Fecha inicio preparación", STAGE_PRECIO_FINAL:"Fecha definición precio", STAGE_DISPONIBLE:"Fecha disponible"}
            form.addRow(label_map.get(self.stage_key,"Fecha"), self.fecha_evento)
            self.comentario=QTextEdit(st.get("comentario") or ""); self.comentario.setMinimumHeight(70)
            if self.stage_key == STAGE_TRASLADO_USA:
                self.transportista=self._combo(load_catalog(F_CATALOG_TRANSPORTISTAS_USA, DEFAULT_TRANSPORTISTAS_USA)); self.transportista.setCurrentText(st.get("proveedor") or "")
                self.naviera_entregada=self._combo(load_catalog(F_CATALOG_NAVIERAS, DEFAULT_NAVIERAS)); self.naviera_entregada.setCurrentText(extra.get("naviera_entregada") or "")
                self.fecha_llegada_yarda=OptionalDateEdit(extra.get("fecha_llegada_yarda") or "", self._date(self.fecha_evento), "Ya llegó")
                form.addRow("Transportista USA:", self.transportista); form.addRow("Naviera entregada:", self.naviera_entregada); form.addRow("Fecha llegada a yarda de consolidación:", self.fecha_llegada_yarda)
                self._add_cost_line(form,"Grúa / traslado USA","TRASLADO_USA","GRUA_TRASLADO_USA",False)
            elif self.stage_key == STAGE_TRANSITO:
                prev=vehicle_stage(v, STAGE_TRASLADO_USA); pextra=prev.get("extra") if isinstance(prev.get("extra"),dict) else {}
                self.naviera=self._combo(load_catalog(F_CATALOG_NAVIERAS, DEFAULT_NAVIERAS)); self.naviera.setCurrentText(extra.get("naviera") or pextra.get("naviera_entregada") or "")
                self.fecha_salida_naviera=OptionalDateEdit(extra.get("fecha_salida_naviera") or "", self._date(self.fecha_evento), "Ya salió")
                self.motivo_extra=QTextEdit(extra.get("motivo_extra") or ""); self.motivo_extra.setMinimumHeight(60)
                form.addRow("Naviera:", self.naviera); form.addRow("Fecha salida de naviera / barco:", self.fecha_salida_naviera); form.addRow("Motivo costo extra:", self.motivo_extra)
                self._add_cost_line(form,"Costo extra tránsito","TRANSITO","COSTO_EXTRA_TRANSITO",False)
            elif self.stage_key == STAGE_ADUANA:
                self.pais=self._combo(load_catalog(F_CATALOG_PAISES, DEFAULT_PAISES_DESTINO)); self.pais.setCurrentText(extra.get("pais") or "EL SALVADOR")
                self.aduana=self._combo(aduanas_for_country(self.pais.currentText())); self.aduana.setCurrentText(extra.get("aduana") or "")
                self.pais.currentTextChanged.connect(self._refresh_aduanas)
                self.fecha_liberacion=OptionalDateEdit(extra.get("fecha_liberacion_aduana") or "", self._date(self.fecha_evento), "Liberada")
                form.addRow("País destino:", self.pais); form.addRow("Aduana:", self.aduana); form.addRow("Fecha liberación aduana:", self.fecha_liberacion)
                self._add_cost_line(form,"Naviera · Grúa interna","ADUANA","NAVIERA_GRUA_INTERNA",True)
                self._add_cost_line(form,"Naviera · Flete","ADUANA","NAVIERA_FLETE",True)
                self._add_cost_line(form,"Naviera · BL","ADUANA","NAVIERA_BL",True)
                self._add_cost_line(form,"Pago de impuestos / IVA DUCA","ADUANA","IMPUESTOS_ADUANA",True)
                self._add_cost_line(form,"Almacenamiento","ADUANA","ALMACENAMIENTO_ADUANA",False)
                self._add_cost_line(form,"Servicio trámite aduanal","ADUANA","TRAMITE_ADUANAL",True)
            elif self.stage_key == STAGE_PREPARACION:
                btnrow=QHBoxLayout(); self.btn_hide_leg=QPushButton("Ocultar legalización"); self.btn_hide_leg.setObjectName("ghost"); self.btn_hide_taller=QPushButton("Ocultar taller"); self.btn_hide_taller.setObjectName("ghost"); btnrow.addWidget(self.btn_hide_leg); btnrow.addWidget(self.btn_hide_taller); btnrow.addStretch(1); main.addLayout(btnrow)
                self.gleg=QGroupBox("Legalización"); fl=QFormLayout(self.gleg); main.addWidget(self.gleg)
                self.emision_pedido=OptionalDateEdit(extra.get("emision_pedido") or "", min_date, "Solicitada")
                self.emision_obtenida=OptionalDateEdit(extra.get("emision_obtenida") or "", min_date, "Obtenida")
                self.cita_pedido=OptionalDateEdit(extra.get("cita_pedido") or "", min_date, "Solicitada")
                self.cita_asignada=OptionalDateEdit(extra.get("cita_asignada") or "", min_date, "Asignada")
                self.placas_ingreso=OptionalDateEdit(extra.get("placas_ingreso") or "", min_date, "Ingresado")
                self.placas_entrega=OptionalDateEdit(extra.get("placas_entrega") or "", min_date, "Entregadas")
                self.legal_fin=OptionalDateEdit(extra.get("legalizacion_fin") or "", min_date, "Finalizada")
                for lab,w in [("Fecha pedimos emisión",self.emision_pedido),("Fecha obtuvimos emisión",self.emision_obtenida),("Fecha pedimos cita",self.cita_pedido),("Fecha asignada cita",self.cita_asignada),("Fecha ingresamos placas",self.placas_ingreso),("Fecha entregaron placas",self.placas_entrega),("Fin legalización",self.legal_fin)]: fl.addRow(lab+":", w)
                self._add_cost_line(fl,"Emisiones","LEGALIZACION","EMISIONES",False); self._add_cost_line(fl,"Cita","LEGALIZACION","CITA",False); self._add_cost_line(fl,"Placas","LEGALIZACION","PLACAS",False)
                self.gtaller=QGroupBox("Taller"); ft=QFormLayout(self.gtaller); main.addWidget(self.gtaller)
                self.taller=self._combo(load_catalog(F_CATALOG_TALLERES, DEFAULT_TALLERES)); self.taller.setCurrentText(extra.get("taller") or "")
                self.taller_ingreso=configure_date_edit(QDateEdit(), extra.get("taller_ingreso") or _today_iso(), min_date)
                self.taller_salida=OptionalDateEdit(extra.get("taller_salida") or "", min_date, "Ya salió")
                self.motivo_taller=QTextEdit(extra.get("motivo_taller") or ""); self.motivo_taller.setMinimumHeight(70)
                ft.addRow("Taller:", self.taller); ft.addRow("Fecha ingreso taller:", self.taller_ingreso); ft.addRow("Fecha salida taller:", self.taller_salida); ft.addRow("Motivo / comentarios:", self.motivo_taller)
                self._add_cost_line(ft,"Enderezado y pintura","TALLER","ENDEREZADO_PINTURA",False); self._add_cost_line(ft,"Servicios mecánicos","TALLER","SERVICIOS_MECANICOS",False); self._add_cost_line(ft,"Grúa local","TALLER","GRUA_LOCAL",False)
                self.repuestos_table=QTableWidget(0,5); self.repuestos_table.setHorizontalHeaderLabels(["Descripción", "Valor USD", "OC", "Comp.", "OC PDF"]); self.repuestos_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.repuestos_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.repuestos_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.repuestos_table.cellDoubleClicked.connect(self._edit_repuesto_row)
                for g in self.vehicle.get("gastos_detallados",[]) or []:
                    if g.get("source")==f"stage:{self.stage_key}" and _norm(g.get("subcategoria"))=="REPUESTO_DETALLE":
                        self.repuestos.append(dict(g))
                brow=QHBoxLayout(); badd=QPushButton("+ Agregar repuesto / gasto"); badd.setObjectName("orange"); bedit=QPushButton("Editar"); bedit.setObjectName("ghost"); bdel=QPushButton("Borrar"); bdel.setObjectName("danger"); badd.clicked.connect(self._add_repuesto_row); bedit.clicked.connect(lambda:self._edit_repuesto_row(self.repuestos_table.currentRow(),0)); bdel.clicked.connect(self._delete_repuesto_row); brow.addWidget(badd); brow.addWidget(bedit); brow.addWidget(bdel); brow.addStretch(1)
                main.addWidget(QLabel("Detalle de repuestos / gastos del taller. Doble click para editar en ventana flotante:")); main.addWidget(self.repuestos_table); main.addLayout(brow); self._refresh_repuestos_table()
                self.btn_hide_leg.clicked.connect(self._toggle_legal); self.btn_hide_taller.clicked.connect(self._toggle_taller)
            elif self.stage_key == STAGE_PRECIO_FINAL:
                self._build_price_section(main, form)
            elif self.stage_key == STAGE_DISPONIBLE:
                info=QLabel(f"Precio cliente final: <b>{_fmt_usd(v.get('precio_venta_usd'))}</b><br>Al confirmar esta etapa el carro caerá automáticamente en Cotizaciones como disponible."); info.setTextFormat(Qt.TextFormat.RichText); info.setStyleSheet("background:#dcfce7;padding:12px;border-radius:10px;color:#14532d;"); main.addWidget(info)
            form.addRow("Comentario:", self.comentario)
            btn=QPushButton("Guardar avance / datos de esta etapa"); btn.setObjectName("orange"); btn.clicked.connect(self.save); lay.addWidget(btn)
        def open_costs(self):
            CostBreakdownDialog(self,self.vehicle).exec()
        def _refresh_aduanas(self):
            cur=self.aduana.currentText(); self.aduana.clear(); self.aduana.addItems(aduanas_for_country(self.pais.currentText())); idx=self.aduana.findText(cur); self.aduana.setCurrentIndex(idx if idx>=0 else 0)
        def _toggle_legal(self):
            self.gleg.setVisible(not self.gleg.isVisible()); self.btn_hide_leg.setText("Mostrar legalización" if not self.gleg.isVisible() else "Ocultar legalización")
        def _toggle_taller(self):
            self.gtaller.setVisible(not self.gtaller.isVisible()); self.btn_hide_taller.setText("Mostrar taller" if not self.gtaller.isVisible() else "Ocultar taller")
        def _refresh_repuestos_table(self):
            self.repuestos_table.setRowCount(len(self.repuestos))
            for r,item in enumerate(self.repuestos):
                vals=[item.get("descripcion",""), _fmt_usd(item.get("monto_usd")), item.get("oc_numero",""), _short_doc_status(item.get("comprobante"), item.get("comprobante_src","")), _short_doc_status(item.get("oc_documento"), item.get("oc_src",""))]
                for c,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.repuestos_table.setItem(r,c,it)
        def _add_repuesto_row(self):
            dlg=WorkshopItemDialog(self)
            if dlg.exec()==QDialog.DialogCode.Accepted:
                d=dlg.data()
                if not d.get("descripcion") or float(d.get("monto_usd") or 0) <= 0:
                    QMessageBox.warning(self,"Validación","La descripción y el valor mayor a cero son obligatorios."); return
                self.repuestos.append(d); self._refresh_repuestos_table()
        def _edit_repuesto_row(self,row,col):
            if row < 0 or row >= len(self.repuestos): return
            dlg=WorkshopItemDialog(self,self.repuestos[row])
            if dlg.exec()==QDialog.DialogCode.Accepted:
                d=dlg.data()
                if not d.get("descripcion") or float(d.get("monto_usd") or 0) <= 0:
                    QMessageBox.warning(self,"Validación","La descripción y el valor mayor a cero son obligatorios."); return
                self.repuestos[row]=d; self._refresh_repuestos_table()
        def _delete_repuesto_row(self):
            row=self.repuestos_table.currentRow()
            if row >=0 and row < len(self.repuestos):
                self.repuestos.pop(row); self._refresh_repuestos_table()
        def _build_price_section(self, main, form):
            btnrow=QHBoxLayout()
            bprev=QPushButton("Visualizar gastos")
            bprev.setObjectName("ghost")
            bprev.setMaximumWidth(145)
            bprev.clicked.connect(self.open_costs)
            btnrow.addWidget(bprev)
            btnrow.addStretch(1)
            main.addLayout(btnrow)
            self.margen=QDoubleSpinBox(); self.margen.setRange(0,100); self.margen.setDecimals(2); self.margen.setSuffix(" %"); self.margen.setValue(float((self.vehicle.get("precio_final") or {}).get("margen_pct") or 15))
            self.iva_duca=MoneyEdit(); self.iva_duca.setRange(0,999999); self.iva_duca.setValue(float((self.vehicle.get("precio_final") or {}).get("iva_duca_usd") or _aduana_impuestos_amount(self.vehicle)))
            self.iva_pct=QDoubleSpinBox(); self.iva_pct.setRange(0,30); self.iva_pct.setDecimals(2); self.iva_pct.setSuffix(" %"); self.iva_pct.setValue(float((self.vehicle.get("precio_final") or {}).get("iva_pct") or 13))
            self.pago_cuenta_pct=QDoubleSpinBox(); self.pago_cuenta_pct.setRange(0,20); self.pago_cuenta_pct.setDecimals(2); self.pago_cuenta_pct.setSuffix(" %"); self.pago_cuenta_pct.setValue(float((self.vehicle.get("precio_final") or {}).get("pago_cuenta_pct") or 1.75))
            self.vts=MoneyEdit(); self.vts.setRange(0,999); self.vts.setValue(float((self.vehicle.get("precio_final") or {}).get("vts_usd") or 2.07))
            self.precio_manual=MoneyEdit(); self.precio_manual.setRange(0,9999999)
            saved_pf=self.vehicle.get("precio_final") or {}
            self.precio_manual.setValue(float(saved_pf.get("precio_redondeado_manual_usd") or saved_pf.get("precio_venta_cliente_usd") or 0))
            loaded=bool(self.vehicle.get("foto_principal")); self.foto_label=QLabel(_legal_doc_flag(self.vehicle.get("foto_principal"), self.foto_path)); self.foto_label.setStyleSheet("font-weight:900;color:#08285a;"); bfoto=QPushButton("Subir foto principal propuesta"); bfoto.setObjectName("ghost"); bfoto.clicked.connect(self._pick_foto)
            frow=QWidget(); fl=QHBoxLayout(frow); fl.setContentsMargins(0,0,0,0); fl.addWidget(self.foto_label); fl.addWidget(bfoto)
            self.calc_lbl=QLabel(); self.calc_lbl.setTextFormat(Qt.TextFormat.RichText); self.calc_lbl.setStyleSheet("background:#fff7ed;border:1px solid #fed7aa;border-radius:10px;padding:12px;color:#08285a;")
            for w in [self.margen,self.iva_pct,self.pago_cuenta_pct]: w.valueChanged.connect(self._refresh_price_calc)
            try: self.iva_duca.textChanged.connect(self._refresh_price_calc); self.vts.textChanged.connect(self._refresh_price_calc); self.precio_manual.textChanged.connect(self._refresh_price_calc)
            except Exception: pass
            form.addRow("Margen deseado:", self.margen); form.addRow("IVA DUCA / impuestos aduana:", self.iva_duca); form.addRow("IVA venta:", self.iva_pct); form.addRow("Pago a cuenta:", self.pago_cuenta_pct); form.addRow("VTS fijo:", self.vts); form.addRow("Redondear precio final cliente:", self.precio_manual); form.addRow("Foto principal:", frow); main.addWidget(self.calc_lbl); self._refresh_price_calc()
        def _pick_foto(self):
            path,_=QFileDialog.getOpenFileName(self,"Selecciona foto principal",str(Path.home()),"Imágenes (*.png *.jpg *.jpeg *.webp);;Todos (*.*)")
            if path: self.foto_path=path; self.foto_label.setText("✅ Cargado")
        def _price_calc(self) -> dict:
            costo=vehicle_total_cost(self.vehicle); margen=self.margen.value()/100.0; venta_neta=round(costo*(1+margen),2); iva_venta=round(venta_neta*(self.iva_pct.value()/100.0),2); iva_duca=self.iva_duca.value(); iva_pagar=max(0.0, round(iva_venta-iva_duca,2)); pago_cuenta=round(venta_neta*(self.pago_cuenta_pct.value()/100.0),2); vts=self.vts.value(); precio_cliente=round(venta_neta+iva_venta+vts,2); utilidad=round(venta_neta-costo,2); utilidad_neta_est=round(utilidad-iva_pagar-pago_cuenta-vts,2)
            return {"costo_total_usd":costo,"margen_pct":self.margen.value(),"venta_neta_usd":venta_neta,"iva_pct":self.iva_pct.value(),"iva_venta_usd":iva_venta,"iva_duca_usd":iva_duca,"iva_pagar_usd":iva_pagar,"pago_cuenta_pct":self.pago_cuenta_pct.value(),"pago_cuenta_usd":pago_cuenta,"vts_usd":vts,"precio_venta_cliente_usd":precio_cliente,"precio_minimo_usd":precio_cliente,"utilidad_bruta_usd":utilidad,"utilidad_neta_estimada_usd":utilidad_neta_est}
        def _refresh_price_calc(self):
            pf=self._price_calc(); manual_txt="<br><b>Precio redondeado/manual aplicado.</b>" if pf.get('precio_redondeado_manual_usd') else ""; self.calc_lbl.setText(f"Costo total: <b>{_fmt_usd(pf['costo_total_usd'])}</b><br>Venta neta: <b>{_fmt_usd(pf['venta_neta_usd'])}</b> · IVA venta: <b>{_fmt_usd(pf['iva_venta_usd'])}</b> · IVA DUCA usado: <b>{_fmt_usd(pf['iva_duca_usd'])}</b> · IVA a pagar: <b>{_fmt_usd(pf['iva_pagar_usd'])}</b><br>Pago a cuenta: <b>{_fmt_usd(pf['pago_cuenta_usd'])}</b> · VTS: <b>{_fmt_usd(pf['vts_usd'])}</b><br>Precio cliente final: <b style='font-size:18px'>{_fmt_usd(pf['precio_venta_cliente_usd'])}</b> · Margen real: <b>{pf['margen_pct']}%</b>{manual_txt}<br>Utilidad bruta contra costo: <b>{_fmt_usd(pf['utilidad_bruta_usd'])}</b> · Utilidad neta estimada interna: <b>{_fmt_usd(pf['utilidad_neta_estimada_usd'])}</b>")
        def _date(self, w):
            if isinstance(w, OptionalDateEdit): return w.date_iso()
            return w.date().toPython().isoformat()
        def _validate_order(self, start, end, msg):
            ds=_parse_date(start); de=_parse_date(end)
            if ds and de and de < ds:
                QMessageBox.warning(self,"Fechas",msg); return False
            return True
        def save(self):
            fecha=self._date(self.fecha_evento); fin=None; proveedor=""; extra={}; cost_items=[]; missing=[]
            if self.stage_key==STAGE_TRASLADO_USA:
                proveedor=self.transportista.currentText(); llegada=self._date(self.fecha_llegada_yarda)
                if not self._validate_order(fecha,llegada,"La llegada a yarda no puede ser menor que la salida de subasta."): return
                fin=llegada; extra={"transportista":proveedor,"naviera_entregada":self.naviera_entregada.currentText(),"fecha_llegada_yarda":llegada}; cost_items,missing=self._collect_costs(fecha, proveedor)
            elif self.stage_key==STAGE_TRANSITO:
                proveedor=self.naviera.currentText(); salida=self._date(self.fecha_salida_naviera)
                if not self._validate_order(fecha,salida,"La salida de naviera no puede ser menor que la llegada a naviera/yarda."): return
                fin=salida; extra={"naviera":proveedor,"fecha_salida_naviera":salida,"motivo_extra":self.motivo_extra.toPlainText()}; cost_items,missing=self._collect_costs(fecha, proveedor)
            elif self.stage_key==STAGE_ADUANA:
                proveedor=f"{self.pais.currentText()} · {self.aduana.currentText()}"; lib=self._date(self.fecha_liberacion)
                if not self._validate_order(fecha,lib,"La liberación de aduana no puede ser menor que la fecha de llegada a aduana."): return
                fin=lib; extra={"pais":self.pais.currentText(),"aduana":self.aduana.currentText(),"fecha_liberacion_aduana":lib}; cost_items,missing=self._collect_costs(fecha, proveedor)
            elif self.stage_key==STAGE_PREPARACION:
                proveedor=self.taller.currentText()
                if not self._validate_order(fecha,self._date(self.emision_pedido),"La fecha de emisión no puede ser anterior al inicio de preparación."): return
                if not self._validate_order(self._date(self.emision_pedido),self._date(self.emision_obtenida),"La fecha obtenida de emisión no puede ser anterior a cuando se pidió."): return
                if not self._validate_order(self._date(self.cita_pedido),self._date(self.cita_asignada),"La cita asignada no puede ser anterior a cuando se pidió."): return
                if not self._validate_order(self._date(self.placas_ingreso),self._date(self.placas_entrega),"La entrega de placas no puede ser anterior al ingreso del trámite."): return
                if not self._validate_order(fecha,self._date(self.taller_ingreso),"El ingreso al taller no puede ser anterior al inicio de preparación."): return
                taller_salida=self._date(self.taller_salida); legal_fin=self._date(self.legal_fin)
                if taller_salida and not self._validate_order(self._date(self.taller_ingreso),taller_salida,"La salida del taller no puede ser anterior al ingreso al taller."): return
                if legal_fin and not self._validate_order(fecha,legal_fin,"El fin de legalización no puede ser anterior al inicio de preparación."): return
                dates=[d for d in [legal_fin,taller_salida] if d]
                fin=max(dates) if len(dates)==2 else None
                extra={"emision_pedido":self._date(self.emision_pedido),"emision_obtenida":self._date(self.emision_obtenida),"cita_pedido":self._date(self.cita_pedido),"cita_asignada":self._date(self.cita_asignada),"placas_ingreso":self._date(self.placas_ingreso),"placas_entrega":self._date(self.placas_entrega),"legalizacion_fin":legal_fin,"taller":self.taller.currentText(),"taller_ingreso":self._date(self.taller_ingreso),"taller_salida":taller_salida,"motivo_taller":self.motivo_taller.toPlainText()}; cost_items,missing=self._collect_costs(fecha, proveedor)
                for item in self.repuestos:
                    if item.get("descripcion") and float(item.get("monto_usd") or 0)>0:
                        cost_items.append({"id": item.get("id"), "categoria":"TALLER","subcategoria":"REPUESTO_DETALLE","descripcion":item.get("descripcion"),"monto_usd":float(item.get("monto_usd") or 0),"proveedor":proveedor,"oc_numero":item.get("oc_numero",""),"fecha":fecha,"comprobante_src":item.get("comprobante_src",""),"oc_src":item.get("oc_src",""),"comprobante":item.get("comprobante",""),"comprobante_nombre":item.get("comprobante_nombre",""),"oc_documento":item.get("oc_documento",""),"oc_documento_nombre":item.get("oc_documento_nombre","")})
            elif self.stage_key==STAGE_PRECIO_FINAL:
                pf=self._price_calc(); proveedor="INTERNO"; extra={"precio_final_guardado":True}; data={"fecha_inicio":fecha,"fecha_fin":fecha,"proveedor":proveedor,"comentario":self.comentario.toPlainText(),"extra":extra,"precio_final":pf,"foto_principal_src":self.foto_path}
                ok,msg=update_vehicle_stage(self.vehicle.get("id"), self.stage_key, data, None, self.user, self.device)
                if not ok: QMessageBox.warning(self,"Validación",msg); return
                QMessageBox.information(self,"Guardado",msg); self.accept(); return
            elif self.stage_key==STAGE_DISPONIBLE:
                fin=fecha; proveedor="VENTAS"; extra={"disponible_para_cotizar":True}; missing=[]
            if missing:
                resp=QMessageBox.question(self,"Gastos en cero", "Estos gastos obligatorios están en cero:\n"+"\n".join(missing)+"\n\n¿Seguro que no aplican o no los tuviste?")
                if resp != QMessageBox.StandardButton.Yes: return
            data={"fecha_inicio":fecha,"fecha_fin":fin,"proveedor":proveedor,"comentario":self.comentario.toPlainText(),"extra":extra,"cost_items":cost_items}
            ok,msg=update_vehicle_stage(self.vehicle.get("id"), self.stage_key, data, None, self.user, self.device)
            if not ok: QMessageBox.warning(self,"Validación",msg); return
            QMessageBox.information(self,"Guardado",msg); self.accept()

    class QuoteEditorDialog(QDialog):
        def __init__(self, parent, user: dict, device: DeviceInfo, vehicle_id: str = "", quote_id: str = "", preset_client: Optional[dict] = None):
            super().__init__(parent); self.user=user; self.device=device; self.vehicle_id=vehicle_id; self.quote_id=quote_id; self.preset_client=preset_client or {}; self._vehicle_ids=[]
            self.setWindowTitle("Cotización leasing"); self.setMinimumSize(980,780); self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint); self._build(); self.refresh_calc()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Cotización leasing", "Primero calcula cuota mensual y en paralelo calcula gastos legales sobre el valor del carro."))
            form=QFormLayout(); lay.addLayout(form)
            self.vehicle_combo=QComboBox(); self.vehicle_combo.setEditable(False)
            for v in load_vehicles():
                ensure_vehicle_runtime_fields(v)
                if v.get("estado_actual")==STAGE_DISPONIBLE and v.get("estado_comercial")==COMM_DISPONIBLE:
                    self._vehicle_ids.append(v.get("id")); self.vehicle_combo.addItem(f"{v.get('codigo')} · {v.get('marca')} {v.get('modelo')} {v.get('anio')} · {_fmt_usd(v.get('precio_venta_usd'))}")
            if self.vehicle_id and self.vehicle_id in self._vehicle_ids:
                self.vehicle_combo.setCurrentIndex(self._vehicle_ids.index(self.vehicle_id)); self.vehicle_combo.setEnabled(False)
            self.nombre=QLineEdit(self.preset_client.get("nombre", "")); self.telefono=QLineEdit(self.preset_client.get("telefono", "")); self.correo=QLineEdit(self.preset_client.get("correo", "")); self.medio=QComboBox(); self.medio.setEditable(True); self.medio.addItems(["WHATSAPP","FACEBOOK","TIKTOK","PAGINA WEB","REFERIDO","LLAMADA","OTRO"]); self.medio.setCurrentText(self.preset_client.get("medio_contacto", "WHATSAPP"))
            self.ingreso=MoneyEdit(); self.ingreso.setRange(0,999999); self.ingreso.setValue(1900)
            self.precio=MoneyEdit(); self.precio.setRange(0,999999); self.prima_pct=QDoubleSpinBox(); self.prima_pct.setRange(0,90); self.prima_pct.setDecimals(2); self.prima_pct.setSuffix(" %"); self.prima_pct.setValue(20)
            self.comision=MoneyEdit(); self.comision.setRange(0,9999); self.comision.setValue(100); self.plazo=QSpinBox(); self.plazo.setRange(1,120); self.plazo.setValue(60)
            self.tasa=QComboBox(); self.tasa.setEditable(True); self.tasa.addItems(load_catalog(F_CATALOG_TASAS_LEASING, DEFAULT_TASAS_LEASING)); self.tasa.setCurrentText("2.50")
            self.seguro=MoneyEdit(); self.seguro.setRange(0,9999); self.seguro.setValue(80); self.gps=MoneyEdit(); self.gps.setRange(0,9999); self.gps.setValue(20); self.iva=QDoubleSpinBox(); self.iva.setRange(0,30); self.iva.setDecimals(2); self.iva.setSuffix(" %"); self.iva.setValue(13)
            self.comentario=QTextEdit(); self.comentario.setMinimumHeight(60)
            for label,w in [("Vehículo:",self.vehicle_combo),("Cliente:",self.nombre),("Teléfono:",self.telefono),("Correo:",self.correo),("Medio contacto:",self.medio),("Ingreso mensual cliente:",self.ingreso),("Precio vehículo:",self.precio),("Prima mínima:",self.prima_pct),("Comisión editable:",self.comision),("Plazo meses:",self.plazo),("Tasa rentabilidad mensual:",self.tasa),("Seguro mensual incluye IVA:",self.seguro),("GPS mensual incluye IVA:",self.gps),("IVA sobre cuota base:",self.iva),("Comentario:",self.comentario)]: form.addRow(label,w)
            boxes=QHBoxLayout(); self.result_cuota=QLabel(); self.result_cuota.setTextFormat(Qt.TextFormat.RichText); self.result_cuota.setWordWrap(True); self.result_cuota.setStyleSheet("background:#eef6ff;border:1px solid #bfdbfe;border-radius:10px;padding:12px;color:#08285a;"); self.result_legal=QLabel(); self.result_legal.setTextFormat(Qt.TextFormat.RichText); self.result_legal.setWordWrap(True); self.result_legal.setStyleSheet("background:#f0fdf4;border:1px solid #86efac;border-radius:10px;padding:12px;color:#14532d;"); boxes.addWidget(self.result_cuota,2); boxes.addWidget(self.result_legal,1); lay.addLayout(boxes)
            self.table=QTableWidget(0,5); self.table.setHorizontalHeaderLabels(["Año","Plazo","Cuota financiamiento","Cuota total","Cuota total con IVA"]); self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); lay.addWidget(self.table)
            self.vehicle_combo.currentIndexChanged.connect(self._vehicle_changed); self._vehicle_changed()
            for w in [self.ingreso,self.precio,self.comision,self.seguro,self.gps]:
                try: w.textChanged.connect(self.refresh_calc)
                except Exception: pass
            self.prima_pct.valueChanged.connect(self.refresh_calc); self.plazo.valueChanged.connect(self.refresh_calc); self.iva.valueChanged.connect(self.refresh_calc); self.tasa.currentTextChanged.connect(self.refresh_calc)
            btns=QHBoxLayout(); bsave=QPushButton("Guardar cotización"); bsave.setObjectName("orange"); bsave.clicked.connect(self.save); btns.addWidget(bsave); btns.addStretch(1); lay.addLayout(btns)
        def _vehicle_changed(self):
            if 0 <= self.vehicle_combo.currentIndex() < len(self._vehicle_ids):
                v=find_vehicle(self._vehicle_ids[self.vehicle_combo.currentIndex()]); self.precio.setValue(float((v or {}).get("precio_venta_usd") or 0))
            if hasattr(self, "result_cuota"):
                self.refresh_calc()
        def _tasa_value(self):
            try: return float(str(self.tasa.currentText()).replace("%","").replace(",","."))
            except Exception: return 2.5
        def _calc_data(self):
            precio=self.precio.value(); pago_inicial=round(precio*self.prima_pct.value()/100.0 + self.comision.value(),2); calc=calculate_leasing(precio,self.ingreso.value(),pago_inicial,self.plazo.value(),self._tasa_value(),self.seguro.value(),self.gps.value(),self.iva.value()); legal=calculate_legal_fees(precio); return calc,legal,pago_inicial
        def refresh_calc(self):
            if not hasattr(self,"result_cuota") or not hasattr(self,"table"): return
            calc,legal,pago=self._calc_data()
            self.result_cuota.setText(f"<b>1) Simulador de cuota</b><br>Monto leasing: <b>{_fmt_usd(calc['monto_leasing'])}</b> · Prima + comisión: <b>{_fmt_usd(pago)}</b><br>Cuota base: <b>{_fmt_usd(calc['cuota_base'])}</b><br>Cuota total con IVA: <b style='font-size:19px'>{_fmt_usd(calc['cuota_total_con_iva'])}</b><br>% ingreso destinado a cuota: <b>{calc['pct_ingreso']}%</b> · Riesgo: <b>{calc['riesgo_texto']} / {calc['riesgo']}</b>")
            self.result_legal.setText(f"<b>2) Gastos legales</b><br>Valor del carro: <b>{_fmt_usd(legal['valor_leasing_vehiculo'])}</b><br>1.5%: <b>{_fmt_usd(legal['valor_pct'])}</b> + Base fija: <b>{_fmt_usd(legal['base_fija'])}</b><br>Subtotal: <b>{_fmt_usd(legal['subtotal'])}</b> · Tope: <b>{_fmt_usd(legal['tope'])}</b><br>Total legal IVA incluido: <b style='font-size:18px'>{_fmt_usd(legal['valor_legales_iva_incluido'])}</b><br>Aplica tope: <b>{legal['aplica_tope']}</b>")
            rows=leasing_table(self.precio.value(),self.ingreso.value(),pago,self._tasa_value(),self.seguro.value(),self.gps.value(),self.iva.value()); self.table.setRowCount(len(rows))
            for r,row in enumerate(rows):
                vals=[row['anio'],row['plazo'],_fmt_usd(row['cuota_base']),_fmt_usd(row['cuota_total_sin_iva']),_fmt_usd(row['cuota_total_con_iva'])]
                for c,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.table.setItem(r,c,it)
        def save(self):
            if not self._vehicle_ids: QMessageBox.warning(self,"Cotización","No hay vehículos disponibles para cotizar."); return
            data={"vehicle_id":self._vehicle_ids[self.vehicle_combo.currentIndex()],"cliente_nombre":self.nombre.text(),"telefono":self.telefono.text(),"correo":self.correo.text(),"medio_contacto":self.medio.currentText(),"ingreso_cliente":self.ingreso.value(),"precio_vehiculo":self.precio.value(),"prima_pct":self.prima_pct.value(),"comision_usd":self.comision.value(),"plazo_meses":self.plazo.value(),"tasa_mensual_pct":self._tasa_value(),"seguro_mensual":self.seguro.value(),"gps_mensual":self.gps.value(),"iva_pct":self.iva.value(),"comentario":self.comentario.toPlainText(),"fecha_cotizacion":date.today().isoformat()}
            ok,msg,qid=create_quote(data,self.user,self.device,self.quote_id)
            if not ok: QMessageBox.warning(self,"Cotización",msg); return
            QMessageBox.information(self,"Cotización",msg); self.quote_id=qid; self.accept()

    class VehicleQuotesDialog(QDialog):
        def __init__(self, parent, vehicle_id: str, user: dict, device: DeviceInfo):
            super().__init__(parent); self.vehicle_id=vehicle_id; self.user=user; self.device=device; self._quote_ids=[]
            self.setWindowTitle("Personas interesadas / cotizaciones del carro"); self.setMinimumSize(980,620); self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint); self._build(); self.refresh()
        def _build(self):
            lay=QVBoxLayout(self); self.header=QLabel(); self.header.setStyleSheet("font-size:20px;font-weight:950;color:#08285a;"); lay.addWidget(self.header)
            row=QHBoxLayout(); bnew=QPushButton("Nueva cotización para este carro"); bnew.setObjectName("orange"); bnew.clicked.connect(self.new_quote); bclose=QPushButton("Cerrar"); bclose.setObjectName("ghost"); bclose.clicked.connect(self.accept); row.addWidget(bnew); row.addStretch(1); row.addWidget(bclose); lay.addLayout(row)
            self.table=QTableWidget(0,9); self.table.setHorizontalHeaderLabels(["Cliente","Teléfono","Medio","Fecha","Última gestión","Días","Estado","Color","Cuota final"]); self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.table.cellDoubleClicked.connect(self.open_quote); lay.addWidget(self.table)
        def refresh(self):
            v=find_vehicle(self.vehicle_id) or {}; self.header.setText(f"{v.get('codigo','')} · {v.get('marca','')} {v.get('modelo','')} {v.get('anio','')} · Precio { _fmt_usd(v.get('precio_venta_usd')) }")
            qs=quotes_for_vehicle(self.vehicle_id); self._quote_ids=[]; self.table.setRowCount(len(qs))
            for r,q in enumerate(qs):
                self._quote_ids.append(q.get("id")); cl=q.get("cliente",{}); le=q.get("leasing",{}); lvl=quote_alert_level(q); vals=[cl.get("nombre"),cl.get("telefono"),cl.get("medio_contacto"),_fmt_date(q.get("fecha_cotizacion")),_fmt_date(q.get("ultima_gestion")),str(quote_days_without_purchase(q)),QUOTE_STATUS_LABELS.get(q.get("estado"),q.get("estado")),lvl,_fmt_usd(le.get("cuota_total_con_iva"))]
                for c,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setBackground(QColor(_quote_alert_color(lvl))); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.table.setItem(r,c,it)
        def new_quote(self):
            dlg=QuoteEditorDialog(self,self.user,self.device,vehicle_id=self.vehicle_id)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def open_quote(self,row,col):
            if row < len(self._quote_ids):
                dlg=QuoteDetailDialog(self,self._quote_ids[row],self.user,self.device); dlg.exec(); self.refresh()

    class CotizacionesPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main=main; self._vehicle_ids=[]; self._quote_ids=[]; self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Cotizaciones", "Vehículos disponibles, clientes interesados y cotizaciones generales sin borrar historial."))
            self.tabs=QTabWidget(); lay.addWidget(self.tabs)
            wv=QWidget(); vl=QVBoxLayout(wv); brow=QHBoxLayout(); bnew=QPushButton("Nueva cotización"); bnew.setObjectName("orange"); bnew.clicked.connect(self.new_quote); brow.addWidget(bnew); brow.addStretch(1); vl.addLayout(brow)
            self.vehicle_table=QTableWidget(0,7); self.vehicle_table.setHorizontalHeaderLabels(["Código","Vehículo","Precio","Costo","Ganancia","Cotizaciones","Días disponible"]); self.vehicle_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.vehicle_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.vehicle_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.vehicle_table.cellDoubleClicked.connect(self.open_vehicle_quotes); vl.addWidget(self.vehicle_table); self.tabs.addTab(wv,"Vehículos disponibles")
            wg=QWidget(); gl=QVBoxLayout(wg); self.quote_table=QTableWidget(0,9); self.quote_table.setHorizontalHeaderLabels(["Cliente","Teléfono","Medio","Vehículo cotizado","Fecha","Última gestión","Días","Estado","Color"]); self.quote_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.quote_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.quote_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.quote_table.cellDoubleClicked.connect(self.open_quote); gl.addWidget(self.quote_table); self.tabs.addTab(wg,"Cotizaciones general")
        def refresh(self):
            vehicles=[v for v in load_vehicles() if ensure_vehicle_runtime_fields(v) and v.get("estado_actual")==STAGE_DISPONIBLE and v.get("estado_comercial")==COMM_DISPONIBLE]
            self._vehicle_ids=[]; self.vehicle_table.setRowCount(len(vehicles))
            for r,v in enumerate(vehicles):
                self._vehicle_ids.append(v.get("id")); vals=[v.get("codigo"),f"{v.get('marca')} {v.get('modelo')} {v.get('anio')}",_fmt_usd(v.get("precio_venta_usd")),_fmt_usd(vehicle_total_cost(v)),_fmt_usd(vehicle_expected_profit(v)),str(len(quotes_for_vehicle(v.get("id")))),str(current_stage_days(v))]
                for c,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.vehicle_table.setItem(r,c,it)
            qs=load_quotes(); self._quote_ids=[]; self.quote_table.setRowCount(len(qs))
            for r,q in enumerate(qs):
                self._quote_ids.append(q.get("id")); cl=q.get("cliente",{}); snap=q.get("vehicle_snapshot",{}); lvl=quote_alert_level(q); vals=[cl.get("nombre"),cl.get("telefono"),cl.get("medio_contacto"),f"{snap.get('anio','')} {snap.get('marca','')} {snap.get('modelo','')} · {q.get('vehicle_code')}",_fmt_date(q.get("fecha_cotizacion")),_fmt_date(q.get("ultima_gestion")),str(quote_days_without_purchase(q)),QUOTE_STATUS_LABELS.get(q.get("estado"),q.get("estado")),lvl]
                for c,val in enumerate(vals):
                    item=QTableWidgetItem(str(val)); item.setBackground(QColor(_quote_alert_color(lvl))); item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable); self.quote_table.setItem(r,c,item)
        def new_quote(self):
            dlg=QuoteEditorDialog(self,self.main.user,self.main.device)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh(); self.main.refresh_all()
        def open_vehicle_quotes(self,row,col):
            if row < len(self._vehicle_ids):
                dlg=VehicleQuotesDialog(self,self._vehicle_ids[row],self.main.user,self.main.device); dlg.exec(); self.refresh(); self.main.refresh_all()
        def open_quote(self,row,col):
            if row < len(self._quote_ids):
                dlg=QuoteDetailDialog(self,self._quote_ids[row],self.main.user,self.main.device); dlg.exec(); self.refresh(); self.main.refresh_all()



# =============================================================================
# LYM AUTO CONTROL V4.2 - PROPUESTAS PDF/EXCEL, UI COTIZACIONES Y REPORTERIA PREMIUM
# =============================================================================
APP_VERSION = "2.1.0_LEASING"


def _to_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return float(default)
        return float(str(value).replace("$", "").replace(",", "").replace("%", "").strip())
    except Exception:
        return float(default)


def _vehicle_display_name(data: dict) -> str:
    """Formato comercial solicitado: nombre del carro primero y año al final."""
    marca = str(data.get("marca") or "").strip()
    modelo = str(data.get("modelo") or "").strip()
    anio = str(data.get("anio") or "").strip()
    parts = [p for p in [marca, modelo, anio] if p]
    return " ".join(parts).strip() or str(data.get("codigo") or "VEHICULO")


def _proposal_output_dir() -> Optional[Path]:
    df = get_data_folder()
    if df is None:
        return None
    out = df / SUB_REPORTES / "PROPUESTAS"
    out.mkdir(parents=True, exist_ok=True)
    return out


def _safe_quote_filename(quote: dict, ext: str) -> str:
    snap = quote.get("vehicle_snapshot", {})
    cl = quote.get("cliente", {})
    base = f"PROPUESTA_LEASING_{_safe_filename(cl.get('nombre','CLIENTE'))}_{_safe_filename(_vehicle_display_name(snap))}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    return base + ext


def _quote_prima_requerida(leasing: dict) -> float:
    return round(_to_float(leasing.get("pago_inicial"), 0), 2)


def _quote_financial_summary(quote: dict) -> dict:
    leasing = quote.get("leasing", {}) or {}
    legal = quote.get("legal", {}) or {}
    return {
        "valor_vehiculo": round(_to_float(leasing.get("precio_vehiculo"), 0), 2),
        "prima_requerida": _quote_prima_requerida(leasing),
        "monto_leasing": round(_to_float(leasing.get("monto_leasing"), 0), 2),
        "plazo": int(_to_float(leasing.get("plazo_meses"), 0)),
        "tasa": round(_to_float(leasing.get("tasa_mensual_pct"), 0), 4),
        "cuota_base": round(_to_float(leasing.get("cuota_base"), 0), 2),
        "cuota_final": round(_to_float(leasing.get("cuota_total_con_iva"), 0), 2),
        "seguro": round(_to_float(leasing.get("seguro_mensual"), 0), 2),
        "gps": round(_to_float(leasing.get("gps_mensual"), 0), 2),
        "legal": round(_to_float(legal.get("valor_legales_iva_incluido"), 0), 2),
        "opcion_compra": round(_to_float(quote.get("opcion_compra_usd"), 500), 2),
        "prima_pct": round(_to_float(leasing.get("prima_pct"), 20), 4),
        "comision": round(_to_float(leasing.get("comision_usd"), 100), 2),
    }


def _split_features(snapshot: dict) -> list[str]:
    raw = str(snapshot.get("caracteristicas") or "").replace(";", "\n")
    items = [x.strip(" •-\t") for x in raw.splitlines() if x.strip(" •-\t")]
    base = []
    if snapshot.get("anio"):
        base.append(f"Año {snapshot.get('anio')}")
    if snapshot.get("millaje") not in (None, ""):
        try:
            base.append(f"{int(snapshot.get('millaje') or 0):,} millas")
        except Exception:
            base.append(f"Millaje {snapshot.get('millaje')}")
    if snapshot.get("color"):
        base.append(f"Color {snapshot.get('color')}")
    full = []
    seen = set()
    for x in base + items:
        k = _norm(x)
        if k and k not in seen:
            full.append(x); seen.add(k)
    if not full:
        full = ["Vehículo revisado por L&M Inversiones", "Aire acondicionado", "Condiciones listas para propuesta"]
    return full[:18]


def create_quote(data: dict, user: dict, device: DeviceInfo, quote_id: str = "") -> tuple[bool, str, str]:
    """Crea o recalcula una cotización sin duplicarla cuando quote_id ya existe."""
    if not user_has_permission(user, PERM_CREATE_QUOTES):
        return False, "No tienes permiso para crear cotizaciones.", ""
    old_quote = find_quote(quote_id) if quote_id else None
    vehicle_id = data.get("vehicle_id") or (old_quote or {}).get("vehicle_id") or ""
    vehicle = find_vehicle(vehicle_id)
    if not vehicle:
        return False, "Vehículo no encontrado.", ""
    ensure_vehicle_runtime_fields(vehicle)
    if vehicle.get("estado_actual") != STAGE_DISPONIBLE:
        return False, "Solo puedes cotizar vehículos que terminaron el proceso de compra y están disponibles.", ""
    if vehicle.get("estado_comercial") == COMM_VENDIDO and not quote_id:
        return False, "Este vehículo ya está vendido. Ofrece otro carro disponible al cliente.", ""
    cliente_nombre = str(data.get("cliente_nombre") or (old_quote or {}).get("cliente", {}).get("nombre") or "").strip()
    telefono = str(data.get("telefono") or (old_quote or {}).get("cliente", {}).get("telefono") or "").strip()
    if not cliente_nombre or not telefono:
        return False, "Nombre y teléfono del cliente son obligatorios.", ""
    precio = round(_to_float(data.get("precio_vehiculo"), float(vehicle.get("precio_venta_usd") or 0)), 2)
    ingreso = round(_to_float(data.get("ingreso_cliente"), (old_quote or {}).get("leasing", {}).get("ingreso_cliente", 0)), 2)
    prima_pct = round(_to_float(data.get("prima_pct"), (old_quote or {}).get("leasing", {}).get("prima_pct", 20)), 4)
    comision = round(_to_float(data.get("comision_usd"), (old_quote or {}).get("leasing", {}).get("comision_usd", 100)), 2)
    pago_inicial = round(_to_float(data.get("pago_inicial"), precio * prima_pct / 100.0 + comision), 2)
    plazo = int(_to_float(data.get("plazo_meses"), (old_quote or {}).get("leasing", {}).get("plazo_meses", 60)))
    tasa = _to_float(data.get("tasa_mensual_pct"), (old_quote or {}).get("leasing", {}).get("tasa_mensual_pct", 2.5))
    seguro = round(_to_float(data.get("seguro_mensual"), (old_quote or {}).get("leasing", {}).get("seguro_mensual", 80)), 2)
    gps = round(_to_float(data.get("gps_mensual"), (old_quote or {}).get("leasing", {}).get("gps_mensual", 20)), 2)
    iva_pct = _to_float(data.get("iva_pct"), (old_quote or {}).get("leasing", {}).get("iva_pct", 13))
    leasing = calculate_leasing(precio, ingreso, pago_inicial, plazo, tasa, seguro, gps, iva_pct)
    legal = calculate_legal_fees(precio)
    quote = old_quote or {"id": uuid.uuid4().hex, "fecha_creacion": _now_iso(), "creado_por": user.get("usuario", "")}
    ensure_quote_runtime_fields(quote)
    old_vehicle = quote.get("vehicle_id")
    if old_vehicle and old_vehicle != vehicle.get("id"):
        quote.setdefault("seguimientos", []).append({"fecha": _now_iso(), "usuario": user.get("usuario", ""), "accion": "REOFERTA", "comentario": f"Se ofreció otro vehículo: {vehicle.get('codigo')}"})
    now = _now_iso()
    quote.update({
        "vehicle_id": vehicle.get("id"),
        "vehicle_code": vehicle.get("codigo"),
        "fecha_cotizacion": data.get("fecha_cotizacion") or quote.get("fecha_cotizacion") or _today_iso(),
        "ultima_gestion": data.get("ultima_gestion") or _today_iso(),
        "estado": data.get("estado") or (quote.get("estado") if quote_id else QUOTE_ENVIADA),
        "cliente": {
            "nombre": cliente_nombre,
            "telefono": telefono,
            "correo": str(data.get("correo") or (old_quote or {}).get("cliente", {}).get("correo") or "").strip(),
            "medio_contacto": _norm(data.get("medio_contacto") or (old_quote or {}).get("cliente", {}).get("medio_contacto") or "WHATSAPP"),
        },
        "vehicle_snapshot": {
            "codigo": vehicle.get("codigo"),
            "marca": vehicle.get("marca"),
            "modelo": vehicle.get("modelo"),
            "anio": vehicle.get("anio"),
            "millaje": vehicle.get("millaje"),
            "color": vehicle.get("color"),
            "precio_venta_usd": precio,
            "foto_principal": vehicle.get("foto_principal"),
            "foto_principal_nombre": vehicle.get("foto_principal_nombre"),
            "caracteristicas": vehicle.get("caracteristicas", ""),
        },
        "leasing": {**leasing, "prima_pct": prima_pct, "comision_usd": comision, "prima_requerida_usd": pago_inicial},
        "legal": legal,
        "opcion_compra_usd": _to_float(data.get("opcion_compra_usd"), (old_quote or {}).get("opcion_compra_usd", 500)),
        "fecha_actualizacion": now,
        "actualizado_por": user.get("usuario", ""),
        "comentario": str(data.get("comentario") or quote.get("comentario") or "").strip(),
    })
    accion = "COTIZACION_RECALCULADA" if quote_id else "COTIZACION_GUARDADA"
    quote.setdefault("seguimientos", []).append({
        "fecha": now,
        "usuario": user.get("usuario", ""),
        "accion": accion,
        "comentario": f"{_vehicle_display_name(quote.get('vehicle_snapshot', {}))} · {plazo} meses · cuota {_fmt_usd(leasing['cuota_total_con_iva'])}",
    })
    if upsert_quote(quote):
        log_audit(accion, user.get("usuario", ""), vehicle.get("codigo", ""), f"Cliente {cliente_nombre}")
        return True, "Cotización actualizada correctamente." if quote_id else "Cotización guardada correctamente.", quote.get("id")
    return False, "No se pudo guardar la cotización.", ""


def generate_quote_proposal_pdf(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada.", None
    out_dir = _proposal_output_dir()
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    except Exception:
        return False, "Para generar PDF instala: pip install reportlab", None
    ensure_quote_runtime_fields(quote)
    cl = quote.get("cliente", {})
    snap = quote.get("vehicle_snapshot", {})
    leasing = quote.get("leasing", {})
    legal = quote.get("legal", {})
    summary = _quote_financial_summary(quote)
    vehicle_name = _vehicle_display_name(snap).upper()
    filename = _safe_quote_filename(quote, ".pdf")
    out = out_dir / filename
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleLYM", parent=styles["Title"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=17, leading=20, textColor=colors.HexColor("#061F4A"), underlineWidth=1))
    styles.add(ParagraphStyle(name="Vehicle", parent=styles["Heading2"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=16, leading=19, textColor=colors.HexColor("#061F4A")))
    styles.add(ParagraphStyle(name="BodyLYM", parent=styles["BodyText"], fontSize=10.5, leading=15, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="SmallLYM", parent=styles["BodyText"], fontSize=8.6, leading=11))
    doc = SimpleDocTemplate(str(out), pagesize=letter, rightMargin=0.55*inch, leftMargin=0.55*inch, topMargin=0.42*inch, bottomMargin=0.45*inch)
    story = []
    logo = ResourceManager.find_logo()
    logo_flow = Image(str(logo), width=0.85*inch, height=0.85*inch) if logo and logo.exists() else Paragraph("<b>L&M</b>", styles["TitleLYM"])
    header = Table([[logo_flow, Paragraph("<b>L&amp;M Inversiones, S.A. de C.V.</b><br/><font size='8'>Ayudando a lograr tus sueños</font>", styles["SmallLYM"])]], colWidths=[1.1*inch, 6.0*inch])
    header.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP"), ("LINEBELOW", (1,0), (1,0), 3, colors.HexColor("#061F4A")), ("BOTTOMPADDING", (0,0), (-1,-1), 8)]))
    story.append(header)
    story.append(Spacer(1, 0.10*inch))
    story.append(Paragraph(f"<b>{html.escape(cl.get('nombre','Cliente').upper())}</b><br/>Presente.", styles["BodyLYM"]))
    story.append(Spacer(1, 0.14*inch))
    story.append(Paragraph("Propuesta de Arrendamiento Vehicular", styles["TitleLYM"]))
    story.append(Spacer(1, 0.14*inch))
    story.append(Paragraph("Por medio de la presente, <b>L&amp;M Inversiones, S.A. de C.V.</b> tiene el agrado de presentarle la propuesta de contrato de arrendamiento (Leasing) opción: <b>RENT A CAR.</b>", styles["BodyLYM"]))
    story.append(Spacer(1, 0.14*inch))
    story.append(Paragraph(f"◆ {html.escape(vehicle_name)} ◆", styles["Vehicle"]))
    story.append(Spacer(1, 0.10*inch))
    story.append(Paragraph("<b>Características destacadas:</b>", styles["Heading3"]))
    feats = _split_features(snap)
    left = feats[0::2]; right = feats[1::2]
    max_rows = max(len(left), len(right))
    rows = []
    for i in range(max_rows):
        rows.append([Paragraph("• " + html.escape(left[i]) if i < len(left) else "", styles["BodyLYM"]), Paragraph("• " + html.escape(right[i]) if i < len(right) else "", styles["BodyLYM"])])
    ft = Table(rows, colWidths=[3.55*inch, 3.55*inch])
    ft.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP"), ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4)]))
    story.append(ft)
    story.append(Spacer(1, 0.14*inch))
    story.append(Paragraph(f"Un vehículo eficiente y equipado para manejar con estilo, seguridad y confianza todos los días. Compartimos ante usted el detalle de plazo y cuota mensual para <b>{html.escape(vehicle_name)}</b>. El cálculo incluye seguro estimado, servicio GPS y una tasa del <b>{summary['tasa']}%</b> mensual.", styles["BodyLYM"]))
    story.append(Spacer(1, 0.12*inch))
    money_rows = [
        [Paragraph("<b>Valor del vehículo</b>", styles["BodyLYM"]), Paragraph(f"<b>{_fmt_usd(summary['valor_vehiculo'])}</b>", styles["BodyLYM"]), Paragraph("<b>Prima requerida</b>", styles["BodyLYM"]), Paragraph(f"<b>{_fmt_usd(summary['prima_requerida'])}</b>", styles["BodyLYM"])],
        [Paragraph("<b>Monto leasing</b>", styles["BodyLYM"]), Paragraph(f"<b>{_fmt_usd(summary['monto_leasing'])}</b>", styles["BodyLYM"]), Paragraph("<b>Costo legal</b>", styles["BodyLYM"]), Paragraph(f"<b>{_fmt_usd(summary['legal'])}</b> incluye IVA", styles["BodyLYM"])],
        [Paragraph("<b>Opción de compra</b>", styles["BodyLYM"]), Paragraph(f"<b>{_fmt_usd(summary['opcion_compra'])}</b>", styles["BodyLYM"]), Paragraph("<b>Prima mínima</b>", styles["BodyLYM"]), Paragraph(f"{summary['prima_pct']}% + comisión {_fmt_usd(summary['comision'])}", styles["BodyLYM"])],
    ]
    mt = Table(money_rows, colWidths=[1.55*inch, 1.85*inch, 1.55*inch, 2.15*inch])
    mt.setStyle(TableStyle([("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E1")), ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#F8FAFC")), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("TOPPADDING", (0,0), (-1,-1), 7), ("BOTTOMPADDING", (0,0), (-1,-1), 7)]))
    story.append(mt)
    story.append(Spacer(1, 0.14*inch))
    selected = Table([
        ["Plazo (meses)", "Tasa rentabilidad (%)", "Cuota base (US$)", "Cuota total mensual incluye seguro, GPS e IVA (US$)"],
        [str(summary["plazo"]), f"{summary['tasa']}%", _fmt_usd(summary["cuota_base"]), _fmt_usd(summary["cuota_final"])],
    ], colWidths=[1.1*inch, 1.6*inch, 1.4*inch, 3.0*inch])
    selected.setStyle(TableStyle([("GRID", (0,0), (-1,-1), 0.75, colors.black), ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#F8FAFC")), ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"), ("FONTNAME", (3,1), (3,1), "Helvetica-Bold"), ("TOPPADDING", (0,0), (-1,-1), 8), ("BOTTOMPADDING", (0,0), (-1,-1), 8)]))
    story.append(selected)
    story.append(Paragraph("<font color='red'><b>La cuota mensual incluye IVA, seguro y GPS.</b></font>", styles["BodyLYM"]))
    story.append(Spacer(1, 0.16*inch))
    story.append(Paragraph("Quedamos atentos a cualquier consulta o duda.", styles["BodyLYM"]))
    story.append(Spacer(1, 0.13*inch))
    story.append(Paragraph(f"San Salvador, {date.today().strftime('%d de %B de %Y')}", styles["BodyLYM"]))
    story.append(PageBreak())
    story.append(header)
    story.append(Spacer(1, 0.18*inch))
    story.append(Paragraph("Atentamente,", styles["BodyLYM"]))
    story.append(Spacer(1, 0.20*inch))
    story.append(Paragraph("<b>Guillermo Moreno</b><br/>L&amp;M Inversiones, S.A. de C.V.<br/>guillermo.moreno@lyminversiones.com<br/>Tel: (503) 7475 5821", styles["BodyLYM"]))
    story.append(Spacer(1, 0.30*inch))
    story.append(Paragraph("X____________________ ___________<br/>Aceptado por cliente", styles["BodyLYM"]))
    story.append(Spacer(1, 0.20*inch))
    story.append(Paragraph("<b>Condiciones y Vigencia de la Oferta</b>", styles["Heading3"]))
    conditions = [
        "Esta propuesta forma parte de una oferta especial de arrendamiento vehicular (leasing) válida por 15 días calendario a partir de la fecha de emisión.",
        f"La tasa de interés mensual del {summary['tasa']}% brinda condiciones accesibles y transparentes para nuestros clientes.",
        "La propuesta se formalizará mediante un contrato de arrendamiento donde se establecerán las condiciones de uso, pagos mensuales y opción de adquisición al finalizar el plazo.",
        "El valor del seguro mostrado es estimado y puede variar según las características y uso del vehículo; el costo final se definirá con base en cotización de aseguradora.",
        "Los gastos administrativos y legales derivados de la formalización del contrato deberán cancelarse al momento de la firma.",
        "Las cuotas mensuales incluyen IVA, seguro y GPS, reflejando el monto total a pagar.",
        "Durante el período de arrendamiento, el cliente deberá mantener el vehículo en buen estado y cumplir con los servicios preventivos o correctivos recomendados.",
        "Esta propuesta tiene carácter informativo y no constituye un compromiso contractual hasta la formalización del contrato de arrendamiento financiero.",
    ]
    for c in conditions:
        story.append(Paragraph("• " + html.escape(c), styles["SmallLYM"]))
        story.append(Spacer(1, 0.04*inch))
    story.append(Spacer(1, 0.18*inch))
    story.append(Paragraph("<para align='center'><font size='9'>Ayudando a lograr tus sueños</font></para>", styles["SmallLYM"]))
    doc.build(story)
    log_audit("GENERAR_PROPUESTA_PDF", (user or {}).get("usuario", ""), quote.get("vehicle_code", ""), out.name)
    return True, "Propuesta PDF generada.", out


def generate_quote_proposal_excel(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada.", None
    out_dir = _proposal_output_dir()
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from openpyxl.drawing.image import Image as XLImage
    except Exception:
        return False, "Para generar Excel instala: pip install openpyxl pillow", None
    ensure_quote_runtime_fields(quote)
    cl = quote.get("cliente", {})
    snap = quote.get("vehicle_snapshot", {})
    leasing = quote.get("leasing", {})
    legal = quote.get("legal", {})
    s = _quote_financial_summary(quote)
    vehicle_name = _vehicle_display_name(snap).upper()
    out = out_dir / _safe_quote_filename(quote, ".xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Propuesta Leasing"
    navy = "08285A"; orange = "F59A13"; blue = "0E3A78"; light = "F8FAFC"
    thin = Side(style="thin", color="1F2937")
    for c in range(1, 9):
        ws.column_dimensions[get_column_letter(c)].width = [15, 18, 18, 18, 18, 18, 18, 18][c-1]
    ws.row_dimensions[1].height = 55
    logo = ResourceManager.find_logo()
    if logo and logo.exists():
        try:
            img = XLImage(str(logo)); img.width = 85; img.height = 85; ws.add_image(img, "A1")
        except Exception:
            ws["A1"] = "L&M"
    ws.merge_cells("B2:H2"); ws["B2"] = "L&M INVERSIONES, S.A. DE C.V."; ws["B2"].font = Font(bold=True, size=14, color=navy); ws["B2"].alignment = Alignment(horizontal="center")
    ws.merge_cells("B3:H3"); ws["B3"] = "Ayudando a lograr tus sueños"; ws["B3"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A5:H5"); ws["A5"] = "PROPUESTA DE ARRENDAMIENTO VEHICULAR"; ws["A5"].font = Font(bold=True, underline="single", size=16, color=navy); ws["A5"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A7:H7"); ws["A7"] = f"◆ {vehicle_name} ◆"; ws["A7"].font = Font(bold=True, size=15, color=navy); ws["A7"].alignment = Alignment(horizontal="center")
    ws["A9"] = "Cliente"; ws["B9"] = cl.get("nombre", "")
    ws["A10"] = "Teléfono"; ws["B10"] = cl.get("telefono", "")
    ws["D9"] = "Fecha"; ws["E9"] = date.today()
    ws["D10"] = "Medio"; ws["E10"] = cl.get("medio_contacto", "")
    for cell in ["A9", "A10", "D9", "D10"]:
        ws[cell].font = Font(bold=True, color=navy)
    ws["E9"].number_format = "dd/mm/yyyy"
    ws.merge_cells("A12:H12"); ws["A12"] = "Resumen para el cliente"; ws["A12"].fill = PatternFill("solid", fgColor=navy); ws["A12"].font = Font(bold=True, color="FFFFFF"); ws["A12"].alignment = Alignment(horizontal="center")
    summary_rows = [
        ["Valor del vehículo", s["valor_vehiculo"], "Prima requerida", s["prima_requerida"], "Monto leasing", s["monto_leasing"]],
        ["Plazo", f"{s['plazo']} meses", "Tasa mensual", f"{s['tasa']}%", "Cuota mensual", s["cuota_final"]],
        ["Costo legal", s["legal"], "Opción compra", s["opcion_compra"], "Seguro + GPS", s["seguro"] + s["gps"]],
    ]
    start = 13
    for r, row in enumerate(summary_rows, start=start):
        for i, val in enumerate(row, start=1):
            ws.cell(r, i).value = val
            ws.cell(r, i).border = Border(top=thin, bottom=thin, left=thin, right=thin)
            ws.cell(r, i).alignment = Alignment(horizontal="center", vertical="center")
            if i in (1,3,5):
                ws.cell(r, i).font = Font(bold=True, color=navy); ws.cell(r, i).fill = PatternFill("solid", fgColor=light)
            else:
                ws.cell(r, i).font = Font(bold=True)
                if isinstance(val, (int,float)): ws.cell(r, i).number_format = '$#,##0.00'
    ws.merge_cells("A18:H18"); ws["A18"] = "Características destacadas"; ws["A18"].fill = PatternFill("solid", fgColor=orange); ws["A18"].font = Font(bold=True, color=navy); ws["A18"].alignment = Alignment(horizontal="center")
    feats = _split_features(snap)
    for idx, feat in enumerate(feats, start=19):
        col = 1 if (idx-19) % 2 == 0 else 5
        row = 19 + (idx-19)//2
        ws.cell(row, col).value = "• " + feat
        ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col+3)
        ws.cell(row, col).alignment = Alignment(wrap_text=True)
    comp_start = 30
    ws.merge_cells(start_row=comp_start, start_column=1, end_row=comp_start, end_column=8)
    ws.cell(comp_start, 1).value = "Tabla comparativa de plazos"
    ws.cell(comp_start, 1).fill = PatternFill("solid", fgColor=navy); ws.cell(comp_start, 1).font = Font(bold=True, color="FFFFFF"); ws.cell(comp_start, 1).alignment = Alignment(horizontal="center")
    headers = ["Año", "Plazo", "Cuota financiamiento", "Cuota total", "Cuota total con IVA"]
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(comp_start+1, c); cell.value = h; cell.fill = PatternFill("solid", fgColor=blue); cell.font = Font(bold=True, color="FFFFFF"); cell.alignment = Alignment(horizontal="center")
    rows = leasing_table(s["valor_vehiculo"], leasing.get("ingreso_cliente", 0), s["prima_requerida"], s["tasa"], s["seguro"], s["gps"], leasing.get("iva_pct", 13))
    for r, item in enumerate(rows, start=comp_start+2):
        values = [item["anio"], item["plazo"], item["cuota_base"], item["cuota_total_sin_iva"], item["cuota_total_con_iva"]]
        for c, val in enumerate(values, start=1):
            cell = ws.cell(r, c); cell.value = val; cell.border = Border(bottom=thin); cell.alignment = Alignment(horizontal="center")
            if c >= 3: cell.number_format = '$#,##0.00'
    note_row = comp_start + 11
    ws.merge_cells(start_row=note_row, start_column=1, end_row=note_row, end_column=8)
    ws.cell(note_row,1).value = "La cuota mensual incluye IVA, seguro y GPS. El IVA mensual se aplica únicamente sobre la cuota base del arrendamiento."
    ws.cell(note_row,1).font = Font(bold=True, color="DC2626"); ws.cell(note_row,1).alignment = Alignment(wrap_text=True)
    ws.page_setup.paperSize = ws.PAPERSIZE_LETTER
    ws.page_setup.orientation = "portrait"
    ws.page_margins.left = 0.35; ws.page_margins.right = 0.35; ws.page_margins.top = 0.45; ws.page_margins.bottom = 0.45
    ws.freeze_panes = "A12"
    wb.save(out)
    log_audit("GENERAR_PROPUESTA_EXCEL", (user or {}).get("usuario", ""), quote.get("vehicle_code", ""), out.name)
    return True, "Propuesta Excel generada.", out


def generate_quote_proposal_files(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, list[Path]]:
    paths: list[Path] = []
    ok_pdf, msg_pdf, pdf = generate_quote_proposal_pdf(quote_id, user)
    if pdf: paths.append(pdf)
    ok_xlsx, msg_xlsx, xlsx = generate_quote_proposal_excel(quote_id, user)
    if xlsx: paths.append(xlsx)
    if ok_pdf and ok_xlsx:
        return True, "Propuesta generada en PDF y Excel.", paths
    return False, f"PDF: {msg_pdf}\nExcel: {msg_xlsx}", paths


def generate_quote_proposal_html(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    """Compatibilidad: ya no se genera HTML de propuesta, ahora genera PDF y Excel."""
    ok, msg, paths = generate_quote_proposal_files(quote_id, user)
    return ok, msg, (paths[0] if paths else None)


def generate_inventory_excel(vehicles: list[dict], user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    df = get_data_folder()
    if df is None:
        return False, "Carpeta del sistema no disponible.", None
    out_dir = df / SUB_REPORTES
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"LYM_INVENTARIO_PREMIUM_{datetime.now().strftime('%Y-%m-%d_%H%M')}.xlsx"
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from openpyxl.drawing.image import Image as XLImage
    except Exception:
        return False, "Para generar Excel instala: pip install openpyxl pillow", None
    wb = Workbook()
    ws = wb.active; ws.title = "Inventario Gerencial"
    wk = wb.create_sheet("Resumen KPI")
    navy="08285A"; orange="F59A13"; blue="0E3A78"; green="10B981"; red="DC2626"; gray="F8FAFC"
    thin = Side(style="thin", color="D9E2EF")
    k = compute_kpis(vehicles)
    for sheet in [ws, wk]:
        sheet.sheet_view.showGridLines = False
    logo = ResourceManager.find_logo()
    if logo and logo.exists():
        try:
            img=XLImage(str(logo)); img.width=90; img.height=90; wk.add_image(img,"A1")
        except Exception: pass
    wk.merge_cells("B2:H2"); wk["B2"]="REPORTE GERENCIAL DE INVENTARIO L&M INVERSIONES"; wk["B2"].font=Font(size=18,bold=True,color=navy); wk["B2"].alignment=Alignment(horizontal="center")
    cards=[("Vehículos",k.get("total",0)),("Activos",k.get("activos",0)),("Disponibles",k.get("disponibles",0)),("Vendidos",sum(1 for v in vehicles if v.get("estado_comercial")==COMM_VENDIDO)),("Capital activo",k.get("capital",0)),("Ganancia esperada",k.get("ganancia_esperada",0)),("Críticos",k.get("criticos",0)),("Compra→Disponible",f"{k.get('prom_compra_disp',0)} días")]
    for idx,(lab,val) in enumerate(cards):
        r=5+(idx//4)*4; c=1+(idx%4)*2
        wk.merge_cells(start_row=r,start_column=c,end_row=r,end_column=c+1); wk.cell(r,c).value=lab; wk.cell(r,c).fill=PatternFill("solid",fgColor=navy); wk.cell(r,c).font=Font(bold=True,color="FFFFFF"); wk.cell(r,c).alignment=Alignment(horizontal="center")
        wk.merge_cells(start_row=r+1,start_column=c,end_row=r+2,end_column=c+1); wk.cell(r+1,c).value=val; wk.cell(r+1,c).fill=PatternFill("solid",fgColor="FFFFFF"); wk.cell(r+1,c).font=Font(bold=True,size=16,color=orange if idx in (4,5) else navy); wk.cell(r+1,c).alignment=Alignment(horizontal="center",vertical="center")
        if isinstance(val,(int,float)) and idx in (4,5): wk.cell(r+1,c).number_format='$#,##0.00'
    for col in range(1,10): wk.column_dimensions[get_column_letter(col)].width=16
    headers=["CV","Lote","Fecha compra","Vehículo","Año","Millaje","Color","Subasta","Estado operativo","Estado comercial","Costo total","Precio cliente","Venta neta","Utilidad","Margen %","Días compra","Días etapa","Alerta","OC Compra","Características"]
    ws.append(headers)
    for c in range(1,len(headers)+1):
        cell=ws.cell(1,c); cell.fill=PatternFill("solid",fgColor=navy); cell.font=Font(bold=True,color="FFFFFF"); cell.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True)
    for v in vehicles:
        ensure_vehicle_runtime_fields(v); pf=v.get("precio_final") or {}; cost=vehicle_total_cost(v); venta_neta=_to_float(pf.get("venta_neta_usd"),0); margen=((venta_neta-cost)/venta_neta if venta_neta else 0)
        row=[v.get("codigo"),v.get("lote"),_parse_date(v.get("fecha_compra")),_vehicle_display_name(v),v.get("anio"),int(v.get("millaje") or 0),v.get("color"),v.get("subasta"),STAGE_META.get(v.get("estado_actual"),{}).get("label",v.get("estado_actual")),v.get("estado_comercial"),cost,_to_float(v.get("precio_venta_usd"),0),venta_neta,vehicle_expected_profit(v),margen,vehicle_days_from_purchase(v),current_stage_days(v),stage_alert_level(v),v.get("oc_compra_numero",""),v.get("caracteristicas","")]
        ws.append(row)
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        alert=row[17].value
        fill = "FEE2E2" if alert=="ROJO" else "FEF9C3" if alert=="AMARILLO" else "DCFCE7" if alert=="OK" else "FFFFFF"
        for cell in row:
            cell.border=Border(bottom=thin); cell.alignment=Alignment(vertical="center",wrap_text=True)
        row[17].fill=PatternFill("solid",fgColor=fill)
    for col in range(1,len(headers)+1):
        ws.column_dimensions[get_column_letter(col)].width = 16 if col not in (4,20) else 34
    for c in [11,12,13,14]:
        for cell in ws.iter_cols(min_col=c,max_col=c,min_row=2,max_row=ws.max_row):
            for x in cell: x.number_format='$#,##0.00'
    for x in ws.iter_cols(min_col=15,max_col=15,min_row=2,max_row=ws.max_row):
        for cell in x: cell.number_format='0.00%'
    ws.freeze_panes="A2"; ws.auto_filter.ref=f"A1:{get_column_letter(len(headers))}{ws.max_row}"
    wb.save(out)
    log_audit("GENERAR_EXCEL_INVENTARIO_PREMIUM", (user or {}).get("usuario", ""), "", out.name)
    return True, "Reporte Excel de inventario premium generado.", out


def generate_vehicle_costs_excel_report(vehicles: Optional[list[dict]] = None, user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    vehicles = vehicles if vehicles is not None else load_vehicles()
    df = get_data_folder()
    if df is None:
        return False, "Carpeta del sistema no disponible.", None
    out_dir = df / SUB_REPORTES
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"LYM_COSTOS_DETALLADOS_VEHICULOS_{datetime.now().strftime('%Y-%m-%d_%H%M')}.xlsx"
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except Exception:
        return False, "Para generar Excel instala: pip install openpyxl", None
    wb=Workbook(); ws=wb.active; ws.title="Costos por vehículo"; det=wb.create_sheet("Detalle de gastos")
    navy="08285A"; orange="F59A13"; green="10B981"; light="F8FAFC"; thin=Side(style="thin",color="D9E2EF")
    summary_headers=["CV","Lote","Fecha adquisición","Vehículo","Valor adquisición","Total USA","Total Aduana/Naviera","Legalización","Taller/Repuestos","Otros","Costo final","Precio cliente","Venta neta","Utilidad","Margen %","Estado"]
    ws.append(summary_headers)
    for cell in ws[1]: cell.fill=PatternFill("solid",fgColor=navy); cell.font=Font(bold=True,color="FFFFFF"); cell.alignment=Alignment(horizontal="center")
    detail_headers=["CV","Fecha","Etapa","Categoría","Subcategoría","Descripción","Proveedor","OC","Monto","Comprobante","Usuario"]
    det.append(detail_headers)
    for cell in det[1]: cell.fill=PatternFill("solid",fgColor=navy); cell.font=Font(bold=True,color="FFFFFF"); cell.alignment=Alignment(horizontal="center")
    for v in vehicles:
        ensure_vehicle_runtime_fields(v); g=v.get("gastos_detallados") or []
        def subtotal(keys):
            total=0.0
            for item in g:
                cat=_norm(item.get("categoria")); sub=_norm(item.get("subcategoria")); stage=_norm(item.get("stage_key"))
                txt=" ".join([cat,sub,stage])
                if any(k in txt for k in keys): total += _to_float(item.get("monto_usd"),0)
            return round(total,2)
        total_usa=subtotal(["USA","GRUA_USA","TRASLADO"])
        total_aduana=subtotal(["ADUANA","NAVIERA","FLETE","BL","IMPUESTOS","ALMACENAMIENTO","TRAMITE"])
        total_legal=subtotal(["LEGALIZACION","EMISION","CITA","PLACAS"])
        total_taller=subtotal(["TALLER","REPUESTO","PINTURA","MECANICA","GRUA_LOCAL"])
        cost=vehicle_total_cost(v); pf=v.get("precio_final") or {}; venta_neta=_to_float(pf.get("venta_neta_usd"),0); margen=(venta_neta-cost)/venta_neta if venta_neta else 0
        ws.append([v.get("codigo"),v.get("lote"),_parse_date(v.get("fecha_compra")),_vehicle_display_name(v),_to_float(v.get("precio_ganado_usd"),0),total_usa,total_aduana,total_legal,total_taller,round(cost-_to_float(v.get("precio_ganado_usd"),0)-total_usa-total_aduana-total_legal-total_taller,2),cost,_to_float(v.get("precio_venta_usd"),0),venta_neta,vehicle_expected_profit(v),margen,v.get("estado_comercial")])
        for item in g:
            det.append([v.get("codigo"),_parse_date(item.get("fecha")),STAGE_META.get(item.get("stage_key"),{}).get("label",item.get("stage_key")),item.get("categoria"),item.get("subcategoria"),item.get("descripcion"),item.get("proveedor"),item.get("oc_numero"),_to_float(item.get("monto_usd"),0),"SI" if item.get("comprobante") else "NO",item.get("usuario")])
    for sh in [ws, det]:
        for row in sh.iter_rows(min_row=2, max_row=sh.max_row):
            for cell in row:
                cell.border=Border(bottom=thin); cell.alignment=Alignment(vertical="center",wrap_text=True)
        sh.freeze_panes="A2"; sh.auto_filter.ref=f"A1:{get_column_letter(sh.max_column)}{sh.max_row}"
        for col in range(1, sh.max_column+1):
            sh.column_dimensions[get_column_letter(col)].width=16 if col not in (4,6) else 34
    for c in range(5,15):
        for row in ws.iter_rows(min_row=2,min_col=c,max_col=c,max_row=ws.max_row):
            for cell in row: cell.number_format='$#,##0.00' if c != 15 else '0.00%'
    for row in det.iter_rows(min_row=2,min_col=9,max_col=9,max_row=det.max_row):
        for cell in row: cell.number_format='$#,##0.00'
    wb.save(out)
    log_audit("GENERAR_EXCEL_COSTOS", (user or {}).get("usuario", ""), "", out.name)
    return True, "Reporte de costos detallados generado.", out


def generate_quotes_excel_report(user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    df=get_data_folder()
    if df is None:
        return False,"Carpeta del sistema no disponible.",None
    out_dir=df/SUB_REPORTES; out_dir.mkdir(parents=True, exist_ok=True)
    out=out_dir/f"LYM_CONTROL_COTIZACIONES_PREMIUM_{datetime.now().strftime('%Y-%m-%d_%H%M')}.xlsx"
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except Exception:
        return False,"Para generar Excel instala: pip install openpyxl",None
    qs=load_quotes(); wb=Workbook(); ws=wb.active; ws.title="Control Cotizaciones"; dash=wb.create_sheet("Dashboard")
    navy="08285A"; orange="F59A13"; green="DCFCE7"; yellow="FEF9C3"; red="FEE2E2"; thin=Side(style="thin",color="D9E2EF")
    total=len(qs); verdes=sum(1 for q in qs if quote_alert_level(q)=="VERDE"); amarillos=sum(1 for q in qs if quote_alert_level(q)=="AMARILLO"); rojos=sum(1 for q in qs if quote_alert_level(q)=="ROJO"); ganadas=sum(1 for q in qs if q.get("estado")==QUOTE_GANADA); reofertar=sum(1 for q in qs if q.get("estado")==QUOTE_REOFERTAR)
    dash.merge_cells("A1:H1"); dash["A1"]="DASHBOARD COMERCIAL - COTIZACIONES LYM"; dash["A1"].fill=PatternFill("solid",fgColor=navy); dash["A1"].font=Font(size=16,bold=True,color="FFFFFF"); dash["A1"].alignment=Alignment(horizontal="center")
    cards=[("Total",total),("Verde",verdes),("Amarillo",amarillos),("Rojo",rojos),("Ganadas",ganadas),("Reofertar",reofertar)]
    for i,(lab,val) in enumerate(cards):
        r=3+(i//3)*3; c=1+(i%3)*2; dash.merge_cells(start_row=r,start_column=c,end_row=r,end_column=c+1); dash.cell(r,c).value=lab; dash.cell(r,c).fill=PatternFill("solid",fgColor=orange if lab in ("Amarillo","Reofertar") else navy); dash.cell(r,c).font=Font(bold=True,color="FFFFFF"); dash.cell(r,c).alignment=Alignment(horizontal="center")
        dash.merge_cells(start_row=r+1,start_column=c,end_row=r+1,end_column=c+1); dash.cell(r+1,c).value=val; dash.cell(r+1,c).font=Font(bold=True,size=18,color=navy); dash.cell(r+1,c).alignment=Alignment(horizontal="center")
    headers=["Cliente","Teléfono","Correo","Medio","Vehículo cotizado","Código","Fecha cotización","Última gestión","Días sin compra/gestión","Color","Estado","Precio vehículo","Prima requerida","Monto leasing","Plazo","Tasa","Cuota final","Legal","Comentario"]
    ws.append(headers)
    for cell in ws[1]: cell.fill=PatternFill("solid",fgColor=navy); cell.font=Font(bold=True,color="FFFFFF"); cell.alignment=Alignment(horizontal="center",wrap_text=True)
    for q in qs:
        cl=q.get("cliente",{}); snap=q.get("vehicle_snapshot",{}); le=q.get("leasing",{}); lg=q.get("legal",{}); lvl=quote_alert_level(q)
        ws.append([cl.get("nombre"),cl.get("telefono"),cl.get("correo"),cl.get("medio_contacto"),_vehicle_display_name(snap),q.get("vehicle_code"),_parse_date(q.get("fecha_cotizacion")),_parse_date(q.get("ultima_gestion")),quote_days_without_purchase(q),lvl,QUOTE_STATUS_LABELS.get(q.get("estado"),q.get("estado")),_to_float(le.get("precio_vehiculo"),0),_quote_prima_requerida(le),_to_float(le.get("monto_leasing"),0),_to_float(le.get("plazo_meses"),0),_to_float(le.get("tasa_mensual_pct"),0),_to_float(le.get("cuota_total_con_iva"),0),_to_float(lg.get("valor_legales_iva_incluido"),0),q.get("comentario","")])
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        lvl=row[9].value; fill={"VERDE":green,"AMARILLO":yellow,"ROJO":red,"CERRADA":"DBEAFE","PERDIDA":"F1F5F9"}.get(lvl,"FFFFFF")
        for cell in row:
            cell.border=Border(bottom=thin); cell.alignment=Alignment(vertical="center",wrap_text=True)
        row[9].fill=PatternFill("solid",fgColor=fill)
    for c in [12,13,14,17,18]:
        for col_cells in ws.iter_cols(min_col=c,max_col=c,min_row=2,max_row=ws.max_row):
            for cell in col_cells: cell.number_format='$#,##0.00'
    for c in [7,8]:
        for col_cells in ws.iter_cols(min_col=c,max_col=c,min_row=2,max_row=ws.max_row):
            for cell in col_cells: cell.number_format='dd/mm/yyyy'
    for col in range(1,len(headers)+1): ws.column_dimensions[get_column_letter(col)].width=16 if col not in (1,5,19) else 30
    ws.freeze_panes="A2"; ws.auto_filter.ref=f"A1:{get_column_letter(len(headers))}{ws.max_row}"
    wb.save(out); log_audit("GENERAR_EXCEL_COTIZACIONES_PREMIUM", (user or {}).get("usuario",""), "", out.name)
    return True,"Reporte Excel de cotizaciones premium generado.",out


def generate_html_report(vehicles: list[dict], user: Optional[dict] = None) -> Optional[Path]:
    df=get_data_folder()
    if df is None: return None
    out_dir=df/SUB_REPORTES; out_dir.mkdir(parents=True, exist_ok=True)
    k=compute_kpis(vehicles); logo=ResourceManager.logo_data_uri(); generated=datetime.now().strftime('%d/%m/%Y %H:%M')
    stage_cards=[]
    for s in STAGES:
        n=sum(1 for v in vehicles if ensure_vehicle_runtime_fields(v) and v.get('estado_actual')==s['key'])
        stage_cards.append(f"<div class='stage'><span style='background:{s['color']}'></span><b>{html.escape(s['label'])}</b><strong>{n}</strong></div>")
    rows=[]
    for v in sorted(vehicles,key=lambda x:(stage_index(ensure_vehicle_runtime_fields(x).get('estado_actual','')),x.get('codigo',''))):
        rows.append(f"<tr><td><b>{html.escape(v.get('codigo',''))}</b></td><td>{html.escape(_vehicle_display_name(v))}</td><td>{html.escape(str(v.get('lote','')))}</td><td>{_fmt_usd(vehicle_total_cost(v))}</td><td>{_fmt_usd(v.get('precio_venta_usd'))}</td><td>{_fmt_usd(vehicle_expected_profit(v))}</td><td>{html.escape(STAGE_META.get(v.get('estado_actual'),{}).get('label',v.get('estado_actual')))}</td><td>{html.escape(str(v.get('estado_comercial','')))}</td><td>{vehicle_days_from_purchase(v)}</td></tr>")
    html_text=f"""<!doctype html><html lang='es'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>LYM Reporte Gerencial</title><style>
:root{{--navy:#08285a;--navy2:#0e3a78;--orange:#f59a13;--bg:#f5f8fc;--ink:#0b172a;--muted:#637083;--line:rgba(8,40,90,.12);--green:#10B981;--red:#DC2626;--amber:#F59E0B;}}
*{{box-sizing:border-box}}body{{margin:0;font-family:Inter,Segoe UI,Arial,sans-serif;background:radial-gradient(circle at 12% 8%,rgba(245,154,19,.16),transparent 28%),radial-gradient(circle at 88% 18%,rgba(8,40,90,.16),transparent 30%),linear-gradient(135deg,#f6f9ff,#eef4fb 55%,#fff8ef);color:var(--ink)}}main{{max-width:1280px;margin:auto;padding:42px 24px 70px}}.hero{{display:grid;grid-template-columns:1.2fr .55fr;gap:24px;align-items:center;min-height:420px}}.logo{{background:rgba(255,255,255,.82);border:1px solid var(--line);border-radius:36px;min-height:340px;display:grid;place-items:center;box-shadow:0 28px 80px rgba(8,40,90,.16)}}.logo img{{max-width:290px;filter:drop-shadow(0 18px 35px rgba(8,40,90,.22))}}.eyebrow{{display:inline-flex;color:var(--orange);font-weight:950;text-transform:uppercase;letter-spacing:.16em;background:rgba(245,154,19,.10);border-radius:99px;padding:8px 13px;font-size:12px}}h1{{font-size:clamp(42px,7vw,82px);line-height:.95;color:var(--navy);letter-spacing:-.06em;margin:18px 0}}h2{{font-size:34px;color:var(--navy);letter-spacing:-.04em}}.lead{{font-size:20px;line-height:1.5;color:#334155;max-width:850px}}.kpis{{display:grid;grid-template-columns:repeat(4,1fr);gap:18px;margin:24px 0}}.card{{background:rgba(255,255,255,.86);border:1px solid var(--line);border-radius:26px;padding:20px;box-shadow:0 20px 55px rgba(8,40,90,.10)}}.label{{color:var(--muted);font-size:12px;text-transform:uppercase;letter-spacing:.09em;font-weight:900}}.value{{font-size:34px;color:var(--navy);font-weight:950;line-height:1.05;margin-top:8px}}.grid2{{display:grid;grid-template-columns:1fr 1fr;gap:20px;margin:24px 0}}.stage{{display:grid;grid-template-columns:16px 1fr 50px;gap:12px;align-items:center;padding:12px 0;border-bottom:1px solid rgba(8,40,90,.08)}}.stage span{{width:14px;height:14px;border-radius:99px}}table{{width:100%;border-collapse:separate;border-spacing:0;background:#fff;border-radius:22px;overflow:hidden;box-shadow:0 18px 46px rgba(8,40,90,.10)}}th{{background:var(--navy);color:#fff;text-align:left;padding:12px;font-size:12px;text-transform:uppercase;letter-spacing:.07em}}td{{padding:12px;border-bottom:1px solid rgba(8,40,90,.08);font-weight:700;color:#263346}}tr:hover td{{background:rgba(245,154,19,.06)}}.note{{background:linear-gradient(135deg,var(--navy),var(--navy2));color:white;border-radius:28px;padding:24px;margin:28px 0;box-shadow:0 24px 70px rgba(8,40,90,.22)}}@media(max-width:980px){{.hero,.grid2{{grid-template-columns:1fr}}.kpis{{grid-template-columns:repeat(2,1fr)}}}}
</style></head><body><main><section class='hero'><div><span class='eyebrow'>Reporte gerencial · Corte {generated}</span><h1>Inventario de vehículos <span style='color:var(--orange)'>L & M</span></h1><p class='lead'>Control operativo y financiero de compra vehicular, costos, capital activo, disponibles para cotización, estados comerciales y rotación del inventario.</p></div><div class='logo'>{'<img src="'+logo+'">' if logo else '<h2>L&M</h2>'}</div></section><section class='kpis'><div class='card'><div class='label'>Vehículos totales</div><div class='value'>{k['total']}</div></div><div class='card'><div class='label'>Activos</div><div class='value'>{k['activos']}</div></div><div class='card'><div class='label'>Disponibles</div><div class='value'>{k['disponibles']}</div></div><div class='card'><div class='label'>Capital activo</div><div class='value'>{_fmt_usd(k['capital'])}</div></div><div class='card'><div class='label'>Ganancia esperada</div><div class='value'>{_fmt_usd(k['ganancia_esperada'])}</div></div><div class='card'><div class='label'>Críticos</div><div class='value'>{k['criticos']}</div></div><div class='card'><div class='label'>Compra→Disponible</div><div class='value'>{k['prom_compra_disp']} días</div></div><div class='card'><div class='label'>Cotizaciones</div><div class='value'>{len(load_quotes())}</div></div></section><section class='grid2'><div class='card'><h2>Estatus operativo</h2>{''.join(stage_cards)}</div><div class='note'><h2 style='color:white'>Lectura ejecutiva</h2><p>Prioriza vehículos con más capital detenido, etapas críticas y unidades disponibles con baja rotación. El costo detallado alimenta la utilidad real y las propuestas de leasing.</p></div></section><section><h2>Detalle general</h2><table><thead><tr><th>CV</th><th>Vehículo</th><th>Lote</th><th>Costo</th><th>Precio</th><th>Utilidad</th><th>Etapa</th><th>Comercial</th><th>Días</th></tr></thead><tbody>{''.join(rows)}</tbody></table></section></main></body></html>"""
    out=out_dir/f"LYM_REPORTE_GERENCIAL_PREMIUM_{datetime.now().strftime('%Y-%m-%d_%H%M')}.html"
    out.write_text(html_text,encoding='utf-8'); log_audit("GENERAR_HTML_GERENCIAL_PREMIUM", (user or {}).get("usuario",""), "", out.name)
    return out


def generate_quotes_html_report(user: Optional[dict] = None) -> Optional[Path]:
    df=get_data_folder()
    if df is None: return None
    out_dir=df/SUB_REPORTES; out_dir.mkdir(parents=True, exist_ok=True)
    qs=load_quotes(); logo=ResourceManager.logo_data_uri(); total=len(qs); verdes=sum(1 for q in qs if quote_alert_level(q)=="VERDE"); amarillos=sum(1 for q in qs if quote_alert_level(q)=="AMARILLO"); rojos=sum(1 for q in qs if quote_alert_level(q)=="ROJO"); reof=sum(1 for q in qs if q.get('estado')==QUOTE_REOFERTAR)
    rows=[]
    for q in qs:
        cl=q.get('cliente',{}); snap=q.get('vehicle_snapshot',{}); le=q.get('leasing',{}); lvl=quote_alert_level(q)
        rows.append(f"<tr class='{lvl.lower()}'><td>{html.escape(cl.get('nombre',''))}</td><td>{html.escape(cl.get('telefono',''))}</td><td>{html.escape(cl.get('medio_contacto',''))}</td><td>{html.escape(_vehicle_display_name(snap))}</td><td>{_fmt_date(q.get('fecha_cotizacion'))}</td><td>{quote_days_without_purchase(q)}</td><td>{html.escape(QUOTE_STATUS_LABELS.get(q.get('estado'),q.get('estado')))}</td><td>{_fmt_usd(le.get('cuota_total_con_iva'))}</td></tr>")
    html_text=f"""<!doctype html><html lang='es'><head><meta charset='utf-8'><title>LYM Cotizaciones</title><style>body{{font-family:Segoe UI,Arial;background:linear-gradient(135deg,#f6f9ff,#fff8ef);margin:0;color:#0b172a}}main{{max-width:1220px;margin:auto;padding:40px 24px}}.hero{{display:grid;grid-template-columns:1fr 260px;gap:24px;align-items:center}}.logo{{background:white;border-radius:28px;display:grid;place-items:center;padding:22px;box-shadow:0 20px 60px #08285a22}}.logo img{{max-width:220px}}h1{{font-size:58px;color:#08285a;line-height:.95}}.kpis{{display:grid;grid-template-columns:repeat(5,1fr);gap:14px}}.card{{background:white;border-radius:22px;padding:18px;box-shadow:0 16px 40px #08285a18}}.label{{font-weight:900;color:#637083;text-transform:uppercase;font-size:12px}}.value{{font-size:34px;font-weight:950;color:#08285a}}table{{width:100%;border-collapse:collapse;margin-top:22px;background:white;border-radius:18px;overflow:hidden;box-shadow:0 16px 40px #08285a18}}th{{background:#08285a;color:white;padding:12px;text-align:left}}td{{padding:11px;border-bottom:1px solid #e2e8f0;font-weight:700}}tr.verde td{{background:#dcfce744}}tr.amarillo td{{background:#fef9c344}}tr.rojo td{{background:#fee2e244}}@media(max-width:900px){{.hero,.kpis{{grid-template-columns:1fr}}}}</style></head><body><main><section class='hero'><div><h1>Control de cotizaciones <span style='color:#f59a13'>L&M</span></h1><p>Seguimiento comercial por cliente, carro cotizado, alerta por días sin compra o gestión y oportunidades para reofertar.</p></div><div class='logo'>{'<img src="'+logo+'">' if logo else '<h2>L&M</h2>'}</div></section><section class='kpis'><div class='card'><div class='label'>Total</div><div class='value'>{total}</div></div><div class='card'><div class='label'>Verde</div><div class='value'>{verdes}</div></div><div class='card'><div class='label'>Amarillo</div><div class='value'>{amarillos}</div></div><div class='card'><div class='label'>Rojo</div><div class='value'>{rojos}</div></div><div class='card'><div class='label'>Reofertar</div><div class='value'>{reof}</div></div></section><table><thead><tr><th>Cliente</th><th>Teléfono</th><th>Medio</th><th>Vehículo</th><th>Fecha</th><th>Días</th><th>Estado</th><th>Cuota</th></tr></thead><tbody>{''.join(rows)}</tbody></table></main></body></html>"""
    out=out_dir/f"LYM_COTIZACIONES_PREMIUM_{datetime.now().strftime('%Y-%m-%d_%H%M')}.html"; out.write_text(html_text,encoding='utf-8')
    log_audit("GENERAR_HTML_COTIZACIONES_PREMIUM", (user or {}).get("usuario",""), "", out.name); return out


if PYSIDE_OK:
    class CollapsibleBox(QGroupBox):
        def __init__(self, title: str, checked: bool = True):
            super().__init__(title)
            self.setCheckable(True); self.setChecked(checked)
            self.inner = QWidget(); self.inner_lay = QVBoxLayout(self.inner); self.inner_lay.setContentsMargins(8,8,8,8)
            outer = QVBoxLayout(self); outer.addWidget(self.inner)
            self.toggled.connect(self.inner.setVisible); self.inner.setVisible(checked)

    class QuoteEditorDialog(QDialog):
        def __init__(self, parent, user: dict, device: DeviceInfo, vehicle_id: str = "", quote_id: str = "", preset_client: Optional[dict] = None):
            super().__init__(parent); self.user=user; self.device=device; self.vehicle_id=vehicle_id; self.quote_id=quote_id; self.preset_client=preset_client or {}; self._vehicle_ids=[]; self._loading=False
            self.quote=find_quote(quote_id) if quote_id else None
            self.setWindowTitle("Cotización leasing")
            self.setMinimumSize(1120,820)
            self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint)
            self._build(); self._load_quote_if_any(); self.refresh_calc()
        def _build(self):
            main=QVBoxLayout(self); main.addWidget(make_title("Cotización leasing", "Calcula en paralelo la cuota mensual y los gastos legales. Puedes recalcular meses sin crear una nueva cotización."))
            scroll=QScrollArea(); scroll.setWidgetResizable(True); content=QWidget(); scroll.setWidget(content); lay=QVBoxLayout(content); main.addWidget(scroll,1)
            sec_client=CollapsibleBox("1) Llenar datos del cliente", True); lay.addWidget(sec_client)
            grid=QGridLayout(); sec_client.inner_lay.addLayout(grid)
            self.nombre=QLineEdit(self.preset_client.get("nombre", "")); self.telefono=QLineEdit(self.preset_client.get("telefono", "")); self.medio=QComboBox(); self.medio.setEditable(True); self.medio.addItems(["WHATSAPP","FACEBOOK","TIKTOK","PAGINA WEB","REFERIDO","LLAMADA","OTRO"]); self.medio.setCurrentText(self.preset_client.get("medio_contacto", "WHATSAPP")); self.correo=QLineEdit(self.preset_client.get("correo", ""))
            labels=[("Nombre del cliente",self.nombre),("Teléfono",self.telefono),("Medio de contacto",self.medio),("Correo",self.correo)]
            for i,(lab,w) in enumerate(labels):
                grid.addWidget(QLabel(lab),0 if i<3 else 2,i if i<3 else 0); grid.addWidget(w,1 if i<3 else 3,i if i<3 else 0)
            sec_car=CollapsibleBox("2) Datos del carro y tarifas", True); lay.addWidget(sec_car)
            form=QFormLayout(); sec_car.inner_lay.addLayout(form)
            self.vehicle_combo=QComboBox(); self.vehicle_combo.setEditable(False)
            available=[]
            for v in load_vehicles():
                ensure_vehicle_runtime_fields(v)
                if v.get("estado_actual")==STAGE_DISPONIBLE and v.get("estado_comercial")==COMM_DISPONIBLE:
                    available.append(v)
            # Si se está recalculando una cotización del mismo vehículo y aún existe, mostrarlo aunque no esté disponible solo para lectura.
            if self.quote and self.quote.get("vehicle_id") and self.quote.get("vehicle_id") not in [x.get("id") for x in available]:
                old=find_vehicle(self.quote.get("vehicle_id"))
                if old: available.insert(0, old)
            for v in available:
                self._vehicle_ids.append(v.get("id")); self.vehicle_combo.addItem(f"{v.get('codigo')} · {_vehicle_display_name(v)} · {_fmt_usd(v.get('precio_venta_usd'))}")
            pick_id=self.vehicle_id or (self.quote or {}).get("vehicle_id")
            if pick_id and pick_id in self._vehicle_ids:
                self.vehicle_combo.setCurrentIndex(self._vehicle_ids.index(pick_id))
            if self.vehicle_id:
                self.vehicle_combo.setEnabled(False)
            self.vehicle_combo.currentIndexChanged.connect(self._vehicle_changed)
            self.ingreso=MoneyEdit(); self.ingreso.setRange(0,999999); self.ingreso.setValue(1900)
            self.precio=MoneyEdit(); self.precio.setRange(0,9999999); self.precio.setValue(0)
            self.prima_pct=QDoubleSpinBox(); self.prima_pct.setRange(0,90); self.prima_pct.setDecimals(2); self.prima_pct.setSuffix(" %"); self.prima_pct.setValue(20)
            self.comision=MoneyEdit(); self.comision.setRange(0,99999); self.comision.setValue(100)
            self.tasa=QComboBox(); self.tasa.setEditable(True); self.tasa.addItems(["1.00","1.80","2.50","3.00"]); self.tasa.setCurrentText("2.50")
            self.seguro=MoneyEdit(); self.seguro.setRange(0,99999); self.seguro.setValue(80)
            self.gps=MoneyEdit(); self.gps.setRange(0,99999); self.gps.setValue(20)
            self.iva=QDoubleSpinBox(); self.iva.setRange(0,30); self.iva.setDecimals(2); self.iva.setSuffix(" %"); self.iva.setValue(13)
            self.opcion_compra=MoneyEdit(); self.opcion_compra.setRange(0,99999); self.opcion_compra.setValue(500)
            for lab,w in [("Vehículo",self.vehicle_combo),("Ingreso mensual cliente",self.ingreso),("Valor del vehículo",self.precio),("Prima mínima",self.prima_pct),("Comisión editable",self.comision),("Tasa rentabilidad mensual",self.tasa),("Seguro mensual incluye IVA",self.seguro),("GPS mensual incluye IVA",self.gps),("IVA sobre cuota base",self.iva),("Opción de compra",self.opcion_compra)]: form.addRow(lab+":",w)
            sec_plazo=CollapsibleBox("3) Plazo, cálculo de cuota y gastos legales", True); lay.addWidget(sec_plazo)
            pgrid=QGridLayout(); sec_plazo.inner_lay.addLayout(pgrid)
            self.plazo=QSpinBox(); self.plazo.setRange(1,120); self.plazo.setValue(60); pgrid.addWidget(QLabel("Plazo elegido (meses):"),0,0); pgrid.addWidget(self.plazo,0,1)
            self.comentario=QTextEdit(); self.comentario.setMinimumHeight(55); pgrid.addWidget(QLabel("Comentario:"),1,0); pgrid.addWidget(self.comentario,1,1,1,3)
            result_row=QHBoxLayout(); sec_plazo.inner_lay.addLayout(result_row)
            self.result_leasing=QLabel(); self.result_leasing.setTextFormat(Qt.TextFormat.RichText); self.result_leasing.setWordWrap(True); self.result_leasing.setStyleSheet("background:#eaf4ff;border:1px solid #bfdbfe;border-radius:12px;padding:14px;color:#08285a;font-weight:800;")
            self.result_legal=QLabel(); self.result_legal.setTextFormat(Qt.TextFormat.RichText); self.result_legal.setWordWrap(True); self.result_legal.setStyleSheet("background:#ecfdf5;border:1px solid #86efac;border-radius:12px;padding:14px;color:#064e3b;font-weight:800;")
            result_row.addWidget(self.result_leasing,2); result_row.addWidget(self.result_legal,1)
            self.table=QTableWidget(0,5); self.table.setHorizontalHeaderLabels(["Año","Plazo","Cuota financiamiento","Cuota total","Cuota total con IVA"]); self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.table.setMinimumHeight(210); sec_plazo.inner_lay.addWidget(self.table)
            btnrow=QHBoxLayout(); main.addLayout(btnrow)
            save=QPushButton("Guardar / recalcular cotización"); save.setObjectName("orange"); save.clicked.connect(self.save); btnrow.addWidget(save)
            close=QPushButton("Cerrar"); close.setObjectName("ghost"); close.clicked.connect(self.reject); btnrow.addWidget(close); btnrow.addStretch(1)
            for w in [self.ingreso,self.precio,self.prima_pct,self.comision,self.plazo,self.seguro,self.gps,self.iva,self.opcion_compra]:
                try: w.valueChanged.connect(self.refresh_calc)
                except Exception: pass
            self.tasa.currentTextChanged.connect(self.refresh_calc)
            QTimer.singleShot(0, self._vehicle_changed)
        def _load_quote_if_any(self):
            if not self.quote: return
            self._loading=True
            cl=self.quote.get("cliente",{}); le=self.quote.get("leasing",{})
            self.nombre.setText(cl.get("nombre", "")); self.telefono.setText(cl.get("telefono", "")); self.correo.setText(cl.get("correo", "")); self.medio.setCurrentText(cl.get("medio_contacto", "WHATSAPP"))
            self.ingreso.setValue(_to_float(le.get("ingreso_cliente"),1900)); self.precio.setValue(_to_float(le.get("precio_vehiculo"),0)); self.prima_pct.setValue(_to_float(le.get("prima_pct"),20)); self.comision.setValue(_to_float(le.get("comision_usd"),100)); self.plazo.setValue(int(_to_float(le.get("plazo_meses"),60))); self.tasa.setCurrentText(str(le.get("tasa_mensual_pct",2.5))); self.seguro.setValue(_to_float(le.get("seguro_mensual"),80)); self.gps.setValue(_to_float(le.get("gps_mensual"),20)); self.iva.setValue(_to_float(le.get("iva_pct"),13)); self.opcion_compra.setValue(_to_float(self.quote.get("opcion_compra_usd"),500)); self.comentario.setPlainText(self.quote.get("comentario", ""))
            self._loading=False
        def _vehicle_changed(self):
            if getattr(self,"_loading",False): return
            idx=self.vehicle_combo.currentIndex()
            if 0 <= idx < len(self._vehicle_ids):
                v=find_vehicle(self._vehicle_ids[idx])
                if v and (not self.quote_id): self.precio.setValue(float(v.get("precio_venta_usd") or 0))
            self.refresh_calc()
        def _calc_data(self):
            precio=self.precio.value(); pago=round(precio*self.prima_pct.value()/100.0 + self.comision.value(),2); tasa=_to_float(self.tasa.currentText(),2.5)
            calc=calculate_leasing(precio,self.ingreso.value(),pago,self.plazo.value(),tasa,self.seguro.value(),self.gps.value(),self.iva.value()); legal=calculate_legal_fees(precio); return calc,legal,pago
        def refresh_calc(self):
            if not hasattr(self,"result_leasing"): return
            calc,legal,pago=self._calc_data()
            self.result_leasing.setText(f"<b>1) Simulador de cuota</b><br>Monto leasing: <b>{_fmt_usd(calc['monto_leasing'])}</b> · Prima requerida: <b>{_fmt_usd(pago)}</b><br>Cuota base: <b>{_fmt_usd(calc['cuota_base'])}</b><br>Cuota total con IVA: <b style='font-size:22px'>{_fmt_usd(calc['cuota_total_con_iva'])}</b><br>% ingreso destinado a cuota: <b>{calc['pct_ingreso']}%</b> · Riesgo: <b>{calc['riesgo_texto']} / {calc['riesgo']}</b>")
            self.result_legal.setText(f"<b>2) Gastos legales</b><br>Valor del carro: <b>{_fmt_usd(legal['valor_leasing_vehiculo'])}</b><br>1.5%: <b>{_fmt_usd(legal['valor_pct'])}</b> + Base fija: <b>{_fmt_usd(legal['base_fija'])}</b><br>Subtotal: <b>{_fmt_usd(legal['subtotal'])}</b> · Tope: <b>{_fmt_usd(legal['tope'])}</b><br>Total legal IVA incluido: <b style='font-size:18px'>{_fmt_usd(legal['valor_legales_iva_incluido'])}</b><br>Aplica tope: <b>{legal['aplica_tope']}</b>")
            rows=leasing_table(self.precio.value(),self.ingreso.value(),pago,_to_float(self.tasa.currentText(),2.5),self.seguro.value(),self.gps.value(),self.iva.value()); self.table.setRowCount(len(rows))
            for r,item in enumerate(rows):
                vals=[item['anio'],item['plazo'],_fmt_usd(item['cuota_base']),_fmt_usd(item['cuota_total_sin_iva']),_fmt_usd(item['cuota_total_con_iva'])]
                for c,val in enumerate(vals): self.table.setItem(r,c,QTableWidgetItem(str(val)))
        def save(self):
            if self.vehicle_combo.currentIndex()<0 or self.vehicle_combo.currentIndex()>=len(self._vehicle_ids): QMessageBox.warning(self,"Validación","No hay vehículo disponible seleccionado."); return
            calc,legal,pago=self._calc_data()
            data={"vehicle_id":self._vehicle_ids[self.vehicle_combo.currentIndex()],"cliente_nombre":self.nombre.text(),"telefono":self.telefono.text(),"correo":self.correo.text(),"medio_contacto":self.medio.currentText(),"ingreso_cliente":self.ingreso.value(),"precio_vehiculo":self.precio.value(),"prima_pct":self.prima_pct.value(),"comision_usd":self.comision.value(),"pago_inicial":pago,"plazo_meses":self.plazo.value(),"tasa_mensual_pct":_to_float(self.tasa.currentText(),2.5),"seguro_mensual":self.seguro.value(),"gps_mensual":self.gps.value(),"iva_pct":self.iva.value(),"opcion_compra_usd":self.opcion_compra.value(),"comentario":self.comentario.toPlainText()}
            ok,msg,qid=create_quote(data,self.user,self.device,self.quote_id)
            if not ok: QMessageBox.warning(self,"Validación",msg); return
            self.quote_id=qid; QMessageBox.information(self,"Guardado",msg); self.accept()

    class QuoteDetailDialog(QDialog):
        def __init__(self,parent,quote_id:str,user:dict,device:DeviceInfo):
            super().__init__(parent); self.quote_id=quote_id; self.user=user; self.device=device
            self.setWindowTitle("Detalle de cotización"); self.setMinimumSize(980,720); self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint); self._build(); self.refresh()
        def _build(self):
            lay=QVBoxLayout(self); self.info=QTextEdit(); self.info.setReadOnly(True); lay.addWidget(self.info)
            row=QHBoxLayout(); lay.addLayout(row)
            for txt,fn,obj in [("Agregar seguimiento",self.add_follow,"orange"),("Editar / recalcular meses",self.edit_quote,"ghost"),("Generar propuesta PDF + Excel",self.generate_proposal,"orange"),("Ofrecer otro carro",self.offer_other,"ghost"),("Marcar compra / vendido",self.mark_sold,"danger")]:
                b=QPushButton(txt); b.setObjectName(obj); b.clicked.connect(fn); row.addWidget(b)
            self.hist=QTextEdit(); self.hist.setReadOnly(True); lay.addWidget(self.hist)
        def refresh(self):
            q=find_quote(self.quote_id)
            if not q: return
            self.q=q; cl=q.get("cliente",{}); snap=q.get("vehicle_snapshot",{}); le=q.get("leasing",{}); lg=q.get("legal",{}); s=_quote_financial_summary(q)
            info=f"""Cliente: {cl.get('nombre','')}\nTeléfono: {cl.get('telefono','')}\nCorreo: {cl.get('correo','')}\nMedio: {cl.get('medio_contacto','')}\n\nVehículo cotizado: {_vehicle_display_name(snap)} · {q.get('vehicle_code')}\nValor del vehículo: {_fmt_usd(s['valor_vehiculo'])}\nPrima requerida: {_fmt_usd(s['prima_requerida'])}\nMonto leasing: {_fmt_usd(s['monto_leasing'])}\nPlazo: {s['plazo']} meses\nTasa: {s['tasa']}%\nCuota final con IVA: {_fmt_usd(s['cuota_final'])}\nGastos legales: {_fmt_usd(s['legal'])}\n\nEstado: {QUOTE_STATUS_LABELS.get(q.get('estado'),q.get('estado'))}\nColor seguimiento: {quote_alert_level(q)}\nDías sin compra/gestión: {quote_days_without_purchase(q)}"""
            self.info.setPlainText(info)
            self.hist.setPlainText("\n".join([f"[{h.get('fecha')}] {h.get('usuario')} · {h.get('accion')} · {h.get('comentario')}" for h in q.get('seguimientos',[])]))
        def add_follow(self):
            estados=list(QUOTE_STATUS_LABELS.keys()); estado,ok=QInputDialog.getItem(self,"Estado","Nuevo estado:",estados,0,False)
            if not ok: return
            com,ok=QInputDialog.getMultiLineText(self,"Seguimiento","Comentario:")
            if not ok: return
            ok,msg=add_quote_followup(self.quote_id,com,estado,self.user); QMessageBox.information(self,"Seguimiento",msg) if ok else QMessageBox.warning(self,"Seguimiento",msg); self.refresh()
        def edit_quote(self):
            dlg=QuoteEditorDialog(self,self.user,self.device,quote_id=self.quote_id)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def generate_proposal(self):
            ok,msg,paths=generate_quote_proposal_files(self.quote_id,self.user)
            if not ok:
                QMessageBox.warning(self,"Propuesta",msg); return
            QMessageBox.information(self,"Propuesta",msg+"\n"+"\n".join(str(p) for p in paths))
            if paths: QDesktopServices.openUrl(QUrl.fromLocalFile(str(paths[0])))
        def offer_other(self):
            cl=self.q.get("cliente",{})
            dlg=QuoteEditorDialog(self,self.user,self.device,preset_client=cl)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def mark_sold(self):
            if QMessageBox.question(self,"Venta","¿Marcar este carro como vendido a este cliente?") != QMessageBox.StandardButton.Yes: return
            ok,msg=mark_quote_won_and_vehicle_sold(self.quote_id,self.user,self.device); QMessageBox.information(self,"Venta",msg) if ok else QMessageBox.warning(self,"Venta",msg); self.refresh()

    class VehicleQuotesDialog(QDialog):
        def __init__(self,parent,vehicle_id:str,user:dict,device:DeviceInfo):
            super().__init__(parent); self.vehicle_id=vehicle_id; self.user=user; self.device=device; self._quote_ids=[]
            self.setWindowTitle("Personas interesadas / cotizaciones del carro"); self.setMinimumSize(980,620); self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint); self._build(); self.refresh()
        def _build(self):
            lay=QVBoxLayout(self); self.header=QLabel(); self.header.setStyleSheet("font-size:20px;font-weight:950;color:#08285a;"); lay.addWidget(self.header)
            row=QHBoxLayout(); bnew=QPushButton("Nueva cotización para este carro"); bnew.setObjectName("orange"); bnew.clicked.connect(self.new_quote); bclose=QPushButton("Cerrar"); bclose.setObjectName("ghost"); bclose.clicked.connect(self.accept); row.addWidget(bnew); row.addStretch(1); row.addWidget(bclose); lay.addLayout(row)
            self.table=QTableWidget(0,9); self.table.setHorizontalHeaderLabels(["Cliente","Teléfono","Medio","Fecha","Última gestión","Días","Estado","Color","Cuota final"]); self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.table.cellDoubleClicked.connect(self.open_quote); lay.addWidget(self.table)
        def refresh(self):
            v=find_vehicle(self.vehicle_id) or {}; self.header.setText(f"{v.get('codigo','')} · {_vehicle_display_name(v)} · Precio { _fmt_usd(v.get('precio_venta_usd')) }")
            qs=quotes_for_vehicle(self.vehicle_id); self._quote_ids=[]; self.table.setRowCount(len(qs))
            for r,q in enumerate(qs):
                self._quote_ids.append(q.get("id")); cl=q.get("cliente",{}); le=q.get("leasing",{}); lvl=quote_alert_level(q); vals=[cl.get("nombre"),cl.get("telefono"),cl.get("medio_contacto"),_fmt_date(q.get("fecha_cotizacion")),_fmt_date(q.get("ultima_gestion")),str(quote_days_without_purchase(q)),QUOTE_STATUS_LABELS.get(q.get("estado"),q.get("estado")),lvl,_fmt_usd(le.get("cuota_total_con_iva"))]
                for c,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setBackground(QColor(_quote_alert_color(lvl))); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.table.setItem(r,c,it)
        def new_quote(self):
            dlg=QuoteEditorDialog(self,self.user,self.device,vehicle_id=self.vehicle_id)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def open_quote(self,row,col):
            if row < len(self._quote_ids):
                dlg=QuoteDetailDialog(self,self._quote_ids[row],self.user,self.device); dlg.exec(); self.refresh()

    class CotizacionesPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main=main; self._vehicle_ids=[]; self._quote_ids=[]; self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Cotizaciones", "Vehículos disponibles, clientes interesados y cotizaciones generales sin borrar historial."))
            self.tabs=QTabWidget(); lay.addWidget(self.tabs)
            wv=QWidget(); vl=QVBoxLayout(wv); brow=QHBoxLayout(); bnew=QPushButton("Nueva cotización"); bnew.setObjectName("orange"); bnew.clicked.connect(self.new_quote); brow.addWidget(bnew); brow.addStretch(1); vl.addLayout(brow)
            self.vehicle_table=QTableWidget(0,7); self.vehicle_table.setHorizontalHeaderLabels(["Código","Vehículo","Precio","Costo","Ganancia","Cotizaciones","Días disponible"]); self.vehicle_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.vehicle_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.vehicle_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.vehicle_table.cellDoubleClicked.connect(self.open_vehicle_quotes); vl.addWidget(self.vehicle_table); self.tabs.addTab(wv,"Vehículos disponibles")
            wg=QWidget(); gl=QVBoxLayout(wg); self.quote_table=QTableWidget(0,9); self.quote_table.setHorizontalHeaderLabels(["Cliente","Teléfono","Medio","Vehículo cotizado","Fecha","Última gestión","Días","Estado","Color"]); self.quote_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.quote_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.quote_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.quote_table.cellDoubleClicked.connect(self.open_quote); gl.addWidget(self.quote_table); self.tabs.addTab(wg,"Cotizaciones general")
        def refresh(self):
            vehicles=[v for v in load_vehicles() if ensure_vehicle_runtime_fields(v) and v.get("estado_actual")==STAGE_DISPONIBLE and v.get("estado_comercial")==COMM_DISPONIBLE]
            self._vehicle_ids=[]; self.vehicle_table.setRowCount(len(vehicles))
            for r,v in enumerate(vehicles):
                self._vehicle_ids.append(v.get("id")); vals=[v.get("codigo"),_vehicle_display_name(v),_fmt_usd(v.get("precio_venta_usd")),_fmt_usd(vehicle_total_cost(v)),_fmt_usd(vehicle_expected_profit(v)),str(len(quotes_for_vehicle(v.get("id")))),str(current_stage_days(v))]
                for c,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.vehicle_table.setItem(r,c,it)
            qs=load_quotes(); self._quote_ids=[]; self.quote_table.setRowCount(len(qs))
            for r,q in enumerate(qs):
                self._quote_ids.append(q.get("id")); cl=q.get("cliente",{}); snap=q.get("vehicle_snapshot",{}); lvl=quote_alert_level(q); vals=[cl.get("nombre"),cl.get("telefono"),cl.get("medio_contacto"),f"{_vehicle_display_name(snap)} · {q.get('vehicle_code')}",_fmt_date(q.get("fecha_cotizacion")),_fmt_date(q.get("ultima_gestion")),str(quote_days_without_purchase(q)),QUOTE_STATUS_LABELS.get(q.get("estado"),q.get("estado")),lvl]
                for c,val in enumerate(vals):
                    item=QTableWidgetItem(str(val)); item.setBackground(QColor(_quote_alert_color(lvl))); item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable); self.quote_table.setItem(r,c,item)
        def new_quote(self):
            dlg=QuoteEditorDialog(self,self.main.user,self.main.device)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh(); self.main.refresh_all()
        def open_vehicle_quotes(self,row,col):
            if row < len(self._vehicle_ids):
                dlg=VehicleQuotesDialog(self,self._vehicle_ids[row],self.main.user,self.main.device); dlg.exec(); self.refresh(); self.main.refresh_all()
        def open_quote(self,row,col):
            if row < len(self._quote_ids):
                dlg=QuoteDetailDialog(self,self._quote_ids[row],self.main.user,self.main.device); dlg.exec(); self.refresh(); self.main.refresh_all()

    class ReporteriaPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main=main; self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Reportería", "Inventario premium, costos detallados, cotizaciones y KPIs comerciales con logo de la empresa."))
            grid=QGridLayout(); lay.addLayout(grid)
            items=[("Excel inventario premium","Inventario gerencial detallado con costos, utilidad, margen y estado comercial.",self.report_excel_inventory),("Excel costos por vehículo","Formato tipo COPART con resumen y detalle completo de costos por carro.",self.report_excel_costs),("HTML gerencial premium","Reporte visual tipo presentación ejecutiva con logo y KPIs.",self.report_html_inventory),("Excel KPI tiempos","Tiempos por etapa y ciclo compra→disponible.",self.report_excel_kpi),("Excel cotizaciones premium","Control comercial con colores verde/amarillo/rojo y cuotas.",self.report_excel_quotes),("HTML cotizaciones premium","Dashboard comercial de clientes, medios y seguimiento.",self.report_html_quotes)]
            for i,(title,desc,fn) in enumerate(items):
                card=QFrame(); card.setStyleSheet("QFrame{background:white;border:1px solid #d9e2ef;border-radius:18px;}"); cl=QVBoxLayout(card); lab=QLabel(title); lab.setStyleSheet("font-size:16pt;font-weight:900;color:#08285a;"); d=QLabel(desc); d.setWordWrap(True); d.setStyleSheet("color:#64748b;"); b=QPushButton("Generar"); b.setObjectName("orange"); b.clicked.connect(fn); cl.addWidget(lab); cl.addWidget(d); cl.addStretch(1); cl.addWidget(b); grid.addWidget(card,i//2,i%2)
            lay.addStretch(1)
        def _after_report(self,path:Path,title:str):
            msg=QMessageBox(self); msg.setWindowTitle(title); msg.setText(f"Reporte generado correctamente:\n{path.name}"); open_btn=msg.addButton("Abrir",QMessageBox.ButtonRole.AcceptRole); copy_btn=msg.addButton("Guardar una copia…",QMessageBox.ButtonRole.ActionRole); msg.addButton("Cerrar",QMessageBox.ButtonRole.RejectRole); msg.exec(); clicked=msg.clickedButton()
            if clicked==open_btn: QDesktopServices.openUrl(QUrl.fromLocalFile(str(path)))
            elif clicked==copy_btn:
                target,_=QFileDialog.getSaveFileName(self,"Guardar copia",str(Path.home()/path.name),"Archivos (*"+path.suffix+")")
                if target:
                    ok,m=copy_report_to(path,Path(target)); QMessageBox.information(self,"Copia",m) if ok else QMessageBox.warning(self,"Copia",m)
        def report_excel_inventory(self):
            ok,msg,out=generate_inventory_excel(load_vehicles(),self.main.user); QMessageBox.warning(self,"Excel",msg) if not ok or not out else self._after_report(out,"Excel inventario premium")
        def report_excel_costs(self):
            ok,msg,out=generate_vehicle_costs_excel_report(load_vehicles(),self.main.user); QMessageBox.warning(self,"Excel costos",msg) if not ok or not out else self._after_report(out,"Excel costos por vehículo")
        def report_html_inventory(self):
            out=generate_html_report(load_vehicles(),self.main.user); self._after_report(out,"HTML gerencial") if out else None
        def report_excel_kpi(self):
            ok,msg,out=generate_kpi_excel_report(load_vehicles(),self.main.user); QMessageBox.warning(self,"Excel KPI",msg) if not ok or not out else self._after_report(out,"Excel KPI")
        def report_excel_quotes(self):
            ok,msg,out=generate_quotes_excel_report(self.main.user); QMessageBox.warning(self,"Excel cotizaciones",msg) if not ok or not out else self._after_report(out,"Excel cotizaciones")
        def report_html_quotes(self):
            out=generate_quotes_html_report(self.main.user); self._after_report(out,"HTML cotizaciones") if out else None




# =============================================================================
# AJUSTES V4.3 - REPORTERIA EXACTA, PROPUESTAS EN CARPETA DEL CARRO Y EXPEDIENTE
# =============================================================================

# El usuario indicó que eliminará las carpetas PLANTILLAS y VEHICULOS.
# El sistema ya no depende de esas carpetas; las plantillas se buscan junto al .py/.exe
# y en la raíz de CONTROL VEHICULOS LYM.
ALL_SUBFOLDERS = [SUB_DATOS, SUB_DOCUMENTOS, SUB_FOTOS, SUB_REPORTES, SUB_RESPALDOS, SUB_TEMP]

INVENTORY_TEMPLATE_CANDIDATES = [
    "Inventario_vehiculos_LYM_JUNIO_2026 ac.xlsx",
    "Inventario_vehiculos_LYM_2026_83_Vehiculos.xlsx",
    "LYM_INVENTARIO_TEMPLATE.xlsx",
]
QUOTES_TEMPLATE_CANDIDATES = [
    "Control_Cotizaciones_LYM.xlsx",
    "LYM_CONTROL_COTIZACIONES_TEMPLATE.xlsx",
]
COSTS_TEMPLATE_CANDIDATES = [
    "COPART INC - COSTOS DE VEHICULOS (JONATHAN).xlsx",
    "COPART_INC_COSTOS_DE_VEHICULOS_JONATHAN.xlsx",
    "LYM_COSTOS_VEHICULOS_TEMPLATE.xlsx",
]
HTML_REPORT_TEMPLATE_CANDIDATES = [
    "LYM_Inversiones_Reporte_Junio_2026_FINAL_LIMPIO.html",
    "LYM_PLANTILLA_REPORTE_GERENCIAL.html",
    F_HTML_TEMPLATE,
]

# Reconfigurar ResourceManager sin carpeta PLANTILLAS.
try:
    ResourceManager.TEMPLATE_CANDIDATES = HTML_REPORT_TEMPLATE_CANDIDATES
    def _lym_search_paths_no_template(cls) -> list[Path]:
        paths: list[Path] = []
        def add(p: Optional[Path]):
            if p and p not in paths:
                paths.append(p)
        app = _app_dir()
        add(app)
        df = get_data_folder()
        if df:
            add(df)
            add(df / SUB_REPORTES)
        return paths
    ResourceManager.search_paths = classmethod(_lym_search_paths_no_template)
except Exception:
    pass


def _vehicle_year_from_code(vehicle_code: str) -> str:
    m = re.search(r"LYM-CV-(\d{4})-", str(vehicle_code or ""))
    return m.group(1) if m else str(date.today().year)


def vehicle_document_folder(vehicle_or_code: Any, create: bool = True) -> Optional[Path]:
    """Carpeta única del carro: documentos, OCs, comprobantes, propuestas y anexos."""
    df = get_data_folder()
    if df is None:
        return None
    if isinstance(vehicle_or_code, dict):
        code = vehicle_or_code.get("codigo") or vehicle_or_code.get("vehicle_code") or "VEHICULO"
    else:
        code = str(vehicle_or_code or "VEHICULO")
    folder = df / SUB_DOCUMENTOS / _vehicle_year_from_code(code) / _safe_filename(code)
    if create:
        folder.mkdir(parents=True, exist_ok=True)
    return folder


def vehicle_proposals_folder(vehicle_or_code: Any, create: bool = True) -> Optional[Path]:
    base = vehicle_document_folder(vehicle_or_code, create=create)
    if base is None:
        return None
    out = base / "PROPUESTAS"
    if create:
        out.mkdir(parents=True, exist_ok=True)
    return out


def _proposal_output_dir(quote: Optional[dict] = None) -> Optional[Path]:
    if quote:
        return vehicle_proposals_folder(quote.get("vehicle_code") or quote.get("vehicle_snapshot", {}).get("codigo") or "VEHICULO")
    df = get_data_folder()
    return None if df is None else df / SUB_REPORTES / "PROPUESTAS"


def _find_template_file(candidates: list[str]) -> Optional[Path]:
    # Búsqueda intencionalmente simple: misma carpeta del .py/.exe y raíz del sistema.
    search: list[Path] = []
    for pth in [_app_dir(), get_data_folder()]:
        if pth and pth not in search:
            search.append(pth)
    for folder in search:
        try:
            if not folder.exists():
                continue
            for name in candidates:
                p = folder / name
                if p.exists() and p.is_file():
                    return p
        except Exception:
            continue
    return None


def _spanish_month_year(d: Optional[date] = None) -> str:
    d = d or date.today()
    names = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    return f"{names[d.month-1]} {d.year}"


def _spanish_date_long(d: Optional[date] = None) -> str:
    d = d or date.today()
    months = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    return f"{d.day:02d} de {months[d.month-1]} de {d.year}"


def _excel_date(value: Any) -> Any:
    dd = _parse_date(value)
    return dd if dd else None


def _copy_template_or_new(template: Optional[Path], out: Path):
    out.parent.mkdir(parents=True, exist_ok=True)
    if template and template.exists():
        shutil.copy2(template, out)
        return True
    return False


def _copy_row_style(ws, src_row: int, dst_row: int, max_col: Optional[int] = None):
    try:
        from copy import copy
        max_col = max_col or ws.max_column
        ws.row_dimensions[dst_row].height = ws.row_dimensions[src_row].height
        for c in range(1, max_col + 1):
            src = ws.cell(src_row, c)
            dst = ws.cell(dst_row, c)
            if src.has_style:
                dst._style = copy(src._style)
            if src.number_format:
                dst.number_format = src.number_format
            if src.alignment:
                dst.alignment = copy(src.alignment)
    except Exception:
        pass


def _clear_row_values(ws, row: int, cols: range):
    # En plantillas con celdas combinadas, los MergedCell no aceptan escritura directa.
    # Solo limpiamos las celdas reales para preservar el diseño original.
    for c in cols:
        try:
            ws.cell(row, c).value = None
        except AttributeError:
            continue


def _report_output_dir() -> Optional[Path]:
    df = get_data_folder()
    if df is None:
        return None
    out = df / SUB_REPORTES
    out.mkdir(parents=True, exist_ok=True)
    return out


def _cost_total_by_keywords(items: list[dict], keys: list[str]) -> float:
    total = 0.0
    for item in items or []:
        txt = " ".join([_norm(item.get("categoria")), _norm(item.get("subcategoria")), _norm(item.get("stage_key")), _norm(item.get("descripcion"))])
        if any(k in txt for k in keys):
            total += _to_float(item.get("monto_usd"), 0)
    return round(total, 2)


def _vehicle_sale_date(v: dict) -> Optional[date]:
    for h in reversed(v.get("historial", []) or []):
        if str(h.get("accion", "")).upper() in ("VENTA", "VENDIDO", "COTIZACION_GANADA", "MARCAR_VENDIDO"):
            d = _parse_date(h.get("fecha"))
            if d:
                return d
    return _parse_date(v.get("fecha_venta"))


def _stage_date(v: dict, stage_key: str, field: str = "fecha_inicio") -> Optional[date]:
    return _parse_date(vehicle_stage(v, stage_key).get(field))


def _doc_label(name: str, rel: str = "") -> str:
    return "✅ " + (name or "Cargado") if rel else "❌ Vacío"


def decrypt_file_to_path(rel_path: str, original_name: str, target_dir: Path) -> Optional[Path]:
    src = rel_to_abs(rel_path)
    if not src or not src.exists():
        return None
    try:
        dec = CryptoManager.decrypt_bytes(src.read_bytes())
        if dec is None:
            return None
        target_dir.mkdir(parents=True, exist_ok=True)
        safe_name = _safe_filename(original_name or src.stem)
        if not Path(safe_name).suffix:
            safe_name += ".pdf"
        out = target_dir / safe_name
        if out.exists():
            out = target_dir / f"{Path(safe_name).stem}_{datetime.now().strftime('%H%M%S')}{Path(safe_name).suffix}"
        out.write_bytes(dec)
        return out
    except Exception:
        return None


def vehicle_document_entries(vehicle: dict) -> list[dict]:
    ensure_vehicle_runtime_fields(vehicle)
    entries: list[dict] = []
    stc = vehicle_stage(vehicle, STAGE_COMPRADO)
    if stc.get("documento"):
        entries.append({
            "stage_key": STAGE_COMPRADO,
            "etapa": STAGE_META[STAGE_COMPRADO]["label"],
            "categoria": "COMPRA",
            "subcategoria": "COMPROBANTE COMPRA",
            "descripcion": "Comprobante de compra del vehículo",
            "monto_usd": stc.get("costo_usd"),
            "proveedor": vehicle.get("subasta", ""),
            "oc_numero": vehicle.get("oc_compra_numero", ""),
            "comprobante": stc.get("documento"),
            "comprobante_nombre": stc.get("documento_nombre") or "comprobante_compra.pdf",
            "oc_documento": vehicle.get("oc_compra_documento", ""),
            "oc_documento_nombre": vehicle.get("oc_compra_documento_nombre", ""),
        })
    for g in vehicle.get("gastos_detallados", []) or []:
        # Evitar duplicar la fila de compra si ya se mostró como comprobante principal.
        if g.get("source") == "purchase" and not g.get("comprobante") and not g.get("oc_documento"):
            continue
        entries.append({
            "stage_key": g.get("stage_key", ""),
            "etapa": STAGE_META.get(g.get("stage_key"), {}).get("label", g.get("stage_key", "")),
            "categoria": g.get("categoria", ""),
            "subcategoria": g.get("subcategoria", ""),
            "descripcion": g.get("descripcion", ""),
            "monto_usd": g.get("monto_usd", 0),
            "proveedor": g.get("proveedor", ""),
            "oc_numero": g.get("oc_numero", ""),
            "comprobante": g.get("comprobante", ""),
            "comprobante_nombre": g.get("comprobante_nombre", ""),
            "oc_documento": g.get("oc_documento", ""),
            "oc_documento_nombre": g.get("oc_documento_nombre", ""),
        })
    return entries


def export_vehicle_documents(vehicle_id: str, target_dir: Path) -> tuple[bool, str, list[Path]]:
    v = find_vehicle(vehicle_id)
    if not v:
        return False, "Vehículo no encontrado.", []
    base = target_dir / _safe_filename(v.get("codigo", "VEHICULO"))
    exported: list[Path] = []
    for idx, item in enumerate(vehicle_document_entries(v), start=1):
        prefix = _safe_filename(f"{idx:02d}_{item.get('etapa','')}_{item.get('subcategoria','')}")
        if item.get("comprobante"):
            out = decrypt_file_to_path(item.get("comprobante", ""), f"{prefix}_COMPROBANTE_{item.get('comprobante_nombre','comprobante.pdf')}", base)
            if out:
                exported.append(out)
        if item.get("oc_documento"):
            out = decrypt_file_to_path(item.get("oc_documento", ""), f"{prefix}_OC_{item.get('oc_documento_nombre','oc.pdf')}", base)
            if out:
                exported.append(out)
    return bool(exported), f"Se exportaron {len(exported)} comprobantes/OC." if exported else "No hay comprobantes u OC para exportar.", exported


def _safe_quote_filename(quote: dict, ext: str) -> str:
    snap = quote.get("vehicle_snapshot", {})
    cl = quote.get("cliente", {})
    base = f"PROPUESTA_LEASING_{_safe_filename(cl.get('nombre','CLIENTE'))}_{_safe_filename(_vehicle_display_name(snap))}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    return base + ext


def generate_quote_proposal_pdf(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada.", None
    out_dir = _proposal_output_dir(quote)
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    except Exception:
        return False, "Para generar PDF instala: pip install reportlab", None
    ensure_quote_runtime_fields(quote)
    cl = quote.get("cliente", {})
    snap = quote.get("vehicle_snapshot", {})
    s = _quote_financial_summary(quote)
    vehicle_name = _vehicle_display_name(snap).upper()
    out = out_dir / _safe_quote_filename(quote, ".pdf")
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="LYMTitle", parent=styles["Title"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=colors.HexColor("#061F4A")))
    styles.add(ParagraphStyle(name="LYMVehicle", parent=styles["Heading2"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor("#061F4A")))
    styles.add(ParagraphStyle(name="LYMBody", parent=styles["BodyText"], alignment=TA_LEFT, fontSize=10.5, leading=15))
    styles.add(ParagraphStyle(name="LYMSmall", parent=styles["BodyText"], fontSize=8.8, leading=11))
    doc = SimpleDocTemplate(str(out), pagesize=letter, rightMargin=0.55*inch, leftMargin=0.55*inch, topMargin=0.42*inch, bottomMargin=0.45*inch)
    story = []
    logo = ResourceManager.find_logo()
    logo_flow = Image(str(logo), width=0.9*inch, height=0.9*inch) if logo and logo.exists() else Paragraph("<b>L&M</b>", styles["LYMTitle"])
    header = Table([[logo_flow, Paragraph("<b>L&amp;M Inversiones, S.A. de C.V.</b><br/><font size='8'>Ayudando a lograr tus sueños</font>", styles["LYMSmall"])]], colWidths=[1.1*inch, 6.0*inch])
    header.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP"), ("LINEBELOW", (1,0), (1,0), 3, colors.HexColor("#061F4A")), ("BOTTOMPADDING", (0,0), (-1,-1), 8)]))
    story.append(header)
    story.append(Spacer(1, 0.10*inch))
    story.append(Paragraph(f"<b>{html.escape(cl.get('nombre','Cliente').upper())}</b><br/>Presente.", styles["LYMBody"]))
    story.append(Spacer(1, 0.14*inch))
    story.append(Paragraph("<u>Propuesta de Arrendamiento Vehicular</u>", styles["LYMTitle"]))
    story.append(Spacer(1, 0.14*inch))
    story.append(Paragraph("Por medio de la presente, <b>L&amp;M Inversiones, S.A. de C.V.</b> tiene el agrado de presentarle la propuesta de contrato de arrendamiento (Leasing) opción: <b>RENT A CAR.</b>", styles["LYMBody"]))
    story.append(Spacer(1, 0.14*inch))
    story.append(Paragraph(f"◆ {html.escape(vehicle_name)} ◆", styles["LYMVehicle"]))
    story.append(Spacer(1, 0.10*inch))
    story.append(Paragraph("<b>Características destacadas:</b>", styles["Heading3"]))
    feats = _split_features(snap)
    left, right = feats[0::2], feats[1::2]
    rows = []
    for i in range(max(len(left), len(right))):
        rows.append([Paragraph("• " + html.escape(left[i]) if i < len(left) else "", styles["LYMBody"]), Paragraph("• " + html.escape(right[i]) if i < len(right) else "", styles["LYMBody"])])
    ft = Table(rows, colWidths=[3.55*inch, 3.55*inch])
    ft.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP"), ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4)]))
    story.append(ft)
    story.append(Spacer(1, 0.14*inch))
    story.append(Paragraph(f"Compartimos ante usted el detalle de prima, monto leasing, plazo y cuota mensual para <b>{html.escape(vehicle_name)}</b>. El cálculo incluye seguro estimado, servicio GPS, IVA y una tasa del <b>{s['tasa']}%</b> mensual.", styles["LYMBody"]))
    story.append(Spacer(1, 0.12*inch))
    summary_rows = [
        ["Valor del vehículo", _fmt_usd(s["valor_vehiculo"]), "Prima requerida", _fmt_usd(s["prima_requerida"])],
        ["Monto leasing", _fmt_usd(s["monto_leasing"]), "Costo legal", _fmt_usd(s["legal"]) + " incluye IVA"],
        ["Plazo", f"{s['plazo']} meses", "Opción de compra", _fmt_usd(s["opcion_compra"])],
        ["Tasa mensual", f"{s['tasa']}%", "Cuota mensual", _fmt_usd(s["cuota_final"])],
    ]
    mt = Table(summary_rows, colWidths=[1.55*inch, 1.85*inch, 1.55*inch, 2.15*inch])
    mt.setStyle(TableStyle([("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E1")), ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#F8FAFC")), ("FONTNAME", (0,0), (0,-1), "Helvetica-Bold"), ("FONTNAME", (2,0), (2,-1), "Helvetica-Bold"), ("FONTNAME", (1,3), (1,3), "Helvetica-Bold"), ("FONTNAME", (3,3), (3,3), "Helvetica-Bold"), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("TOPPADDING", (0,0), (-1,-1), 7), ("BOTTOMPADDING", (0,0), (-1,-1), 7)]))
    story.append(mt)
    story.append(Spacer(1, 0.10*inch))
    story.append(Paragraph("<font color='red'><b>La cuota mensual incluye IVA, seguro y GPS. Prima requerida corresponde a prima mínima más comisión.</b></font>", styles["LYMBody"]))
    story.append(Spacer(1, 0.18*inch))
    story.append(Paragraph("Quedamos atentos a cualquier consulta o duda.", styles["LYMBody"]))
    story.append(Spacer(1, 0.13*inch))
    story.append(Paragraph(f"San Salvador, {_spanish_date_long()}", styles["LYMBody"]))
    story.append(PageBreak())
    story.append(header)
    story.append(Spacer(1, 0.18*inch))
    story.append(Paragraph("Atentamente,", styles["LYMBody"]))
    story.append(Spacer(1, 0.20*inch))
    story.append(Paragraph("<b>Guillermo Moreno</b><br/>L&amp;M Inversiones, S.A. de C.V.<br/>guillermo.moreno@lyminversiones.com<br/>Tel: (503) 7475 5821", styles["LYMBody"]))
    story.append(Spacer(1, 0.30*inch))
    story.append(Paragraph("X____________________ ___________<br/>Aceptado por cliente", styles["LYMBody"]))
    story.append(Spacer(1, 0.20*inch))
    story.append(Paragraph("<b>Condiciones y Vigencia de la Oferta</b>", styles["Heading3"]))
    for c in [
        "Esta propuesta forma parte de una oferta especial de arrendamiento vehicular (leasing) válida por 15 días calendario a partir de la fecha de emisión.",
        f"La tasa de interés mensual del {s['tasa']}% brinda condiciones accesibles y transparentes para nuestros clientes.",
        "La propuesta se formalizará mediante contrato de arrendamiento donde se establecerán condiciones de uso, pagos mensuales y opción de adquisición al finalizar el plazo.",
        "El valor del seguro mostrado es estimado y puede variar según las características y uso del vehículo.",
        "Los gastos administrativos y legales derivados de la formalización deberán cancelarse al momento de la firma.",
        "Las cuotas mensuales incluyen IVA, seguro y GPS, reflejando el monto total a pagar.",
        "Durante el período de arrendamiento, el cliente deberá mantener el vehículo en buen estado y cumplir con los servicios preventivos o correctivos recomendados.",
        "Esta propuesta tiene carácter informativo y no constituye compromiso contractual hasta la formalización del contrato.",
    ]:
        story.append(Paragraph("• " + html.escape(c), styles["LYMSmall"]))
        story.append(Spacer(1, 0.04*inch))
    doc.build(story)
    log_audit("GENERAR_PROPUESTA_PDF", (user or {}).get("usuario", ""), quote.get("vehicle_code", ""), str(out))
    return True, "Propuesta PDF generada en carpeta del carro.", out


def generate_quote_proposal_excel(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada.", None
    out_dir = _proposal_output_dir(quote)
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from openpyxl.drawing.image import Image as XLImage
    except Exception:
        return False, "Para generar Excel instala: pip install openpyxl pillow", None
    cl = quote.get("cliente", {})
    snap = quote.get("vehicle_snapshot", {})
    s = _quote_financial_summary(quote)
    vehicle_name = _vehicle_display_name(snap).upper()
    out = out_dir / _safe_quote_filename(quote, ".xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Propuesta Leasing"
    navy = "08285A"; orange = "F59A13"; light = "F8FAFC"; line = "CBD5E1"
    thin = Side(style="thin", color=line)
    ws.sheet_view.showGridLines = False
    for c, w in enumerate([16, 18, 18, 18, 18, 18, 18, 18], start=1):
        ws.column_dimensions[get_column_letter(c)].width = w
    ws.row_dimensions[1].height = 62
    logo = ResourceManager.find_logo()
    if logo and logo.exists():
        try:
            img = XLImage(str(logo)); img.width = 88; img.height = 88; ws.add_image(img, "A1")
        except Exception:
            ws["A1"] = "L&M"
    ws.merge_cells("B2:H2"); ws["B2"] = "L&M INVERSIONES, S.A. DE C.V."; ws["B2"].font = Font(bold=True, size=14, color=navy); ws["B2"].alignment = Alignment(horizontal="center")
    ws.merge_cells("B3:H3"); ws["B3"] = "Ayudando a lograr tus sueños"; ws["B3"].font = Font(italic=True, color=navy); ws["B3"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A5:H5"); ws["A5"] = "PROPUESTA DE ARRENDAMIENTO VEHICULAR"; ws["A5"].font = Font(bold=True, underline="single", size=16, color=navy); ws["A5"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A7:H7"); ws["A7"] = f"◆ {vehicle_name} ◆"; ws["A7"].font = Font(bold=True, size=15, color=navy); ws["A7"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A9:H9"); ws["A9"] = f"Cliente: {cl.get('nombre','').upper()}    Teléfono: {cl.get('telefono','')}    Medio: {cl.get('medio_contacto','')}    Fecha: {_spanish_date_long()}"; ws["A9"].font = Font(bold=True, color=navy)
    ws.merge_cells("A11:H12"); ws["A11"] = f"Por medio de la presente, L&M Inversiones, S.A. de C.V. tiene el agrado de presentarle la propuesta de contrato de arrendamiento (Leasing) para {vehicle_name}."; ws["A11"].alignment = Alignment(wrap_text=True, vertical="center")
    ws.merge_cells("A14:H14"); ws["A14"] = "RESUMEN PARA EL CLIENTE"; ws["A14"].fill = PatternFill("solid", fgColor=navy); ws["A14"].font = Font(bold=True, color="FFFFFF"); ws["A14"].alignment = Alignment(horizontal="center")
    rows = [
        ["Valor del vehículo", s["valor_vehiculo"], "Prima requerida", s["prima_requerida"]],
        ["Monto leasing", s["monto_leasing"], "Costo legal", s["legal"]],
        ["Plazo", f"{s['plazo']} meses", "Opción de compra", s["opcion_compra"]],
        ["Tasa mensual", f"{s['tasa']}%", "Cuota mensual", s["cuota_final"]],
    ]
    start = 15
    for r, row in enumerate(rows, start=start):
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
        ws.merge_cells(start_row=r, start_column=3, end_row=r, end_column=4)
        ws.merge_cells(start_row=r, start_column=5, end_row=r, end_column=6)
        ws.merge_cells(start_row=r, start_column=7, end_row=r, end_column=8)
        ws.cell(r,1).value = row[0]; ws.cell(r,3).value = row[1]; ws.cell(r,5).value = row[2]; ws.cell(r,7).value = row[3]
        for c in [1,3,5,7]:
            cell = ws.cell(r,c); cell.border = Border(top=thin,bottom=thin,left=thin,right=thin); cell.alignment = Alignment(horizontal="center", vertical="center")
            if c in (1,5):
                cell.fill = PatternFill("solid", fgColor=light); cell.font = Font(bold=True, color=navy)
            else:
                cell.font = Font(bold=True, color=navy)
                if isinstance(cell.value, (int,float)):
                    cell.number_format = '$#,##0.00'
    ws.merge_cells("A20:H20"); ws["A20"] = "CARACTERÍSTICAS DESTACADAS"; ws["A20"].fill = PatternFill("solid", fgColor=orange); ws["A20"].font = Font(bold=True, color=navy); ws["A20"].alignment = Alignment(horizontal="center")
    feats = _split_features(snap)
    for idx, feat in enumerate(feats[:14]):
        row = 21 + idx // 2
        col = 1 if idx % 2 == 0 else 5
        ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col+3)
        ws.cell(row, col).value = "• " + feat
        ws.cell(row, col).alignment = Alignment(wrap_text=True, vertical="top")
    note_row = 30
    ws.merge_cells(start_row=note_row, start_column=1, end_row=note_row+1, end_column=8)
    ws.cell(note_row,1).value = "La cuota mensual incluye IVA, seguro y GPS. Prima requerida corresponde a prima mínima más comisión. Esta propuesta tiene vigencia de 15 días calendario a partir de su emisión."
    ws.cell(note_row,1).font = Font(bold=True, color="DC2626")
    ws.cell(note_row,1).alignment = Alignment(wrap_text=True, vertical="center")
    ws.merge_cells("A34:D34"); ws["A34"] = "Atentamente,\nGuillermo Moreno\nL&M Inversiones, S.A. de C.V."; ws["A34"].alignment = Alignment(wrap_text=True)
    ws.merge_cells("E34:H34"); ws["E34"] = "Aceptado por cliente:\n\nX____________________________"; ws["E34"].alignment = Alignment(wrap_text=True)
    for row in range(1, 38):
        ws.row_dimensions[row].height = 23
    ws.page_setup.paperSize = ws.PAPERSIZE_LETTER
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.page_margins.left = 0.35; ws.page_margins.right = 0.35; ws.page_margins.top = 0.45; ws.page_margins.bottom = 0.45
    wb.save(out)
    log_audit("GENERAR_PROPUESTA_EXCEL", (user or {}).get("usuario", ""), quote.get("vehicle_code", ""), str(out))
    return True, "Propuesta Excel generada en carpeta del carro.", out


def generate_quote_proposal_selected(quote_id: str, formato: str, user: Optional[dict] = None) -> tuple[bool, str, list[Path]]:
    formato = _norm(formato)
    paths: list[Path] = []
    msgs: list[str] = []
    ok_all = True
    if "PDF" in formato:
        ok, msg, p = generate_quote_proposal_pdf(quote_id, user)
        ok_all = ok_all and ok; msgs.append(msg)
        if p: paths.append(p)
    if "EXCEL" in formato or "XLSX" in formato:
        ok, msg, p = generate_quote_proposal_excel(quote_id, user)
        ok_all = ok_all and ok; msgs.append(msg)
        if p: paths.append(p)
    return ok_all and bool(paths), "\n".join(msgs), paths


def generate_quote_proposal_files(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, list[Path]]:
    return generate_quote_proposal_selected(quote_id, "PDF + EXCEL", user)


def generate_inventory_excel(vehicles: list[dict], user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    out_dir = _report_output_dir()
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    out = out_dir / f"Inventario_vehiculos_LYM_{_spanish_month_year().replace(' ','_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    template = _find_template_file(INVENTORY_TEMPLATE_CANDIDATES)
    try:
        from openpyxl import Workbook, load_workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except Exception:
        return False, "Para generar Excel instala: pip install openpyxl", None
    if template:
        shutil.copy2(template, out)
        wb = load_workbook(out)
    else:
        wb = Workbook(); wb.active.title = "Inventario General JUNIO"
    ws = wb["Inventario General JUNIO"] if "Inventario General JUNIO" in wb.sheetnames else wb.active
    start_row = 5
    max_needed = start_row + len(vehicles) + 3
    for r in range(start_row, max_needed):
        if r > ws.max_row:
            ws.append([None] * 26)
        _copy_row_style(ws, start_row, r, 26)
        _clear_row_values(ws, r, range(1, 27))
    for i, v in enumerate(vehicles, start=1):
        ensure_vehicle_runtime_fields(v)
        r = start_row + i - 1
        sale = _vehicle_sale_date(v)
        status = v.get("estado_comercial") if v.get("estado_comercial") == COMM_VENDIDO else STAGE_META.get(v.get("estado_actual"), {}).get("label", v.get("estado_actual"))
        compra = _parse_date(v.get("fecha_compra"))
        disp = _stage_date(v, STAGE_DISPONIBLE, "fecha_inicio") or _stage_date(v, STAGE_DISPONIBLE, "fecha_fin")
        taller_in = _stage_date(v, STAGE_PREPARACION, "taller_ingreso") or _stage_date(v, STAGE_PREPARACION, "fecha_inicio")
        taller_out = _stage_date(v, STAGE_PREPARACION, "taller_salida") or _stage_date(v, STAGE_PREPARACION, "fecha_fin")
        leg_in = _stage_date(v, STAGE_PREPARACION, "placas_ingreso") or _stage_date(v, STAGE_PREPARACION, "fecha_inicio")
        leg_out = _stage_date(v, STAGE_PREPARACION, "legalizacion_fin") or _stage_date(v, STAGE_PREPARACION, "fecha_fin")
        vals = {
            1: i, 2: v.get("lote") or "LOCAL", 3: compra, 4: sale, 5: v.get("marca"), 6: v.get("modelo"), 7: v.get("anio"),
            8: status, 9: v.get("cliente") or v.get("observaciones", ""), 10: STAGE_META.get(v.get("estado_actual"), {}).get("label", v.get("estado_actual")),
            12: _stage_date(v, STAGE_ADUANA, "fecha_inicio"), 13: taller_in, 14: taller_out, 15: leg_in, 16: leg_out, 17: disp,
            18: f"=IF(AND(M{r}<>\"\",N{r}<>\"\"),N{r}-M{r},\"\")",
            19: f"=IF(AND(O{r}<>\"\",P{r}<>\"\"),P{r}-O{r},\"\")",
            20: f"=IF(AND(C{r}<>\"\",Q{r}<>\"\"),Q{r}-C{r},\"\")",
            21: f"=IF(AND(Q{r}<>\"\",D{r}<>\"\"),D{r}-Q{r},\"\")",
            22: f"=IF(AND(C{r}<>\"\",D{r}<>\"\"),D{r}-C{r},\"\")",
            23: f"=IF(D{r}=\"\",TODAY()-C{r},\"\")",
            24: "CERRADO" if sale else stage_alert_level(v), 25: sale.strftime("%B %Y") if sale else "", 26: "LOCAL" if str(v.get("lote", "")).upper() == "LOCAL" else "LOTE",
        }
        for c, val in vals.items():
            ws.cell(r, c).value = val
        for c in [3,4,12,13,14,15,16,17]:
            ws.cell(r, c).number_format = "dd/mm/yyyy"
    try:
        ws.auto_filter.ref = f"A4:Z{start_row + len(vehicles)}"
        ws.freeze_panes = "A5"
    except Exception:
        pass
    wb.save(out)
    log_audit("GENERAR_EXCEL_INVENTARIO_EXACTO", (user or {}).get("usuario", ""), "", out.name)
    return True, "Inventario generado con el formato original de L&M.", out


def generate_quotes_excel_report(user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    out_dir = _report_output_dir()
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    out = out_dir / f"Control_Cotizaciones_LYM_{_spanish_month_year().replace(' ','_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    template = _find_template_file(QUOTES_TEMPLATE_CANDIDATES)
    try:
        from openpyxl import Workbook, load_workbook
    except Exception:
        return False, "Para generar Excel instala: pip install openpyxl", None
    if template:
        shutil.copy2(template, out)
        wb = load_workbook(out)
    else:
        wb = Workbook(); wb.active.title = _spanish_month_year()
    sheet_name = _spanish_month_year()
    ws = wb[sheet_name] if sheet_name in wb.sheetnames else (wb["Junio 2026"] if "Junio 2026" in wb.sheetnames else wb.active)
    start_row = 5
    qs = sorted(load_quotes(), key=lambda q: q.get("fecha_cotizacion", ""))
    for r in range(start_row, start_row + max(len(qs), 1) + 5):
        if r > ws.max_row:
            ws.append([None] * 9)
        _copy_row_style(ws, start_row, r, 9)
        _clear_row_values(ws, r, range(1, 10))
    for idx, q in enumerate(qs, start=1):
        r = start_row + idx - 1
        cl = q.get("cliente", {})
        snap = q.get("vehicle_snapshot", {})
        s = _quote_financial_summary(q)
        ws.cell(r, 1).value = _parse_date(q.get("fecha_cotizacion"))
        ws.cell(r, 1).number_format = "dd/mm/yyyy"
        ws.cell(r, 2).value = f"COT-{idx:03d}"
        ws.cell(r, 3).value = cl.get("nombre", "")
        ws.cell(r, 4).value = " ".join([str(snap.get("marca") or ""), str(snap.get("modelo") or "")]).strip()
        ws.cell(r, 5).value = snap.get("anio", "")
        ws.cell(r, 6).value = cl.get("telefono", "")
        ws.cell(r, 7).value = "LEASING"
        ws.cell(r, 8).value = f"{cl.get('medio_contacto','')} · {QUOTE_STATUS_LABELS.get(q.get('estado'), q.get('estado'))} · {quote_alert_level(q)} · {s['plazo']} meses · prima {_fmt_usd(s['prima_requerida'])} · cuota {_fmt_usd(s['cuota_final'])}"
    try:
        ws.auto_filter.ref = f"A4:H{start_row + len(qs)}"
        ws.freeze_panes = "A5"
    except Exception:
        pass
    wb.save(out)
    log_audit("GENERAR_EXCEL_COTIZACIONES_EXACTO", (user or {}).get("usuario", ""), "", out.name)
    return True, "Control de cotizaciones generado con el formato original de L&M.", out


def generate_vehicle_costs_excel_report(vehicles: Optional[list[dict]] = None, user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    vehicles = vehicles if vehicles is not None else load_vehicles()
    out_dir = _report_output_dir()
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    out = out_dir / f"COPART_INC_COSTOS_DE_VEHICULOS_LYM_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    template = _find_template_file(COSTS_TEMPLATE_CANDIDATES)
    try:
        from openpyxl import Workbook, load_workbook
    except Exception:
        return False, "Para generar Excel instala: pip install openpyxl", None
    if template:
        shutil.copy2(template, out)
        wb = load_workbook(out)
    else:
        wb = Workbook(); wb.active.title = "COPART INC"; wb.create_sheet("COSTO")
    ws = wb["COPART INC"] if "COPART INC" in wb.sheetnames else wb.active
    wc = wb["COSTO"] if "COSTO" in wb.sheetnames else wb.create_sheet("COSTO")
    start_row = 6
    for sh, max_col in [(ws, 54), (wc, 30)]:
        for r in range(start_row, start_row + max(len(vehicles), 1) + 5):
            if r > sh.max_row:
                sh.append([None] * max_col)
            _copy_row_style(sh, start_row, r, max_col)
            _clear_row_values(sh, r, range(1, max_col + 1))
    for idx, v in enumerate(vehicles, start=1):
        ensure_vehicle_runtime_fields(v)
        r = start_row + idx - 1
        g = v.get("gastos_detallados", []) or []
        compra = _to_float(v.get("precio_ganado_usd"), 0)
        traslado = _cost_total_by_keywords(g, ["TRASLADO", "GRUA_TRASLADO", "USA"])
        flete_bl = _cost_total_by_keywords(g, ["FLETE", "BL", "NAVIERA_GRUA", "NAVIERA"])
        impuestos = _cost_total_by_keywords(g, ["IMPUESTOS", "DUCA", "DECLARACION"])
        aduanal = _cost_total_by_keywords(g, ["TRAMITE", "ADUANAL", "ADUANALES"])
        repuestos = _cost_total_by_keywords(g, ["REPUESTO"])
        pintura = _cost_total_by_keywords(g, ["PINTURA", "ENDEREZADO"])
        mecanica = _cost_total_by_keywords(g, ["MECANICO", "MECANICA"])
        grua_local = _cost_total_by_keywords(g, ["GRUA_LOCAL"])
        emisiones = _cost_total_by_keywords(g, ["EMISION"])
        placas = _cost_total_by_keywords(g, ["PLACA", "CITA", "EXPERTICIA"])
        almacen = _cost_total_by_keywords(g, ["ALMACEN"])
        honorarios = _cost_total_by_keywords(g, ["HONORARIO"])
        total_es = flete_bl + impuestos + aduanal + repuestos + pintura + mecanica + grua_local + emisiones + placas + almacen + honorarios
        cost = vehicle_total_cost(v)
        # COPART INC
        values = {1: idx, 2: v.get("lote"), 4: _parse_date(v.get("fecha_compra")), 6: idx, 7: v.get("anio"), 9: v.get("marca"), 13: v.get("modelo"), 14: v.get("color"), 15: v.get("oc_compra_numero", ""), 16: compra, 20: traslado, 25: compra + traslado, 27: flete_bl, 29: impuestos, 31: aduanal, 33: repuestos, 35: pintura, 37: mecanica, 39: grua_local, 41: emisiones, 47: placas, 49: placas, 51: almacen, 52: honorarios, 53: total_es, 54: cost}
        for c, val in values.items():
            ws.cell(r, c).value = val
        for c in [4]: ws.cell(r, c).number_format = "dd/mm/yyyy"
        for c in [16,20,25,27,29,31,33,35,37,39,41,47,49,51,52,53,54]: ws.cell(r,c).number_format = '$#,##0.00'
        # COSTO
        pf = v.get("precio_final") or {}
        venta_neta = _to_float(pf.get("venta_neta_usd"), 0)
        iva_venta = _to_float(pf.get("iva_venta_usd"), 0)
        iva_duca = _to_float(pf.get("iva_duca_usd"), 0)
        iva_pagar = _to_float(pf.get("iva_pagar_usd"), 0)
        pago_cuenta = _to_float(pf.get("pago_cuenta_usd"), 0)
        precio_cliente = _to_float(v.get("precio_venta_usd"), 0)
        leasing_val = precio_cliente
        prima = round(leasing_val * 0.20, 2) if leasing_val else 0
        rowc = {1: v.get("lote"), 3: _parse_date(v.get("fecha_compra")), 5: idx, 6: v.get("anio"), 8: v.get("marca"), 12: v.get("modelo"), 13: v.get("color"), 14: cost, 15: 0.30, 16: cost * 1.30 if cost else None, 17: venta_neta, 18: iva_venta, 19: iva_duca, 20: iva_pagar, 21: pago_cuenta, 22: precio_cliente, 23: precio_cliente*0.15 if precio_cliente else None, 24: precio_cliente*0.30 if precio_cliente else None, 25: leasing_val, 26: precio_cliente, 27: prima, 28: prima + 100 if prima else None, 29: vehicle_expected_profit(v), 30: v.get("observaciones", "")}
        for c, val in rowc.items():
            wc.cell(r, c).value = val
        wc.cell(r, 3).number_format = "dd/mm/yyyy"
        for c in [14,16,17,18,19,20,21,22,23,24,25,26,27,28,29]: wc.cell(r,c).number_format = '$#,##0.00'
        wc.cell(r,15).number_format = '0%'
    for sh in [ws, wc]:
        try:
            sh.freeze_panes = "A6"
            sh.auto_filter.ref = f"A4:{sh.cell(4, sh.max_column).coordinate[0]}{start_row + len(vehicles)}"
        except Exception:
            pass
    wb.save(out)
    log_audit("GENERAR_EXCEL_COSTOS_EXACTO", (user or {}).get("usuario", ""), "", out.name)
    return True, "Costos por vehículo generados con formato tipo COPART.", out


def generate_html_report(vehicles: list[dict], user: Optional[dict] = None) -> Optional[Path]:
    out_dir = _report_output_dir()
    if out_dir is None:
        return None
    template = _find_template_file(HTML_REPORT_TEMPLATE_CANDIDATES)
    k = compute_kpis(vehicles)
    generated = _spanish_date_long()
    rows = []
    for v in vehicles:
        ensure_vehicle_runtime_fields(v)
        rows.append(f"<tr><td>{html.escape(v.get('codigo',''))}</td><td>{html.escape(_vehicle_display_name(v))}</td><td>{html.escape(str(v.get('lote','')))}</td><td>{_fmt_usd(vehicle_total_cost(v))}</td><td>{_fmt_usd(v.get('precio_venta_usd'))}</td><td>{_fmt_usd(vehicle_expected_profit(v))}</td><td>{html.escape(STAGE_META.get(v.get('estado_actual'),{}).get('label',v.get('estado_actual')))}</td><td>{html.escape(str(v.get('estado_comercial','')))}</td><td>{vehicle_days_from_purchase(v)}</td></tr>")
    dynamic_slide = f"""
<section class="slide" id="sistema-actual"><div class="container reveal"><div class="eyebrow">Datos actualizados del sistema · {generated}</div><h2>Resumen actualizado desde LYM AUTO CONTROL</h2><div class="grid kpis"><div class="card kpi-card"><div class="label">Vehículos totales</div><div class="value">{k.get('total',0)}</div></div><div class="card kpi-card"><div class="label">Activos</div><div class="value">{k.get('activos',0)}</div></div><div class="card kpi-card"><div class="label">Disponibles</div><div class="value">{k.get('disponibles',0)}</div></div><div class="card kpi-card"><div class="label">Capital activo</div><div class="value">{_fmt_usd(k.get('capital',0))}</div></div><div class="card kpi-card"><div class="label">Ganancia esperada</div><div class="value">{_fmt_usd(k.get('ganancia_esperada',0))}</div></div><div class="card kpi-card"><div class="label">Cotizaciones</div><div class="value">{len(load_quotes())}</div></div></div></div></section>
<section class="slide" id="detalle-sistema"><div class="container"><div class="section-head"><div><div class="eyebrow">Detalle operativo</div><h2>Inventario actualizado</h2></div><p>Tabla generada automáticamente con los datos actuales del sistema.</p></div><div class="card table-card compact-table"><div class="table-scroll"><table><thead><tr><th>CV</th><th>Vehículo</th><th>Lote</th><th>Costo</th><th>Precio</th><th>Utilidad</th><th>Etapa</th><th>Comercial</th><th>Días</th></tr></thead><tbody>{''.join(rows)}</tbody></table></div></div></div></section>
"""
    if template and template.exists():
        text = template.read_text(encoding="utf-8", errors="ignore")
        text = re.sub(r"Corte\s+\d{1,2}/\d{1,2}/\d{4}", f"Corte {date.today().strftime('%d/%m/%Y')}", text)
        text = re.sub(r"\d+ unidades totales", f"{k.get('total',0)} unidades totales", text, count=1)
        text = re.sub(r"\d+ activas", f"{k.get('activos',0)} activas", text, count=1)
        text = re.sub(r"\d+ disponibles para venta", f"{k.get('disponibles',0)} disponibles para venta", text, count=1)
        if "</main>" in text:
            text = text.replace("</main>", dynamic_slide + "\n</main>")
    else:
        logo = ResourceManager.logo_data_uri()
        text = f"""<!doctype html><html lang="es"><head><meta charset="utf-8"><title>Reporte LYM</title><style>:root{{--navy:#08285a;--orange:#f59a13}}body{{font-family:Segoe UI,Arial;background:#f5f8fc;color:#0b172a;margin:0}}main{{max-width:1180px;margin:auto;padding:36px}}h1{{font-size:56px;color:var(--navy)}}.card{{background:white;border-radius:24px;padding:18px;margin:14px;box-shadow:0 18px 50px #08285a22}}.grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:16px}}table{{width:100%;border-collapse:collapse;background:white}}th{{background:var(--navy);color:white}}td,th{{padding:10px;border-bottom:1px solid #e5e7eb}}.logo img{{max-width:240px}}</style></head><body><main><div class="logo">{'<img src="'+logo+'">' if logo else ''}</div><h1>Reporte Gerencial L&M</h1>{dynamic_slide}</main></body></html>"""
    out = out_dir / f"LYM_Inversiones_Reporte_Gerencial_{datetime.now().strftime('%Y%m%d_%H%M')}.html"
    out.write_text(text, encoding="utf-8")
    log_audit("GENERAR_HTML_GERENCIAL_EXACTO", (user or {}).get("usuario", ""), "", out.name)
    return out


def generate_quotes_html_report(user: Optional[dict] = None) -> Optional[Path]:
    # Se mantiene por compatibilidad, pero ahora solo existe un HTML gerencial.
    return generate_html_report(load_vehicles(), user)


if PYSIDE_OK:
    class VehicleDetailDialog(QDialog):
        def __init__(self, parent, vehicle_id: str, user: dict, device: DeviceInfo):
            super().__init__(parent); self.vehicle_id=vehicle_id; self.user=user; self.device=device; self._cost_entries=[]
            self.setWindowTitle("Expediente del vehículo"); self.setMinimumSize(1280,840); self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint); self._build(); self.refresh()
        def _build(self):
            lay=QVBoxLayout(self)
            self.header=QLabel(); self.header.setStyleSheet("font-size:24px;font-weight:950;color:#08285a;"); lay.addWidget(self.header)
            self.timeline=TimelineWidget(); self.timeline.stageDoubleClicked.connect(self.open_stage_detail); lay.addWidget(self.timeline)
            btns=QHBoxLayout(); lay.addLayout(btns)
            self.next_btn=QPushButton("Pasar a siguiente etapa"); self.next_btn.setObjectName("orange"); self.next_btn.clicked.connect(self.advance_next)
            self.edit_btn=QPushButton("Editar etapa seleccionada"); self.edit_btn.setObjectName("ghost"); self.edit_btn.clicked.connect(self.edit_selected_stage)
            bdoc=QPushButton("Abrir comprobante compra"); bdoc.setObjectName("ghost"); bdoc.clicked.connect(self.open_purchase_doc)
            bfolder=QPushButton("Abrir carpeta del carro"); bfolder.setObjectName("ghost"); bfolder.clicked.connect(self.open_vehicle_folder)
            bexport=QPushButton("Descargar comprobantes y OC"); bexport.setObjectName("orange"); bexport.clicked.connect(self.export_docs)
            for b in [self.next_btn,self.edit_btn,bdoc,bfolder,bexport]: btns.addWidget(b)
            btns.addStretch(1)
            self.tabs=QTabWidget(); lay.addWidget(self.tabs)
            self.summary=QTextEdit(); self.summary.setReadOnly(True); self.tabs.addTab(self.summary,"Resumen")
            self.stage_table=QTableWidget(0,7); self.stage_table.setHorizontalHeaderLabels(["Etapa","Estado","Inicio","Fin","Días","Costo","Dato etapa"]); self.stage_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.stage_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.stage_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.stage_table.cellDoubleClicked.connect(lambda r,c:self.open_stage_detail(STAGE_ORDER[r] if r<len(STAGE_ORDER) else STAGE_COMPRADO)); self.tabs.addTab(self.stage_table,"Línea de tiempo / etapas")
            self.cost_table=QTableWidget(0,8); self.cost_table.setHorizontalHeaderLabels(["Etapa","Categoría","Descripción","Monto","Proveedor","OC","Comprobante","PDF OC"]); self.cost_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.cost_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.cost_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.cost_table.cellDoubleClicked.connect(self.open_cost_doc); self.tabs.addTab(self.cost_table,"Gastos y documentos")
            self.hist=QTextEdit(); self.hist.setReadOnly(True); self.tabs.addTab(self.hist,"Historial")
        def refresh(self):
            self.vehicle=find_vehicle(self.vehicle_id)
            if not self.vehicle: return
            v=self.vehicle; ensure_vehicle_runtime_fields(v)
            self.header.setText(f"{v.get('codigo')} · {_vehicle_display_name(v)} · {STAGE_META.get(v.get('estado_actual'),{}).get('label',v.get('estado_actual'))}")
            self.timeline.set_vehicle(v)
            nxt=next_stage_key(v.get("estado_actual", STAGE_COMPRADO))
            flujo_cerrado = (v.get("estado_comercial") == COMM_VENDIDO) or (v.get("estado_actual") == STAGE_DISPONIBLE)
            self.next_btn.setVisible(not flujo_cerrado and bool(nxt))
            self.next_btn.setEnabled(not flujo_cerrado and bool(nxt))
            self.next_btn.setText(f"Pasar a siguiente etapa: {STAGE_META[nxt]['label']}" if (nxt and not flujo_cerrado) else "Flujo operativo cerrado")
            self.edit_btn.setVisible(not flujo_cerrado)
            self.edit_btn.setEnabled(not flujo_cerrado)
            resumen=f"""Código: {v.get('codigo')}
Vehículo: {_vehicle_display_name(v)}
Millaje: {int(v.get('millaje') or 0):,}
Estado USA: {v.get('estado_usa')}
Subasta: {v.get('subasta')}
Lote: {v.get('lote')}
Precio ganado: {_fmt_usd(v.get('precio_ganado_usd'))}
Fecha compra: {_fmt_date(v.get('fecha_compra'))}
Usuario registró: {v.get('usuario_registro')}

Estado actual: {STAGE_META.get(v.get('estado_actual'),{}).get('label',v.get('estado_actual'))}
Estado comercial: {v.get('estado_comercial')}
Días desde compra: {vehicle_days_from_purchase(v)}
Días en etapa actual: {current_stage_days(v)}
Alerta etapa actual: {stage_alert_level(v)}

Costo total acumulado: {_fmt_usd(vehicle_total_cost(v))}
Precio venta publicado: {_fmt_usd(v.get('precio_venta_usd'))}
Precio mínimo permitido: {_fmt_usd(min_sale_price(v))}
Ganancia esperada: {_fmt_usd(vehicle_expected_profit(v))}
Precio venta real: {_fmt_usd(v.get('precio_venta_real_usd'))}
Cliente: {v.get('cliente','')}

Carpeta del carro:
{vehicle_document_folder(v, create=False) or ''}

Observaciones:
{v.get('observaciones','')}
"""
            self.summary.setPlainText(resumen)
            self.stage_table.setRowCount(len(STAGES))
            for r,s in enumerate(STAGES):
                st=vehicle_stage(v,s["key"]); vals=[s["label"],st.get("status"),_fmt_date(st.get("fecha_inicio")),_fmt_date(st.get("fecha_fin")),str(stage_duration_days(st)),_fmt_usd(st.get("costo_usd")),st.get("proveedor","")]
                for c,val in enumerate(vals): self.stage_table.setItem(r,c,QTableWidgetItem(str(val)))
            self._cost_entries=vehicle_document_entries(v); self.cost_table.setRowCount(len(self._cost_entries))
            for r,g in enumerate(self._cost_entries):
                vals=[g.get("etapa"),g.get("subcategoria") or g.get("categoria"),g.get("descripcion"),_fmt_usd(g.get("monto_usd")),g.get("proveedor",""),g.get("oc_numero",""),_doc_label(g.get("comprobante_nombre",""), g.get("comprobante","")),_doc_label(g.get("oc_documento_nombre",""), g.get("oc_documento",""))]
                for c,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.cost_table.setItem(r,c,it)
            hist="\n".join([f"[{h.get('fecha')}] {h.get('usuario')} · {h.get('accion')} · {h.get('detalle')}" for h in v.get('historial',[])])
            self.hist.setPlainText(hist)
        def advance_next(self):
            v=find_vehicle(self.vehicle_id); nxt=next_stage_key(v.get("estado_actual",STAGE_COMPRADO)) if v else None
            if not v:
                return
            if v.get("estado_comercial") == COMM_VENDIDO or v.get("estado_actual") == STAGE_DISPONIBLE:
                QMessageBox.information(self,"Etapas","El flujo operativo de este vehículo ya está cerrado."); return
            if not nxt: QMessageBox.information(self,"Etapas","El vehículo ya no tiene una etapa siguiente."); return
            dlg=StageUpdateDialog(self,v,nxt,self.user,self.device,mode="advance")
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def selected_stage_key(self) -> str:
            row=self.stage_table.currentRow(); return STAGE_ORDER[row] if 0 <= row < len(STAGE_ORDER) else self.vehicle.get("estado_actual", STAGE_COMPRADO)
        def edit_selected_stage(self):
            if self.vehicle.get("estado_comercial") == COMM_VENDIDO or self.vehicle.get("estado_actual") == STAGE_DISPONIBLE:
                QMessageBox.information(self,"Etapas","El flujo operativo de este vehículo ya está cerrado. Solo puede visualizar el historial."); return
            if not user_can_override_flow(self.user): QMessageBox.warning(self,"Permiso","Solo ADMIN/SUPERVISOR o usuarios con permiso especial pueden editar etapas anteriores o corregir el flujo."); return
            dlg=StageUpdateDialog(self,self.vehicle,self.selected_stage_key(),self.user,self.device,mode="edit")
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def open_stage_detail(self, stage_key: str):
            if self.vehicle.get("estado_comercial") != COMM_VENDIDO and self.vehicle.get("estado_actual") != STAGE_DISPONIBLE and stage_key == next_stage_key(self.vehicle.get("estado_actual",STAGE_COMPRADO)):
                self.advance_next(); return
            st=vehicle_stage(self.vehicle,stage_key)
            QMessageBox.information(self,"Detalle de etapa",f"Etapa: {STAGE_META[stage_key]['label']}\nEstado: {st.get('status')}\nInicio: {_fmt_date(st.get('fecha_inicio'))}\nFin: {_fmt_date(st.get('fecha_fin'))}\nDías: {stage_duration_days(st)}\nCosto: {_fmt_usd(st.get('costo_usd'))}\nDato: {st.get('proveedor','')}\n\nPara corregir usa 'Editar etapa seleccionada'.")
        def open_purchase_doc(self):
            st=vehicle_stage(self.vehicle,STAGE_COMPRADO); p=decrypt_file_to_temp(st.get("documento",""),st.get("documento_nombre","comprobante.pdf"))
            if p: QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
            else: QMessageBox.warning(self,"Documento","No se pudo abrir el comprobante.")
        def open_cost_doc(self,row:int,col:int):
            if row < 0 or row >= len(self._cost_entries): return
            g=self._cost_entries[row]
            if col == 6 and g.get("comprobante"):
                p=decrypt_file_to_temp(g.get("comprobante",""),g.get("comprobante_nombre","comprobante.pdf"))
                if p: QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
                else: QMessageBox.warning(self,"Documento","No se pudo abrir el comprobante.")
            elif col == 7 and g.get("oc_documento"):
                p=decrypt_file_to_temp(g.get("oc_documento",""),g.get("oc_documento_nombre","oc.pdf"))
                if p: QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
                else: QMessageBox.warning(self,"Documento","No se pudo abrir el PDF de OC.")
        def open_vehicle_folder(self):
            folder=vehicle_document_folder(self.vehicle, create=True)
            if folder: QDesktopServices.openUrl(QUrl.fromLocalFile(str(folder)))
        def export_docs(self):
            target=QFileDialog.getExistingDirectory(self,"Selecciona carpeta donde guardar comprobantes y OC",str(Path.home()))
            if not target: return
            ok,msg,paths=export_vehicle_documents(self.vehicle_id,Path(target)); QMessageBox.information(self,"Exportar documentos",msg) if ok else QMessageBox.warning(self,"Exportar documentos",msg)

    class QuoteDetailDialog(QDialog):
        def __init__(self,parent,quote_id:str,user:dict,device:DeviceInfo):
            super().__init__(parent); self.quote_id=quote_id; self.user=user; self.device=device
            self.setWindowTitle("Detalle de cotización"); self.setMinimumSize(980,720); self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint); self._build(); self.refresh()
        def _build(self):
            lay=QVBoxLayout(self); self.info=QTextEdit(); self.info.setReadOnly(True); lay.addWidget(self.info)
            row=QHBoxLayout(); lay.addLayout(row)
            for txt,fn,obj in [("Agregar seguimiento",self.add_follow,"orange"),("Editar / recalcular meses",self.edit_quote,"ghost"),("Generar propuesta",self.generate_proposal,"orange"),("Ofrecer otro carro",self.offer_other,"ghost"),("Marcar compra / vendido",self.mark_sold,"danger")]:
                b=QPushButton(txt); b.setObjectName(obj); b.clicked.connect(fn); row.addWidget(b)
            self.hist=QTextEdit(); self.hist.setReadOnly(True); lay.addWidget(self.hist)
        def refresh(self):
            q=find_quote(self.quote_id)
            if not q: return
            self.q=q; cl=q.get("cliente",{}); snap=q.get("vehicle_snapshot",{}); s=_quote_financial_summary(q)
            info=f"""Cliente: {cl.get('nombre','')}
Teléfono: {cl.get('telefono','')}
Correo: {cl.get('correo','')}
Medio: {cl.get('medio_contacto','')}

Vehículo cotizado: {_vehicle_display_name(snap)} · {q.get('vehicle_code')}
Valor del vehículo: {_fmt_usd(s['valor_vehiculo'])}
Prima requerida: {_fmt_usd(s['prima_requerida'])}
Monto leasing: {_fmt_usd(s['monto_leasing'])}
Plazo: {s['plazo']} meses
Tasa: {s['tasa']}%
Cuota final con IVA: {_fmt_usd(s['cuota_final'])}
Gastos legales: {_fmt_usd(s['legal'])}

Estado: {QUOTE_STATUS_LABELS.get(q.get('estado'),q.get('estado'))}
Color seguimiento: {quote_alert_level(q)}
Días sin compra/gestión: {quote_days_without_purchase(q)}

Carpeta del carro/propuestas:
{vehicle_proposals_folder(q.get('vehicle_code'), create=False) or ''}"""
            self.info.setPlainText(info)
            self.hist.setPlainText("\n".join([f"[{h.get('fecha')}] {h.get('usuario')} · {h.get('accion')} · {h.get('comentario')}" for h in q.get('seguimientos',[])]))
        def add_follow(self):
            estados=list(QUOTE_STATUS_LABELS.keys()); estado,ok=QInputDialog.getItem(self,"Estado","Nuevo estado:",estados,0,False)
            if not ok: return
            com,ok=QInputDialog.getMultiLineText(self,"Seguimiento","Comentario:")
            if not ok: return
            ok,msg=add_quote_followup(self.quote_id,com,estado,self.user); QMessageBox.information(self,"Seguimiento",msg) if ok else QMessageBox.warning(self,"Seguimiento",msg); self.refresh()
        def edit_quote(self):
            dlg=QuoteEditorDialog(self,self.user,self.device,quote_id=self.quote_id)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def generate_proposal(self):
            formato,ok=QInputDialog.getItem(self,"Generar propuesta","Seleccione el formato a generar:",["PDF","Excel","PDF + Excel"],0,False)
            if not ok: return
            ok,msg,paths=generate_quote_proposal_selected(self.quote_id, formato, self.user)
            if not ok:
                QMessageBox.warning(self,"Propuesta",msg); return
            guardar=QMessageBox.question(self,"Propuesta generada",msg+"\n\n¿Deseas guardar una copia en otra ubicación?",QMessageBox.StandardButton.Yes|QMessageBox.StandardButton.No)
            if guardar==QMessageBox.StandardButton.Yes:
                if len(paths)>1:
                    folder=QFileDialog.getExistingDirectory(self,"Selecciona carpeta para guardar la copia",str(Path.home()))
                    if folder:
                        for p in paths: shutil.copy2(p, Path(folder)/p.name)
                        QMessageBox.information(self,"Copia",f"Copia guardada en:\n{folder}")
                else:
                    p=paths[0]; target,_=QFileDialog.getSaveFileName(self,"Guardar copia",str(Path.home()/p.name),"Archivos (*"+p.suffix+")")
                    if target:
                        shutil.copy2(p, target); QMessageBox.information(self,"Copia",f"Copia guardada en:\n{target}")
                return
            msgbox=QMessageBox(self); msgbox.setWindowTitle("Abrir propuesta"); msgbox.setText("¿Qué deseas hacer ahora?")
            pdf_btn=excel_btn=None
            for p in paths:
                if p.suffix.lower()==".pdf" and pdf_btn is None: pdf_btn=msgbox.addButton("Abrir PDF",QMessageBox.ButtonRole.AcceptRole)
                if p.suffix.lower() in (".xlsx",".xlsm") and excel_btn is None: excel_btn=msgbox.addButton("Abrir Excel",QMessageBox.ButtonRole.AcceptRole)
            close_btn=msgbox.addButton("Cerrar",QMessageBox.ButtonRole.RejectRole); msgbox.exec(); clicked=msgbox.clickedButton()
            if clicked==pdf_btn:
                for p in paths:
                    if p.suffix.lower()==".pdf": QDesktopServices.openUrl(QUrl.fromLocalFile(str(p))); break
            elif clicked==excel_btn:
                for p in paths:
                    if p.suffix.lower() in (".xlsx",".xlsm"): QDesktopServices.openUrl(QUrl.fromLocalFile(str(p))); break
        def offer_other(self):
            cl=self.q.get("cliente",{})
            dlg=QuoteEditorDialog(self,self.user,self.device,preset_client=cl)
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def mark_sold(self):
            if QMessageBox.question(self,"Venta","¿Marcar este carro como vendido a este cliente?") != QMessageBox.StandardButton.Yes: return
            ok,msg=mark_quote_won_and_vehicle_sold(self.quote_id,self.user,self.device); QMessageBox.information(self,"Venta",msg) if ok else QMessageBox.warning(self,"Venta",msg); self.refresh()

    class ReporteriaPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main=main; self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Reportería", "Reportes oficiales L&M con formatos originales: inventario, cotizaciones, costos y un único HTML gerencial."))
            grid=QGridLayout(); lay.addLayout(grid)
            items=[
                ("Excel inventario L&M", "Usa exactamente el formato Inventario_vehiculos_LYM_JUNIO_2026 ac.xlsx.", self.report_excel_inventory),
                ("Excel costos por vehículo", "Usa exactamente el formato COPART INC - COSTOS DE VEHICULOS.", self.report_excel_costs),
                ("Excel control de cotizaciones", "Usa exactamente el formato Control_Cotizaciones_LYM, con clientes y seguimiento.", self.report_excel_quotes),
                ("HTML reporte gerencial", "Único reporte HTML con el diseño LYM_Inversiones_Reporte_Junio_2026_FINAL_LIMPIO.", self.report_html_inventory),
            ]
            for i,(title,desc,fn) in enumerate(items):
                card=QFrame(); card.setStyleSheet("QFrame{background:white;border:1px solid #d9e2ef;border-radius:18px;}"); cl=QVBoxLayout(card)
                lab=QLabel(title); lab.setStyleSheet("font-size:16pt;font-weight:900;color:#08285a;"); d=QLabel(desc); d.setWordWrap(True); d.setStyleSheet("color:#64748b;")
                b=QPushButton("Generar"); b.setObjectName("orange"); b.clicked.connect(fn)
                cl.addWidget(lab); cl.addWidget(d); cl.addStretch(1); cl.addWidget(b); grid.addWidget(card,i//2,i%2)
            lay.addStretch(1)
        def _after_report(self,path:Path,title:str):
            msg=QMessageBox(self); msg.setWindowTitle(title); msg.setText(f"Reporte generado correctamente:\n{path.name}\n\n¿Deseas guardar una copia?")
            copy_btn=msg.addButton("Guardar una copia…",QMessageBox.ButtonRole.ActionRole); open_btn=msg.addButton("Abrir",QMessageBox.ButtonRole.AcceptRole); msg.addButton("Cerrar",QMessageBox.ButtonRole.RejectRole); msg.exec(); clicked=msg.clickedButton()
            if clicked==copy_btn:
                target,_=QFileDialog.getSaveFileName(self,"Guardar copia",str(Path.home()/path.name),"Archivos (*"+path.suffix+")")
                if target:
                    ok,m=copy_report_to(path,Path(target)); QMessageBox.information(self,"Copia",m) if ok else QMessageBox.warning(self,"Copia",m)
            elif clicked==open_btn:
                QDesktopServices.openUrl(QUrl.fromLocalFile(str(path)))
        def report_excel_inventory(self):
            ok,msg,out=generate_inventory_excel(load_vehicles(),self.main.user); QMessageBox.warning(self,"Excel",msg) if not ok or not out else self._after_report(out,"Excel inventario L&M")
        def report_excel_costs(self):
            ok,msg,out=generate_vehicle_costs_excel_report(load_vehicles(),self.main.user); QMessageBox.warning(self,"Excel costos",msg) if not ok or not out else self._after_report(out,"Excel costos por vehículo")
        def report_excel_quotes(self):
            ok,msg,out=generate_quotes_excel_report(self.main.user); QMessageBox.warning(self,"Excel cotizaciones",msg) if not ok or not out else self._after_report(out,"Excel control de cotizaciones")
        def report_html_inventory(self):
            out=generate_html_report(load_vehicles(),self.main.user); self._after_report(out,"HTML reporte gerencial") if out else QMessageBox.warning(self,"HTML","No se pudo generar el reporte HTML.")




# =============================================================================
# AJUSTES V4.4 - FECHAS INTELIGENTES, DOCUMENTOS LEGALES, SIDEBAR Y TRACKING
# =============================================================================
APP_VERSION = "2.2.0_LEASING"
ALL_SUBFOLDERS = [SUB_DATOS, SUB_DOCUMENTOS, SUB_FOTOS, SUB_REPORTES, SUB_RESPALDOS, SUB_TEMP]


def _stage_completion_date(vehicle: dict, stage_key: str) -> Optional[date]:
    st = vehicle_stage(vehicle, stage_key)
    extra = st.get("extra") if isinstance(st.get("extra"), dict) else {}
    candidates = []
    if stage_key == STAGE_TRASLADO_USA:
        candidates = [extra.get("fecha_llegada_yarda"), st.get("fecha_fin")]
    elif stage_key == STAGE_TRANSITO:
        candidates = [extra.get("fecha_salida_naviera"), st.get("fecha_fin")]
    elif stage_key == STAGE_ADUANA:
        candidates = [extra.get("fecha_liberacion_aduana"), st.get("fecha_fin")]
    elif stage_key == STAGE_PREPARACION:
        candidates = [extra.get("legalizacion_fin"), extra.get("taller_salida"), st.get("fecha_fin")]
    else:
        candidates = [st.get("fecha_fin"), st.get("fecha_inicio")]
    parsed = [_parse_date(x) for x in candidates if x]
    parsed = [x for x in parsed if x]
    return max(parsed) if parsed else None


def _stage_min_start_date(vehicle: dict, stage_key: str) -> str:
    fc = _parse_date(vehicle.get("fecha_compra")) or date.today()
    idx = stage_index(stage_key)
    if idx > 0:
        prev_key = STAGE_ORDER[idx - 1]
        prev_done = _stage_completion_date(vehicle, prev_key)
        if prev_done:
            return prev_done.isoformat()
        prev_st = vehicle_stage(vehicle, prev_key)
        prev_start = _parse_date(prev_st.get("fecha_inicio"))
        if prev_start:
            return prev_start.isoformat()
    return fc.isoformat()


def _stage_default_start_date(vehicle: dict, stage_key: str) -> str:
    idx = stage_index(stage_key)
    if idx > 0:
        prev_key = STAGE_ORDER[idx - 1]
        prev_done = _stage_completion_date(vehicle, prev_key)
        if prev_done:
            return prev_done.isoformat()
    fc = _parse_date(vehicle.get("fecha_compra")) or date.today()
    return fc.isoformat() if stage_key == STAGE_TRASLADO_USA else ""


def validate_stage_dates(vehicle: dict, stage_key: str, fecha_inicio: date, fecha_fin: Optional[date], user: dict) -> tuple[bool, str]:
    fc = _parse_date(vehicle.get("fecha_compra")) or date.today()
    if fecha_inicio > date.today():
        return False, "No se permiten fechas futuras."
    if fecha_fin and fecha_fin > date.today():
        return False, "No se permiten fechas futuras."
    if fecha_inicio < fc:
        return False, "La fecha de la etapa no puede ser anterior a la fecha de compra."
    if fecha_fin and fecha_fin < fecha_inicio:
        return False, "La fecha final no puede ser anterior a la fecha de inicio."
    idx = stage_index(stage_key)
    if idx > 0:
        prev_key = STAGE_ORDER[idx - 1]
        prev_gate = _stage_completion_date(vehicle, prev_key)
        if prev_gate and fecha_inicio < prev_gate:
            return False, f"La etapa {STAGE_META[stage_key]['label']} no puede iniciar antes del fin lógico de {STAGE_META[prev_key]['label']} ({_fmt_date(prev_gate)})."
        prev_start = _parse_date(vehicle_stage(vehicle, prev_key).get("fecha_inicio"))
        if prev_start and fecha_inicio < prev_start:
            return False, f"La etapa {STAGE_META[stage_key]['label']} no puede iniciar antes de {STAGE_META[prev_key]['label']} ({_fmt_date(prev_start)})."
    return True, "OK"


def add_vehicle_legal_paper(vehicle_id: str, tipo: str, descripcion: str, src_path: str, user: dict, device: DeviceInfo) -> tuple[bool, str]:
    v = find_vehicle(vehicle_id)
    if not v:
        return False, "Vehículo no encontrado."
    ensure_vehicle_runtime_fields(v)
    p = Path(src_path or "")
    if not p.exists():
        return False, "Debes seleccionar un archivo válido."
    tipo_n = _norm(tipo or "PAPEL_LEGAL")
    rel, name = store_document_named(p, v.get("codigo", "VEHICULO"), "PAPEL_LEGAL", f"{v.get('codigo')}_PAPEL_LEGAL_{tipo_n}_{datetime.now().strftime('%Y%m%d_%H%M%S')}", p.name)
    if not rel:
        return False, "No se pudo guardar el papel legal."
    item = {"id": uuid.uuid4().hex, "tipo": tipo_n, "descripcion": descripcion.strip(), "documento": rel, "documento_nombre": name, "fecha": _today_iso(), "usuario": user.get("usuario", ""), "fecha_registro": _now_iso()}
    v.setdefault("papeles_legales", []).append(item)
    v.setdefault("historial", []).append({"fecha": _now_iso(), "usuario": user.get("usuario", ""), "computadora": device.computer_name, "accion": "SUBIR_PAPEL_LEGAL", "detalle": f"{tipo_n} · {name}"})
    if save_vehicle(v):
        return True, "Papel legal guardado correctamente."
    return False, "No se pudo guardar el vehículo."


_lym_v43_vehicle_document_entries = vehicle_document_entries

def vehicle_document_entries(vehicle: dict) -> list[dict]:
    entries = _lym_v43_vehicle_document_entries(vehicle)
    for p in vehicle.get("papeles_legales", []) or []:
        entries.append({
            "stage_key": "PAPELES_LEGALES",
            "etapa": "Papeles legales",
            "categoria": "LEGAL",
            "subcategoria": p.get("tipo", "PAPEL_LEGAL"),
            "descripcion": p.get("descripcion", "Papel legal del vehículo"),
            "monto_usd": 0,
            "proveedor": "",
            "oc_numero": "",
            "comprobante": p.get("documento", ""),
            "comprobante_nombre": p.get("documento_nombre", "papel_legal.pdf"),
            "oc_documento": "",
            "oc_documento_nombre": "",
            "legal_id": p.get("id", ""),
        })
    return entries


if PYSIDE_OK:
    class CostBreakdownDialog(QDialog):
        def __init__(self, parent, vehicle: dict):
            super().__init__(parent)
            self.vehicle = vehicle
            self._entries = vehicle_document_entries(vehicle)
            self.setWindowTitle(f"Visualizar gastos y documentos · {vehicle.get('codigo','')}")
            self.setMinimumSize(1120, 680)
            self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint)
            lay = QVBoxLayout(self)
            head = QLabel(f"<b>{vehicle.get('codigo')}</b> · {_vehicle_display_name(vehicle)} · Total gastos: <b>{_fmt_usd(vehicle_total_cost(vehicle))}</b><br><span style='color:#64748b'>Doble click en Comprobante u OC PDF para abrirlo.</span>")
            head.setTextFormat(Qt.TextFormat.RichText); head.setStyleSheet("font-size:18px;color:#08285a;padding:8px;")
            lay.addWidget(head)
            self.table = QTableWidget(0, 9)
            self.table.setHorizontalHeaderLabels(["Etapa","Categoría","Subcategoría","Descripción","Fecha","Proveedor","Monto","Comprobante","OC PDF"])
            self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
            self.table.cellDoubleClicked.connect(self.open_doc)
            lay.addWidget(self.table)
            self.refresh()
            row=QHBoxLayout(); bex=QPushButton("Descargar toda la documentación"); bex.setObjectName("orange"); bex.clicked.connect(self.export_all); bclose=QPushButton("Cerrar"); bclose.setObjectName("ghost"); bclose.clicked.connect(self.accept); row.addWidget(bex); row.addStretch(1); row.addWidget(bclose); lay.addLayout(row)
        def refresh(self):
            self._entries = vehicle_document_entries(self.vehicle)
            self.table.setRowCount(len(self._entries))
            for r,g in enumerate(self._entries):
                vals=[g.get("etapa"), g.get("categoria"), g.get("subcategoria"), g.get("descripcion"), _fmt_date(g.get("fecha")), g.get("proveedor",""), _fmt_usd(g.get("monto_usd")), _doc_label(g.get("comprobante_nombre",""), g.get("comprobante","")), _doc_label(g.get("oc_documento_nombre",""), g.get("oc_documento",""))]
                for c,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.table.setItem(r,c,it)
        def open_doc(self,row:int,col:int):
            if row < 0 or row >= len(self._entries): return
            g=self._entries[row]
            if col == 7 and g.get("comprobante"):
                p=decrypt_file_to_temp(g.get("comprobante",""),g.get("comprobante_nombre","documento.pdf"))
                if p: QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
                else: QMessageBox.warning(self,"Documento","No se pudo abrir el comprobante/documento.")
            elif col == 8 and g.get("oc_documento"):
                p=decrypt_file_to_temp(g.get("oc_documento",""),g.get("oc_documento_nombre","oc.pdf"))
                if p: QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
                else: QMessageBox.warning(self,"Documento","No se pudo abrir el PDF de OC.")
        def export_all(self):
            target=QFileDialog.getExistingDirectory(self,"Selecciona carpeta donde guardar comprobantes y OC",str(Path.home()))
            if not target: return
            ok,msg,paths=export_vehicle_documents(self.vehicle.get("id"),Path(target))
            QMessageBox.information(self,"Exportar documentos",msg) if ok else QMessageBox.warning(self,"Exportar documentos",msg)

    _LYM_V43_VehicleDetailDialog = VehicleDetailDialog
    class VehicleDetailDialog(_LYM_V43_VehicleDetailDialog):
        def _build(self):
            super()._build()
            self.legal_tab = QWidget(); ll=QVBoxLayout(self.legal_tab)
            top=QHBoxLayout(); badd=QPushButton("Subir papel legal"); badd.setObjectName("orange"); badd.clicked.connect(self.add_legal_paper); bex=QPushButton("Descargar comprobantes, OC y papeles legales"); bex.setObjectName("ghost"); bex.clicked.connect(self.export_docs); top.addWidget(badd); top.addWidget(bex); top.addStretch(1); ll.addLayout(top)
            self.legal_table=QTableWidget(0,4); self.legal_table.setHorizontalHeaderLabels(["Tipo","Descripción","Fecha","Documento"]); self.legal_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.legal_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); self.legal_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); self.legal_table.cellDoubleClicked.connect(self.open_legal_doc); ll.addWidget(self.legal_table)
            self.tabs.addTab(self.legal_tab,"Papeles legales")
        def refresh(self):
            super().refresh()
            v=getattr(self,"vehicle",None)
            if not v: return
            self.header.setText(f"{v.get('codigo')} · {_vehicle_display_name(v)} · {STAGE_META.get(v.get('estado_actual'),{}).get('label',v.get('estado_actual'))} · Días desde compra: {vehicle_days_from_purchase(v)}")
            papers=v.get("papeles_legales",[]) or []
            self.legal_table.setRowCount(len(papers))
            for r,p in enumerate(papers):
                vals=[p.get("tipo",""),p.get("descripcion",""),_fmt_date(p.get("fecha")),_doc_label(p.get("documento_nombre",""),p.get("documento",""))]
                for c,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.legal_table.setItem(r,c,it)
        def advance_next(self):
            v=find_vehicle(self.vehicle_id); current=v.get("estado_actual",STAGE_COMPRADO) if v else STAGE_COMPRADO
            ok,msg=can_advance_from_stage(v,current) if v else (False,"Vehículo no encontrado.")
            if not ok:
                QMessageBox.warning(self,"No se puede avanzar",msg); return
            nxt=next_stage_key(current)
            if not nxt: QMessageBox.information(self,"Etapas","El vehículo ya no tiene una etapa siguiente."); return
            dlg=StageUpdateDialog(self,v,nxt,self.user,self.device,mode="advance")
            if dlg.exec()==QDialog.DialogCode.Accepted: self.refresh()
        def add_legal_paper(self):
            tipos=["TARJETA_CIRCULACION","PLACAS","EMISION_GASES","EXPERTICIA","CONTRATO","DUCA","OTRO"]
            tipo,ok=QInputDialog.getItem(self,"Papel legal","Tipo de documento:",tipos,0,False)
            if not ok: return
            desc,ok=QInputDialog.getText(self,"Papel legal","Descripción breve:")
            if not ok: return
            path,_=QFileDialog.getOpenFileName(self,"Selecciona papel legal",str(Path.home()),"Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if not path: return
            ok,msg=add_vehicle_legal_paper(self.vehicle_id,tipo,desc,path,self.user,self.device)
            QMessageBox.information(self,"Papel legal",msg) if ok else QMessageBox.warning(self,"Papel legal",msg)
            self.refresh()
        def open_legal_doc(self,row:int,col:int):
            if col != 3: return
            papers=(self.vehicle or {}).get("papeles_legales",[]) or []
            if row < 0 or row >= len(papers): return
            p=decrypt_file_to_temp(papers[row].get("documento",""),papers[row].get("documento_nombre","papel_legal.pdf"))
            if p: QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
            else: QMessageBox.warning(self,"Documento","No se pudo abrir el papel legal.")

    class MainWindow(QMainWindow):
        def __init__(self, user: dict, device: DeviceInfo):
            super().__init__(); self.user=user; self.device=device; self.setWindowTitle(f"{APP_NAME} · {user.get('usuario')} · v{APP_VERSION}"); self.resize(1480,900); self.setWindowIcon(QIcon(str(ResourceManager.find_logo() or ""))); self._build(); self.refresh_all()
        def _build(self):
            central=QWidget(); self.setCentralWidget(central); main=QHBoxLayout(central); main.setContentsMargins(0,0,0,0)
            side=QFrame(); side.setFixedWidth(265); side.setStyleSheet("QFrame{background:#08285a;} QLabel{color:white;} QPushButton{background:transparent;color:white;text-align:left;padding:11px 14px;border-radius:0;font-weight:850;} QPushButton:hover{background:#0e3a78;border-left:5px solid #f59a13;}")
            sl=QVBoxLayout(side); sl.setContentsMargins(14,20,14,16)
            logo=QLabel(); logo.setAlignment(Qt.AlignmentFlag.AlignCenter); pix=QPixmap(str(ResourceManager.find_logo() or ""))
            if not pix.isNull(): logo.setPixmap(pix.scaled(145,145,Qt.AspectRatioMode.KeepAspectRatio,Qt.TransformationMode.SmoothTransformation))
            else: logo.setText("L&M"); logo.setStyleSheet("font-size:30px;font-weight:950;color:#f59a13;")
            sl.addWidget(logo); title=QLabel("LYM AUTO CONTROL"); title.setAlignment(Qt.AlignmentFlag.AlignCenter); title.setStyleSheet("font-weight:950;color:#ffcf7a;font-size:14pt;"); sub=QLabel(f"Rol: {self.user.get('rol')}"); sub.setAlignment(Qt.AlignmentFlag.AlignCenter); sub.setStyleSheet("font-weight:700;color:#d7e6ff;font-size:10pt;"); sl.addWidget(title); sl.addWidget(sub)
            self.stack=QStackedWidget(); self.pages=[]; self._section_widgets=[]
            self.dashboard=DashboardPage(self); self.purchase=PurchasePage(self); self.inventory=InventoryPage(self); self.cotizaciones=CotizacionesPage(self); self.reporteria=ReporteriaPage(self); self.catalogos=CatalogosPage(self)
            def add_page(container_layout, name, widget):
                idx=self.stack.addWidget(widget); self.pages.append(widget); b=QPushButton(name); b.clicked.connect(lambda _,i=idx:self.stack.setCurrentIndex(i)); container_layout.addWidget(b); return b
            add_page(sl,"🏠  Inicio",self.dashboard)
            def collapsible_section(title, entries, opened=True):
                header=QPushButton(("▾  " if opened else "▸  ")+title); header.setStyleSheet("color:#ffcf7a;font-weight:950;margin-top:10px;border-bottom:1px solid rgba(255,255,255,.14);")
                box=QWidget(); bl=QVBoxLayout(box); bl.setContentsMargins(0,0,0,0); bl.setSpacing(0)
                for name,widget in entries: add_page(bl,name,widget)
                box.setVisible(opened)
                def toggle():
                    box.setVisible(not box.isVisible()); header.setText(("▾  " if box.isVisible() else "▸  ")+title)
                header.clicked.connect(toggle); sl.addWidget(header); sl.addWidget(box); self._section_widgets.append((header,box))
            collapsible_section("COMPRA VEHICULAR", [("➕  Nueva compra",self.purchase),("🚗  Inventario / CV",self.inventory)], True)
            collapsible_section("COMERCIAL", [("💬  Cotizaciones",self.cotizaciones)], True)
            collapsible_section("REPORTERÍA", [("📊  Reportería",self.reporteria),("📚  Catálogos",self.catalogos)], True)
            if user_has_permission(self.user, PERM_CONFIG):
                bconf=QPushButton("⚙️  Configuración"); bconf.clicked.connect(self.open_config); sl.addWidget(bconf)
            sl.addStretch(1); bclose=QPushButton("🚪  Cerrar"); bclose.clicked.connect(self.close); sl.addWidget(bclose); main.addWidget(side); main.addWidget(self.stack,1); self.setStatusBar(QStatusBar()); self.statusBar().showMessage(f"Usuario: {self.user.get('usuario')} · Carpeta: {get_data_folder()}")
        def refresh_all(self):
            bootstrap_system(); self.dashboard.refresh(); self.inventory.refresh(); self.purchase.refresh_catalogs(); self.cotizaciones.refresh(); self.catalogos.refresh_all()
        def open_vehicle_detail(self, vehicle_id: str):
            dlg=VehicleDetailDialog(self,vehicle_id,self.user,self.device); dlg.exec(); self.refresh_all()
        def open_config(self):
            dlg=ConfigDialog(self,self.user,self.device); dlg.exec(); self.refresh_all()


# =============================================================================
# AJUSTES V4.5 - GANANCIA REAL, PRECIO EDITABLE, VENTA CON REGALÍA Y REPORTERÍA DESDE CERO
# =============================================================================
APP_VERSION = "2.3.0_LEASING"

LYM_NAVY = "08285A"
LYM_NAVY2 = "0E3A78"
LYM_ORANGE = "F59A13"
LYM_GOLD = "FFC000"
LYM_GREEN = "10B981"
LYM_RED = "DC2626"
LYM_AMBER = "F59E0B"
LYM_BLUE = "2563EB"
LYM_LIGHT = "EFF4F9"
LYM_BORDER = "B7C9E2"


def calculate_price_final_breakdown(costo_total: float, margen_pct: float = 30.0, iva_duca_usd: float = 0.0, iva_pct: float = 13.0,
                                    pago_cuenta_pct: float = 1.75, vts_usd: float = 2.07,
                                    precio_cliente_manual_usd: float = 0.0, precio_minimo_manual_usd: float = 0.0,
                                    regalia_usd: float = 0.0) -> dict:
    """Cálculo oficial L&M para precio final y ganancia real estimada.

    Regla corregida: el cliente paga el precio final completo; internamente se resta solo el IVA a pagar
    después de aplicar el IVA DUCA, pago a cuenta, VTS y cualquier regalía. Luego se compara contra costo.
    """
    costo = round(_to_float(costo_total, 0), 2)
    margen_pct = round(_to_float(margen_pct, 0), 4)
    iva_pct = round(_to_float(iva_pct, 13), 4)
    pago_cuenta_pct = round(_to_float(pago_cuenta_pct, 1.75), 4)
    iva_duca = round(max(0.0, _to_float(iva_duca_usd, 0)), 2)
    vts = round(max(0.0, _to_float(vts_usd, 2.07)), 2)
    regalia = round(max(0.0, _to_float(regalia_usd, 0)), 2)
    suggested_venta_neta = round(costo * (1 + margen_pct / 100.0), 2)
    suggested_iva = round(suggested_venta_neta * (iva_pct / 100.0), 2)
    suggested_price = round(suggested_venta_neta + suggested_iva + vts, 2)
    manual = round(_to_float(precio_cliente_manual_usd, 0), 2)
    if manual > 0:
        precio_cliente = manual
        venta_neta = round(max(0.0, (precio_cliente - vts) / (1 + iva_pct / 100.0)), 2) if iva_pct > -100 else 0.0
        price_source = "MANUAL_REDONDEADO"
    else:
        precio_cliente = suggested_price
        venta_neta = suggested_venta_neta
        price_source = "MARGEN_DESEADO"
    iva_venta = round(venta_neta * (iva_pct / 100.0), 2)
    iva_pagar = round(max(0.0, iva_venta - iva_duca), 2)
    pago_cuenta = round(venta_neta * (pago_cuenta_pct / 100.0), 2)
    utilidad_bruta = round(venta_neta - costo, 2)
    total_descuentos = round(iva_pagar + pago_cuenta + vts + regalia, 2)
    ingreso_neto_despues_impuestos = round(precio_cliente - iva_pagar - pago_cuenta - vts - regalia, 2)
    ganancia_real = round(ingreso_neto_despues_impuestos - costo, 2)
    margen_bruto_pct = round((utilidad_bruta / costo * 100.0), 2) if costo else 0.0
    margen_real_pct = round((ganancia_real / costo * 100.0), 2) if costo else 0.0
    precio_minimo_manual = round(_to_float(precio_minimo_manual_usd, 0), 2)
    precio_minimo = precio_minimo_manual if precio_minimo_manual > 0 else precio_cliente
    return {
        "costo_total_usd": costo,
        "margen_pct": margen_pct,
        "margen_bruto_pct": margen_bruto_pct,
        "venta_neta_usd": venta_neta,
        "venta_neta_sugerida_usd": suggested_venta_neta,
        "precio_sugerido_por_margen_usd": suggested_price,
        "iva_pct": iva_pct,
        "iva_venta_usd": iva_venta,
        "iva_duca_usd": iva_duca,
        "iva_pagar_usd": iva_pagar,
        "pago_cuenta_pct": pago_cuenta_pct,
        "pago_cuenta_usd": pago_cuenta,
        "vts_usd": vts,
        "regalia_usd": regalia,
        "total_descuentos_internos_usd": total_descuentos,
        "ingreso_neto_despues_impuestos_usd": ingreso_neto_despues_impuestos,
        "precio_venta_cliente_usd": precio_cliente,
        "precio_redondeado_manual_usd": manual,
        "precio_minimo_usd": precio_minimo,
        "utilidad_bruta_usd": utilidad_bruta,
        "ganancia_real_estimada_usd": ganancia_real,
        "utilidad_neta_estimada_usd": ganancia_real,
        "margen_real_estimado_pct": margen_real_pct,
        "price_source": price_source,
    }


def vehicle_profit_summary(vehicle: dict, sale_price: Optional[float] = None, regalia_usd: float = 0.0) -> dict:
    ensure_vehicle_runtime_fields(vehicle)
    costo = vehicle_total_cost(vehicle)
    pf = vehicle.get("precio_final") or {}
    price = _to_float(sale_price, 0) or _to_float(vehicle.get("precio_venta_usd"), 0) or _to_float(pf.get("precio_venta_cliente_usd"), 0)
    return calculate_price_final_breakdown(
        costo,
        _to_float(pf.get("margen_pct"), 30),
        _to_float(pf.get("iva_duca_usd"), 0),
        _to_float(pf.get("iva_pct"), 13),
        _to_float(pf.get("pago_cuenta_pct"), 1.75),
        _to_float(pf.get("vts_usd"), 2.07),
        price,
        _to_float(pf.get("precio_minimo_usd"), 0),
        regalia_usd,
    )


def vehicle_expected_profit(vehicle: dict) -> float:
    try:
        vd = vehicle.get("venta_detalle") or {}
        if vd.get("ganancia_real_final_usd") not in (None, ""):
            return round(_to_float(vd.get("ganancia_real_final_usd"), 0), 2)
        pf = vehicle.get("precio_final") or {}
        if pf.get("ganancia_real_estimada_usd") not in (None, ""):
            return round(_to_float(pf.get("ganancia_real_estimada_usd"), 0), 2)
        return round(vehicle_profit_summary(vehicle).get("ganancia_real_estimada_usd", 0), 2)
    except Exception:
        return 0.0


def register_vehicle_price_change(vehicle: dict, old_price: float, new_price: float, user: Optional[dict], device: Optional[DeviceInfo] = None) -> None:
    old_price = round(_to_float(old_price, 0), 2)
    new_price = round(_to_float(new_price, 0), 2)
    if abs(old_price - new_price) < 0.01:
        return
    now = _now_iso()
    vehicle.setdefault("precio_historial", []).append({
        "fecha": now,
        "usuario": (user or {}).get("usuario", ""),
        "precio_anterior_usd": old_price,
        "precio_nuevo_usd": new_price,
        "comentario": f"Precio actualizado de {_fmt_usd(old_price)} a {_fmt_usd(new_price)}",
    })
    vehicle.setdefault("historial", []).append({
        "fecha": now,
        "usuario": (user or {}).get("usuario", ""),
        "computadora": getattr(device, "computer_name", ""),
        "accion": "CAMBIO_PRECIO_CLIENTE",
        "detalle": f"Precio anterior {_fmt_usd(old_price)} · nuevo {_fmt_usd(new_price)}",
    })


def sync_active_quotes_after_price_change(vehicle: dict, old_price: float, new_price: float, user: Optional[dict]) -> int:
    if abs(round(_to_float(old_price, 0), 2) - round(_to_float(new_price, 0), 2)) < 0.01:
        return 0
    quotes = load_quotes()
    changed = 0
    now = _now_iso()
    notice = f"El precio del vehículo cambió de {_fmt_usd(old_price)} a {_fmt_usd(new_price)} el {_fmt_date(date.today())}. Recalcula o avisa al cliente antes de enviar propuesta."
    for q in quotes:
        if q.get("vehicle_id") != vehicle.get("id"):
            continue
        if q.get("estado") in (QUOTE_GANADA, QUOTE_PERDIDA):
            continue
        le = q.get("leasing", {}) or {}
        prima_pct = _to_float(le.get("prima_pct"), 20)
        comision = _to_float(le.get("comision_usd"), 100)
        pago = round(new_price * prima_pct / 100.0 + comision, 2)
        calc = calculate_leasing(new_price, _to_float(le.get("ingreso_cliente"), 0), pago, int(_to_float(le.get("plazo_meses"), 60)), _to_float(le.get("tasa_mensual_pct"), 2.5), _to_float(le.get("seguro_mensual"), 80), _to_float(le.get("gps_mensual"), 20), _to_float(le.get("iva_pct"), 13))
        q["leasing"] = {**calc, "prima_pct": prima_pct, "comision_usd": comision}
        q["legal"] = calculate_legal_fees(new_price)
        snap = q.get("vehicle_snapshot", {}) or {}
        snap["precio_venta_usd"] = new_price
        q["vehicle_snapshot"] = snap
        q["ultimo_aviso_precio"] = notice
        q.setdefault("avisos_precio", []).append({"fecha": now, "precio_anterior_usd": old_price, "precio_nuevo_usd": new_price, "comentario": notice})
        q.setdefault("seguimientos", []).append({"fecha": now, "usuario": (user or {}).get("usuario", ""), "accion": "PRECIO_ACTUALIZADO", "comentario": notice})
        q["fecha_actualizacion"] = now
        changed += 1
    if changed:
        save_quotes(quotes)
    return changed


def _quote_status_counts(quotes: Optional[list[dict]] = None) -> dict:
    quotes = quotes if quotes is not None else load_quotes()
    counts = {"VERDE": 0, "AMARILLO": 0, "ROJO": 0, "GANADAS": 0, "REOFERTAR": 0, "PERDIDAS": 0}
    for q in quotes:
        st = q.get("estado")
        if st == QUOTE_GANADA:
            counts["GANADAS"] += 1
        elif st == QUOTE_REOFERTAR:
            counts["REOFERTAR"] += 1
            counts[quote_alert_level(q)] = counts.get(quote_alert_level(q), 0) + 1
        elif st == QUOTE_PERDIDA:
            counts["PERDIDAS"] += 1
        else:
            counts[quote_alert_level(q)] = counts.get(quote_alert_level(q), 0) + 1
    counts["TOTAL"] = len(quotes)
    return counts


def _report_month_key(value: Any) -> str:
    d = _parse_date(value)
    if not d:
        return "Sin fecha"
    return f"{d.year}-{d.month:02d}"


def _counter_by(items: list[Any], keyfunc) -> dict:
    out: dict[str, int] = {}
    for item in items:
        key = str(keyfunc(item) or "Sin dato")
        out[key] = out.get(key, 0) + 1
    return out


def _cost_buckets(vehicle: dict) -> dict:
    ensure_vehicle_runtime_fields(vehicle)
    g = vehicle.get("gastos_detallados", []) or []
    return {
        "compra": _to_float(vehicle.get("precio_ganado_usd"), 0),
        "comision_bancaria": _cost_total_by_keywords(g, ["COMISION", "BANCARIA"]),
        "grua_usa": _cost_total_by_keywords(g, ["GRUA_TRASLADO", "TRASLADO", "USA"]),
        "pago_complementario": _cost_total_by_keywords(g, ["COMPLEMENTARIO"]),
        "fedex": _cost_total_by_keywords(g, ["FEDEX"]),
        "flete_bl": _cost_total_by_keywords(g, ["FLETE", "BL", "NAVIERA_GRUA", "NAVIERA"]),
        "declaracion": _cost_total_by_keywords(g, ["DECLARACION", "DUCA", "IMPUESTOS"]),
        "servicios_aduanales": _cost_total_by_keywords(g, ["TRAMITE", "ADUANAL", "ADUANALES"]),
        "repuestos": _cost_total_by_keywords(g, ["REPUESTO"]),
        "pintura": _cost_total_by_keywords(g, ["PINTURA", "ENDEREZADO"]),
        "mecanica": _cost_total_by_keywords(g, ["MECANICO", "MECANICA"]),
        "grua_local": _cost_total_by_keywords(g, ["GRUA_LOCAL"]),
        "emisiones": _cost_total_by_keywords(g, ["EMISION"]),
        "tapiceria": _cost_total_by_keywords(g, ["TAPICERIA"]),
        "limpieza": _cost_total_by_keywords(g, ["LIMPIEZA", "PULIDA"]),
        "experticias": _cost_total_by_keywords(g, ["EXPERTICIA"]),
        "placas": _cost_total_by_keywords(g, ["PLACA", "CITA", "TARJETA"]),
        "almacenadora": _cost_total_by_keywords(g, ["ALMACEN"]),
        "honorarios": _cost_total_by_keywords(g, ["HONORARIO"]),
    }


def _report_alerts(vehicles: list[dict], quotes: Optional[list[dict]] = None) -> list[str]:
    alerts = []
    for v in vehicles:
        ensure_vehicle_runtime_fields(v)
        code = v.get("codigo", "")
        if vehicle_total_cost(v) <= 0:
            alerts.append(f"{code}: costo total en cero o incompleto.")
        if v.get("estado_actual") == STAGE_DISPONIBLE and _to_float(v.get("precio_venta_usd"), 0) <= 0:
            alerts.append(f"{code}: disponible sin precio cliente final.")
        if _to_float(v.get("precio_venta_usd"), 0) and _to_float(v.get("precio_minimo_usd"), 0) and _to_float(v.get("precio_venta_usd"), 0) < _to_float(v.get("precio_minimo_usd"), 0):
            alerts.append(f"{code}: precio publicado menor al mínimo aceptable.")
        for st in STAGE_ORDER:
            s = vehicle_stage(v, st)
            ini = _parse_date(s.get("fecha_inicio")); fin = _parse_date(s.get("fecha_fin"))
            if ini and fin and fin < ini:
                alerts.append(f"{code}: fechas invertidas en {STAGE_META.get(st,{}).get('label',st)}.")
    for q in quotes if quotes is not None else load_quotes():
        if quote_alert_level(q) == "ROJO" and q.get("estado") not in (QUOTE_GANADA, QUOTE_PERDIDA):
            alerts.append(f"Cotización roja: {q.get('cliente',{}).get('nombre','')} · {q.get('vehicle_code','')}.")
    return alerts[:30]


def _openpyxl_common():
    try:
        from openpyxl import Workbook
        from openpyxl.chart import BarChart, LineChart, PieChart, Reference
        from openpyxl.chart.pie_chart import DoughnutChart
        from openpyxl.drawing.image import Image as XLImage
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
        return Workbook, BarChart, LineChart, PieChart, Reference, XLImage, Alignment, Border, Font, PatternFill, Side, get_column_letter
    except Exception:
        return None


def _style_range(ws, cell_range: str, fill: Optional[str] = None, font_color: str = "000000", bold: bool = False,
                 size: int = 11, align: str = "center", border: bool = True):
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    side = Side(style="thin", color=LYM_BORDER)
    for row in ws[cell_range]:
        for cell in row:
            if fill:
                cell.fill = PatternFill("solid", fgColor=fill)
            cell.font = Font(bold=bold, color=font_color, size=size)
            cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
            if border:
                cell.border = Border(left=side, right=side, top=side, bottom=side)


def _merge_title(ws, rng: str, text: str, font_size: int = 16, fill: str = LYM_NAVY, color: str = "FFFFFF"):
    ws.merge_cells(rng)
    cell = ws[rng.split(":")[0]]
    cell.value = text
    _style_range(ws, rng, fill, color, True, font_size, "center")


def _set_widths(ws, widths: dict[int, float]):
    from openpyxl.utils import get_column_letter
    for col, width in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = width


def _fmt_xlsx_money(ws, cols: list[int], start: int, end: int):
    for r in range(start, end + 1):
        for c in cols:
            ws.cell(r, c).number_format = '$#,##0.00'


def _vehicle_report_row(v: dict, idx: int) -> dict:
    ensure_vehicle_runtime_fields(v)
    sale = _vehicle_sale_date(v)
    compra = _parse_date(v.get("fecha_compra"))
    disp = _stage_date(v, STAGE_DISPONIBLE, "fecha_inicio") or _stage_date(v, STAGE_DISPONIBLE, "fecha_fin")
    prep = vehicle_stage(v, STAGE_PREPARACION).get("extra", {}) or {}
    taller_in = _parse_date(prep.get("taller_ingreso")) or _stage_date(v, STAGE_PREPARACION, "fecha_inicio")
    taller_out = _parse_date(prep.get("taller_salida")) or _stage_date(v, STAGE_PREPARACION, "fecha_fin")
    leg_in = _parse_date(prep.get("placas_ingreso")) or _parse_date(prep.get("emision_pedido")) or _stage_date(v, STAGE_PREPARACION, "fecha_inicio")
    leg_out = _parse_date(prep.get("legalizacion_fin")) or _parse_date(prep.get("placas_entrega")) or _stage_date(v, STAGE_PREPARACION, "fecha_fin")
    stage_label = STAGE_META.get(v.get("estado_actual"), {}).get("label", v.get("estado_actual"))
    commercial = v.get("estado_comercial", "")
    status = "VENDIDO" if commercial == COMM_VENDIDO else stage_label
    return {"no": idx,"lot": v.get("lote") or "LOCAL","fecha_compra": compra,"fecha_venta": sale,"marca": v.get("marca"),"modelo": v.get("modelo"),"anio": v.get("anio"),"estatus": status,"observaciones": v.get("cliente") or v.get("observaciones") or "","summary": stage_label,"llegada_es": _stage_date(v, STAGE_ADUANA, "fecha_inicio"),"taller_in": taller_in,"taller_out": taller_out,"legal_in": leg_in,"legal_out": leg_out,"disponible": disp,"dias_taller": (taller_out - taller_in).days if taller_in and taller_out else None,"dias_legal": (leg_out - leg_in).days if leg_in and leg_out else None,"dias_compra_disp": (disp - compra).days if compra and disp else None,"dias_disp_venta": (sale - disp).days if sale and disp else None,"dias_compra_venta": (sale - compra).days if sale and compra else None,"dias_actuales": vehicle_days_from_purchase(v),"alerta": "CERRADO" if sale else stage_alert_level(v),"mes_venta": sale.strftime("%Y-%m") if sale else "","tipo": "LOCAL" if str(v.get("lote", "")).upper() == "LOCAL" else "LOTE"}


def generate_inventory_excel(vehicles: list[dict], user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    out_dir = _report_output_dir()
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    api = _openpyxl_common()
    if api is None:
        return False, "Para generar Excel instala: pip install openpyxl", None
    Workbook, BarChart, LineChart, PieChart, Reference, XLImage, Alignment, Border, Font, PatternFill, Side, get_column_letter = api
    vehicles = [ensure_vehicle_runtime_fields(v) for v in vehicles]
    out = out_dir / f"Inventario_vehiculos_LYM_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    wb = Workbook()
    ws = wb.active; ws.title = "Inventario General"
    dash = wb.create_sheet("Dashboard Gerencial")
    hist = wb.create_sheet("Historial Disponible")
    _merge_title(ws, "A1:Z1", "INVENTARIO Y ESTATUS DE VEHÍCULOS GENERAL", 16)
    ws.merge_cells("A2:Z2"); ws["A2"] = f"Corte {_spanish_date_long()} · Generado desde LYM AUTO CONTROL · Solo datos registrados en el sistema"; _style_range(ws,"A2:Z2",LYM_LIGHT,LYM_NAVY,True,11,"center")
    headers3 = {"A3":"No.","B3":"LOT","C3":"FECHA DE COMPRA","D3":"FECHA DE VENTA","E3":"VEHÍCULO","H3":"ESTATUS","I3":"OBSERVACIONES","J3":"SUMMARY","L3":"NUEVAS FECHAS DEL PROCESO","R3":"HISTORIAL DE TIEMPOS"}
    for c,v in headers3.items(): ws[c]=v
    for rng in ["A3:A4","B3:B4","C3:C4","D3:D4","E3:G3","H3:H4","I3:I4","J3:K3","L3:Q3","R3:Z3"]:
        try: ws.merge_cells(rng)
        except Exception: pass
    row4 = [None,None,None,None,"MARCA","MODELO","AÑO",None,None,"ETAPA ACTUAL","DÍAS ETAPA","LLEGADA ADUANA/ES","ING. TALLER","SAL. TALLER","INICIO LEGAL","FIN LEGAL","DISPONIBLE","DÍAS TALLER","DÍAS LEGAL","COMPRA→DISP","DISP→VENTA","COMPRA→VENTA","DÍAS ACTUALES","ALERTA","MES VENTA","TIPO"]
    for i,v in enumerate(row4, start=1):
        if v is not None:
            ws.cell(4,i).value=v
    _style_range(ws,"A3:Z4",LYM_NAVY,"FFFFFF",True,10,"center")
    records = [_vehicle_report_row(v, i) for i,v in enumerate(vehicles, 1)]
    start=5
    keys=["no","lot","fecha_compra","fecha_venta","marca","modelo","anio","estatus","observaciones","summary","dias_actuales","llegada_es","taller_in","taller_out","legal_in","legal_out","disponible","dias_taller","dias_legal","dias_compra_disp","dias_disp_venta","dias_compra_venta","dias_actuales","alerta","mes_venta","tipo"]
    for r, rec in enumerate(records, start=start):
        for c,k in enumerate(keys, start=1): ws.cell(r,c).value = rec.get(k)
        for c in [3,4,12,13,14,15,16,17]: ws.cell(r,c).number_format = "dd/mm/yyyy"
        alert = rec.get("alerta")
        fill = "DCFCE7" if alert in ("OK","VERDE","CERRADO") else ("FEF3C7" if alert in ("ALERTA","AMARILLO") else "FEE2E2")
        _style_range(ws, f"A{r}:Z{r}", fill, "0B172A", False, 10, "center")
    end = max(start, start + len(records) - 1)
    _set_widths(ws,{1:6,2:13,3:14,4:14,5:15,6:20,7:8,8:16,9:32,10:18,11:10,12:16,13:14,14:14,15:14,16:14,17:14,18:10,19:10,20:13,21:13,22:13,23:12,24:12,25:12,26:10})
    ws.freeze_panes="A5"; ws.auto_filter.ref=f"A4:Z{end}"
    _merge_title(dash,"A1:N1","REPORTE GERENCIAL DE INVENTARIO DE VEHÍCULOS",18)
    dash.merge_cells("A2:N2"); dash["A2"] = f"Corte {_spanish_date_long()} · Reporte generado desde datos reales del sistema"; _style_range(dash,"A2:N2",LYM_LIGHT,LYM_NAVY,True,11,"center")
    sold=[v for v in vehicles if v.get("estado_comercial")==COMM_VENDIDO]
    activos=[v for v in vehicles if v.get("estado_comercial")!=COMM_VENDIDO]
    disponibles=[v for v in vehicles if v.get("estado_actual")==STAGE_DISPONIBLE and v.get("estado_comercial")==COMM_DISPONIBLE]
    sale_days=[]
    for v in sold:
        fc=_parse_date(v.get("fecha_compra")); fv=_vehicle_sale_date(v)
        if fc and fv: sale_days.append((fv-fc).days)
    avg_sale=round(sum(sale_days)/len(sale_days),1) if sale_days else 0
    quotes=load_quotes(); qcounts=_quote_status_counts(quotes)
    kpis=[("Total unidades",len(vehicles)),("Inventario activo",len(activos)),("Vendidos total",len(sold)),("En tránsito",sum(1 for v in activos if v.get("estado_actual")==STAGE_TRANSITO)),("Disponibles",len(disponibles)),("Prom. días venta",avg_sale),("Capital activo",sum(vehicle_total_cost(v) for v in activos)),("Ganancia esperada",sum(vehicle_expected_profit(v) for v in vehicles)),("Cotizaciones",len(quotes))]
    for i,(lab,val) in enumerate(kpis, start=1): dash.cell(4,i).value=lab; dash.cell(5,i).value=val
    _style_range(dash,"A4:I4",LYM_NAVY,"FFFFFF",True,10,"center"); _style_range(dash,"A5:I5","FFFFFF",LYM_NAVY,True,14,"center")
    for c in [7,8]: dash.cell(5,c).number_format='$#,##0.00'
    status_counts=_counter_by(activos, lambda v: STAGE_META.get(v.get("estado_actual"),{}).get("label",v.get("estado_actual")))
    general_counts=_counter_by(vehicles, lambda v: "VENDIDO" if v.get("estado_comercial")==COMM_VENDIDO else STAGE_META.get(v.get("estado_actual"),{}).get("label",v.get("estado_actual")))
    dash["A8"]="ESTATUS ACTIVO — SIN VENDIDOS"; dash["H8"]="ESTATUS GENERAL — INCLUYE VENDIDOS"; _style_range(dash,"A8:C8",LYM_NAVY,"FFFFFF",True,11,"center"); _style_range(dash,"H8:J8",LYM_NAVY,"FFFFFF",True,11,"center")
    dash["A9"]="Estatus"; dash["B9"]="Cantidad"; dash["C9"]="% activo"; dash["H9"]="Estatus"; dash["I9"]="Cantidad"; dash["J9"]="% total"; _style_range(dash,"A9:C9",LYM_ORANGE,LYM_NAVY,True,10,"center"); _style_range(dash,"H9:J9",LYM_ORANGE,LYM_NAVY,True,10,"center")
    row=10
    for st,n in sorted(status_counts.items(), key=lambda x:x[1], reverse=True): dash.cell(row,1).value=st; dash.cell(row,2).value=n; dash.cell(row,3).value=(n/len(activos) if activos else 0); dash.cell(row,3).number_format='0.0%'; row+=1
    row2=10
    for st,n in sorted(general_counts.items(), key=lambda x:x[1], reverse=True): dash.cell(row2,8).value=st; dash.cell(row2,9).value=n; dash.cell(row2,10).value=(n/len(vehicles) if vehicles else 0); dash.cell(row2,10).number_format='0.0%'; row2+=1
    maxrow=max(row,row2)-1
    if maxrow>=10:
        _style_range(dash,f"A10:C{maxrow}","FFFFFF","0B172A",False,10,"center"); _style_range(dash,f"H10:J{maxrow}","FFFFFF","0B172A",False,10,"center")
        pie=PieChart(); pie.title="Estatus activo"; pie.add_data(Reference(dash,min_col=2,min_row=9,max_row=row-1),titles_from_data=True); pie.set_categories(Reference(dash,min_col=1,min_row=10,max_row=row-1)); dash.add_chart(pie,"D8")
        bar=BarChart(); bar.title="Estatus general"; bar.add_data(Reference(dash,min_col=9,min_row=9,max_row=row2-1),titles_from_data=True); bar.set_categories(Reference(dash,min_col=8,min_row=10,max_row=row2-1)); bar.height=7; bar.width=13; dash.add_chart(bar,"K8")
    sales_month=_counter_by(sold, lambda v: _report_month_key(_vehicle_sale_date(v)))
    q_month=_counter_by(quotes, lambda q: _report_month_key(q.get("fecha_cotizacion")))
    dash["A25"]="VENTAS POR MES"; dash["D25"]="COTIZACIONES POR MES"; _style_range(dash,"A25:B25",LYM_NAVY,"FFFFFF",True,11,"center"); _style_range(dash,"D25:E25",LYM_NAVY,"FFFFFF",True,11,"center")
    dash["A26"]="Mes"; dash["B26"]="Ventas"; dash["D26"]="Mes"; dash["E26"]="Cotizaciones"; _style_range(dash,"A26:B26",LYM_ORANGE,LYM_NAVY,True,10,"center"); _style_range(dash,"D26:E26",LYM_ORANGE,LYM_NAVY,True,10,"center")
    for i,(m,n) in enumerate(sorted(sales_month.items()), start=27): dash.cell(i,1).value=m; dash.cell(i,2).value=n
    for i,(m,n) in enumerate(sorted(q_month.items()), start=27): dash.cell(i,4).value=m; dash.cell(i,5).value=n
    if sales_month:
        ch=LineChart(); ch.title="Ventas por mes"; ch.add_data(Reference(dash,min_col=2,min_row=26,max_row=26+len(sales_month)),titles_from_data=True); ch.set_categories(Reference(dash,min_col=1,min_row=27,max_row=26+len(sales_month))); ch.height=7; ch.width=12; dash.add_chart(ch,"G25")
    if q_month:
        ch2=BarChart(); ch2.title="Cotizaciones por mes"; ch2.add_data(Reference(dash,min_col=5,min_row=26,max_row=26+len(q_month)),titles_from_data=True); ch2.set_categories(Reference(dash,min_col=4,min_row=27,max_row=26+len(q_month))); ch2.height=7; ch2.width=12; dash.add_chart(ch2,"G42")
    alerts=_report_alerts(vehicles, quotes); dash["A45"]="ALERTAS / VALIDACIONES DEL REPORTE"; _style_range(dash,"A45:F45",LYM_RED,"FFFFFF",True,11,"center")
    for i,a in enumerate(alerts or ["Sin alertas críticas detectadas."], start=46): dash.cell(i,1).value=a
    _merge_title(hist,"A1:L1","HISTORIAL DISPONIBLE PARA LA VENTA — COMPARATIVO ACTUAL",16)
    hist.merge_cells("A2:L2"); hist["A2"]=f"Fecha de corte {_spanish_date_long()} · Disponibles actuales: {len(disponibles)}"; _style_range(hist,"A2:L2",LYM_LIGHT,LYM_NAVY,True,11,"center")
    headers=["Detalle","No.","Lot","Marca","Modelo","Año","Fecha compra","Días desde compra","Precio cliente","Costo","Ganancia esperada","Observación"]
    for c,h in enumerate(headers,1): hist.cell(4,c).value=h
    _style_range(hist,"A4:L4",LYM_NAVY,"FFFFFF",True,10,"center")
    for r,v in enumerate(disponibles,start=5):
        hist.cell(r,1).value="Disponible actual"; hist.cell(r,2).value=r-4; hist.cell(r,3).value=v.get("lote"); hist.cell(r,4).value=v.get("marca"); hist.cell(r,5).value=v.get("modelo"); hist.cell(r,6).value=v.get("anio"); hist.cell(r,7).value=_parse_date(v.get("fecha_compra")); hist.cell(r,8).value=vehicle_days_from_purchase(v); hist.cell(r,9).value=_to_float(v.get("precio_venta_usd"),0); hist.cell(r,10).value=vehicle_total_cost(v); hist.cell(r,11).value=vehicle_expected_profit(v); hist.cell(r,12).value=v.get("observaciones","")
        hist.cell(r,7).number_format="dd/mm/yyyy"; hist.cell(r,9).number_format='$#,##0.00'; hist.cell(r,10).number_format='$#,##0.00'; hist.cell(r,11).number_format='$#,##0.00'
    if disponibles:
        _style_range(hist,f"A5:L{4+len(disponibles)}","FFFFFF","0B172A",False,10,"center")
    for sh in [ws,dash,hist]:
        _set_widths(sh,{i:14 for i in range(1,27)}); sh.sheet_view.showGridLines=False
    wb.save(out)
    log_audit("GENERAR_EXCEL_INVENTARIO_DESDE_CERO", (user or {}).get("usuario", ""), "", out.name)
    return True, "Inventario generado desde cero con Dashboard Gerencial, gráficos, KPIs, validaciones e historial disponible.", out


def generate_vehicle_costs_excel_report(vehicles: Optional[list[dict]] = None, user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    vehicles = [ensure_vehicle_runtime_fields(v) for v in (vehicles if vehicles is not None else load_vehicles())]
    out_dir = _report_output_dir()
    if out_dir is None: return False, "Carpeta del sistema no disponible.", None
    api = _openpyxl_common()
    if api is None: return False, "Para generar Excel instala: pip install openpyxl", None
    Workbook, BarChart, LineChart, PieChart, Reference, XLImage, Alignment, Border, Font, PatternFill, Side, get_column_letter = api
    out = out_dir / f"COPART_INC_COSTOS_DE_VEHICULOS_LYM_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    wb=Workbook(); ws=wb.active; ws.title="COPART INC"; wc=wb.create_sheet("COSTO")
    _merge_title(ws,"B1:BB1","L&M INVERSIONES, S.A. DE C.V.",16)
    _merge_title(ws,"B2:BB2","VEHÍCULOS IMPORTADOS",14,LYM_ORANGE,LYM_NAVY)
    _merge_title(ws,"B3:BB3","COMPRA DE VEHÍCULOS COPART INC",13,LYM_NAVY,"FFFFFF")
    headers=["No.","# LOTE","","FECHA DE ADQUISICION","CLASE","#","AÑO","PLACA","MARCA","TIPO","No. DE CHASIS","No. DE MOTOR","MODELO","COLOR","OC","VALOR DE ADQUISICIÓN COPART","OC","Comisión Bancaria","OC","GRÚA INTERNA USA","OC","PAGO COMPLEMENTARIO","OC","FEDEX","TOTAL USA","OC","GRÚA, FLETE Y BL","OC","DECLARACIÓN DE MERCANCÍA","OC","SERVICIOS ADUANALES","OC","REPUESTOS","OC","ENDEREZADO Y PINTURA","OC","SERVICIOS MECÁNICOS","OC","GRÚA LOCAL","OC","EMISIÓN DE GASES","OC","TAPICERÍA","OC","LIMPIEZA Y PULIDA","OC","EXPERTICIAS","OC","TARJETA DE CIRCULACIÓN Y PLACAS","OC","ALMACENADORA","HONORARIOS","TOTAL ES","COSTO"]
    for c,h in enumerate(headers,1): ws.cell(4,c).value=h
    ws.merge_cells("D5:Y5"); ws["D5"]="PAGOS USA"; ws.merge_cells("Z5:BA5"); ws["Z5"]="PAGOS EL SALVADOR"; _style_range(ws,"A4:BB5",LYM_NAVY,"FFFFFF",True,9,"center")
    wc.merge_cells("A1:AD1"); wc["A1"]="L&M INVERSIONES, S.A. DE C.V."; _style_range(wc,"A1:AD1",LYM_NAVY,"FFFFFF",True,16,"center")
    wc.merge_cells("A2:AD2"); wc["A2"]="VEHÍCULOS IMPORTADOS / COMPRAS COPART INC"; _style_range(wc,"A2:AD2",LYM_ORANGE,LYM_NAVY,True,13,"center")
    cost_headers=["# LOTE","AÑO COMPRA","FECHA DE ADQUISICION","CLASE","#","AÑO","PLACA","MARCA","TIPO","No. DE CHASIS","No. DE MOTOR","MODELO","COLOR","COSTO DE VEHICULO","MARGEN 30%","COSTO DE VENTA","VENTA NETA","IVA VENTA","IVA DUCA","IVA A PAGAR","PAGO A CUENTA","PRECIO DE VENTA","REF 15%","REF 30%","LEASING","CONTADO","PRIMA (20%)","PRIMA CON COMISION","UTILIDAD GENERADA","OBSERVACIONES"]
    for c,h in enumerate(cost_headers,1): wc.cell(4,c).value=h
    _style_range(wc,"A4:AD5",LYM_NAVY,"FFFFFF",True,9,"center")
    start=6
    for idx,v in enumerate(vehicles,start=1):
        r=start+idx-1; b=_cost_buckets(v); total_usa=round(b['comision_bancaria']+b['grua_usa']+b['pago_complementario']+b['fedex'],2); total_es=round(b['flete_bl']+b['declaracion']+b['servicios_aduanales']+b['repuestos']+b['pintura']+b['mecanica']+b['grua_local']+b['emisiones']+b['tapiceria']+b['limpieza']+b['experticias']+b['placas']+b['almacenadora']+b['honorarios'],2); cost=vehicle_total_cost(v)
        row={1:idx,2:v.get('lote'),4:_parse_date(v.get('fecha_compra')),6:idx,7:v.get('anio'),9:v.get('marca'),13:v.get('modelo'),14:v.get('color'),15:v.get('oc_compra_numero'),16:b['compra'],18:b['comision_bancaria'],20:b['grua_usa'],22:b['pago_complementario'],24:b['fedex'],25:total_usa,27:b['flete_bl'],29:b['declaracion'],31:b['servicios_aduanales'],33:b['repuestos'],35:b['pintura'],37:b['mecanica'],39:b['grua_local'],41:b['emisiones'],43:b['tapiceria'],45:b['limpieza'],47:b['experticias'],49:b['placas'],51:b['almacenadora'],52:b['honorarios'],53:total_es,54:cost}
        for c,val in row.items(): ws.cell(r,c).value=val
        ws.cell(r,4).number_format='dd/mm/yyyy'; _fmt_xlsx_money(ws,[16,18,20,22,24,25,27,29,31,33,35,37,39,41,43,45,47,49,51,52,53,54],r,r)
        pf=vehicle_profit_summary(v); precio=_to_float(v.get('precio_venta_usd'),0) or pf.get('precio_venta_cliente_usd',0); prima=round(precio*0.20,2) if precio else 0
        rowc={1:v.get('lote'),2:_parse_date(v.get('fecha_compra')).year if _parse_date(v.get('fecha_compra')) else '',3:_parse_date(v.get('fecha_compra')),5:idx,6:v.get('anio'),8:v.get('marca'),12:v.get('modelo'),13:v.get('color'),14:cost,15:0.30,16:round(cost*1.30,2),17:pf.get('venta_neta_usd'),18:pf.get('iva_venta_usd'),19:pf.get('iva_duca_usd'),20:pf.get('iva_pagar_usd'),21:pf.get('pago_cuenta_usd'),22:precio,23:round(precio*0.15,2) if precio else 0,24:round(precio*0.30,2) if precio else 0,25:precio,26:precio,27:prima,28:round(prima+100,2) if prima else 0,29:vehicle_expected_profit(v),30:v.get('observaciones','')}
        for c,val in rowc.items(): wc.cell(r,c).value=val
        wc.cell(r,3).number_format='dd/mm/yyyy'; wc.cell(r,15).number_format='0%'; _fmt_xlsx_money(wc,[14,16,17,18,19,20,21,22,23,24,25,26,27,28,29],r,r)
    end=max(start,start+len(vehicles)-1)
    _style_range(ws,f"A6:BB{end}","FFFFFF","0B172A",False,9,"center"); _style_range(wc,f"A6:AD{end}","FFFFFF","0B172A",False,9,"center")
    for sh in [ws,wc]:
        sh.freeze_panes='A6'; sh.auto_filter.ref=f"A4:{get_column_letter(sh.max_column)}{end}"; sh.sheet_view.showGridLines=False; _set_widths(sh,{i:14 for i in range(1,sh.max_column+1)})
    wb.save(out)
    log_audit("GENERAR_EXCEL_COSTOS_DESDE_CERO", (user or {}).get("usuario", ""), "", out.name)
    return True, "Costos por vehículo generados desde cero con hojas COPART INC y COSTO.", out


def generate_quotes_excel_report(user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    out_dir = _report_output_dir()
    if out_dir is None: return False, "Carpeta del sistema no disponible.", None
    api = _openpyxl_common()
    if api is None: return False, "Para generar Excel instala: pip install openpyxl", None
    Workbook, BarChart, LineChart, PieChart, Reference, XLImage, Alignment, Border, Font, PatternFill, Side, get_column_letter = api
    quotes=sorted(load_quotes(), key=lambda q: q.get('fecha_cotizacion',''))
    out=out_dir / f"Control_Cotizaciones_LYM_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    wb=Workbook(); ws=wb.active; ws.title="Control Cotizaciones"; dash=wb.create_sheet("Dashboard Cotizaciones")
    _merge_title(ws,"A1:M2","CONTROL DE COTIZACIONES - L&M INVERSIONES",16)
    headers=["FECHA","CORRELATIVO","NOMBRE DEL CLIENTE","VEHÍCULO COTIZADO","AÑO","TELÉFONO","MEDIO","TIPO DE NEGOCIO","OBSERVACIONES","ESTATUS","COLOR","CUOTA","PRIMA REQUERIDA"]
    for c,h in enumerate(headers,1): ws.cell(4,c).value=h
    _style_range(ws,"A4:M4",LYM_NAVY,"FFFFFF",True,10,"center")
    for idx,q in enumerate(quotes,start=1):
        r=4+idx; cl=q.get('cliente',{}); snap=q.get('vehicle_snapshot',{}); s=_quote_financial_summary(q); lvl=quote_alert_level(q)
        vals=[_parse_date(q.get('fecha_cotizacion')),f"COT-{idx:03d}",cl.get('nombre',''),f"{snap.get('marca','')} {snap.get('modelo','')}",snap.get('anio',''),cl.get('telefono',''),cl.get('medio_contacto',''),"LEASING",q.get('comentario',''),QUOTE_STATUS_LABELS.get(q.get('estado'),q.get('estado')),lvl,s['cuota_final'],s['prima_requerida']]
        for c,val in enumerate(vals,1): ws.cell(r,c).value=val
        ws.cell(r,1).number_format='dd/mm/yyyy'; ws.cell(r,12).number_format='$#,##0.00'; ws.cell(r,13).number_format='$#,##0.00'
        fill={'VERDE':'DCFCE7','AMARILLO':'FEF3C7','ROJO':'FEE2E2','CERRADA':'DBEAFE','PERDIDA':'F1F5F9'}.get(lvl,'FFFFFF')
        _style_range(ws,f"A{r}:M{r}",fill,"0B172A",False,10,"center")
    end=max(5,4+len(quotes)); ws.freeze_panes='A5'; ws.auto_filter.ref=f"A4:M{end}"; _set_widths(ws,{1:13,2:13,3:28,4:24,5:8,6:15,7:16,8:16,9:36,10:22,11:12,12:14,13:16})
    _merge_title(dash,"A1:N1","DASHBOARD GERENCIAL DE COTIZACIONES",18)
    dash.merge_cells("A2:N2"); dash["A2"]=f"Corte {_spanish_date_long()} · Seguimiento comercial por cliente, medio, color y vehículo cotizado"; _style_range(dash,"A2:N2",LYM_LIGHT,LYM_NAVY,True,11,"center")
    counts=_quote_status_counts(quotes); medios=_counter_by(quotes, lambda q:q.get('cliente',{}).get('medio_contacto','SIN MEDIO')); meses=_counter_by(quotes, lambda q:_report_month_key(q.get('fecha_cotizacion'))); carros=_counter_by(quotes, lambda q:_vehicle_display_name(q.get('vehicle_snapshot',{})))
    kpis=[('Total',counts['TOTAL']),('Verde',counts.get('VERDE',0)),('Amarillo',counts.get('AMARILLO',0)),('Rojo',counts.get('ROJO',0)),('Ganadas',counts.get('GANADAS',0)),('Reofertar',counts.get('REOFERTAR',0)),('Perdidas',counts.get('PERDIDAS',0))]
    for i,(lab,val) in enumerate(kpis,1): dash.cell(4,i).value=lab; dash.cell(5,i).value=val
    _style_range(dash,"A4:G4",LYM_NAVY,"FFFFFF",True,10,"center"); _style_range(dash,"A5:G5","FFFFFF",LYM_NAVY,True,16,"center")
    tables=[('A8','Cotizaciones por mes',meses),('D8','Medios de contacto',medios),('G8','Colores seguimiento',{k:counts.get(k,0) for k in ['VERDE','AMARILLO','ROJO','GANADAS','REOFERTAR','PERDIDAS']}),('J8','Carros más cotizados',dict(sorted(carros.items(), key=lambda x:x[1], reverse=True)[:10]))]
    for start_cell,title,data in tables:
        col=dash[start_cell].column; row=dash[start_cell].row
        dash.cell(row,col).value=title; dash.cell(row+1,col).value='Concepto'; dash.cell(row+1,col+1).value='Cantidad'
        _style_range(dash,f"{get_column_letter(col)}{row}:{get_column_letter(col+1)}{row}",LYM_NAVY,"FFFFFF",True,11,"center"); _style_range(dash,f"{get_column_letter(col)}{row+1}:{get_column_letter(col+1)}{row+1}",LYM_ORANGE,LYM_NAVY,True,10,"center")
        for i,(k,v) in enumerate(sorted(data.items(), key=lambda x:x[1], reverse=True),start=row+2): dash.cell(i,col).value=k; dash.cell(i,col+1).value=v
        if data: _style_range(dash,f"{get_column_letter(col)}{row+2}:{get_column_letter(col+1)}{row+1+len(data)}","FFFFFF","0B172A",False,10,"center")
    if meses:
        ch=BarChart(); ch.title='Cotizaciones por mes'; ch.add_data(Reference(dash,min_col=2,min_row=9,max_row=9+len(meses)),titles_from_data=True); ch.set_categories(Reference(dash,min_col=1,min_row=10,max_row=9+len(meses))); ch.height=7; ch.width=12; dash.add_chart(ch,'A23')
    if medios:
        pie=PieChart(); pie.title='Medios de contacto'; pie.add_data(Reference(dash,min_col=5,min_row=9,max_row=9+len(medios)),titles_from_data=True); pie.set_categories(Reference(dash,min_col=4,min_row=10,max_row=9+len(medios))); pie.height=7; pie.width=10; dash.add_chart(pie,'F23')
    if carros:
        bar=BarChart(); bar.title='Carros más cotizados'; bar.add_data(Reference(dash,min_col=11,min_row=9,max_row=9+min(len(carros),10)),titles_from_data=True); bar.set_categories(Reference(dash,min_col=10,min_row=10,max_row=9+min(len(carros),10))); bar.height=7; bar.width=13; dash.add_chart(bar,'K23')
    alerts=_report_alerts(load_vehicles(), quotes); dash["A42"]="ALERTAS COMERCIALES"; _style_range(dash,"A42:H42",LYM_RED,"FFFFFF",True,11,"center")
    for i,a in enumerate(alerts or ["Sin alertas comerciales críticas."], start=43): dash.cell(i,1).value=a
    for sh in [ws,dash]: sh.sheet_view.showGridLines=False; _set_widths(sh,{i:16 for i in range(1,15)})
    wb.save(out)
    log_audit("GENERAR_EXCEL_COTIZACIONES_DESDE_CERO", (user or {}).get("usuario", ""), "", out.name)
    return True, "Control de cotizaciones generado desde cero con dashboard, gráficos y colores verde/amarillo/rojo.", out


def generate_html_report(vehicles: list[dict], user: Optional[dict] = None) -> Optional[Path]:
    out_dir = _report_output_dir()
    if out_dir is None: return None
    vehicles=[ensure_vehicle_runtime_fields(v) for v in vehicles]
    quotes=load_quotes(); qcounts=_quote_status_counts(quotes); k=compute_kpis(vehicles); alerts=_report_alerts(vehicles, quotes); logo=ResourceManager.logo_data_uri(); template=_find_template_file(HTML_REPORT_TEMPLATE_CANDIDATES)
    style="""<style>:root{--navy:#08285a;--navy2:#0e3a78;--orange:#f59a13;--bg:#f5f8fc;--ink:#0b172a;--muted:#637083;--line:rgba(8,40,90,.12);--green:#10B981;--red:#DC2626;--amber:#F59E0B;--blue:#2563EB}*{box-sizing:border-box}body{margin:0;font-family:Inter,Segoe UI,Arial,sans-serif;color:var(--ink);background:linear-gradient(135deg,#f6f9ff,#eef4fb 55%,#fff8ef)}.slide{min-height:100vh;padding:58px clamp(22px,5vw,86px);display:flex;align-items:center}.container{width:min(1220px,100%);margin:auto}.eyebrow{display:inline-flex;color:var(--orange);text-transform:uppercase;font-size:12px;letter-spacing:.16em;font-weight:900;background:rgba(245,154,19,.10);border-radius:999px;padding:8px 13px}h1{font-size:clamp(40px,7vw,82px);line-height:.94;color:var(--navy);letter-spacing:-.06em}h2{font-size:clamp(30px,4vw,56px);color:var(--navy);letter-spacing:-.045em}.lead{font-size:20px;line-height:1.45;color:#334155;max-width:880px}.hero-grid{display:grid;grid-template-columns:1.1fr .65fr;gap:34px;align-items:center}.logo-card{min-height:420px;border-radius:40px;background:rgba(255,255,255,.86);border:1px solid var(--line);display:grid;place-items:center;box-shadow:0 30px 90px rgba(8,40,90,.18)}.logo-card img{width:min(330px,78%)}.kpis{display:grid;grid-template-columns:repeat(4,1fr);gap:18px}.card{background:rgba(255,255,255,.86);border:1px solid var(--line);border-radius:26px;padding:20px;box-shadow:0 20px 55px rgba(8,40,90,.10)}.label{color:var(--muted);font-size:12px;text-transform:uppercase;letter-spacing:.08em;font-weight:900}.value{font-size:38px;color:var(--navy);font-weight:950;line-height:1;margin-top:8px}.grid2{display:grid;grid-template-columns:1fr 1fr;gap:22px}table{width:100%;border-collapse:separate;border-spacing:0;background:#fff;border-radius:22px;overflow:hidden;box-shadow:0 18px 46px rgba(8,40,90,.10)}th{background:var(--navy);color:white;text-align:left;padding:12px;font-size:12px;text-transform:uppercase}td{padding:11px;border-bottom:1px solid rgba(8,40,90,.08);font-weight:700}.bar{height:28px;border-radius:999px;background:#edf3fa;overflow:hidden}.fill{height:100%;background:linear-gradient(90deg,var(--orange),var(--navy));border-radius:999px}.note{background:linear-gradient(135deg,var(--navy),var(--navy2));color:white;border-radius:28px;padding:24px}.note h2{color:white}.red{color:var(--red)}@media(max-width:980px){.hero-grid,.grid2,.kpis{grid-template-columns:1fr}.slide{padding:38px 18px}}</style>"""
    if template and template.exists():
        t=template.read_text(encoding='utf-8', errors='ignore'); m=re.search(r"<style>.*?</style>", t, flags=re.S|re.I)
        if m: style=m.group(0)
    status_counts=_counter_by([v for v in vehicles if v.get('estado_comercial')!=COMM_VENDIDO], lambda v: STAGE_META.get(v.get('estado_actual'),{}).get('label',v.get('estado_actual')))
    medios=_counter_by(quotes, lambda q:q.get('cliente',{}).get('medio_contacto','SIN MEDIO')); q_month=_counter_by(quotes, lambda q:_report_month_key(q.get('fecha_cotizacion'))); top_cars=dict(sorted(_counter_by(quotes, lambda q:_vehicle_display_name(q.get('vehicle_snapshot',{}))).items(), key=lambda x:x[1], reverse=True)[:8])
    def table_rows(d):
        total=max(sum(d.values()),1)
        return ''.join(f"<tr><td>{html.escape(str(k))}</td><td>{v}</td><td><div class='bar'><div class='fill' style='width:{(v/total*100):.1f}%'></div></div></td></tr>" for k,v in sorted(d.items(), key=lambda x:x[1], reverse=True))
    inv_rows=''.join(f"<tr><td>{html.escape(v.get('codigo',''))}</td><td>{html.escape(_vehicle_display_name(v))}</td><td>{html.escape(str(v.get('lote','')))}</td><td>{_fmt_usd(vehicle_total_cost(v))}</td><td>{_fmt_usd(v.get('precio_venta_usd'))}</td><td>{_fmt_usd(vehicle_expected_profit(v))}</td><td>{vehicle_days_from_purchase(v)}</td><td>{html.escape(str(v.get('estado_comercial','')))}</td></tr>" for v in vehicles[:120])
    html_text=f"""<!doctype html><html lang='es'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>L&M Inversiones · Reporte Gerencial</title>{style}</head><body><main class='deck'>
<section class='slide'><div class='container hero-grid'><div><div class='eyebrow'>Reporte gerencial · Corte {_spanish_date_long()}</div><h1>Inventario y comercial <span style='color:var(--orange)'>L & M</span></h1><p class='lead'>Presentación ejecutiva generada automáticamente desde LYM AUTO CONTROL. Incluye inventario, costos, utilidad real estimada, tiempos, cotizaciones, medios de contacto y alertas de riesgo.</p></div><div class='logo-card'>{'<img src="'+logo+'">' if logo else '<h2>L&M</h2>'}</div></div></section>
<section class='slide'><div class='container'><div class='eyebrow'>KPIs del sistema</div><h2>Dashboard gerencial</h2><div class='kpis'><div class='card'><div class='label'>Vehículos totales</div><div class='value'>{k.get('total',0)}</div></div><div class='card'><div class='label'>Activos</div><div class='value'>{k.get('activos',0)}</div></div><div class='card'><div class='label'>Disponibles</div><div class='value'>{k.get('disponibles',0)}</div></div><div class='card'><div class='label'>Vendidos</div><div class='value'>{k.get('vendidos',0)}</div></div><div class='card'><div class='label'>Capital activo</div><div class='value'>{_fmt_usd(k.get('capital',0))}</div></div><div class='card'><div class='label'>Ganancia esperada</div><div class='value'>{_fmt_usd(k.get('ganancia_esperada',0))}</div></div><div class='card'><div class='label'>Cotizaciones</div><div class='value'>{qcounts.get('TOTAL',0)}</div></div><div class='card'><div class='label'>Cotizaciones rojas</div><div class='value red'>{qcounts.get('ROJO',0)}</div></div></div></div></section>
<section class='slide'><div class='container grid2'><div class='card'><h2>Estatus operativo</h2><table><thead><tr><th>Etapa</th><th>Cantidad</th><th>%</th></tr></thead><tbody>{table_rows(status_counts)}</tbody></table></div><div class='card'><h2>Medios de contacto</h2><table><thead><tr><th>Medio</th><th>Cotizaciones</th><th>%</th></tr></thead><tbody>{table_rows(medios)}</tbody></table></div></div></section>
<section class='slide'><div class='container grid2'><div class='card'><h2>Cotizaciones por mes</h2><table><thead><tr><th>Mes</th><th>Cotizaciones</th><th>%</th></tr></thead><tbody>{table_rows(q_month)}</tbody></table></div><div class='card'><h2>Carros más cotizados</h2><table><thead><tr><th>Vehículo</th><th>Cotizaciones</th><th>%</th></tr></thead><tbody>{table_rows(top_cars)}</tbody></table></div></div></section>
<section class='slide'><div class='container'><div class='eyebrow'>Validación inteligente</div><h2>Alertas para evitar pérdidas</h2><div class='note'><h2>Revisión automática</h2><ul>{''.join('<li>'+html.escape(a)+'</li>' for a in (alerts or ['Sin alertas críticas detectadas.']))}</ul></div></div></section>
<section class='slide'><div class='container'><div class='eyebrow'>Detalle</div><h2>Inventario actualizado</h2><table><thead><tr><th>CV</th><th>Vehículo</th><th>Lote</th><th>Costo</th><th>Precio</th><th>Ganancia</th><th>Días</th><th>Comercial</th></tr></thead><tbody>{inv_rows}</tbody></table></div></section>
</main></body></html>"""
    out=out_dir / f"LYM_Inversiones_Reporte_Gerencial_{datetime.now().strftime('%Y%m%d_%H%M')}.html"; out.write_text(html_text, encoding='utf-8'); log_audit("GENERAR_HTML_GERENCIAL_DINAMICO", (user or {}).get("usuario", ""), "", out.name); return out


def generate_quotes_html_report(user: Optional[dict] = None) -> Optional[Path]:
    return generate_html_report(load_vehicles(), user)


if PYSIDE_OK:
    class PriceFinalReviewDialog(QDialog):
        def __init__(self, parent, vehicle: dict, pf: dict):
            super().__init__(parent); self.vehicle=vehicle; self.pf=pf
            self.setWindowTitle("Revisión final de precio y ganancia"); self.setMinimumSize(880,650); self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint)
            lay=QVBoxLayout(self)
            title=QLabel(f"{vehicle.get('codigo')} · {_vehicle_display_name(vehicle)}"); title.setStyleSheet("font-size:22px;font-weight:950;color:#08285a;"); lay.addWidget(title)
            note=QLabel("Revisa el desglose completo antes de aprobar el precio. La ganancia real estimada usa: precio final cliente - IVA a pagar - pago a cuenta - VTS - regalía - costo total."); note.setWordWrap(True); note.setStyleSheet("background:#eaf4ff;border:1px solid #bfdbfe;border-radius:12px;padding:12px;color:#08285a;font-weight:800;"); lay.addWidget(note)
            self.table=QTableWidget(0,3); self.table.setHorizontalHeaderLabels(["Concepto","Valor","Explicación"]); self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); lay.addWidget(self.table,1)
            rows=[("Costo total del vehículo", _fmt_usd(pf['costo_total_usd']), "Suma real de compra + gastos + documentos."),("Margen bruto aplicado", f"{pf['margen_pct']}%", "Margen usado para sugerir la venta neta."),("Venta neta", _fmt_usd(pf['venta_neta_usd']), "Base antes de IVA venta."),("IVA venta generado", _fmt_usd(pf['iva_venta_usd']), "IVA calculado sobre venta neta."),("IVA DUCA/crédito usado", _fmt_usd(pf['iva_duca_usd']), "Crédito que resta al IVA venta."),("IVA a pagar", _fmt_usd(pf['iva_pagar_usd']), "IVA venta - IVA DUCA."),("Pago a cuenta", _fmt_usd(pf['pago_cuenta_usd']), f"{pf['pago_cuenta_pct']}% sobre venta neta."),("VTS", _fmt_usd(pf['vts_usd']), "Valor fijo configurado."),("Precio final cliente", _fmt_usd(pf['precio_venta_cliente_usd']), "Lo que verá y pagará el cliente."),("Precio mínimo aceptable", _fmt_usd(pf['precio_minimo_usd']), "Mínimo permitido al cerrar venta."),("Ingreso neto después de impuestos", _fmt_usd(pf['ingreso_neto_despues_impuestos_usd']), "Precio cliente - IVA a pagar - pago a cuenta - VTS."),("Ganancia real estimada", _fmt_usd(pf['ganancia_real_estimada_usd']), "Ingreso neto después de impuestos - costo total."),("Margen real estimado", f"{pf['margen_real_estimado_pct']}%", "Ganancia real estimada / costo total.")]
            self.table.setRowCount(len(rows))
            for r,row in enumerate(rows):
                for c,val in enumerate(row):
                    it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.table.setItem(r,c,it)
            warn="" if pf['ganancia_real_estimada_usd'] >= 0 else "⚠️ Esta operación proyecta pérdida."
            self.summary=QLabel(f"Ganancia real estimada: <b>{_fmt_usd(pf['ganancia_real_estimada_usd'])}</b> · Margen real: <b>{pf['margen_real_estimado_pct']}%</b> {warn}"); self.summary.setTextFormat(Qt.TextFormat.RichText); self.summary.setStyleSheet("font-size:16px;background:#fff7ed;border:1px solid #fed7aa;border-radius:12px;padding:14px;color:#08285a;"); lay.addWidget(self.summary)
            row=QHBoxLayout(); lay.addLayout(row); ok=QPushButton("Aprobar y guardar precio"); ok.setObjectName("orange"); ok.clicked.connect(self.accept); no=QPushButton("Rechazar / volver a editar"); no.setObjectName("danger"); no.clicked.connect(self.reject); row.addWidget(ok); row.addWidget(no); row.addStretch(1)

    def _patched_build_price_section(self, main, form):
        btnrow=QHBoxLayout(); bprev=QPushButton("Visualizar gastos"); bprev.setObjectName("ghost"); bprev.setMaximumWidth(145); bprev.clicked.connect(self.open_costs); btnrow.addWidget(bprev); btnrow.addStretch(1); main.addLayout(btnrow)
        saved_pf=self.vehicle.get("precio_final") or {}
        self.margen=QDoubleSpinBox(); self.margen.setRange(0,100); self.margen.setDecimals(2); self.margen.setSuffix(" %"); self.margen.setValue(float(saved_pf.get("margen_pct") or 30))
        self.iva_duca=MoneyEdit(); self.iva_duca.setRange(0,999999); self.iva_duca.setValue(float(saved_pf.get("iva_duca_usd") or _aduana_impuestos_amount(self.vehicle)))
        self.iva_pct=QDoubleSpinBox(); self.iva_pct.setRange(0,30); self.iva_pct.setDecimals(2); self.iva_pct.setSuffix(" %"); self.iva_pct.setValue(float(saved_pf.get("iva_pct") or 13))
        self.pago_cuenta_pct=QDoubleSpinBox(); self.pago_cuenta_pct.setRange(0,20); self.pago_cuenta_pct.setDecimals(2); self.pago_cuenta_pct.setSuffix(" %"); self.pago_cuenta_pct.setValue(float(saved_pf.get("pago_cuenta_pct") or 1.75))
        self.vts=MoneyEdit(); self.vts.setRange(0,999); self.vts.setValue(float(saved_pf.get("vts_usd") or 2.07))
        self.precio_manual=MoneyEdit(); self.precio_manual.setRange(0,9999999); self.precio_manual.setValue(float(saved_pf.get("precio_redondeado_manual_usd") or saved_pf.get("precio_venta_cliente_usd") or 0))
        self.precio_minimo_manual=MoneyEdit(); self.precio_minimo_manual.setRange(0,9999999); self.precio_minimo_manual.setValue(float(saved_pf.get("precio_minimo_usd") or 0))
        self.foto_label=QLabel(_legal_doc_flag(self.vehicle.get("foto_principal"), self.foto_path)); self.foto_label.setStyleSheet("font-weight:900;color:#08285a;"); bfoto=QPushButton("Subir foto principal propuesta"); bfoto.setObjectName("ghost"); bfoto.clicked.connect(self._pick_foto)
        frow=QWidget(); fl=QHBoxLayout(frow); fl.setContentsMargins(0,0,0,0); fl.addWidget(self.foto_label); fl.addWidget(bfoto)
        self.calc_lbl=QLabel(); self.calc_lbl.setTextFormat(Qt.TextFormat.RichText); self.calc_lbl.setStyleSheet("background:#fff7ed;border:1px solid #fed7aa;border-radius:10px;padding:12px;color:#08285a;")
        self.detail_btn=QPushButton("Ver pantalla final de cálculo"); self.detail_btn.setObjectName("ghost"); self.detail_btn.clicked.connect(lambda: PriceFinalReviewDialog(self,self.vehicle,self._price_calc()).exec())
        for w in [self.margen,self.iva_pct,self.pago_cuenta_pct]: w.valueChanged.connect(self._refresh_price_calc)
        try: self.iva_duca.textChanged.connect(self._refresh_price_calc); self.vts.textChanged.connect(self._refresh_price_calc); self.precio_manual.textChanged.connect(self._refresh_price_calc); self.precio_minimo_manual.textChanged.connect(self._refresh_price_calc)
        except Exception: pass
        form.addRow("Margen bruto deseado:", self.margen); form.addRow("IVA DUCA / impuestos aduana:", self.iva_duca); form.addRow("IVA venta:", self.iva_pct); form.addRow("Pago a cuenta:", self.pago_cuenta_pct); form.addRow("VTS fijo:", self.vts); form.addRow("Precio final cliente / redondeado:", self.precio_manual); form.addRow("Precio mínimo aceptable:", self.precio_minimo_manual); form.addRow("Foto principal:", frow); main.addWidget(self.calc_lbl); main.addWidget(self.detail_btn); self._refresh_price_calc()

    def _patched_price_calc(self) -> dict:
        return calculate_price_final_breakdown(vehicle_total_cost(self.vehicle), self.margen.value(), self.iva_duca.value(), self.iva_pct.value(), self.pago_cuenta_pct.value(), self.vts.value(), self.precio_manual.value(), self.precio_minimo_manual.value())

    def _patched_refresh_price_calc(self):
        pf=self._price_calc(); manual_txt="<br><b>Precio manual/redondeado aplicado.</b>" if pf.get('precio_redondeado_manual_usd') else ""
        self.calc_lbl.setText(f"Costo total: <b>{_fmt_usd(pf['costo_total_usd'])}</b><br>Venta neta: <b>{_fmt_usd(pf['venta_neta_usd'])}</b> · IVA venta: <b>{_fmt_usd(pf['iva_venta_usd'])}</b> · IVA DUCA usado: <b>{_fmt_usd(pf['iva_duca_usd'])}</b> · IVA a pagar: <b>{_fmt_usd(pf['iva_pagar_usd'])}</b><br>Pago a cuenta: <b>{_fmt_usd(pf['pago_cuenta_usd'])}</b> · VTS: <b>{_fmt_usd(pf['vts_usd'])}</b><br>Precio cliente final: <b style='font-size:20px'>{_fmt_usd(pf['precio_venta_cliente_usd'])}</b> · Precio mínimo: <b>{_fmt_usd(pf['precio_minimo_usd'])}</b>{manual_txt}<br>Ganancia real estimada: <b>{_fmt_usd(pf['ganancia_real_estimada_usd'])}</b> · Margen real estimado: <b>{pf['margen_real_estimado_pct']}%</b><br><span style='color:#64748b'>Fórmula: precio final - IVA a pagar - pago a cuenta - VTS - costo total.</span>")

    _ORIG_STAGE_SAVE = StageUpdateDialog.save
    def _patched_stage_save(self):
        if getattr(self, 'stage_key', None) != STAGE_PRECIO_FINAL:
            return _ORIG_STAGE_SAVE(self)
        fecha=self._date(self.fecha_evento); pf=self._price_calc(); review=PriceFinalReviewDialog(self,self.vehicle,pf)
        if review.exec()!=QDialog.DialogCode.Accepted:
            QMessageBox.information(self,"Precio final","No se guardó el precio. Puedes ajustar los valores y volver a aprobar."); return
        old_price=_to_float(self.vehicle.get("precio_venta_usd"),0)
        data={"fecha_inicio":fecha,"fecha_fin":fecha,"proveedor":"INTERNO","comentario":self.comentario.toPlainText(),"extra":{"precio_final_guardado":True},"precio_final":pf,"foto_principal_src":self.foto_path}
        ok,msg=update_vehicle_stage(self.vehicle.get("id"), self.stage_key, data, None, self.user, self.device)
        if not ok: QMessageBox.warning(self,"Validación",msg); return
        updated=find_vehicle(self.vehicle.get("id")) or self.vehicle; new_price=_to_float(updated.get("precio_venta_usd"),0)
        if abs(old_price-new_price)>=0.01:
            register_vehicle_price_change(updated, old_price, new_price, self.user, self.device); changed=sync_active_quotes_after_price_change(updated, old_price, new_price, self.user); save_vehicle(updated); msg += f"\n\nPrecio actualizado de {_fmt_usd(old_price)} a {_fmt_usd(new_price)}. Cotizaciones activas recalculadas: {changed}."
        QMessageBox.information(self,"Guardado",msg); self.accept(); return

    StageUpdateDialog._build_price_section = _patched_build_price_section
    StageUpdateDialog._price_calc = _patched_price_calc
    StageUpdateDialog._refresh_price_calc = _patched_refresh_price_calc
    StageUpdateDialog.save = _patched_stage_save


def mark_quote_won_and_vehicle_sold(quote_id: str, user: dict, device: DeviceInfo, precio_vendido_usd: Optional[float] = None, regalia_descripcion: str = "", regalia_usd: float = 0.0) -> tuple[bool, str]:
    if not user_has_permission(user, PERM_CLOSE_QUOTES): return False, "No tienes permiso para cerrar ventas."
    quotes = load_quotes(); idx = next((i for i,q in enumerate(quotes) if q.get("id") == quote_id), -1)
    if idx < 0: return False, "Cotización no encontrada."
    quote = quotes[idx]; vehicle = find_vehicle(quote.get("vehicle_id", ""))
    if not vehicle: return False, "Vehículo no encontrado."
    ensure_vehicle_runtime_fields(vehicle)
    precio = round(_to_float(precio_vendido_usd, 0) or _to_float(quote.get("leasing",{}).get("precio_vehiculo"),0) or _to_float(vehicle.get("precio_venta_usd"),0), 2)
    min_price = round(_to_float(vehicle.get("precio_minimo_usd"),0),2)
    if min_price and precio < min_price and not user_can_override_flow(user): return False, f"El precio vendido {_fmt_usd(precio)} es menor al precio mínimo aceptable {_fmt_usd(min_price)}. Requiere autorización."
    regalia = round(max(0.0, _to_float(regalia_usd, 0)), 2); calc = vehicle_profit_summary(vehicle, precio, regalia)
    if calc.get("ganancia_real_estimada_usd",0) < 0 and not user_can_override_flow(user): return False, "La venta proyecta pérdida. Requiere autorización de administración/supervisión."
    now = _now_iso(); venta_detalle = {**calc, "precio_vendido_cliente_usd": precio, "regalia_descripcion": regalia_descripcion.strip(), "regalia_usd": regalia, "fecha_venta": date.today().isoformat(), "ganancia_real_final_usd": calc.get("ganancia_real_estimada_usd",0), "margen_real_final_pct": calc.get("margen_real_estimado_pct",0)}
    quote["estado"] = QUOTE_GANADA; quote["fecha_cierre"] = date.today().isoformat(); quote["ultima_gestion"] = date.today().isoformat(); quote["venta_detalle"] = venta_detalle
    quote.setdefault("seguimientos", []).append({"fecha": now, "usuario": user.get("usuario", ""), "accion": "VENTA_CERRADA", "comentario": f"Venta cerrada por {_fmt_usd(precio)}. Regalía: {_fmt_usd(regalia)}. Ganancia real final: {_fmt_usd(venta_detalle['ganancia_real_final_usd'])}."})
    quotes[idx] = quote
    for i,q in enumerate(quotes):
        if q.get("vehicle_id") == vehicle.get("id") and q.get("id") != quote_id and q.get("estado") not in (QUOTE_GANADA, QUOTE_PERDIDA):
            q["estado"] = QUOTE_REOFERTAR; q["ultima_gestion"] = date.today().isoformat(); q.setdefault("seguimientos", []).append({"fecha": now, "usuario": user.get("usuario", ""), "accion": "VEHICULO_VENDIDO", "comentario": "El carro cotizado se vendió. Cliente queda para reofertar otro disponible."}); quotes[i]=q
    vehicle["estado_comercial"] = COMM_VENDIDO; vehicle["fecha_venta"] = date.today().isoformat(); vehicle["venta_quote_id"] = quote_id; vehicle["cliente"] = quote.get("cliente",{}).get("nombre", ""); vehicle["precio_venta_real_usd"] = precio; vehicle["venta_detalle"] = venta_detalle
    vehicle.setdefault("historial", []).append({"fecha": now, "usuario": user.get("usuario", ""), "computadora": device.computer_name, "accion": "VENTA_CERRADA", "detalle": f"Venta a {vehicle.get('cliente')} por {_fmt_usd(precio)} · regalía {_fmt_usd(regalia)} · ganancia {_fmt_usd(venta_detalle['ganancia_real_final_usd'])}"})
    ok1=save_vehicle(vehicle); ok2=save_quotes(quotes)
    if ok1 and ok2:
        log_audit("VENTA_CERRADA", user.get("usuario", ""), vehicle.get("codigo", ""), f"{vehicle.get('cliente','')} · {_fmt_usd(precio)}")
        return True, f"Venta cerrada. Ganancia real final: {_fmt_usd(venta_detalle['ganancia_real_final_usd'])}. Margen real final: {venta_detalle['margen_real_final_pct']}%."
    return False, "No se pudo cerrar la venta correctamente."


if PYSIDE_OK:
    def _patched_quote_detail_refresh(self):
        q=find_quote(self.quote_id)
        if not q: return
        self.q=q; cl=q.get("cliente",{}); snap=q.get("vehicle_snapshot",{}); s=_quote_financial_summary(q); avisos=q.get("avisos_precio",[]) or []; aviso_txt=""
        if avisos:
            last=avisos[-1]; aviso_txt=f"\n\n⚠ AVISO DE PRECIO:\nEl precio antes era {_fmt_usd(last.get('precio_anterior_usd'))} y ahora es {_fmt_usd(last.get('precio_nuevo_usd'))}. Fecha: {_fmt_date(last.get('fecha'))}."
        venta=q.get("venta_detalle") or {}; venta_txt=""
        if venta:
            venta_txt=f"\n\nVENTA CERRADA:\nPrecio vendido: {_fmt_usd(venta.get('precio_vendido_cliente_usd'))}\nRegalía: {_fmt_usd(venta.get('regalia_usd'))} · {venta.get('regalia_descripcion','')}\nGanancia real final: {_fmt_usd(venta.get('ganancia_real_final_usd'))}\nMargen real final: {venta.get('margen_real_final_pct',0)}%"
        info=f"""Cliente: {cl.get('nombre','')}
Teléfono: {cl.get('telefono','')}
Correo: {cl.get('correo','')}
Medio: {cl.get('medio_contacto','')}

Vehículo cotizado: {_vehicle_display_name(snap)} · {q.get('vehicle_code')}
Valor del vehículo: {_fmt_usd(s['valor_vehiculo'])}
Prima requerida: {_fmt_usd(s['prima_requerida'])}
Monto leasing: {_fmt_usd(s['monto_leasing'])}
Plazo: {s['plazo']} meses
Tasa: {s['tasa']}%
Cuota final con IVA: {_fmt_usd(s['cuota_final'])}
Gastos legales: {_fmt_usd(s['legal'])}

Estado: {QUOTE_STATUS_LABELS.get(q.get('estado'),q.get('estado'))}
Color seguimiento: {quote_alert_level(q)}
Días sin compra/gestión: {quote_days_without_purchase(q)}{aviso_txt}{venta_txt}

Carpeta del carro/propuestas:
{vehicle_proposals_folder(q.get('vehicle_code'), create=False) or ''}"""
        self.info.setPlainText(info); self.hist.setPlainText("\n".join([f"[{h.get('fecha')}] {h.get('usuario')} · {h.get('accion')} · {h.get('comentario')}" for h in q.get('seguimientos',[])]))

    def _patched_quote_mark_sold(self):
        q=getattr(self, 'q', None) or find_quote(self.quote_id)
        if not q: return
        v=find_vehicle(q.get('vehicle_id','')); default=_to_float(q.get('leasing',{}).get('precio_vehiculo'),0) or (_to_float(v.get('precio_venta_usd'),0) if v else 0)
        precio,ok=QInputDialog.getDouble(self,"Precio vendido","Precio real al que se vendió el vehículo:",default,0,9999999,2)
        if not ok: return
        desc,ok=QInputDialog.getText(self,"Regalía","Descripción de regalía / descuento / cortesía (dejar vacío si no aplica):")
        if not ok: return
        regalia,ok=QInputDialog.getDouble(self,"Valor regalía","Valor USD de la regalía a restar a la ganancia:",0,0,999999,2)
        if not ok: return
        if regalia>0 and not str(desc).strip(): QMessageBox.warning(self,"Regalía","Si hay valor de regalía, debes escribir una descripción."); return
        if QMessageBox.question(self,"Venta",f"¿Confirmas venta por {_fmt_usd(precio)}?\nRegalía: {_fmt_usd(regalia)}") != QMessageBox.StandardButton.Yes: return
        ok,msg=mark_quote_won_and_vehicle_sold(self.quote_id,self.user,self.device,precio,desc,regalia); QMessageBox.information(self,"Venta",msg) if ok else QMessageBox.warning(self,"Venta",msg); self.refresh()

    try:
        QuoteDetailDialog.refresh = _patched_quote_detail_refresh
        QuoteDetailDialog.mark_sold = _patched_quote_mark_sold
    except Exception:
        pass



# =============================================================================
# AJUSTES V4.7 - CALIDAD PROFESIONAL: VENTAS, LOCAL, DOCUMENTOS, REPORTERÍA INVENTARIO
# =============================================================================
APP_VERSION = "2.4.0_LEASING"

TIPO_COMPRA_LOTE = "LOTE"
TIPO_COMPRA_LOCAL = "LOCAL"
TIPO_COMPRA_IMPORTACION = "SERVICIO DE IMPORTACION"


def _v46_money(v: Any) -> float:
    return round(_to_float(v, 0), 2)


def _v46_stage_extra(vehicle: dict, stage_key: str) -> dict:
    st = vehicle_stage(vehicle, stage_key)
    return st.get("extra") if isinstance(st.get("extra"), dict) else {}


def _v46_date(value: Any) -> Optional[date]:
    return _parse_date(value)


def _v46_stage_start(vehicle: dict, stage_key: str) -> Optional[date]:
    return _parse_date(vehicle_stage(vehicle, stage_key).get("fecha_inicio"))


def _v46_stage_end(vehicle: dict, stage_key: str) -> Optional[date]:
    return _parse_date(vehicle_stage(vehicle, stage_key).get("fecha_fin"))


def _v46_date_from_extra(vehicle: dict, stage_key: str, key: str) -> Optional[date]:
    return _parse_date((_v46_stage_extra(vehicle, stage_key) or {}).get(key))


def _v46_tipo_compra(vehicle: dict) -> str:
    if bool(vehicle.get("servicio_importacion")) or _norm(vehicle.get("tipo_compra")) == _norm(TIPO_COMPRA_IMPORTACION):
        return TIPO_COMPRA_IMPORTACION
    if _norm(vehicle.get("tipo_compra")) == "LOCAL" or _norm(vehicle.get("lote")) == "LOCAL" or _norm(vehicle.get("subasta")) == "LOCAL":
        return TIPO_COMPRA_LOCAL
    return TIPO_COMPRA_LOTE


def _v46_report_status(vehicle: dict) -> str:
    ensure_vehicle_runtime_fields(vehicle)
    if vehicle.get("estado_comercial") == COMM_VENDIDO:
        return "VENDIDO"
    if vehicle.get("estado_comercial") == COMM_APARTADO:
        return "RESERVADO"
    if _v46_tipo_compra(vehicle) == TIPO_COMPRA_IMPORTACION:
        return "SERVICIO DE IMPORTACION"
    st = vehicle.get("estado_actual", STAGE_COMPRADO)
    if st in (STAGE_COMPRADO, STAGE_TRASLADO_USA, STAGE_TRANSITO):
        return "TRANSITO"
    if st == STAGE_ADUANA:
        return "ADUANA"
    if st == STAGE_PREPARACION:
        extra = _v46_stage_extra(vehicle, STAGE_PREPARACION)
        legal_done = bool(extra.get("legalizacion_fin") or extra.get("placas_entrega") or extra.get("numero_placa"))
        taller_done = bool(extra.get("taller_salida"))
        if not legal_done:
            return "LEGALIZACION"
        if not taller_done:
            return "REPARACION"
        return "REPARACION"
    if st in (STAGE_PRECIO_FINAL, STAGE_DISPONIBLE):
        return "DISPONIBLE VENTA"
    return STAGE_META.get(st, {}).get("label", st).upper()


def _v46_status_fill(status: str) -> str:
    s = _norm(status)
    if "VENDIDO" in s:
        return "92D050"      # verde vivo como el archivo original
    if "DISPONIBLE" in s:
        return "00B050"      # verde fuerte
    if "ADUANA" in s:
        return "F4B183"      # naranja claro
    if "LEGAL" in s:
        return "D9EAD3"      # verde tenue
    if "REPARACION" in s or "TALLER" in s:
        return "BDD7EE"      # azul claro
    if "TRANSITO" in s:
        return "FFD966"      # amarillo
    if "RESERVADO" in s:
        return "C6E0B4"
    if "IMPORTACION" in s:
        return "9BC2E6"
    return "F2F2F2"


def _v46_month_name(d: Optional[date]) -> str:
    if not d:
        return ""
    names = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    return f"{names[d.month-1]} {d.year}"


def _v46_vehicle_dates(vehicle: dict) -> dict:
    extra_t = _v46_stage_extra(vehicle, STAGE_TRASLADO_USA)
    extra_tr = _v46_stage_extra(vehicle, STAGE_TRANSITO)
    extra_ad = _v46_stage_extra(vehicle, STAGE_ADUANA)
    extra_pr = _v46_stage_extra(vehicle, STAGE_PREPARACION)
    compra = _parse_date(vehicle.get("fecha_compra"))
    traslado = _v46_stage_start(vehicle, STAGE_TRASLADO_USA)
    yarda = _parse_date(extra_t.get("fecha_llegada_yarda")) or _v46_stage_end(vehicle, STAGE_TRASLADO_USA)
    transito = _v46_stage_start(vehicle, STAGE_TRANSITO) or yarda
    aduana = _v46_stage_start(vehicle, STAGE_ADUANA)
    legal = _parse_date(extra_pr.get("emision_pedido")) or _parse_date(extra_pr.get("placas_ingreso")) or _v46_stage_start(vehicle, STAGE_PREPARACION)
    taller = _parse_date(extra_pr.get("taller_ingreso")) or _v46_stage_start(vehicle, STAGE_PREPARACION)
    disponible = _v46_stage_start(vehicle, STAGE_DISPONIBLE) or _v46_stage_end(vehicle, STAGE_PRECIO_FINAL) or _v46_stage_end(vehicle, STAGE_DISPONIBLE)
    venta = _vehicle_sale_date(vehicle)
    aduana_fin = _parse_date(extra_ad.get("fecha_liberacion_aduana")) or _v46_stage_end(vehicle, STAGE_ADUANA)
    legal_fin = _parse_date(extra_pr.get("legalizacion_fin")) or _parse_date(extra_pr.get("placas_entrega"))
    taller_fin = _parse_date(extra_pr.get("taller_salida"))
    return {"compra": compra, "traslado": traslado, "yarda": yarda, "transito": transito, "aduana": aduana, "aduana_fin": aduana_fin, "legal": legal, "legal_fin": legal_fin, "taller": taller, "taller_fin": taller_fin, "disponible": disponible, "venta": venta}


def _v46_days_between(a: Optional[date], b: Optional[date]) -> Optional[int]:
    if a and b:
        return max(0, (b - a).days)
    return None


def _v46_stage_days(vehicle: dict) -> dict:
    d = _v46_vehicle_dates(vehicle)
    return {
        "traslado": _v46_days_between(d["traslado"] or d["compra"], d["yarda"]),
        "transito": _v46_days_between(d["transito"] or d["yarda"], d["aduana"]),
        "aduana": _v46_days_between(d["aduana"], d["aduana_fin"]),
        "legal": _v46_days_between(d["legal"], d["legal_fin"]),
        "taller": _v46_days_between(d["taller"], d["taller_fin"]),
        "compra_disp": _v46_days_between(d["compra"], d["disponible"]),
        "compra_venta": _v46_days_between(d["compra"], d["venta"]),
        "sin_vender": _v46_days_between(d["disponible"], date.today()) if not d["venta"] else None,
    }


def _v46_sales_monthly(vehicles: list[dict]) -> dict[str, list[dict]]:
    out: dict[str, list[dict]] = {}
    for v in vehicles:
        sale = _vehicle_sale_date(v)
        if not sale:
            continue
        key = _v46_month_name(sale)
        out.setdefault(key, []).append(v)
    return out


def _v46_status_summary(vehicles: list[dict]) -> list[tuple[str, int]]:
    counts: dict[str, int] = {}
    for v in vehicles:
        st = _v46_report_status(v)
        counts[st] = counts.get(st, 0) + 1
    sold_months = _v46_sales_monthly(vehicles)
    order = ["TRANSITO", "REPARACION", "DISPONIBLE VENTA", "LEGALIZACION", "ADUANA", "SERVICIO DE IMPORTACION", "RESERVADO"]
    rows = [(x, counts.get(x, 0)) for x in order if counts.get(x, 0)]
    for month, vs in sorted(sold_months.items()):
        rows.append(("VENDIDO " + month.upper(), len(vs)))
    rows.append(("VENDIDO TOTAL", sum(1 for v in vehicles if v.get("estado_comercial") == COMM_VENDIDO)))
    rows.append(("INVENTARIO TOTAL", len(vehicles)))
    return rows


def _v46_apply_table_hardening(table: Any):
    try:
        table.setMouseTracking(False)
        table.viewport().setMouseTracking(False)
        table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        table.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        table.setDragEnabled(False)
        table.setAlternatingRowColors(False)
    except Exception:
        pass
    return table


# --- Compra LOCAL / Servicio de importación ---------------------------------
_ORIG_CREATE_VEHICLE_PURCHASE_V46 = create_vehicle_purchase

def create_vehicle_purchase(data: dict, comprobante_src: Path, user: dict, device: DeviceInfo) -> tuple[bool, str, str]:
    tipo = _norm(data.get("tipo_compra") or TIPO_COMPRA_LOTE)
    is_local = tipo == "LOCAL"
    is_import_service = tipo == _norm(TIPO_COMPRA_IMPORTACION) or bool(data.get("servicio_importacion"))
    if is_local:
        data = {**data, "lote": "LOCAL", "subasta": "LOCAL", "estado_usa": data.get("pais_compra") or data.get("estado_usa") or "LOCAL"}
    ok, msg, vid = _ORIG_CREATE_VEHICLE_PURCHASE_V46(data, comprobante_src, user, device)
    if not ok or not vid:
        return ok, msg, vid
    v = find_vehicle(vid)
    if not v:
        return ok, msg, vid
    ensure_vehicle_runtime_fields(v)
    v["tipo_compra"] = TIPO_COMPRA_LOCAL if is_local else (TIPO_COMPRA_IMPORTACION if is_import_service else TIPO_COMPRA_LOTE)
    v["pais_compra"] = data.get("pais_compra") or data.get("estado_usa", "")
    v["servicio_importacion"] = bool(is_import_service)
    if is_local:
        fecha = data.get("fecha_compra") or date.today().isoformat()
        v["estado_actual"] = STAGE_DISPONIBLE
        v["estado_comercial"] = COMM_DISPONIBLE
        if _to_float(data.get("precio_venta_local_usd"), 0) > 0:
            precio = _to_float(data.get("precio_venta_local_usd"), 0)
            v["precio_venta_usd"] = precio
            v.setdefault("precio_final", {})["precio_venta_cliente_usd"] = precio
            v.setdefault("precio_final", {})["precio_minimo_usd"] = _to_float(data.get("precio_minimo_local_usd"), 0) or precio
        for sk in STAGE_ORDER:
            st = vehicle_stage(v, sk)
            if sk == STAGE_COMPRADO:
                st.update({"status": "COMPLETADO", "fecha_inicio": fecha, "fecha_fin": fecha})
            elif sk == STAGE_DISPONIBLE:
                st.update({"status": "COMPLETADO", "fecha_inicio": fecha, "fecha_fin": fecha, "proveedor": "COMPRA LOCAL", "comentario": "Compra local registrada. Se omiten traslado USA, tránsito y aduana."})
            elif sk not in (STAGE_PRECIO_FINAL,):
                st.update({"status": "OMITIDO", "fecha_inicio": None, "fecha_fin": None, "comentario": "No aplica por compra local."})
        v.setdefault("historial", []).append({"fecha": _now_iso(), "usuario": user.get("usuario", ""), "computadora": device.computer_name, "accion": "COMPRA_LOCAL", "detalle": "Compra local registrada y disponible para la venta."})
        save_vehicle(v)
        msg += "\nCompra LOCAL: se omitió el flujo USA/Aduana y el vehículo quedó disponible."
    elif is_import_service:
        v.setdefault("historial", []).append({"fecha": _now_iso(), "usuario": user.get("usuario", ""), "computadora": device.computer_name, "accion": "SERVICIO_IMPORTACION", "detalle": "Vehículo marcado como servicio de importación."})
        save_vehicle(v)
    return True, msg, vid


# --- Reporte Inventario desde cero con diseño L&M ----------------------------
def _v46_xl_styles():
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    thin = Side(style="thin", color="808080")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    return Alignment, Border, Font, PatternFill, Side, border


def _v46_cell(ws, row: int, col: int, value: Any = None, fill: Optional[str] = None, color: str = "000000", bold: bool = False, size: int = 10, align: str = "center"):
    Alignment, Border, Font, PatternFill, Side, border = _v46_xl_styles()
    c = ws.cell(row, col)
    if value is not None:
        c.value = value
    c.font = Font(name="Calibri", size=size, bold=bold, color=color)
    c.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
    c.border = border
    if fill:
        c.fill = PatternFill("solid", fgColor=fill)
    return c


def _v46_merge(ws, rng: str, value: str, fill: str = LYM_NAVY, color: str = "FFFFFF", size: int = 11):
    ws.merge_cells(rng)
    c = ws[rng.split(":")[0]]
    c.value = value
    _style_range(ws, rng, fill, color, True, size, "center")


def _v46_fill_status_cell(cell, status: str):
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    side = Side(style="thin", color="808080")
    cell.fill = PatternFill("solid", fgColor=_v46_status_fill(status))
    cell.font = Font(name="Calibri", size=10, bold=True, color="000000")
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = Border(left=side, right=side, top=side, bottom=side)


def generate_inventory_excel(vehicles: list[dict], user: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    out_dir = _report_output_dir()
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    try:
        from openpyxl import Workbook
        from openpyxl.chart import BarChart, LineChart, PieChart, Reference
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except Exception:
        return False, "Para generar Excel instala: pip install openpyxl", None
    vehicles = [ensure_vehicle_runtime_fields(v) for v in vehicles]
    out = out_dir / f"Inventario_vehiculos_LYM_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Inventario General"
    dash = wb.create_sheet("Dashboard Gerencial")
    detail = wb.create_sheet("Reportes Detalle")
    hist = wb.create_sheet("Historial Disponible")

    # Hoja 1: Inventario General con formato del ejemplo y sin relleno verde en toda la fila.
    ws.sheet_view.showGridLines = False
    for col in range(1, 30):
        ws.column_dimensions[get_column_letter(col)].width = 13
    custom_widths = {1:5,2:13,3:13,4:13,5:14,6:18,7:8,8:18,9:32,10:20,11:8,12:15,13:16,14:15,15:15,16:15,17:15,18:15,19:10,20:10,21:10,22:10,23:10,24:12,25:12,26:12,27:12,28:13,29:12}
    for c,w in custom_widths.items():
        ws.column_dimensions[get_column_letter(c)].width = w
    ws.row_dimensions[2].height = 22
    ws.row_dimensions[3].height = 25
    ws.row_dimensions[4].height = 30
    _v46_merge(ws, "A2:AC2", f"INVENTARIO Y ESTATUS DE VEHÍCULOS GENERAL { _spanish_month_year().upper() }", LYM_ORANGE, "000000", 12)
    headers = {"A3":"No.","B3":"LOT","C3":"FECHA DE COMPRA","D3":"FECHA DE VENTA","E3":"VEHÍCULO","H3":"ESTATUS","I3":"OBSERVACIONES","J3":"SUMMARY","L3":"FECHAS LINEALES DEL PROCESO","S3":"DÍAS POR ETAPA Y TOTALES"}
    for cell, value in headers.items():
        ws[cell] = value
    for rng in ["A3:A4","B3:B4","C3:C4","D3:D4","E3:G3","H3:H4","I3:I4","J3:K3","L3:R3","S3:AC3"]:
        try: ws.merge_cells(rng)
        except Exception: pass
    row4 = {5:"MARCA",6:"MODELO",7:"AÑO",10:"ESTATUS",11:"CANTIDAD",12:"FECHA TRASLADO GRÚA",13:"LLEGADA YARDA",14:"INICIO TRÁNSITO",15:"LLEGADA ADUANA",16:"INICIO LEGALIZACIÓN",17:"INGRESO TALLER",18:"DISPONIBLE VENTA",19:"DÍAS TRASLADO",20:"DÍAS TRÁNSITO",21:"DÍAS ADUANA",22:"DÍAS LEGALIZACIÓN",23:"DÍAS TALLER",24:"COMPRA → DISPONIBLE",25:"COMPRA → VENTA",26:"DÍAS SIN VENDER",27:"SEMÁFORO",28:"MES VENTA",29:"TIPO COMPRA"}
    for c, val in row4.items(): ws.cell(4, c).value = val
    _style_range(ws, "A3:AC4", LYM_NAVY, "FFFFFF", True, 9, "center")
    summary_rows = _v46_status_summary(vehicles)
    for i, (label, count) in enumerate(summary_rows[:28], start=5):
        ws.cell(i, 10).value = label
        ws.cell(i, 11).value = count
        _v46_fill_status_cell(ws.cell(i,10), label)
        _v46_cell(ws, i, 11, count, "F2F2F2", "000000", True, 10, "center")

    for idx, v in enumerate(vehicles, start=1):
        r = 5 + idx - 1
        d = _v46_vehicle_dates(v)
        days = _v46_stage_days(v)
        status = _v46_report_status(v)
        sale = d["venta"]
        rowvals = {1: idx, 2: v.get("lote") or ("LOCAL" if _v46_tipo_compra(v)==TIPO_COMPRA_LOCAL else ""), 3: d["compra"], 4: sale, 5: v.get("marca"), 6: v.get("modelo"), 7: v.get("anio"), 8: status, 9: v.get("cliente") or v.get("observaciones", ""), 12: d["traslado"], 13: d["yarda"], 14: d["transito"], 15: d["aduana"], 16: d["legal"], 17: d["taller"], 18: d["disponible"], 19: days["traslado"], 20: days["transito"], 21: days["aduana"], 22: days["legal"], 23: days["taller"], 24: days["compra_disp"], 25: days["compra_venta"], 26: days["sin_vender"], 27: "CERRADO" if sale else stage_alert_level(v), 28: _v46_month_name(sale), 29: _v46_tipo_compra(v)}
        for c, val in rowvals.items():
            cell = _v46_cell(ws, r, c, val, None, "000000", False, 10, "center")
            if c in (3,4,12,13,14,15,16,17,18): cell.number_format = "dd/mm/yyyy"
        # Celdas grises claras para lectura, status único con color.
        for c in list(range(1, 10)) + list(range(12, 30)):
            ws.cell(r,c).fill = PatternFill("solid", fgColor="F8FAFC")
        _v46_fill_status_cell(ws.cell(r, 8), status)
        if sale: _v46_fill_status_cell(ws.cell(r, 27), "VENDIDO")
        elif stage_alert_level(v) == "ROJO": _v46_fill_status_cell(ws.cell(r, 27), "CRÍTICO")
        elif stage_alert_level(v) == "AMARILLO": _v46_fill_status_cell(ws.cell(r, 27), "ALERTA")
        else: _v46_fill_status_cell(ws.cell(r, 27), "OK")
    # Mantener el bloque visual completo: si el resumen (J:K) tiene más líneas que el inventario,
    # se rellenan las celdas vacías con gris claro y bordes para evitar filas/celdas en blanco.
    end_row = max(5, 4 + len(vehicles), 4 + len(summary_rows))
    for rr in range(5, end_row + 1):
        for cc in range(1, 30):
            cell = ws.cell(rr, cc)
            if cell.value is None:
                _v46_cell(ws, rr, cc, "", "F8FAFC", "000000", False, 10, "center")
    ws.freeze_panes = "A5"
    ws.auto_filter.ref = f"A4:AC{end_row}"

    # Dashboard Gerencial con KPI y gráficas como el ejemplo.
    dash.sheet_view.showGridLines = False
    for c in range(1, 12): dash.column_dimensions[get_column_letter(c)].width = 18
    sold = [v for v in vehicles if v.get("estado_comercial") == COMM_VENDIDO]
    active = [v for v in vehicles if v.get("estado_comercial") != COMM_VENDIDO and v.get("estado_actual") != STAGE_ANULADO]
    disponibles = [v for v in vehicles if _v46_report_status(v) == "DISPONIBLE VENTA"]
    sale_days = []
    for v in sold:
        d = _v46_vehicle_dates(v)
        if d["compra"] and d["venta"]: sale_days.append((d["venta"] - d["compra"]).days)
    avg_sale = round(sum(sale_days)/len(sale_days), 1) if sale_days else 0
    marca_counts = _counter_by(sold, lambda v: v.get("marca"))
    marca_mas = ", ".join([k for k,n in sorted(marca_counts.items(), key=lambda x: x[1], reverse=True)[:2]]) if marca_counts else ""
    brand_speed: dict[str, list[int]] = {}
    for v in sold:
        d = _v46_vehicle_dates(v)
        if d["compra"] and d["venta"]: brand_speed.setdefault(str(v.get("marca") or ""), []).append((d["venta"] - d["compra"]).days)
    marca_fast = min(brand_speed.items(), key=lambda kv: sum(kv[1])/len(kv[1]))[0] if brand_speed else ""
    criticos = sum(1 for v in active if (vehicle_days_from_purchase(v) or 0) >= 120 or stage_alert_level(v) == "ROJO")
    kpis = [("Total unidades", len(vehicles)), ("Inventario activo", len(active)), ("Vendidos total", len(sold)), ("En tránsito", sum(1 for v in active if _v46_report_status(v)=="TRANSITO")), ("Disponibles", len(disponibles)), ("Prom. días venta", avg_sale), ("Marca más vendida", marca_mas), ("Marca más rápida", marca_fast), ("Críticos +120", criticos)]
    for i, (lab, val) in enumerate(kpis, start=1):
        _v46_cell(dash, 4, i, lab, "FFF2CC", "FFFFFF", True, 10, "center")
        _v46_cell(dash, 5, i, val, "FFF2CC", LYM_NAVY, True, 11, "center")
    status_active = _counter_by(active, _v46_report_status)
    status_general = _counter_by(vehicles, _v46_report_status)
    dash.merge_cells("A8:C8"); dash["A8"]="ESTATUS ACTIVO — 100% SIN VENDIDOS"; _style_range(dash,"A8:C8","548235","FFFFFF",True,10,"center")
    dash.merge_cells("H8:J8"); dash["H8"]="ESTATUS GENERAL — 100% INCLUYE VENDIDOS"; _style_range(dash,"H8:J8","9E480E","FFFFFF",True,10,"center")
    for col, text in enumerate(["Estatus", "Cantidad", "% activo"], 1): _v46_cell(dash, 9, col, text, LYM_NAVY, "FFFFFF", True, 10, "center")
    for col, text in enumerate(["Estatus", "Cantidad", "% total"], 8): _v46_cell(dash, 9, col, text, LYM_NAVY, "FFFFFF", True, 10, "center")
    r = 10
    for st, n in sorted(status_active.items(), key=lambda x: x[1], reverse=True):
        _v46_cell(dash, r, 1, st); _v46_cell(dash, r, 2, n); _v46_cell(dash, r, 3, n/len(active) if active else 0); dash.cell(r,3).number_format='0.0%'; r += 1
    r2 = 10
    for st, n in sorted(status_general.items(), key=lambda x: x[1], reverse=True):
        _v46_cell(dash, r2, 8, st); _v46_cell(dash, r2, 9, n); _v46_cell(dash, r2, 10, n/len(vehicles) if vehicles else 0); dash.cell(r2,10).number_format='0.0%'; r2 += 1
    # Gráficas tipo dona para que el Dashboard se vea igual que la presentación gerencial HTML.
    if r > 10:
        chart = DoughnutChart()
        chart.title = "Estatus activo sin vendidos"
        chart.holeSize = 58
        chart.firstSliceAng = 270
        chart.add_data(Reference(dash, min_col=2, min_row=9, max_row=r-1), titles_from_data=True)
        chart.set_categories(Reference(dash, min_col=1, min_row=10, max_row=r-1))
        chart.height = 7
        chart.width = 13
        dash.add_chart(chart, "D8")
    if r2 > 10:
        chart2 = DoughnutChart()
        chart2.title = "Estatus general"
        chart2.holeSize = 58
        chart2.firstSliceAng = 270
        chart2.add_data(Reference(dash, min_col=9, min_row=9, max_row=r2-1), titles_from_data=True)
        chart2.set_categories(Reference(dash, min_col=8, min_row=10, max_row=r2-1))
        chart2.height = 7
        chart2.width = 13
        dash.add_chart(chart2, "K8")
    sales_by_month = _v46_sales_monthly(vehicles)
    dash.merge_cells("A25:F25"); dash["A25"] = "VENTAS POR MES, % Y VARIACIÓN MENSUAL"; _style_range(dash,"A25:F25",LYM_NAVY,"FFFFFF",True,10,"center")
    for c,h in enumerate(["Mes","Vendidos","% vendidos","Variación vs mes anterior","Tendencia","Marcas vendidas"],1): _v46_cell(dash,26,c,h,LYM_NAVY,"FFFFFF",True,10,"center")
    prev = None; row = 27
    for month, vs in sorted(sales_by_month.items(), key=lambda kv: _parse_date(kv[1][0].get("fecha_venta")) or date.today()):
        n=len(vs); pct=n/len(sold) if sold else 0; var="Sin comparación" if prev is None else ((n-prev)/prev if prev else None); tendencia="Sin comparación" if prev is None else ("Subió" if n>prev else ("Bajó" if n<prev else "Se mantuvo")); marcas=_counter_by(vs, lambda v:v.get('marca'))
        vals=[month.title(), n, pct, var if isinstance(var,(int,float)) else var, tendencia, ", ".join(f"{m} ({c})" for m,c in sorted(marcas.items()))]
        for c,val in enumerate(vals,1): _v46_cell(dash,row,c,val)
        dash.cell(row,3).number_format='0.0%'
        if isinstance(var,(int,float)): dash.cell(row,4).number_format='0.0%'
        prev=n; row += 1
    if row > 27:
        line=LineChart(); line.title="Unidades vendidas por mes"; line.add_data(Reference(dash,min_col=2,min_row=26,max_row=row-1),titles_from_data=True); line.set_categories(Reference(dash,min_col=1,min_row=27,max_row=row-1)); line.height=9; line.width=16; dash.add_chart(line,"H25")

    # Detalle de tiempos por fase para auditoría de números.
    detail.sheet_view.showGridLines = False
    _v46_merge(detail, "A1:M1", "DETALLE DE TIEMPOS POR VEHÍCULO", LYM_NAVY, "FFFFFF", 14)
    det_headers = ["Código", "Vehículo", "Compra", "Traslado grúa", "Yarda", "Tránsito", "Aduana", "Legalización", "Taller", "Disponible", "Venta", "Días compra→disp", "Días compra→venta"]
    for c,h in enumerate(det_headers,1): _v46_cell(detail,3,c,h,LYM_NAVY,"FFFFFF",True,10,"center")
    for idx,v in enumerate(vehicles,start=4):
        d=_v46_vehicle_dates(v); dy=_v46_stage_days(v); vals=[v.get('codigo'), _vehicle_display_name(v), d['compra'], d['traslado'], d['yarda'], d['transito'], d['aduana'], d['legal'], d['taller'], d['disponible'], d['venta'], dy['compra_disp'], dy['compra_venta']]
        for c,val in enumerate(vals,1):
            cell=_v46_cell(detail,idx,c,val, "F8FAFC")
            if c in range(3,12): cell.number_format='dd/mm/yyyy'
    detail.freeze_panes='A4'
    for c in range(1,14): detail.column_dimensions[get_column_letter(c)].width=16

    # Historial Disponible semanal derivado de las fechas del sistema.
    hist.sheet_view.showGridLines = False
    _v46_merge(hist, "A1:L1", "HISTORIAL DISPONIBLE PARA LA VENTA", LYM_NAVY, "FFFFFF", 14)
    hist.merge_cells("A2:L2"); hist["A2"] = "Semanas calculadas desde las fechas de disponible y venta registradas en el sistema."; _style_range(hist,"A2:L2",LYM_LIGHT,LYM_NAVY,True,10,"center")
    for c,h in enumerate(["Semana", "Desde", "Hasta", "Disponibles", "Vendidos acumulados", "Capital disponible", "Ganancia esperada", "Vehículos disponibles", "Observaciones"],1): _v46_cell(hist,4,c,h,LYM_NAVY,"FFFFFF",True,10,"center")
    all_dates = [d for v in vehicles for d in [_v46_vehicle_dates(v)["compra"], _v46_vehicle_dates(v)["disponible"], _v46_vehicle_dates(v)["venta"]] if d]
    start_week = min(all_dates) if all_dates else date.today()
    start_week = start_week - timedelta(days=start_week.weekday())
    today = date.today()
    row = 5; week_no=1
    cur = start_week
    while cur <= today:
        end = cur + timedelta(days=6)
        avail=[]
        for v in vehicles:
            d=_v46_vehicle_dates(v)
            if d["disponible"] and d["disponible"] <= end and (not d["venta"] or d["venta"] > cur):
                avail.append(v)
        vals=[f"Semana {week_no}", cur, end, len(avail), sum(1 for v in vehicles if (_v46_vehicle_dates(v)["venta"] and _v46_vehicle_dates(v)["venta"] <= end)), sum(_to_float(v.get("precio_venta_usd"),0) for v in avail), sum(vehicle_expected_profit(v) for v in avail), ", ".join((v.get("modelo") or v.get("codigo") or "") for v in avail[:8]), ""]
        for c,val in enumerate(vals,1):
            cell=_v46_cell(hist,row,c,val,"F8FAFC")
            if c in (2,3): cell.number_format='dd/mm/yyyy'
            if c in (6,7): cell.number_format='$#,##0.00'
        cur += timedelta(days=7); row += 1; week_no += 1
    hist.freeze_panes='A5'
    for c in range(1,13): hist.column_dimensions[get_column_letter(c)].width=18
    wb.save(out)
    log_audit("GENERAR_EXCEL_INVENTARIO_V46", (user or {}).get("usuario", ""), "", out.name)
    return True, "Inventario generado desde cero con formato L&M, Dashboard Gerencial, Reportes Detalle e Historial Disponible semanal.", out


# --- Documentos, índice PDF y totalización -----------------------------------
_ORIG_VEHICLE_DOCUMENT_ENTRIES_V46 = vehicle_document_entries

def vehicle_document_entries(vehicle: dict) -> list[dict]:
    entries = list(_ORIG_VEHICLE_DOCUMENT_ENTRIES_V46(vehicle) or [])
    venta = vehicle.get("venta_detalle") or {}
    if venta.get("comprobante_venta"):
        entries.append({"stage_key": "VENTA", "etapa": "Venta", "categoria": "VENTA", "subcategoria": "COMPROBANTE VENTA", "descripcion": "Comprobante PDF de venta", "monto_usd": venta.get("precio_vendido_cliente_usd", 0), "proveedor": venta.get("cliente", vehicle.get("cliente", "")), "oc_numero": "", "comprobante": venta.get("comprobante_venta"), "comprobante_nombre": venta.get("comprobante_venta_nombre", "comprobante_venta.pdf"), "oc_documento": "", "oc_documento_nombre": ""})
    if venta.get("foto_cliente"):
        entries.append({"stage_key": "VENTA", "etapa": "Venta", "categoria": "VENTA", "subcategoria": "FOTO CLIENTE", "descripcion": "Foto del cliente con el vehículo", "monto_usd": 0, "proveedor": venta.get("cliente", vehicle.get("cliente", "")), "oc_numero": "", "comprobante": venta.get("foto_cliente"), "comprobante_nombre": venta.get("foto_cliente_nombre", "foto_cliente.jpg"), "oc_documento": "", "oc_documento_nombre": ""})
    for p in vehicle.get("papeles_legales", []) or []:
        entries.append({"stage_key": "LEGAL", "etapa": "Papeles legales", "categoria": "LEGAL", "subcategoria": p.get("tipo", "PAPEL LEGAL"), "descripcion": p.get("descripcion", ""), "monto_usd": 0, "proveedor": "", "oc_numero": "", "comprobante": p.get("documento", ""), "comprobante_nombre": p.get("documento_nombre", "papel_legal.pdf"), "oc_documento": "", "oc_documento_nombre": ""})
    return entries


def _v46_write_doc_index_pdf(vehicle: dict, target: Path, exported: list[Path]) -> Optional[Path]:
    target.mkdir(parents=True, exist_ok=True)
    out = target / f"INDICE_DOCUMENTOS_{_safe_filename(vehicle.get('codigo','VEHICULO'))}.pdf"
    expected = vehicle_document_entries(vehicle)
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        doc = SimpleDocTemplate(str(out), pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet(); title = ParagraphStyle("LymTitle", parent=styles["Title"], textColor=colors.HexColor("#08285A"))
        story = [Paragraph(f"Índice de documentación · {vehicle.get('codigo','')}", title), Paragraph(_vehicle_display_name(vehicle), styles["Heading2"]), Spacer(1, 10)]
        data = [["#", "Etapa", "Documento", "Comprobante", "OC"]]
        for i, e in enumerate(expected, 1):
            data.append([i, e.get("etapa",""), e.get("subcategoria",""), "SUBIDO" if e.get("comprobante") else "FALTA", "SUBIDO" if e.get("oc_documento") else "FALTA"])
        tbl = Table(data, repeatRows=1, colWidths=[28, 105, 195, 80, 70])
        tbl.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#08285A")), ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("GRID", (0,0), (-1,-1), 0.35, colors.grey), ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("FONTSIZE", (0,0), (-1,-1), 8)]))
        story.append(tbl); story.append(Spacer(1, 12)); story.append(Paragraph(f"Archivos exportados: {len(exported)}", styles["Normal"]))
        doc.build(story)
        return out
    except Exception:
        txt = out.with_suffix(".txt")
        lines = [f"INDICE DOCUMENTOS {vehicle.get('codigo','')}", _vehicle_display_name(vehicle), ""]
        for i,e in enumerate(expected,1):
            lines.append(f"{i}. {e.get('etapa')} · {e.get('subcategoria')} · Comprobante: {'SUBIDO' if e.get('comprobante') else 'FALTA'} · OC: {'SUBIDO' if e.get('oc_documento') else 'FALTA'}")
        txt.write_text("\n".join(lines), encoding="utf-8")
        return txt


def export_vehicle_documents(vehicle_id: str, target_dir: Path) -> tuple[bool, str, list[Path]]:
    v = find_vehicle(vehicle_id)
    if not v:
        return False, "Vehículo no encontrado.", []
    base = target_dir / _safe_filename(v.get("codigo", "VEHICULO"))
    exported: list[Path] = []
    for idx, item in enumerate(vehicle_document_entries(v), start=1):
        prefix = _safe_filename(f"{idx:02d}_{item.get('etapa','')}_{item.get('subcategoria','')}")
        if item.get("comprobante"):
            out = decrypt_file_to_path(item.get("comprobante", ""), f"{prefix}_COMPROBANTE_{item.get('comprobante_nombre','comprobante.pdf')}", base)
            if out: exported.append(out)
        if item.get("oc_documento"):
            out = decrypt_file_to_path(item.get("oc_documento", ""), f"{prefix}_OC_{item.get('oc_documento_nombre','oc.pdf')}", base)
            if out: exported.append(out)
    idx = _v46_write_doc_index_pdf(v, base, exported)
    if idx: exported.append(idx)
    return True, f"Se exportaron {len(exported)} archivos. Incluye índice/desglose de documentos y faltantes.", exported


# --- Venta avanzada con comprobante, foto, fecha y protección contra doble venta
_ORIG_MARK_QUOTE_WON_V46 = mark_quote_won_and_vehicle_sold

def mark_quote_won_and_vehicle_sold(quote_id: str, user: dict, device: DeviceInfo, precio_vendido_usd: Optional[float] = None, regalia_descripcion: str = "", regalia_usd: float = 0.0, fecha_venta: Optional[str] = None, comprobante_venta_src: str = "", foto_cliente_src: str = "") -> tuple[bool, str]:
    if not user_has_permission(user, PERM_CLOSE_QUOTES):
        return False, "No tienes permiso para cerrar ventas."
    quotes = load_quotes(); idx = next((i for i,q in enumerate(quotes) if q.get("id") == quote_id), -1)
    if idx < 0: return False, "Cotización no encontrada."
    quote = quotes[idx]; vehicle = find_vehicle(quote.get("vehicle_id", ""))
    if not vehicle: return False, "Vehículo no encontrado."
    ensure_vehicle_runtime_fields(vehicle)
    if vehicle.get("estado_comercial") == COMM_VENDIDO or vehicle.get("venta_detalle"):
        return False, "Este vehículo ya fue vendido. No se puede cerrar una segunda venta."
    precio = round(_to_float(precio_vendido_usd, 0) or _to_float(quote.get("leasing",{}).get("precio_vehiculo"),0) or _to_float(vehicle.get("precio_venta_usd"),0), 2)
    min_price = round(_to_float(vehicle.get("precio_minimo_usd"),0),2)
    if min_price and precio < min_price and not user_can_override_flow(user):
        return False, f"El precio vendido {_fmt_usd(precio)} es menor al precio mínimo aceptable {_fmt_usd(min_price)}. Requiere autorización."
    regalia = round(max(0.0, _to_float(regalia_usd, 0)), 2)
    calc = vehicle_profit_summary(vehicle, precio, regalia)
    if calc.get("ganancia_real_estimada_usd",0) < 0:
        return False, "La venta proyecta pérdida. No se puede cerrar sin revisar precio, costos o autorización administrativa."
    fventa = _parse_date(fecha_venta) or date.today()
    comp_rel = comp_name = foto_rel = foto_name = ""
    if comprobante_venta_src:
        comp_rel, comp_name = store_document_named(Path(comprobante_venta_src), vehicle.get("codigo", "VEHICULO"), "VENTA", f"{vehicle.get('codigo')}_COMPROBANTE_VENTA_{fventa.isoformat()}", Path(comprobante_venta_src).name)
    if foto_cliente_src:
        foto_rel, foto_name = store_document_named(Path(foto_cliente_src), vehicle.get("codigo", "VEHICULO"), "FOTO_CLIENTE", f"{vehicle.get('codigo')}_FOTO_CLIENTE_{fventa.isoformat()}", Path(foto_cliente_src).name)
    now = _now_iso(); venta_detalle = {**calc, "precio_vendido_cliente_usd": precio, "regalia_descripcion": regalia_descripcion.strip(), "regalia_usd": regalia, "fecha_venta": fventa.isoformat(), "ganancia_real_final_usd": calc.get("ganancia_real_estimada_usd",0), "margen_real_final_pct": calc.get("margen_real_estimado_pct",0), "cliente": quote.get("cliente",{}).get("nombre", ""), "comprobante_venta": comp_rel, "comprobante_venta_nombre": comp_name, "foto_cliente": foto_rel, "foto_cliente_nombre": foto_name}
    quote["estado"] = QUOTE_GANADA; quote["fecha_cierre"] = fventa.isoformat(); quote["ultima_gestion"] = fventa.isoformat(); quote["venta_detalle"] = venta_detalle
    quote.setdefault("seguimientos", []).append({"fecha": now, "usuario": user.get("usuario", ""), "accion": "VENTA_CERRADA", "comentario": f"Venta cerrada por {_fmt_usd(precio)}. Regalía: {_fmt_usd(regalia)}. Ganancia real final: {_fmt_usd(venta_detalle['ganancia_real_final_usd'])}."})
    quotes[idx] = quote
    for i,q in enumerate(quotes):
        if q.get("vehicle_id") == vehicle.get("id") and q.get("id") != quote_id and q.get("estado") not in (QUOTE_GANADA, QUOTE_PERDIDA):
            q["estado"] = QUOTE_REOFERTAR; q["ultima_gestion"] = fventa.isoformat(); q.setdefault("seguimientos", []).append({"fecha": now, "usuario": user.get("usuario", ""), "accion": "VEHICULO_VENDIDO", "comentario": "El carro cotizado se vendió. Cliente queda para reofertar otro disponible."}); quotes[i]=q
    vehicle["estado_comercial"] = COMM_VENDIDO; vehicle["fecha_venta"] = fventa.isoformat(); vehicle["venta_quote_id"] = quote_id; vehicle["cliente"] = quote.get("cliente",{}).get("nombre", ""); vehicle["precio_venta_real_usd"] = precio; vehicle["venta_detalle"] = venta_detalle
    vehicle.setdefault("historial", []).append({"fecha": now, "usuario": user.get("usuario", ""), "computadora": device.computer_name, "accion": "VENTA_CERRADA", "detalle": f"Venta a {vehicle.get('cliente')} por {_fmt_usd(precio)} · regalía {_fmt_usd(regalia)} · ganancia {_fmt_usd(venta_detalle['ganancia_real_final_usd'])}"})
    ok1=save_vehicle(vehicle); ok2=save_quotes(quotes)
    if ok1 and ok2:
        log_audit("VENTA_CERRADA", user.get("usuario", ""), vehicle.get("codigo", ""), f"{vehicle.get('cliente','')} · {_fmt_usd(precio)}")
        return True, f"Venta cerrada. Ganancia real final: {_fmt_usd(venta_detalle['ganancia_real_final_usd'])}. Margen real final: {venta_detalle['margen_real_final_pct']}%."
    return False, "No se pudo cerrar la venta correctamente."


if PYSIDE_OK:
    class SaleClosingDialog(QDialog):
        def __init__(self, parent, quote: dict, vehicle: dict):
            super().__init__(parent); self.quote=quote; self.vehicle=vehicle; self.comp_path=""; self.photo_path=""
            self.setWindowTitle("Cerrar venta del vehículo"); self.setMinimumSize(760, 520); self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Cerrar venta", "Registra precio real, fecha de venta, regalía y documentos de respaldo."))
            form=QFormLayout(); lay.addLayout(form)
            default_price = _to_float(self.quote.get("leasing",{}).get("precio_vehiculo"),0) or _to_float(self.vehicle.get("precio_venta_usd"),0)
            self.fecha=configure_date_edit(QDateEdit(), date.today().isoformat(), self.vehicle.get("fecha_compra"))
            self.precio=MoneyEdit(); self.precio.setRange(0,9999999); self.precio.setValue(default_price)
            self.regalia_desc=QLineEdit(); self.regalia_desc.setPlaceholderText("Ejemplo: polarizado, mantenimiento, descuento comercial")
            self.regalia=MoneyEdit(); self.regalia.setRange(0,999999); self.regalia.setValue(0)
            self.comp_lbl=QLineEdit(); self.comp_lbl.setReadOnly(True); bcomp=QPushButton("Subir comprobante PDF venta"); bcomp.setObjectName("ghost"); bcomp.clicked.connect(self.pick_comp)
            self.photo_lbl=QLineEdit(); self.photo_lbl.setReadOnly(True); bphoto=QPushButton("Subir foto cliente con carro"); bphoto.setObjectName("ghost"); bphoto.clicked.connect(self.pick_photo)
            row1=QWidget(); l1=QHBoxLayout(row1); l1.setContentsMargins(0,0,0,0); l1.addWidget(self.comp_lbl); l1.addWidget(bcomp)
            row2=QWidget(); l2=QHBoxLayout(row2); l2.setContentsMargins(0,0,0,0); l2.addWidget(self.photo_lbl); l2.addWidget(bphoto)
            form.addRow("Fecha venta:", self.fecha); form.addRow("Precio real vendido:", self.precio); form.addRow("Descripción regalía:", self.regalia_desc); form.addRow("Valor regalía:", self.regalia); form.addRow("Comprobante venta:", row1); form.addRow("Foto cliente:", row2)
            self.preview=QLabel(); self.preview.setTextFormat(Qt.TextFormat.RichText); self.preview.setStyleSheet("background:#fff7ed;border:1px solid #fed7aa;border-radius:10px;padding:12px;color:#08285a;"); lay.addWidget(self.preview)
            for w in [self.precio, self.regalia]:
                try: w.textChanged.connect(self.refresh_preview)
                except Exception: pass
            self.refresh_preview()
            btns=QHBoxLayout(); lay.addLayout(btns); ok=QPushButton("Cerrar venta"); ok.setObjectName("orange"); ok.clicked.connect(self.accept); no=QPushButton("Cancelar"); no.setObjectName("ghost"); no.clicked.connect(self.reject); btns.addWidget(ok); btns.addWidget(no); btns.addStretch(1)
        def pick_comp(self):
            path,_=QFileDialog.getOpenFileName(self,"Selecciona comprobante de venta",str(Path.home()),"PDF (*.pdf);;Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if path: self.comp_path=path; self.comp_lbl.setText(path)
        def pick_photo(self):
            path,_=QFileDialog.getOpenFileName(self,"Selecciona foto del cliente con el carro",str(Path.home()),"Imágenes (*.png *.jpg *.jpeg *.webp);;Todos (*.*)")
            if path: self.photo_path=path; self.photo_lbl.setText(path)
        def refresh_preview(self):
            calc=vehicle_profit_summary(self.vehicle, self.precio.value(), self.regalia.value())
            self.preview.setText(f"Precio vendido: <b>{_fmt_usd(self.precio.value())}</b><br>Regalía: <b>{_fmt_usd(self.regalia.value())}</b><br>Ganancia real estimada: <b>{_fmt_usd(calc.get('ganancia_real_estimada_usd'))}</b> · Margen real: <b>{calc.get('margen_real_estimado_pct')}%</b>")
        def values(self):
            return {"fecha_venta": self.fecha.date().toPython().isoformat(), "precio": self.precio.value(), "regalia_desc": self.regalia_desc.text(), "regalia": self.regalia.value(), "comprobante": self.comp_path, "foto": self.photo_path}

    _ORIG_STAGE_BUILD_V46 = StageUpdateDialog._build
    _ORIG_STAGE_SAVE_V46 = StageUpdateDialog.save
    def _v46_stage_build(self):
        _ORIG_STAGE_BUILD_V46(self)
        if getattr(self, 'stage_key', None) == STAGE_PREPARACION:
            extra = _v46_stage_extra(self.vehicle, STAGE_PREPARACION)
            row = QWidget(); lay = QHBoxLayout(row); lay.setContentsMargins(0,0,0,0)
            lab = QLabel("Número de placa asignada:"); lab.setStyleSheet("font-weight:800;color:#08285a;")
            self.numero_placa = QLineEdit(extra.get("numero_placa", "")); self.numero_placa.setPlaceholderText("Ejemplo: P123ABC")
            lay.addWidget(lab); lay.addWidget(self.numero_placa)
            try:
                self.layout().insertWidget(max(0, self.layout().count()-1), row)
            except Exception:
                self.layout().addWidget(row)
    def _v46_stage_save(self):
        if getattr(self, 'stage_key', None) != STAGE_PREPARACION:
            return _ORIG_STAGE_SAVE_V46(self)
        fecha=self._date(self.fecha_evento); proveedor=self.taller.currentText(); missing=[]
        if not self._validate_order(fecha,self._date(self.emision_pedido),"La fecha de emisión no puede ser anterior al inicio de preparación."): return
        if not self._validate_order(self._date(self.emision_pedido),self._date(self.emision_obtenida),"La fecha obtenida de emisión no puede ser anterior a cuando se pidió."): return
        if not self._validate_order(self._date(self.cita_pedido),self._date(self.cita_asignada),"La cita asignada no puede ser anterior a cuando se pidió."): return
        if not self._validate_order(self._date(self.placas_ingreso),self._date(self.placas_entrega),"La entrega de placas no puede ser anterior al ingreso del trámite."): return
        if not self._validate_order(fecha,self._date(self.taller_ingreso),"El ingreso al taller no puede ser anterior al inicio de preparación."): return
        taller_salida=self._date(self.taller_salida); legal_fin=self._date(self.legal_fin)
        if taller_salida and not self._validate_order(self._date(self.taller_ingreso),taller_salida,"La salida del taller no puede ser anterior al ingreso al taller."): return
        if legal_fin and not self._validate_order(fecha,legal_fin,"El fin de legalización no puede ser anterior al inicio de preparación."): return
        dates=[d for d in [legal_fin,taller_salida] if d]
        fin=max(dates) if len(dates)==2 else None
        extra={"emision_pedido":self._date(self.emision_pedido),"emision_obtenida":self._date(self.emision_obtenida),"cita_pedido":self._date(self.cita_pedido),"cita_asignada":self._date(self.cita_asignada),"placas_ingreso":self._date(self.placas_ingreso),"placas_entrega":self._date(self.placas_entrega),"legalizacion_fin":legal_fin,"numero_placa":getattr(self,"numero_placa",QLineEdit()).text().strip(),"taller":self.taller.currentText(),"taller_ingreso":self._date(self.taller_ingreso),"taller_salida":taller_salida,"motivo_taller":self.motivo_taller.toPlainText()}
        cost_items, missing = self._collect_costs(fecha, proveedor)
        for item in getattr(self,"repuestos",[]):
            if item.get("descripcion") and float(item.get("monto_usd") or 0)>0:
                cost_items.append({"id": item.get("id"), "categoria":"TALLER","subcategoria":"REPUESTO_DETALLE","descripcion":item.get("descripcion"),"monto_usd":float(item.get("monto_usd") or 0),"proveedor":proveedor,"oc_numero":item.get("oc_numero",""),"fecha":fecha,"comprobante_src":item.get("comprobante_src",""),"oc_src":item.get("oc_src",""),"comprobante":item.get("comprobante",""),"comprobante_nombre":item.get("comprobante_nombre",""),"oc_documento":item.get("oc_documento",""),"oc_documento_nombre":item.get("oc_documento_nombre","")})
        if missing:
            resp=QMessageBox.question(self,"Gastos en cero", "Estos gastos obligatorios están en cero:\n"+"\n".join(missing)+"\n\n¿Seguro que no aplican o no los tuviste?")
            if resp != QMessageBox.StandardButton.Yes: return
        data={"fecha_inicio":fecha,"fecha_fin":fin,"proveedor":proveedor,"comentario":self.comentario.toPlainText(),"extra":extra,"cost_items":cost_items}
        ok,msg=update_vehicle_stage(self.vehicle.get("id"), self.stage_key, data, None, self.user, self.device)
        if not ok: QMessageBox.warning(self,"Validación",msg); return
        QMessageBox.information(self,"Guardado",msg); self.accept()
    StageUpdateDialog._build = _v46_stage_build
    StageUpdateDialog.save = _v46_stage_save

    _ORIG_VEHICLE_DETAIL_REFRESH_V46 = VehicleDetailDialog.refresh
    def _v46_vehicle_detail_refresh(self):
        _ORIG_VEHICLE_DETAIL_REFRESH_V46(self)
        try:
            entries = getattr(self, "_cost_entries", None) or vehicle_document_entries(self.vehicle)
            total = sum(_to_float(e.get("monto_usd"), 0) for e in entries)
            r = self.cost_table.rowCount(); self.cost_table.insertRow(r)
            vals=["", "", "TOTAL DOCUMENTADO", _fmt_usd(total), "", "", "", ""]
            for c,val in enumerate(vals):
                it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); it.setBackground(QColor("#FFF2CC")); self.cost_table.setItem(r,c,it)
            _v46_apply_table_hardening(self.cost_table); _v46_apply_table_hardening(self.stage_table)
        except Exception:
            pass
    VehicleDetailDialog.refresh = _v46_vehicle_detail_refresh

    _ORIG_QUOTE_MARK_SOLD_V46 = QuoteDetailDialog.mark_sold
    def _v46_quote_mark_sold(self):
        q=getattr(self, 'q', None) or find_quote(self.quote_id)
        if not q: return
        v=find_vehicle(q.get('vehicle_id',''))
        if not v: QMessageBox.warning(self,"Venta","Vehículo no encontrado."); return
        if v.get("estado_comercial") == COMM_VENDIDO or v.get("venta_detalle"):
            QMessageBox.warning(self,"Venta","Este vehículo ya fue vendido. No se puede volver a vender."); return
        dlg=SaleClosingDialog(self,q,v)
        if dlg.exec()!=QDialog.DialogCode.Accepted: return
        vals=dlg.values()
        ok,msg=mark_quote_won_and_vehicle_sold(self.quote_id,self.user,self.device,vals["precio"],vals["regalia_desc"],vals["regalia"],vals["fecha_venta"],vals["comprobante"],vals["foto"])
        QMessageBox.information(self,"Venta",msg) if ok else QMessageBox.warning(self,"Venta",msg)
        self.refresh()
    QuoteDetailDialog.mark_sold = _v46_quote_mark_sold

    _ORIG_QUOTE_GENERATE_PROPOSAL_V46 = QuoteDetailDialog.generate_proposal
    def _v46_quote_generate_proposal(self):
        dlg=QDialog(self); dlg.setWindowTitle("Generar propuesta"); dlg.setMinimumWidth(460); lay=QVBoxLayout(dlg)
        lay.addWidget(make_title("Generar propuesta", "Elige formato y si deseas usar el valor actual de la cotización."))
        form=QFormLayout(); lay.addLayout(form)
        formato=QComboBox(); formato.addItems(["PDF", "Excel", "PDF + Excel"]); form.addRow("Formato:", formato)
        nota=QLabel("La propuesta se guardará dentro de la carpeta del carro. Después podrás abrirla o guardar una copia."); nota.setWordWrap(True); lay.addWidget(nota)
        row=QHBoxLayout(); lay.addLayout(row); ok=QPushButton("Generar"); ok.setObjectName("orange"); ok.clicked.connect(dlg.accept); no=QPushButton("Cancelar"); no.setObjectName("ghost"); no.clicked.connect(dlg.reject); row.addWidget(ok); row.addWidget(no); row.addStretch(1)
        if dlg.exec()!=QDialog.DialogCode.Accepted: return
        ok,msg,paths=generate_quote_proposal_selected(self.quote_id, formato.currentText(), self.user)
        if not ok:
            QMessageBox.warning(self,"Propuesta",msg); return
        guardar=QMessageBox.question(self,"Propuesta generada",msg+"\n\n¿Deseas guardar una copia en otra ubicación?",QMessageBox.StandardButton.Yes|QMessageBox.StandardButton.No)
        if guardar==QMessageBox.StandardButton.Yes:
            folder=QFileDialog.getExistingDirectory(self,"Selecciona carpeta para guardar copia",str(Path.home()))
            if folder:
                for p in paths: shutil.copy2(p, Path(folder)/p.name)
                QMessageBox.information(self,"Copia",f"Copia guardada en:\n{folder}")
            return
        msgbox=QMessageBox(self); msgbox.setWindowTitle("Abrir propuesta"); msgbox.setText("¿Qué deseas hacer ahora?")
        pdf_btn=excel_btn=None
        for p in paths:
            if p.suffix.lower()==".pdf" and pdf_btn is None: pdf_btn=msgbox.addButton("Abrir PDF",QMessageBox.ButtonRole.AcceptRole)
            if p.suffix.lower() in (".xlsx",".xlsm") and excel_btn is None: excel_btn=msgbox.addButton("Abrir Excel",QMessageBox.ButtonRole.AcceptRole)
        msgbox.addButton("Cerrar",QMessageBox.ButtonRole.RejectRole); msgbox.exec(); clicked=msgbox.clickedButton()
        if clicked==pdf_btn:
            for p in paths:
                if p.suffix.lower()==".pdf": QDesktopServices.openUrl(QUrl.fromLocalFile(str(p))); break
        elif clicked==excel_btn:
            for p in paths:
                if p.suffix.lower() in (".xlsx",".xlsm"): QDesktopServices.openUrl(QUrl.fromLocalFile(str(p))); break
    QuoteDetailDialog.generate_proposal = _v46_quote_generate_proposal

    class PurchasePage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main=main; self.comprobante_path=""; self.oc_path=""; self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Compra vehicular · Nueva compra", "Compra USA, compra LOCAL o servicio de importación."))
            scroll=QScrollArea(); scroll.setWidgetResizable(True); content=QWidget(); scroll.setWidget(content); body=QVBoxLayout(content)
            g1=QGroupBox("Datos del vehículo"); f1=QFormLayout(g1)
            self.marca=QComboBox(); self.marca.setEditable(True); self.modelo=QLineEdit(); self.anio=QSpinBox(); self.anio.setRange(1980,date.today().year+1); self.anio.setValue(date.today().year); self.millaje=QSpinBox(); self.millaje.setRange(0,999999); self.color=QLineEdit(); self.tipo=QComboBox(); self.tipo.setEditable(True); self.tipo.addItems(["SEDAN","SUV","PICKUP","VAN","CAMIONETA","DEPORTIVO","OTRO"]); self.caracteristicas=QTextEdit(); self.caracteristicas.setMinimumHeight(110)
            for lab,w in [("Marca:",self.marca),("Modelo:",self.modelo),("Año:",self.anio),("Millaje:",self.millaje),("Color:",self.color),("Tipo:",self.tipo),("Características propuesta:",self.caracteristicas)]: f1.addRow(lab,w)
            g2=QGroupBox("Tipo de compra y costos iniciales"); f2=QFormLayout(g2)
            self.tipo_compra=QComboBox(); self.tipo_compra.addItems(["LOTE USA","LOCAL","SERVICIO DE IMPORTACION"]); self.pais_compra=QComboBox(); self.pais_compra.setEditable(True); self.pais_compra.addItems(load_catalog(F_CATALOG_PAISES, DEFAULT_PAISES_DESTINO)); self.estado_usa=QComboBox(); self.subasta=QComboBox(); self.subasta.setEditable(True); self.lote=QLineEdit(); self.precio=MoneyEdit(); self.precio.setRange(0,9999999); self.precio_publicado_local=MoneyEdit(); self.precio_publicado_local.setRange(0,9999999); self.precio_minimo_local=MoneyEdit(); self.precio_minimo_local.setRange(0,9999999); self.fecha=configure_date_edit(QDateEdit()); self.fecha.setDate(QDate.currentDate())
            for lab,w in [("Tipo compra:",self.tipo_compra),("País compra/local:",self.pais_compra),("Estado USA:",self.estado_usa),("Subasta:",self.subasta),("Número lote:",self.lote),("Costo / precio compra USD:",self.precio),("Precio venta local opcional:",self.precio_publicado_local),("Precio mínimo local opcional:",self.precio_minimo_local),("Fecha compra:",self.fecha)]: f2.addRow(lab,w)
            g3=QGroupBox("Comprobantes y OC"); f3=QFormLayout(g3)
            self.comp_label=QLineEdit(); self.comp_label.setReadOnly(True); bcomp=QPushButton("Subir comprobante"); bcomp.setObjectName("ghost"); bcomp.clicked.connect(self.pick_comprobante); r1=QWidget(); l1=QHBoxLayout(r1); l1.setContentsMargins(0,0,0,0); l1.addWidget(self.comp_label); l1.addWidget(bcomp)
            self.oc_num=QLineEdit(); self.oc_label=QLineEdit(); self.oc_label.setReadOnly(True); boc=QPushButton("Subir PDF OC"); boc.setObjectName("ghost"); boc.clicked.connect(self.pick_oc); r2=QWidget(); l2=QHBoxLayout(r2); l2.setContentsMargins(0,0,0,0); l2.addWidget(self.oc_label); l2.addWidget(boc)
            self.obs=QTextEdit(); self.obs.setMinimumHeight(80)
            f3.addRow("Comprobante:",r1); f3.addRow("Número OC:",self.oc_num); f3.addRow("PDF OC:",r2); f3.addRow("Observaciones:",self.obs)
            body.addWidget(g1); body.addWidget(g2); body.addWidget(g3); btn=QPushButton("Crear compra"); btn.setObjectName("orange"); btn.clicked.connect(self.save_purchase); body.addWidget(btn); lay.addWidget(scroll); self.refresh_catalogs()
        def refresh_catalogs(self):
            self.marca.clear(); self.marca.addItems(load_catalog(F_CATALOG_MARCAS, DEFAULT_MARCAS)); self.subasta.clear(); self.subasta.addItems(load_catalog(F_CATALOG_SUBASTAS, DEFAULT_SUBASTAS)); self.estado_usa.clear(); self.estado_usa.addItems(load_catalog(F_CATALOG_ESTADOS_USA, US_STATES))
        def pick_comprobante(self):
            path,_=QFileDialog.getOpenFileName(self,"Selecciona comprobante",str(Path.home()),"Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if path: self.comprobante_path=path; self.comp_label.setText(path)
        def pick_oc(self):
            path,_=QFileDialog.getOpenFileName(self,"Selecciona PDF de OC",str(Path.home()),"PDF (*.pdf);;Todos (*.*)")
            if path: self.oc_path=path; self.oc_label.setText(path)
        def save_purchase(self):
            if not self.comprobante_path:
                QMessageBox.warning(self,"Validación","Debes subir el comprobante de compra."); return
            tipo_text=self.tipo_compra.currentText(); is_local=_norm(tipo_text)=="LOCAL"; is_servicio=_norm(tipo_text)==_norm(TIPO_COMPRA_IMPORTACION)
            data={"marca":self.marca.currentText(),"modelo":self.modelo.text(),"anio":self.anio.value(),"millaje":self.millaje.value(),"color":self.color.text(),"tipo":self.tipo.currentText(),"estado_usa":self.estado_usa.currentText(),"subasta":self.subasta.currentText(),"lote":self.lote.text(),"precio_ganado_usd":self.precio.value(),"fecha_compra":self.fecha.date().toPython().isoformat(),"observaciones":self.obs.toPlainText(),"caracteristicas":self.caracteristicas.toPlainText(),"oc_compra_numero":self.oc_num.text(),"oc_compra_src":self.oc_path,"tipo_compra":TIPO_COMPRA_LOCAL if is_local else (TIPO_COMPRA_IMPORTACION if is_servicio else TIPO_COMPRA_LOTE),"pais_compra":self.pais_compra.currentText(),"servicio_importacion":is_servicio,"precio_venta_local_usd":self.precio_publicado_local.value(),"precio_minimo_local_usd":self.precio_minimo_local.value()}
            ok,msg,vid=create_vehicle_purchase(data,Path(self.comprobante_path),self.main.user,self.main.device)
            if not ok: QMessageBox.warning(self,"Validación",msg); return
            QMessageBox.information(self,"Compra creada",msg); self.main.refresh_all(); self.main.open_vehicle_detail(vid)

    class SalesPage(QWidget):
        def __init__(self, main):
            super().__init__(); self.main=main; self._ids=[]; self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Ventas", "Carros vendidos con historial, documentos, cotizaciones, ganancia y fecha de venta."))
            self.table=QTableWidget(0,9); self.table.setHorizontalHeaderLabels(["Código","Vehículo","Cliente","Fecha venta","Precio vendido","Ganancia real","Margen","Comprobante","Días compra→venta"]); self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); _v46_apply_table_hardening(self.table); self.table.cellDoubleClicked.connect(self.open_vehicle); lay.addWidget(self.table)
        def refresh(self):
            sold=[v for v in load_vehicles() if ensure_vehicle_runtime_fields(v) and v.get("estado_comercial")==COMM_VENDIDO]
            self._ids=[]; self.table.setRowCount(len(sold))
            for r,v in enumerate(sold):
                self._ids.append(v.get("id")); venta=v.get("venta_detalle",{}) or {}; vals=[v.get("codigo"),_vehicle_display_name(v),v.get("cliente",""),_fmt_date(v.get("fecha_venta")),_fmt_usd(v.get("precio_venta_real_usd")),_fmt_usd(venta.get("ganancia_real_final_usd")),str(venta.get("margen_real_final_pct",0))+"%","✅" if venta.get("comprobante_venta") else "❌",str(_v46_stage_days(v).get("compra_venta") or "")]
                for c,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.table.setItem(r,c,it)
        def open_vehicle(self,row,col):
            if 0 <= row < len(self._ids): self.main.open_vehicle_detail(self._ids[row]); self.refresh()

    _ORIG_COT_PAGE = CotizacionesPage
    class CotizacionesPage(_ORIG_COT_PAGE):
        def _build(self):
            super()._build()
            wc=QWidget(); cl=QVBoxLayout(wc); self.client_table=QTableWidget(0,9); self.client_table.setHorizontalHeaderLabels(["Cliente","Teléfono","Medio principal","Ingreso","Cotizaciones","Primera cotización","Días desde primera","Último status","Días último status"]); self.client_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); _v46_apply_table_hardening(self.client_table); cl.addWidget(self.client_table); self.tabs.addTab(wc,"Base de datos clientes")
        def refresh(self):
            super().refresh()
            clients: dict[str, dict] = {}
            for q in load_quotes():
                cl=q.get("cliente",{}) or {}; key=_norm(cl.get("telefono") or cl.get("nombre") or q.get("id"))
                if key not in clients:
                    clients[key]={"nombre":cl.get("nombre",""),"telefono":cl.get("telefono",""),"medio":cl.get("medio_contacto",""),"ingreso":_to_float((q.get("leasing") or {}).get("ingreso_cliente"),0),"quotes":[],"first":None,"last":None}
                clients[key]["quotes"].append(q)
                fc=_parse_date(q.get("fecha_cotizacion")); lg=_parse_date(q.get("ultima_gestion") or q.get("fecha_cotizacion"))
                if fc and (not clients[key]["first"] or fc<clients[key]["first"]): clients[key]["first"]=fc
                if lg and (not clients[key]["last"] or lg>clients[key]["last"]): clients[key]["last"]=lg
            rows=list(clients.values()); self.client_table.setRowCount(len(rows))
            for r,cdata in enumerate(rows):
                last=cdata.get("last"); first=cdata.get("first"); vals=[cdata["nombre"],cdata["telefono"],cdata["medio"],_fmt_usd(cdata["ingreso"]),len(cdata["quotes"]),_fmt_date(first),(date.today()-first).days if first else "",_fmt_date(last),(date.today()-last).days if last else ""]
                for col,val in enumerate(vals):
                    it=QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.client_table.setItem(r,col,it)
            _v46_apply_table_hardening(self.client_table)

    class MainWindow(QMainWindow):
        def __init__(self, user: dict, device: DeviceInfo):
            super().__init__(); self.user=user; self.device=device; self.setWindowTitle(f"{APP_NAME} · {user.get('usuario')} · v{APP_VERSION}"); self.resize(1480,900); self.setWindowIcon(QIcon(str(ResourceManager.find_logo() or ""))); self._build(); self.refresh_all()
        def _build(self):
            central=QWidget(); self.setCentralWidget(central); main=QHBoxLayout(central); main.setContentsMargins(0,0,0,0)
            side=QFrame(); side.setFixedWidth(265); side.setStyleSheet("QFrame{background:#08285a;} QLabel{color:white;} QPushButton{background:transparent;color:white;text-align:left;padding:11px 14px;border-radius:0;font-weight:850;} QPushButton:hover{background:#0e3a78;border-left:5px solid #f59a13;}")
            sl=QVBoxLayout(side); sl.setContentsMargins(14,20,14,16); logo=QLabel(); logo.setAlignment(Qt.AlignmentFlag.AlignCenter); pix=QPixmap(str(ResourceManager.find_logo() or ""))
            if not pix.isNull(): logo.setPixmap(pix.scaled(145,145,Qt.AspectRatioMode.KeepAspectRatio,Qt.TransformationMode.SmoothTransformation))
            else: logo.setText("L&M"); logo.setStyleSheet("font-size:30px;font-weight:950;color:#f59a13;")
            sl.addWidget(logo); title=QLabel("LYM AUTO CONTROL"); title.setAlignment(Qt.AlignmentFlag.AlignCenter); title.setStyleSheet("font-weight:950;color:#ffcf7a;font-size:14pt;"); sub=QLabel(f"Rol: {self.user.get('rol')}"); sub.setAlignment(Qt.AlignmentFlag.AlignCenter); sub.setStyleSheet("font-weight:700;color:#d7e6ff;font-size:10pt;"); sl.addWidget(title); sl.addWidget(sub)
            self.stack=QStackedWidget(); self.pages=[]; self._section_widgets=[]
            self.dashboard=DashboardPage(self); self.purchase=PurchasePage(self); self.inventory=InventoryPage(self); self.sales=SalesPage(self); self.cotizaciones=CotizacionesPage(self); self.reporteria=ReporteriaPage(self); self.catalogos=CatalogosPage(self)
            def add_page(container_layout, name, widget):
                idx=self.stack.addWidget(widget); self.pages.append(widget); b=QPushButton(name); b.clicked.connect(lambda _,i=idx:self.stack.setCurrentIndex(i)); container_layout.addWidget(b); return b
            add_page(sl,"🏠  Inicio",self.dashboard)
            def collapsible_section(title, entries, opened=True):
                header=QPushButton(("▾  " if opened else "▸  ")+title); header.setStyleSheet("color:#ffcf7a;font-weight:950;margin-top:10px;border-bottom:1px solid rgba(255,255,255,.14);")
                box=QWidget(); bl=QVBoxLayout(box); bl.setContentsMargins(0,0,0,0); bl.setSpacing(0)
                for name,widget in entries: add_page(bl,name,widget)
                box.setVisible(opened)
                def toggle(): box.setVisible(not box.isVisible()); header.setText(("▾  " if box.isVisible() else "▸  ")+title)
                header.clicked.connect(toggle); sl.addWidget(header); sl.addWidget(box); self._section_widgets.append((header,box))
            collapsible_section("COMPRA VEHICULAR", [("➕  Nueva compra",self.purchase),("🚗  Inventario / CV",self.inventory),("✅  Ventas",self.sales)], True)
            collapsible_section("COMERCIAL", [("💬  Cotizaciones",self.cotizaciones)], True)
            collapsible_section("REPORTERÍA", [("📊  Reportería",self.reporteria),("📚  Catálogos",self.catalogos)], True)
            if user_has_permission(self.user, PERM_CONFIG):
                bconf=QPushButton("⚙️  Configuración"); bconf.clicked.connect(self.open_config); sl.addWidget(bconf)
            sl.addStretch(1); bclose=QPushButton("🚪  Cerrar"); bclose.clicked.connect(self.close); sl.addWidget(bclose); main.addWidget(side); main.addWidget(self.stack,1); self.setStatusBar(QStatusBar()); self.statusBar().showMessage(f"Usuario: {self.user.get('usuario')} · Carpeta: {get_data_folder()}")
        def refresh_all(self):
            bootstrap_system(); self.dashboard.refresh(); self.inventory.refresh(); self.sales.refresh(); self.purchase.refresh_catalogs(); self.cotizaciones.refresh(); self.catalogos.refresh_all()
        def open_vehicle_detail(self, vehicle_id: str):
            dlg=VehicleDetailDialog(self,vehicle_id,self.user,self.device); dlg.exec(); self.refresh_all()
        def open_config(self):
            dlg=ConfigDialog(self,self.user,self.device); dlg.exec(); self.refresh_all()




# =============================================================================
# AJUSTES V4.8 - PROPUESTA PROFESIONAL, CATALOGOS COMERCIALES Y COSTOS SIN DUPLICADOS
# =============================================================================
APP_VERSION = "2.5.0_LEASING"

F_CATALOG_FIRMAS_PROPUESTA = "firmas_propuesta.json"
F_CATALOG_CONDICIONES_LEASING = "condiciones_leasing.json"
F_CATALOG_OPCIONES_VENTA = "opciones_venta.json"

DEFAULT_OPCIONES_VENTA = ["USO PERSONAL", "RENTA CAR", "PLATAFORMA"]
DEFAULT_FIRMAS_PROPUESTA = [
    {"nombre": "Guillermo Moreno", "telefono": "(503) 7475 5821", "correo": "guillermo.moreno@lyminversiones.com", "usuario": ""},
]
DEFAULT_CONDICIONES_LEASING = [
    "Esta propuesta forma parte de una oferta especial de arrendamiento vehicular (leasing) válida por 15 días calendario a partir de la fecha de emisión.",
    "La tasa de interés mensual del {tasa}% corresponde a la cotización solicitada y brinda condiciones accesibles y transparentes para nuestros clientes.",
    "La propuesta se formalizará mediante un contrato de arrendamiento donde se establecerán las condiciones de uso, pagos mensuales y opción de adquisición al finalizar el plazo.",
    "El valor del seguro mostrado es estimado y puede variar según las características y el uso que le dé al vehículo; el costo final se definirá con base en la cotización emitida por la aseguradora correspondiente.",
    "Los gastos administrativos y legales derivados de la formalización del contrato deberán ser cancelados al momento de la firma.",
    "Las cuotas mensuales incluyen IVA, seguro y GPS, reflejando el monto total a pagar. No se aplican cobros adicionales fuera de los valores detallados en esta propuesta.",
    "Durante el período de arrendamiento, el cliente deberá mantener el vehículo en buen estado y cumplir con los servicios preventivos o correctivos recomendados.",
    "Esta propuesta tiene carácter informativo y no constituye compromiso contractual hasta la formalización del contrato de arrendamiento financiero.",
]

_ORIG_BOOTSTRAP_V48 = bootstrap_system

def bootstrap_system() -> None:
    _ORIG_BOOTSTRAP_V48()
    # Catálogos nuevos: configuraciones por defecto, firmas, condiciones y opciones de venta.
    try:
        for name, default in [
            (F_CATALOG_OPCIONES_VENTA, DEFAULT_OPCIONES_VENTA),
            (F_CATALOG_CONDICIONES_LEASING, DEFAULT_CONDICIONES_LEASING),
            (F_CATALOG_FIRMAS_PROPUESTA, DEFAULT_FIRMAS_PROPUESTA),
        ]:
            p = system_file(name)
            if p and not p.exists():
                _write_json_file(p, default)
    except Exception:
        pass


def _load_json_list_catalog(filename: str, default: list) -> list:
    data = _read_json_system(filename, default)
    if not isinstance(data, list):
        data = list(default)
    if not data:
        data = list(default)
    return data


def load_sale_options() -> list[str]:
    vals = load_catalog(F_CATALOG_OPCIONES_VENTA, DEFAULT_OPCIONES_VENTA)
    # Mantener orden funcional mínimo.
    for v in DEFAULT_OPCIONES_VENTA:
        if v not in vals:
            vals.append(v)
    return vals


def load_quote_conditions() -> list[str]:
    data = _load_json_list_catalog(F_CATALOG_CONDICIONES_LEASING, DEFAULT_CONDICIONES_LEASING)
    clean = [str(x).strip() for x in data if str(x).strip()]
    if not clean:
        clean = DEFAULT_CONDICIONES_LEASING.copy()
    return clean


def save_quote_conditions(conditions: list[str]) -> bool:
    clean = [str(x).strip() for x in conditions if str(x).strip()]
    return _write_json_system(F_CATALOG_CONDICIONES_LEASING, clean or DEFAULT_CONDICIONES_LEASING.copy())


def load_quote_signatures() -> list[dict]:
    data = _load_json_list_catalog(F_CATALOG_FIRMAS_PROPUESTA, DEFAULT_FIRMAS_PROPUESTA)
    clean = []
    for raw in data:
        if isinstance(raw, dict):
            nombre = str(raw.get("nombre") or "").strip()
            telefono = str(raw.get("telefono") or "").strip()
            correo = str(raw.get("correo") or "").strip()
            usuario = str(raw.get("usuario") or "").strip()
        else:
            nombre, telefono, correo, usuario = str(raw).strip(), "", "", ""
        if nombre:
            clean.append({"nombre": nombre, "telefono": telefono, "correo": correo, "usuario": usuario})
    if not clean:
        clean = DEFAULT_FIRMAS_PROPUESTA.copy()
    return clean


def save_quote_signatures(signatures: list[dict]) -> bool:
    clean = []
    for raw in signatures or []:
        if not isinstance(raw, dict):
            continue
        nombre = str(raw.get("nombre") or "").strip()
        telefono = str(raw.get("telefono") or "").strip()
        correo = str(raw.get("correo") or "").strip()
        usuario = str(raw.get("usuario") or "").strip()
        if nombre:
            clean.append({"nombre": nombre, "telefono": telefono, "correo": correo, "usuario": usuario})
    return _write_json_system(F_CATALOG_FIRMAS_PROPUESTA, clean or DEFAULT_FIRMAS_PROPUESTA.copy())


def default_signature_for_user(user: Optional[dict] = None) -> dict:
    signatures = load_quote_signatures()
    uname = str((user or {}).get("usuario") or "").strip().lower()
    if uname:
        for sig in signatures:
            if str(sig.get("usuario") or "").strip().lower() == uname or str(sig.get("nombre") or "").strip().lower() == uname:
                return sig
    return signatures[0] if signatures else DEFAULT_FIRMAS_PROPUESTA[0]


def calculate_purchase_option_from_quote(quote: dict) -> float:
    cuota = _to_float((quote.get("leasing") or {}).get("cuota_total_con_iva"), 0)
    if cuota > 0 and cuota < 500:
        return round(cuota, 2)
    return 500.0


# Costeo corregido: no duplicar precio ganado cuando existe gasto detallado + etapa comprado.
def vehicle_total_cost(vehicle: dict) -> float:
    ensure_vehicle_runtime_fields(vehicle)
    costs = vehicle.get("gastos_detallados")
    if isinstance(costs, list) and costs:
        total = 0.0
        purchase_counted = False
        purchase_amount = _to_float(vehicle.get("precio_ganado_usd"), 0)
        for g in costs:
            source = str(g.get("source") or "").lower()
            cat = _norm(g.get("categoria"))
            sub = _norm(g.get("subcategoria"))
            is_purchase = source == "purchase" or source == "stage:comprado" or (cat == "COMPRA" and sub in ("PRECIO_GANADO", "COMPROBANTE COMPRA", "COMPROBANTE_COMPRA"))
            monto = _to_float(g.get("monto_usd"), 0)
            if is_purchase:
                if not purchase_counted:
                    total += purchase_amount if purchase_amount > 0 else monto
                    purchase_counted = True
                continue
            total += monto
        if not purchase_counted and purchase_amount > 0:
            total += purchase_amount
        return round(total, 2)
    total = _to_float(vehicle.get("precio_ganado_usd"), 0)
    for st_key, st in (vehicle.get("etapas", {}) or {}).items():
        if st_key == STAGE_COMPRADO:
            continue
        total += _to_float(st.get("costo_usd"), 0)
    for g in vehicle.get("gastos_extra", []) or []:
        total += _to_float(g.get("monto_usd"), 0)
    return round(total, 2)


def _quote_prima_requerida(leasing: dict) -> float:
    precio = _to_float(leasing.get("precio_vehiculo"), 0)
    prima_pct = _to_float(leasing.get("prima_pct"), 20)
    comision = _to_float(leasing.get("comision_usd"), 0)
    return round(_to_float(leasing.get("prima_requerida_usd"), precio * prima_pct / 100.0 + comision), 2)


def _quote_financial_summary(quote: dict) -> dict:
    leasing = quote.get("leasing", {}) or {}
    legal = quote.get("legal", {}) or {}
    auto_oc = calculate_purchase_option_from_quote(quote)
    return {
        "valor_vehiculo": round(_to_float(leasing.get("precio_vehiculo"), 0), 2),
        "prima_requerida": _quote_prima_requerida(leasing),
        "monto_leasing": round(_to_float(leasing.get("monto_leasing"), 0), 2),
        "plazo": int(_to_float(leasing.get("plazo_meses"), 0)),
        "tasa": round(_to_float(leasing.get("tasa_mensual_pct"), 0), 4),
        "cuota_base": round(_to_float(leasing.get("cuota_base"), 0), 2),
        "cuota_final": round(_to_float(leasing.get("cuota_total_con_iva"), 0), 2),
        "seguro": round(_to_float(leasing.get("seguro_mensual"), 0), 2),
        "gps": round(_to_float(leasing.get("gps_mensual"), 0), 2),
        "legal": round(_to_float(legal.get("valor_legales_iva_incluido"), 0), 2),
        "opcion_compra": round(_to_float(quote.get("opcion_compra_usd"), auto_oc), 2),
        "prima_pct": round(_to_float(leasing.get("prima_pct"), 20), 4),
        "comision": round(_to_float(leasing.get("comision_usd"), 100), 2),
    }


_ORIG_CREATE_QUOTE_V48 = create_quote

def create_quote(data: dict, user: dict, device: DeviceInfo, quote_id: str = "") -> tuple[bool, str, str]:
    # Validar que prima siempre sea 20% + comisión si el usuario no manda pago_inicial explícito.
    precio_tmp = _to_float(data.get("precio_vehiculo"), 0)
    prima_pct_tmp = _to_float(data.get("prima_pct"), 20)
    comision_tmp = _to_float(data.get("comision_usd"), 100)
    if data.get("pago_inicial") in (None, ""):
        data = {**data, "pago_inicial": round(precio_tmp * prima_pct_tmp / 100.0 + comision_tmp, 2)}
    ok, msg, qid = _ORIG_CREATE_QUOTE_V48(data, user, device, quote_id)
    if not ok or not qid:
        return ok, msg, qid
    q = find_quote(qid)
    if not q:
        return ok, msg, qid
    # Opción venta y opción de compra por regla de negocio.
    q["opcion_venta"] = _norm(data.get("opcion_venta") or q.get("opcion_venta") or "RENTA CAR")
    if data.get("opcion_compra_usd") in (None, "") or _to_float(data.get("opcion_compra_usd"), 0) <= 0:
        q["opcion_compra_usd"] = calculate_purchase_option_from_quote(q)
    else:
        q["opcion_compra_usd"] = round(_to_float(data.get("opcion_compra_usd"), calculate_purchase_option_from_quote(q)), 2)
    # Asegurar que la prima requerida guardada incluya la comisión.
    le = q.setdefault("leasing", {})
    le["prima_requerida_usd"] = round(_to_float(le.get("precio_vehiculo"), 0) * _to_float(le.get("prima_pct"), 20) / 100.0 + _to_float(le.get("comision_usd"), 0), 2)
    upsert_quote(q)
    return ok, msg, qid


def _proposal_vehicle_photo_path(snapshot: dict) -> Optional[Path]:
    rel = snapshot.get("foto_principal") or ""
    name = snapshot.get("foto_principal_nombre") or "foto_vehiculo.jpg"
    if not rel:
        return None
    return decrypt_file_to_temp(rel, name)


def _proposal_footer(canvas, doc):
    canvas.saveState()
    try:
        canvas.setFont("Helvetica", 9)
        canvas.setFillColorRGB(0.05, 0.15, 0.35)
        canvas.drawCentredString(doc.pagesize[0] / 2.0, 0.28 * 72, "Ayudando a lograr tus sueños")
        canvas.drawRightString(doc.pagesize[0] - 0.55 * 72, 0.28 * 72, str(canvas.getPageNumber()))
    finally:
        canvas.restoreState()


def _proposal_conditions_for_quote(quote: dict) -> list[str]:
    s = _quote_financial_summary(quote)
    out = []
    for line in load_quote_conditions():
        try:
            out.append(line.format(tasa=s["tasa"], opcion_venta=quote.get("opcion_venta", "RENTA CAR")))
        except Exception:
            out.append(line)
    return out


def generate_quote_proposal_pdf(quote_id: str, user: Optional[dict] = None, options: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada.", None
    out_dir = _proposal_output_dir(quote)
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether
    except Exception:
        return False, "Para generar PDF instala: pip install reportlab", None
    ensure_quote_runtime_fields(quote)
    options = options or {}
    signer = options.get("firma") or default_signature_for_user(user)
    if options.get("opcion_venta"):
        quote["opcion_venta"] = _norm(options.get("opcion_venta"))
    if options.get("opcion_compra_usd") not in (None, ""):
        quote["opcion_compra_usd"] = round(_to_float(options.get("opcion_compra_usd"), calculate_purchase_option_from_quote(quote)), 2)
    else:
        quote.setdefault("opcion_compra_usd", calculate_purchase_option_from_quote(quote))
    upsert_quote(quote)
    cl = quote.get("cliente", {})
    snap = quote.get("vehicle_snapshot", {})
    s = _quote_financial_summary(quote)
    vehicle_name = _vehicle_display_name(snap).upper()
    out = out_dir / _safe_quote_filename(quote, ".pdf")
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="LYMTitle48", parent=styles["Title"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=16, leading=19, textColor=colors.HexColor("#061F4A")))
    styles.add(ParagraphStyle(name="LYMVehicle48", parent=styles["Heading2"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor("#061F4A")))
    styles.add(ParagraphStyle(name="LYMBody48", parent=styles["BodyText"], alignment=TA_LEFT, fontSize=10.2, leading=14.2))
    styles.add(ParagraphStyle(name="LYMSmall48", parent=styles["BodyText"], fontSize=8.6, leading=10.8))
    styles.add(ParagraphStyle(name="LYMRed48", parent=styles["BodyText"], fontSize=9.5, leading=12, textColor=colors.red, fontName="Helvetica-Bold"))
    doc = SimpleDocTemplate(str(out), pagesize=letter, rightMargin=0.55*inch, leftMargin=0.55*inch, topMargin=0.38*inch, bottomMargin=0.48*inch)
    story = []
    logo = ResourceManager.find_logo()
    logo_flow = Image(str(logo), width=0.88*inch, height=0.88*inch) if logo and logo.exists() else Paragraph("<b>L&M</b>", styles["LYMTitle48"])
    header = Table([[logo_flow, ""]], colWidths=[1.05*inch, 6.1*inch])
    header.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP"), ("LINEBELOW", (1,0), (1,0), 3, colors.HexColor("#061F4A")), ("BOTTOMPADDING", (0,0), (-1,-1), 8)]))
    story.append(header)
    story.append(Spacer(1, 0.09*inch))
    story.append(Paragraph(f"<b>{html.escape(cl.get('nombre','Cliente'))}</b><br/>Presente.", styles["LYMBody48"]))
    story.append(Spacer(1, 0.12*inch))
    story.append(Paragraph("<u>Propuesta de Arrendamiento Vehicular</u>", styles["LYMTitle48"]))
    story.append(Spacer(1, 0.12*inch))
    opcion_venta = quote.get("opcion_venta") or "RENTA CAR"
    story.append(Paragraph(f"Por medio de la presente, <b>L&amp;M Inversiones, S.A. de C.V.</b> tiene el agrado de presentarle la propuesta de contrato de arrendamiento (Leasing) opción: <b>{html.escape(opcion_venta)}</b>.", styles["LYMBody48"]))
    story.append(Spacer(1, 0.12*inch))
    story.append(Paragraph(f"◆ {html.escape(vehicle_name)} ◆", styles["LYMVehicle48"]))
    photo_path = _proposal_vehicle_photo_path(snap)
    if photo_path and photo_path.exists():
        try:
            car_img = Image(str(photo_path), width=2.35*inch, height=1.55*inch)
            photo_table = Table([["", car_img, ""]], colWidths=[2.25*inch, 2.35*inch, 2.25*inch])
            photo_table.setStyle(TableStyle([("ALIGN", (1,0), (1,0), "CENTER"), ("VALIGN", (0,0), (-1,-1), "MIDDLE")]))
            story.append(photo_table)
            story.append(Spacer(1, 0.07*inch))
        except Exception:
            pass
    story.append(Paragraph("<b>Características destacadas:</b>", styles["Heading3"]))
    feats = _split_features(snap)
    left, right = feats[0::2], feats[1::2]
    feat_rows = []
    for i in range(max(len(left), len(right))):
        feat_rows.append([Paragraph("• " + html.escape(left[i]) if i < len(left) else "", styles["LYMBody48"]), Paragraph("• " + html.escape(right[i]) if i < len(right) else "", styles["LYMBody48"])])
    ft = Table(feat_rows, colWidths=[3.55*inch, 3.55*inch])
    ft.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP"), ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4)]))
    story.append(ft)
    story.append(Spacer(1, 0.12*inch))
    story.append(Paragraph(f"Compartimos ante usted el detalle de prima, monto leasing, plazo y cuota mensual para <b>{html.escape(vehicle_name)}</b>. El cálculo incluye seguro estimado, servicio GPS, IVA y una tasa del <b>{s['tasa']}%</b> mensual.", styles["LYMBody48"]))
    story.append(Spacer(1, 0.10*inch))
    summary_line = Table([[Paragraph(f"<b>MONTO LEASING:</b> <u>{_fmt_usd(s['monto_leasing'])}</u>", styles["LYMBody48"]), Paragraph(f"<b>Costo Legal:</b> <u>{_fmt_usd(s['legal'])}</u> <font size='8'>Incluye IVA</font>", styles["LYMBody48"]), Paragraph(f"<b>Opción de compra:</b> <u>{_fmt_usd(s['opcion_compra'])}</u>", styles["LYMBody48"])]], colWidths=[2.35*inch, 2.35*inch, 2.35*inch])
    summary_line.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("ALIGN", (0,0), (-1,-1), "LEFT")]))
    story.append(summary_line)
    # Resumen claro para el cliente: valor del vehículo y prima requerida separada de la tabla de cuota.
    client_rows = [
        [Paragraph("<b>Valor del vehículo</b>", styles["LYMBody48"]), Paragraph(f"<b>{_fmt_usd(s['valor_vehiculo'])}</b>", styles["LYMBody48"]), Paragraph("<b>Prima requerida</b>", styles["LYMBody48"]), Paragraph(f"<b>{_fmt_usd(s['prima_requerida'])}</b>", styles["LYMBody48"])],
    ]
    client_tbl = Table(client_rows, colWidths=[1.65*inch, 1.55*inch, 1.65*inch, 2.25*inch])
    client_tbl.setStyle(TableStyle([("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#64748B")), ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#F8FAFC")), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("TOPPADDING", (0,0), (-1,-1), 6), ("BOTTOMPADDING", (0,0), (-1,-1), 6)]))
    story.append(Spacer(1, 0.06*inch)); story.append(client_tbl); story.append(Spacer(1, 0.08*inch))
    quote_tbl = Table([
        ["Plazo\n(meses)", "Tasa rentabilidad\n(%)", "Cuota base\n(US$)", "Cuota total mensual incluye seguro\nGPS e IVA\n(US$)"],
        [str(s["plazo"]), f"{s['tasa']}%", _fmt_usd(s["cuota_base"]).replace("$ ", ""), _fmt_usd(s["cuota_final"]).replace("$ ", "")],
    ], colWidths=[1.05*inch, 1.55*inch, 1.35*inch, 3.15*inch])
    quote_tbl.setStyle(TableStyle([("GRID", (0,0), (-1,-1), 0.65, colors.black), ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"), ("FONTNAME", (3,1), (3,1), "Helvetica-Bold"), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("ALIGN", (0,0), (-1,-1), "LEFT"), ("BACKGROUND", (0,0), (-1,0), colors.white), ("FONTSIZE", (0,0), (-1,-1), 9.2), ("TOPPADDING", (0,0), (-1,-1), 7), ("BOTTOMPADDING", (0,0), (-1,-1), 7)]))
    story.append(quote_tbl)
    story.append(Paragraph("<u>La cuota mensual incluye IVA, seguro y GPS.</u>", styles["LYMRed48"]))
    story.append(Spacer(1, 0.12*inch))
    story.append(Paragraph("Quedamos atentos a cualquier consulta o duda.", styles["LYMBody48"]))
    story.append(Spacer(1, 0.12*inch))
    story.append(Paragraph(f"San Salvador, {_spanish_date_long(date.today())}", styles["LYMBody48"]))
    story.append(PageBreak())
    story.append(Paragraph("Atentamente,", styles["LYMBody48"]))
    sig_left = Paragraph(f"<br/><b>{html.escape(signer.get('nombre',''))}</b><br/>L&amp;M Inversiones, S.A. de C.V.<br/>{html.escape(signer.get('correo',''))}<br/>Tel: {html.escape(signer.get('telefono',''))}<br/><br/>X__________________________<br/><font size='8'>Firma asesor</font>", styles["LYMBody48"])
    sig_right = Paragraph(f"<br/><b>{html.escape(cl.get('nombre','Cliente'))}</b><br/>Aceptado por cliente<br/><br/><br/>X__________________________<br/><font size='8'>Firma cliente</font>", styles["LYMBody48"])
    sig_tbl = Table([[sig_left, sig_right]], colWidths=[3.45*inch, 3.45*inch])
    sig_tbl.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP"), ("ALIGN", (0,0), (-1,-1), "LEFT"), ("TOPPADDING", (0,0), (-1,-1), 16)]))
    story.append(sig_tbl)
    story.append(Spacer(1, 0.24*inch))
    story.append(Paragraph("<b>Condiciones y Vigencia de la Oferta</b>", styles["LYMTitle48"]))
    conds = _proposal_conditions_for_quote(quote)
    for c in conds:
        story.append(Paragraph("• " + html.escape(c), styles["LYMBody48"]))
        story.append(Spacer(1, 0.055*inch))
    doc.build(story, onFirstPage=_proposal_footer, onLaterPages=_proposal_footer)
    quote.setdefault("propuestas", []).append({"fecha": _now_iso(), "tipo": "PDF", "archivo": str(out), "firma": signer, "opcion_venta": quote.get("opcion_venta"), "opcion_compra_usd": quote.get("opcion_compra_usd"), "usuario": (user or {}).get("usuario", "")})
    upsert_quote(quote)
    log_audit("GENERAR_PROPUESTA_PDF", (user or {}).get("usuario", ""), quote.get("vehicle_code", ""), str(out))
    return True, "Propuesta PDF generada en carpeta del carro.", out


def generate_quote_proposal_excel(quote_id: str, user: Optional[dict] = None, options: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada.", None
    out_dir = _proposal_output_dir(quote)
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.drawing.image import Image as XLImage
        from openpyxl.utils import get_column_letter
    except Exception:
        return False, "Para generar Excel instala: pip install openpyxl pillow", None
    ensure_quote_runtime_fields(quote)
    options = options or {}
    signer = options.get("firma") or default_signature_for_user(user)
    if options.get("opcion_venta"):
        quote["opcion_venta"] = _norm(options.get("opcion_venta"))
    if options.get("opcion_compra_usd") not in (None, ""):
        quote["opcion_compra_usd"] = round(_to_float(options.get("opcion_compra_usd"), calculate_purchase_option_from_quote(quote)), 2)
    else:
        quote.setdefault("opcion_compra_usd", calculate_purchase_option_from_quote(quote))
    upsert_quote(quote)
    cl = quote.get("cliente", {})
    snap = quote.get("vehicle_snapshot", {})
    s = _quote_financial_summary(quote)
    vehicle_name = _vehicle_display_name(snap).upper()
    out = out_dir / _safe_quote_filename(quote, ".xlsx")
    wb = Workbook(); ws = wb.active; ws.title = "Propuesta Leasing"
    ws.sheet_view.showGridLines = False
    for c in range(1, 9): ws.column_dimensions[get_column_letter(c)].width = [13, 14, 14, 14, 14, 14, 14, 14][c-1]
    navy = "061F4A"; orange = "F59A13"; light = "F8FAFC"; thin = Side(style="thin", color="111111"); border = Border(left=thin, right=thin, top=thin, bottom=thin)
    def cell(r,c,v=None,bold=False,size=11,color="000000",fill=None,align="left"):
        x=ws.cell(r,c); x.value=v; x.font=Font(name="Calibri",size=size,bold=bold,color=color); x.alignment=Alignment(horizontal=align,vertical="center",wrap_text=True); x.border=border
        if fill: x.fill=PatternFill("solid",fgColor=fill)
        return x
    logo=ResourceManager.find_logo()
    if logo and logo.exists():
        try:
            img=XLImage(str(logo)); img.width=72; img.height=72; ws.add_image(img,"A1")
        except Exception: pass
    ws.merge_cells("B4:H4"); cell(4,2,"",fill=navy)
    ws.merge_cells("A6:H6"); cell(6,1,f"{cl.get('nombre','Cliente')}\nPresente.",bold=True,size=12)
    ws.merge_cells("A9:H9"); cell(9,1,"Propuesta de Arrendamiento Vehicular",bold=True,size=16,color=navy,align="center")
    ws.merge_cells("A11:H12"); cell(11,1,f"Por medio de la presente, L&M Inversiones, S.A. de C.V. tiene el agrado de presentarle la propuesta de contrato de arrendamiento (Leasing) opción: {quote.get('opcion_venta','RENTA CAR')}.",size=11)
    ws.merge_cells("A14:H14"); cell(14,1,f"◆ {vehicle_name} ◆",bold=True,size=16,color=navy,align="center")
    photo_path=_proposal_vehicle_photo_path(snap)
    if photo_path and photo_path.exists():
        try:
            img=XLImage(str(photo_path)); img.width=220; img.height=140; ws.add_image(img,"C15")
        except Exception: pass
        feat_start=23
    else:
        feat_start=16
    ws.merge_cells(start_row=feat_start,end_row=feat_start,start_column=1,end_column=8); cell(feat_start,1,"Características destacadas:",bold=True,size=13,color=navy)
    feats=_split_features(snap)
    for i,feat in enumerate(feats[:18]):
        row=feat_start+1+i//2; col=1 if i%2==0 else 5
        ws.merge_cells(start_row=row,start_column=col,end_row=row,end_column=col+3)
        cell(row,col,"• "+feat,size=10)
    finance_row=feat_start+12
    ws.merge_cells(start_row=finance_row,end_row=finance_row+1,start_column=1,end_column=8); cell(finance_row,1,f"Compartimos ante usted el detalle de prima, monto leasing, plazo y cuota mensual para {vehicle_name}. El cálculo incluye seguro estimado, servicio GPS, IVA y una tasa del {s['tasa']}% mensual.",size=10)
    r=finance_row+3
    rows=[["MONTO LEASING",s["monto_leasing"],"Costo Legal",s["legal"],"Opción de compra",s["opcion_compra"]],["Valor del vehículo",s["valor_vehiculo"],"Prima requerida",s["prima_requerida"],"Cuota mensual",s["cuota_final"]]]
    for rr,row in enumerate(rows,start=r):
        for cc,val in enumerate(row,start=1):
            cell(rr,cc,val,bold=cc%2==1,fill=light if rr==r else None)
    r+=4
    headers=["Plazo (meses)","Tasa rentabilidad (%)","Cuota base (US$)","Cuota total mensual incluye seguro, GPS e IVA (US$)"]
    for cc,h in enumerate(headers,start=1): cell(r,cc,h,bold=True,color="FFFFFF",fill=navy,align="center")
    vals=[s["plazo"],s["tasa"],s["cuota_base"],s["cuota_final"]]
    for cc,v in enumerate(vals,start=1): cell(r+1,cc,v,align="center")
    ws.merge_cells(start_row=r+3,end_row=r+3,start_column=1,end_column=8); cell(r+3,1,"La cuota mensual incluye IVA, seguro y GPS.",bold=True,color="DC2626")
    ws.merge_cells(start_row=r+5,end_row=r+7,start_column=1,end_column=4); cell(r+5,1,f"Atentamente,\n{signer.get('nombre','')}\nL&M Inversiones, S.A. de C.V.\n{signer.get('correo','')}\nTel: {signer.get('telefono','')}\n\nX____________________________\nFirma asesor",size=10)
    ws.merge_cells(start_row=r+5,end_row=r+7,start_column=5,end_column=8); cell(r+5,5,f"Aceptado por cliente:\n{cl.get('nombre','Cliente')}\n\nX____________________________\nFirma cliente",size=10)
    cond_r=r+9; ws.merge_cells(start_row=cond_r,end_row=cond_r,start_column=1,end_column=8); cell(cond_r,1,"Condiciones y Vigencia de la Oferta",bold=True,size=14,color=navy,align="center")
    for i,cond in enumerate(_proposal_conditions_for_quote(quote),start=cond_r+1):
        ws.merge_cells(start_row=i,end_row=i,start_column=1,end_column=8); cell(i,1,"• "+cond,size=9)
    ws.merge_cells(start_row=cond_r+11,end_row=cond_r+11,start_column=1,end_column=8); cell(cond_r+11,1,"Ayudando a lograr tus sueños",size=10,color=navy,align="center")
    for rr in range(1,cond_r+12): ws.row_dimensions[rr].height=22
    ws.page_setup.orientation="portrait"; ws.page_setup.paperSize=ws.PAPERSIZE_LETTER; ws.page_setup.fitToWidth=1; ws.page_setup.fitToHeight=0
    wb.save(out)
    quote.setdefault("propuestas", []).append({"fecha": _now_iso(), "tipo": "EXCEL", "archivo": str(out), "firma": signer, "opcion_venta": quote.get("opcion_venta"), "opcion_compra_usd": quote.get("opcion_compra_usd"), "usuario": (user or {}).get("usuario", "")})
    upsert_quote(quote)
    log_audit("GENERAR_PROPUESTA_EXCEL", (user or {}).get("usuario", ""), quote.get("vehicle_code", ""), str(out))
    return True, "Propuesta Excel generada en carpeta del carro.", out


def generate_quote_proposal_selected(quote_id: str, formato: str, user: Optional[dict] = None, options: Optional[dict] = None) -> tuple[bool, str, list[Path]]:
    formato = _norm(formato)
    paths: list[Path] = []
    msgs: list[str] = []
    ok_all = True
    if "PDF" in formato:
        ok, msg, p = generate_quote_proposal_pdf(quote_id, user, options)
        ok_all = ok_all and ok; msgs.append(msg)
        if p: paths.append(p)
    if "EXCEL" in formato or "XLSX" in formato:
        ok, msg, p = generate_quote_proposal_excel(quote_id, user, options)
        ok_all = ok_all and ok; msgs.append(msg)
        if p: paths.append(p)
    return ok_all and bool(paths), "\n".join(msgs), paths


def generate_quote_proposal_files(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, list[Path]]:
    return generate_quote_proposal_selected(quote_id, "PDF + EXCEL", user, None)


if PYSIDE_OK:
    class ProposalOptionsDialog(QDialog):
        def __init__(self, parent, quote: dict, user: dict):
            super().__init__(parent); self.quote=quote; self.user=user; self.setWindowTitle("Generar propuesta"); self.setMinimumWidth(560); self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Opciones de propuesta", "Selecciona formato, firma, opción de venta y opción de compra antes de generar."))
            f=QFormLayout(); lay.addLayout(f)
            self.formato=QComboBox(); self.formato.addItems(["PDF + EXCEL", "PDF", "EXCEL"])
            self.firmas=load_quote_signatures(); self.firma_combo=QComboBox();
            for sig in self.firmas: self.firma_combo.addItem(f"{sig.get('nombre')} · {sig.get('telefono')} · {sig.get('correo')}")
            default=default_signature_for_user(self.user)
            for i,sig in enumerate(self.firmas):
                if sig.get("nombre") == default.get("nombre"): self.firma_combo.setCurrentIndex(i); break
            self.opcion_venta=QComboBox(); self.opcion_venta.setEditable(True); self.opcion_venta.addItems(load_sale_options()); self.opcion_venta.setCurrentText(self.quote.get("opcion_venta") or "RENTA CAR")
            self.opcion_compra=MoneyEdit(); self.opcion_compra.setRange(0,999999); self.opcion_compra.setValue(_to_float(self.quote.get("opcion_compra_usd"), calculate_purchase_option_from_quote(self.quote)))
            f.addRow("Formato:", self.formato); f.addRow("Quién hace la propuesta:", self.firma_combo); f.addRow("Opción de venta:", self.opcion_venta); f.addRow("Opción de compra:", self.opcion_compra)
            row=QHBoxLayout(); lay.addLayout(row); b=QPushButton("Generar"); b.setObjectName("orange"); b.clicked.connect(self.accept); c=QPushButton("Cancelar"); c.setObjectName("ghost"); c.clicked.connect(self.reject); row.addWidget(b); row.addWidget(c); row.addStretch(1)
        def data(self):
            idx=self.firma_combo.currentIndex(); sig=self.firmas[idx] if 0 <= idx < len(self.firmas) else default_signature_for_user(self.user)
            return {"formato": self.formato.currentText(), "firma": sig, "opcion_venta": self.opcion_venta.currentText(), "opcion_compra_usd": self.opcion_compra.value()}

    _ORIG_QUOTE_EDITOR_V48 = QuoteEditorDialog
    class QuoteEditorDialog(_ORIG_QUOTE_EDITOR_V48):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._opcion_manual = False
            self._opcion_auto_update = False
            try:
                panel=QFrame(); panel.setStyleSheet("QFrame{background:#fff7ed;border:1px solid #fed7aa;border-radius:12px;}"); pl=QFormLayout(panel)
                self.opcion_venta=QComboBox(); self.opcion_venta.setEditable(True); self.opcion_venta.addItems(load_sale_options())
                if getattr(self,"quote",None): self.opcion_venta.setCurrentText(self.quote.get("opcion_venta") or "RENTA CAR")
                else: self.opcion_venta.setCurrentText("RENTA CAR")
                pl.addRow("Opción de venta:", self.opcion_venta)
                self.layout().insertWidget(2, panel)
                self.opcion_compra.valueChanged.connect(self._mark_opcion_manual)
                QTimer.singleShot(0, self.refresh_calc)
            except Exception:
                pass
        def _mark_opcion_manual(self, *args):
            if not getattr(self,"_opcion_auto_update",False):
                self._opcion_manual=True
        def refresh_calc(self):
            try:
                calc, legal, pago = self._calc_data()
                auto_oc = min(500.0, round(_to_float(calc.get("cuota_total_con_iva"), 0), 2)) if _to_float(calc.get("cuota_total_con_iva"),0) > 0 else 500.0
                if hasattr(self,"opcion_compra") and not getattr(self,"_opcion_manual",False):
                    self._opcion_auto_update=True
                    self.opcion_compra.setValue(auto_oc)
                    self._opcion_auto_update=False
            except Exception:
                pass
            return super().refresh_calc()
        def save(self):
            if self.vehicle_combo.currentIndex()<0 or self.vehicle_combo.currentIndex()>=len(self._vehicle_ids):
                QMessageBox.warning(self,"Validación","No hay vehículo disponible seleccionado."); return
            calc,legal,pago=self._calc_data()
            data={"vehicle_id":self._vehicle_ids[self.vehicle_combo.currentIndex()],"cliente_nombre":self.nombre.text(),"telefono":self.telefono.text(),"correo":self.correo.text(),"medio_contacto":self.medio.currentText(),"ingreso_cliente":self.ingreso.value(),"precio_vehiculo":self.precio.value(),"prima_pct":self.prima_pct.value(),"comision_usd":self.comision.value(),"pago_inicial":pago,"plazo_meses":self.plazo.value(),"tasa_mensual_pct":_to_float(self.tasa.currentText(),2.5),"seguro_mensual":self.seguro.value(),"gps_mensual":self.gps.value(),"iva_pct":self.iva.value(),"opcion_compra_usd":self.opcion_compra.value(),"opcion_venta":getattr(self,"opcion_venta",QComboBox()).currentText() if hasattr(self,"opcion_venta") else "RENTA CAR","comentario":self.comentario.toPlainText()}
            ok,msg,qid=create_quote(data,self.user,self.device,self.quote_id)
            if not ok: QMessageBox.warning(self,"Validación",msg); return
            self.quote_id=qid; QMessageBox.information(self,"Guardado",msg); self.accept()

    _ORIG_QUOTE_DETAIL_V48 = QuoteDetailDialog
    class QuoteDetailDialog(_ORIG_QUOTE_DETAIL_V48):
        def refresh(self):
            super().refresh()
            try:
                q=find_quote(self.quote_id); s=_quote_financial_summary(q or {})
                if q:
                    extra=f"\nOpción de venta: {q.get('opcion_venta','RENTA CAR')}\nOpción de compra: {_fmt_usd(s['opcion_compra'])}\nPrima incluye comisión: {_fmt_usd(s['prima_requerida'])}"
                    self.info.setPlainText(self.info.toPlainText()+extra)
            except Exception:
                pass
        def generate_proposal(self):
            q=find_quote(self.quote_id)
            if not q: QMessageBox.warning(self,"Propuesta","Cotización no encontrada."); return
            dlg=ProposalOptionsDialog(self,q,self.user)
            if dlg.exec()!=QDialog.DialogCode.Accepted: return
            data=dlg.data(); formato=data.pop("formato")
            ok,msg,paths=generate_quote_proposal_selected(self.quote_id,formato,self.user,data)
            if not ok:
                QMessageBox.warning(self,"Propuesta",msg); return
            box=QMessageBox(self); box.setWindowTitle("Propuesta generada"); box.setText(msg+"\n\n"+"\n".join(p.name for p in paths))
            pdf_btn=None; xls_btn=None
            for p in paths:
                if p.suffix.lower()==".pdf" and pdf_btn is None: pdf_btn=box.addButton("Abrir PDF",QMessageBox.ButtonRole.AcceptRole)
                if p.suffix.lower() in (".xlsx",".xls") and xls_btn is None: xls_btn=box.addButton("Abrir Excel",QMessageBox.ButtonRole.ActionRole)
            copy_btn=box.addButton("Guardar copia…",QMessageBox.ButtonRole.ActionRole); box.addButton("Cerrar",QMessageBox.ButtonRole.RejectRole); box.exec(); clicked=box.clickedButton()
            if clicked==pdf_btn:
                for p in paths:
                    if p.suffix.lower()==".pdf": QDesktopServices.openUrl(QUrl.fromLocalFile(str(p))); break
            elif clicked==xls_btn:
                for p in paths:
                    if p.suffix.lower() in (".xlsx",".xls"): QDesktopServices.openUrl(QUrl.fromLocalFile(str(p))); break
            elif clicked==copy_btn:
                folder=QFileDialog.getExistingDirectory(self,"Selecciona carpeta para guardar copia",str(Path.home()))
                if folder:
                    for p in paths:
                        try: shutil.copy2(p, Path(folder)/p.name)
                        except Exception: pass
                    QMessageBox.information(self,"Copia","Copia guardada correctamente.")
            self.refresh()

    _ORIG_CATALOGOS_V48 = CatalogosPage
    class CatalogosPage(_ORIG_CATALOGOS_V48):
        def _build(self):
            super()._build()
            try:
                self.tabs.addTab(self._build_sale_options_tab(), "Opciones venta")
                self.tabs.addTab(self._build_signatures_tab(), "Firmas propuestas")
                self.tabs.addTab(self._build_conditions_tab(), "Condiciones")
            except Exception:
                pass
        def _build_sale_options_tab(self):
            return self._build_simple_tab("Opciones venta", F_CATALOG_OPCIONES_VENTA, DEFAULT_OPCIONES_VENTA)
        def _build_signatures_tab(self):
            w=QWidget(); lay=QVBoxLayout(w); row=QHBoxLayout(); lay.addLayout(row)
            badd=QPushButton("Agregar firma"); badd.setObjectName("orange"); bedit=QPushButton("Editar"); bedit.setObjectName("ghost"); bdel=QPushButton("Borrar"); bdel.setObjectName("danger"); row.addWidget(badd); row.addWidget(bedit); row.addWidget(bdel); row.addStretch(1)
            self.sign_table=QTableWidget(0,4); self.sign_table.setHorizontalHeaderLabels(["Nombre y apellido","Teléfono","Correo","Usuario asociado"]); self.sign_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); _v46_apply_table_hardening(self.sign_table); lay.addWidget(self.sign_table)
            badd.clicked.connect(self.add_signature); bedit.clicked.connect(self.edit_signature); bdel.clicked.connect(self.delete_signature); return w
        def _signature_dialog(self, data=None):
            data=data or {}; dlg=QDialog(self); dlg.setWindowTitle("Firma propuesta"); fl=QFormLayout(dlg); nombre=QLineEdit(data.get("nombre", "")); tel=QLineEdit(data.get("telefono", "")); correo=QLineEdit(data.get("correo", "")); usuario=QLineEdit(data.get("usuario", "")); fl.addRow("Nombre y apellido:",nombre); fl.addRow("Teléfono:",tel); fl.addRow("Correo:",correo); fl.addRow("Usuario asociado:",usuario); row=QHBoxLayout(); b=QPushButton("Guardar"); b.setObjectName("orange"); c=QPushButton("Cancelar"); c.setObjectName("ghost"); row.addWidget(b); row.addWidget(c); fl.addRow(row); b.clicked.connect(dlg.accept); c.clicked.connect(dlg.reject)
            if dlg.exec()!=QDialog.DialogCode.Accepted: return None
            return {"nombre":nombre.text().strip(),"telefono":tel.text().strip(),"correo":correo.text().strip(),"usuario":usuario.text().strip()}
        def add_signature(self):
            data=self._signature_dialog();
            if data and data.get("nombre"):
                vals=load_quote_signatures(); vals.append(data); save_quote_signatures(vals); self.refresh_signatures()
        def edit_signature(self):
            row=self.sign_table.currentRow(); vals=load_quote_signatures()
            if row<0 or row>=len(vals): QMessageBox.information(self,"Firmas","Selecciona una firma."); return
            data=self._signature_dialog(vals[row])
            if data and data.get("nombre"):
                vals[row]=data; save_quote_signatures(vals); self.refresh_signatures()
        def delete_signature(self):
            row=self.sign_table.currentRow(); vals=load_quote_signatures()
            if row<0 or row>=len(vals): QMessageBox.information(self,"Firmas","Selecciona una firma."); return
            if QMessageBox.question(self,"Borrar","¿Borrar firma seleccionada?") != QMessageBox.StandardButton.Yes: return
            vals.pop(row); save_quote_signatures(vals); self.refresh_signatures()
        def refresh_signatures(self):
            if not hasattr(self,"sign_table"): return
            vals=load_quote_signatures(); self.sign_table.setRowCount(len(vals))
            for r,sig in enumerate(vals):
                for c,k in enumerate(["nombre","telefono","correo","usuario"]): self.sign_table.setItem(r,c,QTableWidgetItem(str(sig.get(k,""))))
            _v46_apply_table_hardening(self.sign_table)
        def _build_conditions_tab(self):
            w=QWidget(); lay=QVBoxLayout(w); tip=QLabel("Cada línea es un párrafo nuevo. Puedes usar {tasa} para insertar la tasa de la cotización y {opcion_venta} para insertar la opción de venta."); tip.setWordWrap(True); lay.addWidget(tip)
            self.conditions_text=QTextEdit(); lay.addWidget(self.conditions_text); row=QHBoxLayout(); lay.addLayout(row); b=QPushButton("Guardar condiciones"); b.setObjectName("orange"); b.clicked.connect(self.save_conditions); row.addWidget(b); row.addStretch(1); return w
        def save_conditions(self):
            lines=[x.strip() for x in self.conditions_text.toPlainText().splitlines() if x.strip()]
            ok=save_quote_conditions(lines); QMessageBox.information(self,"Condiciones","Condiciones guardadas." if ok else "No se pudieron guardar.")
        def refresh_all(self):
            super().refresh_all()
            try:
                self.refresh_signatures()
                if hasattr(self,"conditions_text"): self.conditions_text.setPlainText("\n".join(load_quote_conditions()))
                if hasattr(self,"simple_tabs") and "Opciones venta" in self.simple_tabs: self.refresh_simple("Opciones venta")
            except Exception:
                pass



# =============================================================================
# AJUSTES V4.9 - PROPUESTA PDF 2 PAGINAS + WORD PROFESIONAL, SIN EXCEL EN PROPUESTAS
# =============================================================================
APP_VERSION = "2.6.0_LEASING"

# Excel inventario usa graficas tipo dona; esta importacion global evita NameError en openpyxl.
try:
    from openpyxl.chart.pie_chart import DoughnutChart
except Exception:
    try:
        from openpyxl.chart import DoughnutChart  # type: ignore
    except Exception:
        DoughnutChart = None  # type: ignore


def _lym_escape(value: Any) -> str:
    return html.escape(str(value or ""))


def _proposal_signer(options: Optional[dict], user: Optional[dict]) -> dict:
    options = options or {}
    sig = options.get("firma") if isinstance(options.get("firma"), dict) else None
    if sig:
        return sig
    return default_signature_for_user(user)


def _apply_proposal_options_to_quote(quote: dict, options: Optional[dict], user: Optional[dict] = None) -> tuple[dict, dict]:
    options = options or {}
    signer = _proposal_signer(options, user)
    if options.get("opcion_venta"):
        quote["opcion_venta"] = _norm(options.get("opcion_venta"))
    if options.get("opcion_compra_usd") not in (None, ""):
        quote["opcion_compra_usd"] = round(_to_float(options.get("opcion_compra_usd"), calculate_purchase_option_from_quote(quote)), 2)
    else:
        quote.setdefault("opcion_compra_usd", calculate_purchase_option_from_quote(quote))
    upsert_quote(quote)
    return quote, signer


def _proposal_register_file(quote: dict, tipo: str, path: Path, signer: dict, user: Optional[dict]) -> None:
    quote.setdefault("propuestas", []).append({
        "fecha": _now_iso(),
        "tipo": tipo,
        "archivo": str(path),
        "path": str(path),
        "nombre": path.name,
        "firma": signer,
        "opcion_venta": quote.get("opcion_venta"),
        "opcion_compra_usd": quote.get("opcion_compra_usd"),
        "usuario": (user or {}).get("usuario", ""),
    })
    quote.setdefault("seguimientos", []).append({
        "fecha": _now_iso(),
        "usuario": (user or {}).get("usuario", ""),
        "accion": f"PROPUESTA_{tipo}_GENERADA",
        "comentario": f"Se generó propuesta {tipo}: {path.name}",
    })
    upsert_quote(quote)


def _proposal_header_table_pdf(styles, logo_flow):
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import Table, TableStyle
    header = Table([[logo_flow, ""]], colWidths=[1.05*inch, 6.1*inch])
    header.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LINEBELOW", (1,0), (1,0), 3, colors.HexColor("#061F4A")),
        ("BOTTOMPADDING", (0,0), (-1,-1), 6),
        ("TOPPADDING", (0,0), (-1,-1), 0),
    ]))
    return header


def _proposal_footer_v49(canvas, doc):
    canvas.saveState()
    try:
        canvas.setFont("Times-Roman", 10)
        canvas.setFillColorRGB(0.15, 0.20, 0.30)
        canvas.drawCentredString(doc.pagesize[0] / 2.0, 0.28 * 72, "Ayudando a lograr tus sueños")
        canvas.setFont("Helvetica", 9)
        canvas.drawRightString(doc.pagesize[0] - 0.45 * 72, 0.28 * 72, str(canvas.getPageNumber()))
    finally:
        canvas.restoreState()


def generate_quote_proposal_pdf(quote_id: str, user: Optional[dict] = None, options: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada.", None
    out_dir = _proposal_output_dir(quote)
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    except Exception:
        return False, "Para generar PDF instala: pip install reportlab", None

    ensure_quote_runtime_fields(quote)
    quote, signer = _apply_proposal_options_to_quote(quote, options, user)
    cl = quote.get("cliente", {}) or {}
    snap = quote.get("vehicle_snapshot", {}) or {}
    s = _quote_financial_summary(quote)
    vehicle_name = _vehicle_display_name(snap).upper()
    opcion_venta = quote.get("opcion_venta") or "RENTA CAR"
    out = out_dir / _safe_quote_filename(quote, ".pdf")

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="LYMTitle49", parent=styles["Title"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=15.2, leading=18, textColor=colors.HexColor("#061F4A")))
    styles.add(ParagraphStyle(name="LYMVehicle49", parent=styles["Heading2"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=15.2, leading=18.2, textColor=colors.HexColor("#061F4A")))
    styles.add(ParagraphStyle(name="LYMBody49", parent=styles["BodyText"], alignment=TA_LEFT, fontSize=9.7, leading=12.6))
    styles.add(ParagraphStyle(name="LYMFeature49", parent=styles["BodyText"], alignment=TA_LEFT, fontSize=9.1, leading=11.5))
    styles.add(ParagraphStyle(name="LYMSmall49", parent=styles["BodyText"], fontSize=8.35, leading=9.75))
    styles.add(ParagraphStyle(name="LYMRed49", parent=styles["BodyText"], fontSize=9.0, leading=11, textColor=colors.red, fontName="Helvetica-Bold"))

    doc = SimpleDocTemplate(str(out), pagesize=letter, rightMargin=0.50*inch, leftMargin=0.50*inch, topMargin=0.32*inch, bottomMargin=0.48*inch)
    story = []
    logo = ResourceManager.find_logo()
    logo_flow = Image(str(logo), width=0.82*inch, height=0.82*inch) if logo and logo.exists() else Paragraph("<b>L&amp;M</b>", styles["LYMTitle49"])
    header = _proposal_header_table_pdf(styles, logo_flow)

    # PAGINA 1: sin foto del carro, igual al modelo ejecutivo de 2 paginas.
    story.append(header)
    story.append(Spacer(1, 0.05*inch))
    story.append(Paragraph(f"<b>{_lym_escape(cl.get('nombre','Cliente')).upper()}</b><br/>Presente.", styles["LYMBody49"]))
    story.append(Spacer(1, 0.08*inch))
    story.append(Paragraph("<u>Propuesta de Arrendamiento Vehicular</u>", styles["LYMTitle49"]))
    story.append(Spacer(1, 0.08*inch))
    story.append(Paragraph(f"Por medio de la presente, <b>L&amp;M Inversiones, S.A. de C.V.</b> tiene el agrado de presentarle la propuesta de contrato de arrendamiento (Leasing) opción: <b>{_lym_escape(opcion_venta)}</b>.", styles["LYMBody49"]))
    story.append(Spacer(1, 0.08*inch))
    story.append(Paragraph(f"◆ {_lym_escape(vehicle_name)} ◆", styles["LYMVehicle49"]))
    story.append(Spacer(1, 0.05*inch))
    story.append(Paragraph("<b>Características destacadas:</b>", styles["Heading3"]))

    feats = _split_features(snap)
    # Mantener dos columnas como el ejemplo; si hay muchas lineas, se compacta para conservar 2 paginas.
    left, right = feats[0::2], feats[1::2]
    feat_rows = []
    for i in range(max(len(left), len(right))):
        feat_rows.append([
            Paragraph("• " + _lym_escape(left[i]) if i < len(left) else "", styles["LYMFeature49"]),
            Paragraph("• " + _lym_escape(right[i]) if i < len(right) else "", styles["LYMFeature49"]),
        ])
    ft = Table(feat_rows, colWidths=[3.45*inch, 3.45*inch])
    ft.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 2),
        ("RIGHTPADDING", (0,0), (-1,-1), 4),
        ("TOPPADDING", (0,0), (-1,-1), 1.2),
        ("BOTTOMPADDING", (0,0), (-1,-1), 1.2),
    ]))
    story.append(ft)
    story.append(Spacer(1, 0.08*inch))
    story.append(Paragraph(f"Compartimos ante usted el detalle de prima, monto leasing, plazo y cuota mensual para <b>{_lym_escape(vehicle_name)}</b>. El cálculo incluye seguro estimado, servicio GPS, IVA y una tasa del <b>{s['tasa']}%</b> mensual.", styles["LYMBody49"]))
    story.append(Spacer(1, 0.08*inch))

    summary_line = Table([[
        Paragraph(f"<b>MONTO LEASING:</b> <u>{_fmt_usd(s['monto_leasing'])}</u>", styles["LYMBody49"]),
        Paragraph(f"<b>Costo Legal:</b> <u>{_fmt_usd(s['legal'])}</u> <font size='7'>Incluye IVA</font>", styles["LYMBody49"]),
        Paragraph(f"<b>Opción de compra:</b> <u>{_fmt_usd(s['opcion_compra'])}</u>", styles["LYMBody49"]),
    ]], colWidths=[2.30*inch, 2.30*inch, 2.30*inch])
    summary_line.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("ALIGN", (0,0), (-1,-1), "LEFT"), ("BOTTOMPADDING", (0,0), (-1,-1), 3)]))
    story.append(summary_line)

    client_rows = [[Paragraph("<b>Valor del vehículo</b>", styles["LYMBody49"]), Paragraph(f"<b>{_fmt_usd(s['valor_vehiculo'])}</b>", styles["LYMBody49"]), Paragraph("<b>Prima requerida</b>", styles["LYMBody49"]), Paragraph(f"<b>{_fmt_usd(s['prima_requerida'])}</b>", styles["LYMBody49"])]]
    client_tbl = Table(client_rows, colWidths=[1.5*inch, 2.0*inch, 1.5*inch, 1.9*inch])
    client_tbl.setStyle(TableStyle([("GRID", (0,0), (-1,-1), 0.45, colors.HexColor("#94A3B8")), ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#F8FAFC")), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("TOPPADDING", (0,0), (-1,-1), 5), ("BOTTOMPADDING", (0,0), (-1,-1), 5)]))
    story.append(client_tbl)
    story.append(Spacer(1, 0.07*inch))

    quote_tbl = Table([
        ["Plazo\n(meses)", "Tasa rentabilidad\n(%)", "Cuota base\n(US$)", "Cuota total mensual incluye seguro, GPS e IVA\n(US$)"],
        [str(s["plazo"]), f"{s['tasa']}%", _fmt_usd(s["cuota_base"]), _fmt_usd(s["cuota_final"])],
    ], colWidths=[1.05*inch, 1.55*inch, 1.35*inch, 2.95*inch])
    quote_tbl.setStyle(TableStyle([("GRID", (0,0), (-1,-1), 0.65, colors.black), ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"), ("FONTNAME", (3,1), (3,1), "Helvetica-Bold"), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("ALIGN", (0,0), (-1,-1), "LEFT"), ("FONTSIZE", (0,0), (-1,-1), 8.9), ("TOPPADDING", (0,0), (-1,-1), 6), ("BOTTOMPADDING", (0,0), (-1,-1), 6)]))
    story.append(quote_tbl)
    story.append(Paragraph("<u>La cuota mensual incluye IVA, seguro y GPS.</u>", styles["LYMRed49"]))
    story.append(Spacer(1, 0.10*inch))
    story.append(Paragraph("Quedamos atentos a cualquier consulta o duda.", styles["LYMBody49"]))
    story.append(Spacer(1, 0.10*inch))
    story.append(Paragraph(f"San Salvador, {_spanish_date_long(date.today())}", styles["LYMBody49"]))

    story.append(PageBreak())

    # PAGINA 2: firmas y condiciones. La foto del carro va aqui, a la derecha de condiciones.
    sig_left = Paragraph(f"Atentamente,<br/><br/><b>{_lym_escape(signer.get('nombre',''))}</b><br/>L&amp;M Inversiones, S.A. de C.V.<br/>{_lym_escape(signer.get('correo',''))}<br/>Tel: {_lym_escape(signer.get('telefono',''))}<br/><br/>X__________________________<br/><font size='8'>Firma asesor</font>", styles["LYMBody49"])
    sig_right = Paragraph(f"<br/><b>{_lym_escape(cl.get('nombre','Cliente'))}</b><br/>Aceptado por cliente<br/><br/><br/>X__________________________<br/><font size='8'>Firma cliente</font>", styles["LYMBody49"])
    sig_tbl = Table([[sig_left, sig_right]], colWidths=[3.45*inch, 3.45*inch])
    sig_tbl.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP"), ("TOPPADDING", (0,0), (-1,-1), 0), ("BOTTOMPADDING", (0,0), (-1,-1), 8)]))
    story.append(sig_tbl)
    story.append(Spacer(1, 0.05*inch))
    story.append(Paragraph("<u><b>Condiciones y Vigencia de la Oferta</b></u>", styles["LYMTitle49"]))
    story.append(Spacer(1, 0.06*inch))

    conds = _proposal_conditions_for_quote(quote)
    cond_flow = []
    for c in conds:
        cond_flow.append(Paragraph("• " + _lym_escape(c), styles["LYMSmall49"]))
        cond_flow.append(Spacer(1, 0.035*inch))
    photo_path = _proposal_vehicle_photo_path(snap)
    photo_cell = ""
    if photo_path and photo_path.exists():
        try:
            photo_cell = Image(str(photo_path), width=2.65*inch, height=1.75*inch)
        except Exception:
            photo_cell = ""
    cond_table = Table([[cond_flow, photo_cell]], colWidths=[4.10*inch, 2.80*inch])
    cond_table.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP"), ("ALIGN", (1,0), (1,0), "CENTER"), ("LEFTPADDING", (0,0), (-1,-1), 0), ("RIGHTPADDING", (0,0), (-1,-1), 0), ("TOPPADDING", (0,0), (-1,-1), 0)]))
    story.append(cond_table)

    doc.build(story, onFirstPage=_proposal_footer_v49, onLaterPages=_proposal_footer_v49)
    _proposal_register_file(quote, "PDF", out, signer, user)
    log_audit("GENERAR_PROPUESTA_PDF", (user or {}).get("usuario", ""), quote.get("vehicle_code", ""), str(out))
    return True, "Propuesta PDF generada en carpeta del carro.", out


def _docx_add_run_bold(paragraph, text: str):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def _docx_set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text or ""))
    r.bold = bold
    return p


def _docx_money(value: float) -> str:
    return _fmt_usd(value)


def generate_quote_proposal_docx(quote_id: str, user: Optional[dict] = None, options: Optional[dict] = None) -> tuple[bool, str, Optional[Path]]:
    quote = find_quote(quote_id)
    if not quote:
        return False, "Cotización no encontrada.", None
    out_dir = _proposal_output_dir(quote)
    if out_dir is None:
        return False, "Carpeta del sistema no disponible.", None
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return False, "Para generar Word instala: pip install python-docx", None

    ensure_quote_runtime_fields(quote)
    quote, signer = _apply_proposal_options_to_quote(quote, options, user)
    cl = quote.get("cliente", {}) or {}
    snap = quote.get("vehicle_snapshot", {}) or {}
    s = _quote_financial_summary(quote)
    vehicle_name = _vehicle_display_name(snap).upper()
    opcion_venta = quote.get("opcion_venta") or "RENTA CAR"
    out = out_dir / _safe_quote_filename(quote, ".docx")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.35); sec.bottom_margin = Inches(0.45); sec.left_margin = Inches(0.55); sec.right_margin = Inches(0.55)
    styles = doc.styles
    styles['Normal'].font.name = 'Arial'; styles['Normal'].font.size = Pt(10)

    def add_blue_line(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '18')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '061F4A')
        pBdr.append(bottom); pPr.append(pBdr)

    def add_center(text, size=14, bold=True, underline=False):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = bold; r.underline = underline; r.font.size = Pt(size); r.font.color.rgb = RGBColor(6,31,74)
        return p

    # Header con logo y linea azul, sin texto extra al lado del logo.
    hdr = doc.add_table(rows=1, cols=2); hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.columns[0].width = Inches(0.95); hdr.columns[1].width = Inches(6.0)
    logo = ResourceManager.find_logo()
    if logo and logo.exists():
        try:
            hdr.cell(0,0).paragraphs[0].add_run().add_picture(str(logo), width=Inches(0.78))
        except Exception:
            hdr.cell(0,0).text = 'L&M'
    add_blue_line(hdr.cell(0,1).paragraphs[0])

    p = doc.add_paragraph(); _docx_add_run_bold(p, str(cl.get('nombre','Cliente')).upper()); p.add_run('\nPresente.')
    add_center('Propuesta de Arrendamiento Vehicular', 15, True, True)
    p = doc.add_paragraph()
    p.add_run('Por medio de la presente, ')
    _docx_add_run_bold(p, 'L&M Inversiones, S.A. de C.V.')
    p.add_run(f' tiene el agrado de presentarle la propuesta de contrato de arrendamiento (Leasing) opción: ')
    _docx_add_run_bold(p, str(opcion_venta))
    p.add_run('.')
    add_center(f'◆ {vehicle_name} ◆', 14, True, False)

    p = doc.add_paragraph(); _docx_add_run_bold(p, 'Características destacadas:')
    feats = _split_features(snap)
    ft = doc.add_table(rows=max(1, (len(feats)+1)//2), cols=2); ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, feat in enumerate(feats):
        cell = ft.cell(idx//2, idx%2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        para = cell.paragraphs[0]
        para.style = doc.styles['Normal']
        para.add_run('• ' + str(feat))
    p = doc.add_paragraph()
    p.add_run('Compartimos ante usted el detalle de prima, monto leasing, plazo y cuota mensual para ')
    _docx_add_run_bold(p, vehicle_name)
    p.add_run(f'. El cálculo incluye seguro estimado, servicio GPS, IVA y una tasa del {s["tasa"]}% mensual.')

    line = doc.add_table(rows=1, cols=3); line.alignment = WD_TABLE_ALIGNMENT.CENTER
    vals = [("MONTO LEASING", s['monto_leasing']), ("Costo Legal", s['legal']), ("Opción de compra", s['opcion_compra'])]
    for i,(lab,val) in enumerate(vals):
        para = line.cell(0,i).paragraphs[0]
        _docx_add_run_bold(para, f'{lab}: '); para.add_run(_docx_money(val))
    money = doc.add_table(rows=1, cols=4); money.alignment = WD_TABLE_ALIGNMENT.CENTER
    _docx_set_cell_text(money.cell(0,0), 'Valor del vehículo', True); _docx_set_cell_text(money.cell(0,1), _docx_money(s['valor_vehiculo']), True)
    _docx_set_cell_text(money.cell(0,2), 'Prima requerida', True); _docx_set_cell_text(money.cell(0,3), _docx_money(s['prima_requerida']), True)

    qt = doc.add_table(rows=2, cols=4); qt.style = 'Table Grid'; qt.alignment = WD_TABLE_ALIGNMENT.CENTER
    heads = ['Plazo\n(meses)', 'Tasa rentabilidad\n(%)', 'Cuota base\n(US$)', 'Cuota total mensual incluye seguro, GPS e IVA\n(US$)']
    vals = [str(s['plazo']), f"{s['tasa']}%", _docx_money(s['cuota_base']), _docx_money(s['cuota_final'])]
    for i,h in enumerate(heads): _docx_set_cell_text(qt.cell(0,i), h, True)
    for i,v in enumerate(vals): _docx_set_cell_text(qt.cell(1,i), v, i==3)
    p = doc.add_paragraph(); r = p.add_run('La cuota mensual incluye IVA, seguro y GPS.'); r.bold=True; r.underline=True; r.font.color.rgb = RGBColor(220,38,38)
    doc.add_paragraph('Quedamos atentos a cualquier consulta o duda.')
    doc.add_paragraph(f'San Salvador, {_spanish_date_long(date.today())}')

    doc.add_page_break()
    sig = doc.add_table(rows=1, cols=2); sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = sig.cell(0,0); right = sig.cell(0,1)
    left.text = ''
    lp = left.paragraphs[0]; lp.add_run('Atentamente,\n\n'); _docx_add_run_bold(lp, signer.get('nombre','')); lp.add_run(f"\nL&M Inversiones, S.A. de C.V.\n{signer.get('correo','')}\nTel: {signer.get('telefono','')}\n\nX__________________________\nFirma asesor")
    right.text = ''
    rp = right.paragraphs[0]; _docx_add_run_bold(rp, str(cl.get('nombre','Cliente'))); rp.add_run('\nAceptado por cliente\n\n\nX__________________________\nFirma cliente')

    add_center('Condiciones y Vigencia de la Oferta', 13, True, True)
    cond_table = doc.add_table(rows=1, cols=2); cond_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cond_cell, photo_cell = cond_table.cell(0,0), cond_table.cell(0,1)
    cond_cell.text = ''
    for cond in _proposal_conditions_for_quote(quote):
        p = cond_cell.add_paragraph('• ' + str(cond))
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs: r.font.size = Pt(8.5)
    photo_path = _proposal_vehicle_photo_path(snap)
    if photo_path and photo_path.exists():
        try:
            pp = photo_cell.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.add_run().add_picture(str(photo_path), width=Inches(2.65))
        except Exception:
            pass
    # Footer
    footer = sec.footer.paragraphs[0]
    footer.text = 'Ayudando a lograr tus sueños'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(out))
    _proposal_register_file(quote, "WORD", out, signer, user)
    log_audit("GENERAR_PROPUESTA_WORD", (user or {}).get("usuario", ""), quote.get("vehicle_code", ""), str(out))
    return True, "Propuesta Word generada en carpeta del carro.", out


def generate_quote_proposal_selected(quote_id: str, formato: str, user: Optional[dict] = None, options: Optional[dict] = None) -> tuple[bool, str, list[Path]]:
    formato_norm = _norm(formato)
    paths: list[Path] = []
    msgs: list[str] = []
    ok_all = True
    if "PDF" in formato_norm:
        ok, msg, p = generate_quote_proposal_pdf(quote_id, user, options)
        ok_all = ok_all and ok; msgs.append(msg)
        if p: paths.append(p)
    # Desde V4.9 ya no se genera Excel de propuesta; WORD reemplaza a Excel. EXCEL legacy se mapea a WORD.
    if any(x in formato_norm for x in ("WORD", "DOCX", "EXCEL", "XLSX")):
        ok, msg, p = generate_quote_proposal_docx(quote_id, user, options)
        ok_all = ok_all and ok; msgs.append(msg)
        if p: paths.append(p)
    return ok_all and bool(paths), "\n".join(msgs), paths


def generate_quote_proposal_files(quote_id: str, user: Optional[dict] = None) -> tuple[bool, str, list[Path]]:
    return generate_quote_proposal_selected(quote_id, "PDF + WORD", user, None)


if PYSIDE_OK:
    # Mejorar busqueda de vehiculo disponible en cotizaciones: maximo 10 visibles + busqueda por texto.
    _ORIG_QUOTE_EDITOR_V49 = QuoteEditorDialog
    class QuoteEditorDialog(_ORIG_QUOTE_EDITOR_V49):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            try:
                self.vehicle_combo.setMaxVisibleItems(10)
                self.vehicle_combo.setEditable(True)
                self.vehicle_combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
                from PySide6.QtWidgets import QCompleter
                items = [self.vehicle_combo.itemText(i) for i in range(self.vehicle_combo.count())]
                comp = QCompleter(items, self.vehicle_combo)
                comp.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
                comp.setFilterMode(Qt.MatchFlag.MatchContains)
                self.vehicle_combo.setCompleter(comp)
            except Exception:
                pass

    class ProposalOptionsDialog(QDialog):
        def __init__(self, parent, quote: dict, user: dict):
            super().__init__(parent); self.quote=quote; self.user=user; self.setWindowTitle("Generar propuesta"); self.setMinimumWidth(580); self._build()
        def _build(self):
            lay=QVBoxLayout(self); lay.addWidget(make_title("Opciones de propuesta", "La propuesta se puede generar en PDF, Word o ambos. Excel ya no se usa para propuestas."))
            f=QFormLayout(); lay.addLayout(f)
            self.formato=QComboBox(); self.formato.addItems(["PDF + WORD", "PDF", "WORD"])
            self.firmas=load_quote_signatures(); self.firma_combo=QComboBox()
            for sig in self.firmas:
                self.firma_combo.addItem(f"{sig.get('nombre')} · {sig.get('telefono')} · {sig.get('correo')}")
            default=default_signature_for_user(self.user)
            for i,sig in enumerate(self.firmas):
                if sig.get("nombre") == default.get("nombre"):
                    self.firma_combo.setCurrentIndex(i); break
            self.opcion_venta=QComboBox(); self.opcion_venta.setEditable(True); self.opcion_venta.addItems(load_sale_options()); self.opcion_venta.setCurrentText(self.quote.get("opcion_venta") or "RENTA CAR")
            self.opcion_compra=MoneyEdit(); self.opcion_compra.setRange(0,999999); self.opcion_compra.setValue(_to_float(self.quote.get("opcion_compra_usd"), calculate_purchase_option_from_quote(self.quote)))
            f.addRow("Formato:", self.formato); f.addRow("Quién hace la propuesta:", self.firma_combo); f.addRow("Opción de venta:", self.opcion_venta); f.addRow("Opción de compra:", self.opcion_compra)
            row=QHBoxLayout(); lay.addLayout(row); b=QPushButton("Generar"); b.setObjectName("orange"); b.clicked.connect(self.accept); c=QPushButton("Cancelar"); c.setObjectName("ghost"); c.clicked.connect(self.reject); row.addWidget(b); row.addWidget(c); row.addStretch(1)
        def data(self):
            idx=self.firma_combo.currentIndex(); sig=self.firmas[idx] if 0 <= idx < len(self.firmas) else default_signature_for_user(self.user)
            return {"formato": self.formato.currentText(), "firma": sig, "opcion_venta": self.opcion_venta.currentText(), "opcion_compra_usd": self.opcion_compra.value()}

    _ORIG_QUOTE_DETAIL_V49 = QuoteDetailDialog
    class QuoteDetailDialog(_ORIG_QUOTE_DETAIL_V49):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            try:
                for btn in self.findChildren(QPushButton):
                    if btn.text().strip().lower().startswith("generar propuesta"):
                        btn.setText("Generar propuesta PDF / Word")
                extra_row = QHBoxLayout()
                bver = QPushButton("Ver propuestas guardadas"); bver.setObjectName("ghost"); bver.clicked.connect(self.view_saved_proposals)
                bsin = QPushButton("Cliente sin interés"); bsin.setObjectName("danger"); bsin.clicked.connect(self.mark_no_interest)
                extra_row.addWidget(bver); extra_row.addWidget(bsin); extra_row.addStretch(1)
                self.layout().insertLayout(2, extra_row)
            except Exception:
                pass
        def generate_proposal(self):
            q=find_quote(self.quote_id)
            if not q:
                QMessageBox.warning(self,"Propuesta","Cotización no encontrada."); return
            dlg=ProposalOptionsDialog(self,q,self.user)
            if dlg.exec()!=QDialog.DialogCode.Accepted: return
            data=dlg.data(); formato=data.pop("formato")
            ok,msg,paths=generate_quote_proposal_selected(self.quote_id,formato,self.user,data)
            if not ok:
                QMessageBox.warning(self,"Propuesta",msg); return
            box=QMessageBox(self); box.setWindowTitle("Propuesta generada"); box.setText(msg+"\n\n"+"\n".join(p.name for p in paths))
            pdf_btn=word_btn=None
            for p in paths:
                if p.suffix.lower()==".pdf" and pdf_btn is None: pdf_btn=box.addButton("Abrir PDF",QMessageBox.ButtonRole.AcceptRole)
                if p.suffix.lower()==".docx" and word_btn is None: word_btn=box.addButton("Abrir Word",QMessageBox.ButtonRole.ActionRole)
            copy_btn=box.addButton("Guardar copia…",QMessageBox.ButtonRole.ActionRole); box.addButton("Cerrar",QMessageBox.ButtonRole.RejectRole); box.exec(); clicked=box.clickedButton()
            if clicked==pdf_btn:
                for p in paths:
                    if p.suffix.lower()==".pdf": QDesktopServices.openUrl(QUrl.fromLocalFile(str(p))); break
            elif clicked==word_btn:
                for p in paths:
                    if p.suffix.lower()==".docx": QDesktopServices.openUrl(QUrl.fromLocalFile(str(p))); break
            elif clicked==copy_btn:
                folder=QFileDialog.getExistingDirectory(self,"Selecciona carpeta para guardar copia",str(Path.home()))
                if folder:
                    for p in paths:
                        try: shutil.copy2(p, Path(folder)/p.name)
                        except Exception: pass
                    QMessageBox.information(self,"Copia","Copia guardada correctamente.")
            self.refresh()
        def view_saved_proposals(self):
            q=find_quote(self.quote_id) or {}
            props=q.get("propuestas",[]) or []
            if not props:
                QMessageBox.information(self,"Propuestas","Este cliente aún no tiene propuestas guardadas."); return
            dlg=QDialog(self); dlg.setWindowTitle("Propuestas guardadas"); dlg.setMinimumSize(820,420); lay=QVBoxLayout(dlg)
            tbl=QTableWidget(0,4); tbl.setHorizontalHeaderLabels(["Fecha","Tipo","Archivo","Ruta"]); tbl.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); tbl.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows); tbl.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers); lay.addWidget(tbl)
            paths=[]; tbl.setRowCount(len(props))
            for r,p in enumerate(props):
                path=Path(str(p.get("archivo") or p.get("path") or "")); paths.append(path)
                vals=[_fmt_date(p.get("fecha")), p.get("tipo",""), p.get("nombre") or path.name, str(path)]
                for c,v in enumerate(vals): tbl.setItem(r,c,QTableWidgetItem(str(v)))
            row=QHBoxLayout(); lay.addLayout(row)
            bopen=QPushButton("Abrir seleccionado"); bopen.setObjectName("orange"); bclose=QPushButton("Cerrar"); bclose.setObjectName("ghost"); row.addWidget(bopen); row.addWidget(bclose); row.addStretch(1)
            def open_sel():
                rr=tbl.currentRow()
                if rr<0 or rr>=len(paths): QMessageBox.information(dlg,"Abrir","Selecciona una propuesta."); return
                p=paths[rr]
                if not p.exists(): QMessageBox.warning(dlg,"Abrir",f"No existe el archivo:\n{p}"); return
                QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
            bopen.clicked.connect(open_sel); bclose.clicked.connect(dlg.reject); tbl.cellDoubleClicked.connect(lambda r,c: open_sel())
            dlg.exec()
        def mark_no_interest(self):
            q=find_quote(self.quote_id)
            if not q: return
            if QMessageBox.question(self,"Cliente sin interés","¿Marcar este cliente como sin interés y detener seguimiento?") != QMessageBox.StandardButton.Yes: return
            q["estado"] = QUOTE_PERDIDA
            q["ultima_gestion"] = date.today().isoformat()
            q.setdefault("seguimientos", []).append({"fecha": _now_iso(), "usuario": self.user.get("usuario", ""), "accion": "CLIENTE_SIN_INTERES", "comentario": "Cliente marcado sin interés. No requiere seguimiento hasta reactivarlo."})
            upsert_quote(q)
            self.refresh()
            QMessageBox.information(self,"Cliente sin interés","Cliente marcado como perdido/sin interés.")



# =============================================================================
# AJUSTES V5.0 - INSTRUCCIONES WORD: COMPRA LOCAL, CARACTERISTICAS, PROPUESTAS Y COTIZACIONES
# =============================================================================
# Este bloque corrige lo solicitado en INSTRUCCIONES PARA MEJORAR LYM:
# - Caracteristicas por lineas desde ventana flotante, no QTextEdit directo.
# - Catalogo Info vehiculos / Marcas.
# - Compra LOCAL oculta Estado USA/Subasta y ya no usa precio venta local ni precio minimo local.
# - Compra LOCAL no pasa automaticamente a Disponible: queda en Precio Final para definir precio y foto.
# - Se elimina el registro de compra dentro de gastos detallados para no duplicar costo.
# - Cotizaciones con fecha disponible, ordenadas de mas reciente a mas antigua.
# - Historial de propuestas PDF/Word visible y nueva propuesta despues de recalcular meses.
# - Boton Cliente sin interes y reactivacion.
# - Selector de vehiculos con busqueda y maximo 10 visibles.

# Asegurar DoughnutChart a nivel global para reportes de inventario.
try:
    from openpyxl.chart.pie_chart import DoughnutChart as DoughnutChart
except Exception:
    try:
        from openpyxl.chart import DoughnutChart as DoughnutChart  # type: ignore
    except Exception:
        DoughnutChart = None  # type: ignore

F_CATALOG_INFO_MARCAS = F_CATALOG_MARCAS


def _v50_is_purchase_cost_item(g: dict) -> bool:
    source = str(g.get("source") or "").lower()
    cat = _norm(g.get("categoria"))
    sub = _norm(g.get("subcategoria"))
    return source in ("purchase", "stage:comprado") or (cat == "COMPRA") or sub in ("PRECIO_GANADO", "COMPROBANTE COMPRA", "COMPROBANTE_COMPRA", "PRECIO GANADO")


def _v50_remove_purchase_cost_items(vehicle: dict) -> bool:
    """Evita que el precio ganado aparezca como gasto detallado y se duplique con el costo base."""
    gastos = vehicle.get("gastos_detallados")
    if not isinstance(gastos, list):
        return False
    new_items = [g for g in gastos if not _v50_is_purchase_cost_item(g)]
    if len(new_items) != len(gastos):
        vehicle["gastos_detallados"] = new_items
        return True
    return False


_ORIG_CREATE_VEHICLE_PURCHASE_V50 = create_vehicle_purchase

def create_vehicle_purchase(data: dict, comprobante_src: Path, user: dict, device: DeviceInfo) -> tuple[bool, str, str]:
    tipo = _norm(data.get("tipo_compra") or TIPO_COMPRA_LOTE)
    is_local = tipo == "LOCAL"
    # Compra LOCAL: no debe llenar flujo USA/Subasta. Se fuerza LOCAL solamente como referencia.
    clean_data = dict(data)
    if is_local:
        clean_data["estado_usa"] = ""
        clean_data["subasta"] = "LOCAL"
        clean_data["pais_compra"] = data.get("pais_compra") or "LOCAL"
        clean_data.pop("precio_venta_local_usd", None)
        clean_data.pop("precio_minimo_local_usd", None)
    ok, msg, vid = _ORIG_CREATE_VEHICLE_PURCHASE_V50(clean_data, comprobante_src, user, device)
    if not ok or not vid:
        return ok, msg, vid
    v = find_vehicle(vid)
    if not v:
        return ok, msg, vid
    ensure_vehicle_runtime_fields(v)
    removed = _v50_remove_purchase_cost_items(v)
    if is_local:
        fecha = clean_data.get("fecha_compra") or date.today().isoformat()
        # Revertir el comportamiento anterior que lo dejaba disponible automatico.
        v["tipo_compra"] = TIPO_COMPRA_LOCAL
        v["estado_actual"] = STAGE_PRECIO_FINAL
        v["estado_comercial"] = COMM_NO_DISPONIBLE
        v["subasta"] = "LOCAL"
        v["estado_usa"] = ""
        v["pais_compra"] = clean_data.get("pais_compra") or "LOCAL"
        v["precio_venta_usd"] = 0.0
        v["precio_minimo_usd"] = 0.0
        v["precio_final"] = {}
        for sk in STAGE_ORDER:
            st = vehicle_stage(v, sk)
            if sk == STAGE_COMPRADO:
                st.update({"status": "COMPLETADO", "fecha_inicio": fecha, "fecha_fin": fecha, "comentario": "Compra local registrada."})
            elif sk in (STAGE_TRASLADO_USA, STAGE_TRANSITO, STAGE_ADUANA, STAGE_PREPARACION):
                st.update({"status": "OMITIDO", "fecha_inicio": None, "fecha_fin": None, "costo_usd": 0, "comentario": "No aplica por compra local."})
            elif sk == STAGE_PRECIO_FINAL:
                st.update({"status": "EN PROCESO", "fecha_inicio": fecha, "fecha_fin": None, "costo_usd": 0, "comentario": "Compra local lista para definir precio final antes de poner disponible."})
            elif sk == STAGE_DISPONIBLE:
                st.update({"status": "PENDIENTE", "fecha_inicio": None, "fecha_fin": None, "costo_usd": 0, "comentario": "Pendiente de aprobación de precio final."})
        v.setdefault("historial", []).append({"fecha": _now_iso(), "usuario": user.get("usuario", ""), "computadora": getattr(device, 'computer_name', ''), "accion": "COMPRA_LOCAL_PRECIO_PENDIENTE", "detalle": "Compra local registrada. Debe definir precio final y luego marcar disponible."})
        msg = msg.split("\nCompra LOCAL:")[0] + "\nCompra LOCAL: se omitió USA/Aduana, pero queda pendiente definir precio final antes de estar disponible."
    if removed:
        v.setdefault("historial", []).append({"fecha": _now_iso(), "usuario": user.get("usuario", ""), "computadora": getattr(device, 'computer_name', ''), "accion": "AJUSTE_COSTO_COMPRA", "detalle": "Se retiró el precio ganado del historial de gastos para evitar duplicar el costo base de compra."})
    save_vehicle(v)
    return True, msg, vid


# Refuerzo adicional: si por datos viejos existe compra en gastos_detallados, no se suma doble.
_ORIG_VEHICLE_TOTAL_COST_V50 = vehicle_total_cost

def vehicle_total_cost(vehicle: dict) -> float:
    ensure_vehicle_runtime_fields(vehicle)
    gastos = vehicle.get("gastos_detallados")
    if isinstance(gastos, list) and any(_v50_is_purchase_cost_item(g) for g in gastos):
        # No modifica en lectura; solo calcula evitando duplicados.
        cloned = dict(vehicle)
        cloned["gastos_detallados"] = [g for g in gastos if not _v50_is_purchase_cost_item(g)]
        base = _to_float(vehicle.get("precio_ganado_usd"), 0)
        total = base
        for g in cloned.get("gastos_detallados", []) or []:
            total += _to_float(g.get("monto_usd"), 0)
        # Incluir gastos_extra si existieran y no están dentro de gastos_detallados.
        for g in cloned.get("gastos_extra", []) or []:
            total += _to_float(g.get("monto_usd"), 0)
        return round(total, 2)
    return _ORIG_VEHICLE_TOTAL_COST_V50(vehicle)


def _v50_available_date(vehicle: dict) -> str:
    ensure_vehicle_runtime_fields(vehicle)
    d = None
    try:
        d = _v46_stage_start(vehicle, STAGE_DISPONIBLE) or _v46_stage_end(vehicle, STAGE_PRECIO_FINAL) or _v46_stage_end(vehicle, STAGE_DISPONIBLE)
    except Exception:
        d = None
    return d.isoformat() if isinstance(d, date) else ""


if PYSIDE_OK:
    class FeatureListDialog(QDialog):
        def __init__(self, parent=None, initial_text: str = ""):
            super().__init__(parent)
            self.setWindowTitle("Agregar características")
            self.setMinimumSize(620, 420)
            self._build(initial_text)
        def _build(self, initial_text: str):
            lay = QVBoxLayout(self)
            title = make_title("Agregar características", "Cada línea será una característica independiente en la propuesta.")
            lay.addWidget(title)
            self.list = QListWidget(); lay.addWidget(self.list)
            for line in str(initial_text or "").replace(";", "\n").splitlines():
                line = line.strip(" •-\t")
                if line:
                    self.list.addItem(line)
            row = QHBoxLayout(); lay.addLayout(row)
            badd = QPushButton("Agregar línea"); badd.setObjectName("orange"); badd.clicked.connect(self.add_line)
            bedit = QPushButton("Editar línea"); bedit.setObjectName("ghost"); bedit.clicked.connect(self.edit_line)
            bdel = QPushButton("Quitar línea"); bdel.setObjectName("danger"); bdel.clicked.connect(self.delete_line)
            row.addWidget(badd); row.addWidget(bedit); row.addWidget(bdel); row.addStretch(1)
            row2 = QHBoxLayout(); lay.addLayout(row2)
            ok = QPushButton("Guardar características"); ok.setObjectName("orange"); ok.clicked.connect(self.accept)
            cancel = QPushButton("Cancelar"); cancel.setObjectName("ghost"); cancel.clicked.connect(self.reject)
            row2.addWidget(ok); row2.addWidget(cancel); row2.addStretch(1)
        def add_line(self):
            txt, ok = QInputDialog.getText(self, "Agregar característica", "Escribe una característica:")
            if ok and txt.strip():
                self.list.addItem(txt.strip(" •-\t"))
        def edit_line(self):
            item = self.list.currentItem()
            if not item:
                QMessageBox.information(self, "Características", "Selecciona una línea."); return
            txt, ok = QInputDialog.getText(self, "Editar característica", "Característica:", text=item.text())
            if ok and txt.strip():
                item.setText(txt.strip(" •-\t"))
        def delete_line(self):
            row = self.list.currentRow()
            if row >= 0:
                self.list.takeItem(row)
        def text(self) -> str:
            return "\n".join(self.list.item(i).text().strip() for i in range(self.list.count()) if self.list.item(i).text().strip())

    class PurchasePage(QWidget):
        def __init__(self, main):
            super().__init__()
            self.main = main
            self.comprobante_path = ""
            self.oc_path = ""
            self._features_text = ""
            self._build()
        def _build(self):
            lay = QVBoxLayout(self)
            lay.addWidget(make_title("Compra vehicular · Nueva compra", "Compra USA, compra LOCAL o servicio de importación."))
            scroll = QScrollArea(); scroll.setWidgetResizable(True); content = QWidget(); scroll.setWidget(content); body = QVBoxLayout(content)
            g1 = QGroupBox("Datos del vehículo"); f1 = QFormLayout(g1)
            self.marca = QComboBox(); self.marca.setEditable(True)
            self.modelo = QLineEdit()
            self.anio = QSpinBox(); self.anio.setRange(1980, date.today().year + 1); self.anio.setValue(date.today().year)
            self.millaje = QSpinBox(); self.millaje.setRange(0, 999999)
            self.color = QLineEdit()
            self.tipo = QComboBox(); self.tipo.setEditable(True); self.tipo.addItems(["SEDAN", "SUV", "PICKUP", "VAN", "CAMIONETA", "DEPORTIVO", "OTRO"])
            self.feature_preview = QTextEdit(); self.feature_preview.setReadOnly(True); self.feature_preview.setMinimumHeight(85); self.feature_preview.setPlaceholderText("Sin características. Usa el botón para agregar líneas.")
            bfeat = QPushButton("Agregar características"); bfeat.setObjectName("orange"); bfeat.clicked.connect(self.open_features)
            featrow = QWidget(); fl = QHBoxLayout(featrow); fl.setContentsMargins(0,0,0,0); fl.addWidget(self.feature_preview); fl.addWidget(bfeat)
            for lab, w in [("Marca:", self.marca), ("Modelo:", self.modelo), ("Año:", self.anio), ("Millaje:", self.millaje), ("Color:", self.color), ("Tipo:", self.tipo), ("Características propuesta:", featrow)]:
                f1.addRow(lab, w)
            g2 = QGroupBox("Tipo de compra y costos iniciales"); self.f2 = QFormLayout(g2)
            self.tipo_compra = QComboBox(); self.tipo_compra.addItems(["LOTE USA", "LOCAL", "SERVICIO DE IMPORTACION"]); self.tipo_compra.currentTextChanged.connect(self._toggle_tipo_compra)
            self.pais_compra = QComboBox(); self.pais_compra.setEditable(True); self.pais_compra.addItems(load_catalog(F_CATALOG_PAISES, DEFAULT_PAISES_DESTINO))
            self.estado_usa = QComboBox()
            self.subasta = QComboBox(); self.subasta.setEditable(True)
            self.lote = QLineEdit()
            self.precio = MoneyEdit(); self.precio.setRange(0, 9999999)
            self.fecha = configure_date_edit(QDateEdit()); self.fecha.setDate(QDate.currentDate())
            self._form_rows = {}
            for lab, w, key in [
                ("Tipo compra:", self.tipo_compra, "tipo"),
                ("País compra/local:", self.pais_compra, "pais"),
                ("Estado USA:", self.estado_usa, "estado_usa"),
                ("Subasta:", self.subasta, "subasta"),
                ("Número lote:", self.lote, "lote"),
                ("Costo / precio compra USD:", self.precio, "precio"),
                ("Fecha compra:", self.fecha, "fecha"),
            ]:
                label = QLabel(lab); self.f2.addRow(label, w); self._form_rows[key] = (label, w)
            g3 = QGroupBox("Comprobantes y OC"); f3 = QFormLayout(g3)
            self.comp_label = QLineEdit(); self.comp_label.setReadOnly(True)
            bcomp = QPushButton("Subir comprobante"); bcomp.setObjectName("ghost"); bcomp.clicked.connect(self.pick_comprobante)
            r1 = QWidget(); l1 = QHBoxLayout(r1); l1.setContentsMargins(0,0,0,0); l1.addWidget(self.comp_label); l1.addWidget(bcomp)
            self.oc_num = QLineEdit(); self.oc_label = QLineEdit(); self.oc_label.setReadOnly(True)
            boc = QPushButton("Subir PDF OC"); boc.setObjectName("ghost"); boc.clicked.connect(self.pick_oc)
            r2 = QWidget(); l2 = QHBoxLayout(r2); l2.setContentsMargins(0,0,0,0); l2.addWidget(self.oc_label); l2.addWidget(boc)
            self.obs = QTextEdit(); self.obs.setMinimumHeight(80)
            f3.addRow("Comprobante:", r1); f3.addRow("Número OC:", self.oc_num); f3.addRow("PDF OC:", r2); f3.addRow("Observaciones:", self.obs)
            body.addWidget(g1); body.addWidget(g2); body.addWidget(g3)
            btn = QPushButton("Crear compra"); btn.setObjectName("orange"); btn.clicked.connect(self.save_purchase); body.addWidget(btn)
            lay.addWidget(scroll); self.refresh_catalogs(); self._toggle_tipo_compra(self.tipo_compra.currentText())
        def refresh_catalogs(self):
            self.marca.clear(); self.marca.addItems(load_catalog(F_CATALOG_INFO_MARCAS, DEFAULT_MARCAS))
            self.subasta.clear(); self.subasta.addItems(load_catalog(F_CATALOG_SUBASTAS, DEFAULT_SUBASTAS))
            self.estado_usa.clear(); self.estado_usa.addItems(load_catalog(F_CATALOG_ESTADOS_USA, US_STATES))
        def _toggle_tipo_compra(self, text: str):
            is_local = _norm(text) == "LOCAL"
            for key in ("estado_usa", "subasta"):
                lab, w = self._form_rows.get(key, (None, None))
                if lab: lab.setVisible(not is_local)
                if w: w.setVisible(not is_local)
            if is_local:
                self.subasta.setCurrentText("LOCAL")
                self.estado_usa.setCurrentText("")
                if not self.lote.text().strip():
                    self.lote.setPlaceholderText("Opcional para compra local")
            else:
                self.lote.setPlaceholderText("")
        def open_features(self):
            dlg = FeatureListDialog(self, self._features_text)
            if dlg.exec() == QDialog.DialogCode.Accepted:
                self._features_text = dlg.text()
                self.feature_preview.setPlainText(self._features_text)
        def pick_comprobante(self):
            path, _ = QFileDialog.getOpenFileName(self, "Selecciona comprobante", str(Path.home()), "Documentos (*.pdf *.png *.jpg *.jpeg);;Todos (*.*)")
            if path:
                self.comprobante_path = path; self.comp_label.setText(path)
        def pick_oc(self):
            path, _ = QFileDialog.getOpenFileName(self, "Selecciona PDF de OC", str(Path.home()), "PDF (*.pdf);;Todos (*.*)")
            if path:
                self.oc_path = path; self.oc_label.setText(path)
        def save_purchase(self):
            if not self.comprobante_path:
                QMessageBox.warning(self, "Validación", "Debes subir el comprobante de compra."); return
            tipo_text = self.tipo_compra.currentText(); is_local = _norm(tipo_text) == "LOCAL"; is_servicio = _norm(tipo_text) == _norm(TIPO_COMPRA_IMPORTACION)
            data = {
                "marca": self.marca.currentText(),
                "modelo": self.modelo.text(),
                "anio": self.anio.value(),
                "millaje": self.millaje.value(),
                "color": self.color.text(),
                "tipo": self.tipo.currentText(),
                "estado_usa": "" if is_local else self.estado_usa.currentText(),
                "subasta": "LOCAL" if is_local else self.subasta.currentText(),
                "lote": self.lote.text() or ("LOCAL" if is_local else ""),
                "precio_ganado_usd": self.precio.value(),
                "fecha_compra": self.fecha.date().toPython().isoformat(),
                "observaciones": self.obs.toPlainText(),
                "caracteristicas": self._features_text,
                "oc_compra_numero": self.oc_num.text(),
                "oc_compra_src": self.oc_path,
                "tipo_compra": TIPO_COMPRA_LOCAL if is_local else (TIPO_COMPRA_IMPORTACION if is_servicio else TIPO_COMPRA_LOTE),
                "pais_compra": self.pais_compra.currentText(),
                "servicio_importacion": is_servicio,
            }
            ok, msg, vid = create_vehicle_purchase(data, Path(self.comprobante_path), self.main.user, self.main.device)
            if not ok:
                QMessageBox.warning(self, "Validación", msg); return
            QMessageBox.information(self, "Compra creada", msg)
            self.main.refresh_all(); self.main.open_vehicle_detail(vid)

    _ORIG_CATALOGOS_PAGE_V50 = CatalogosPage
    class CatalogosPage(_ORIG_CATALOGOS_PAGE_V50):
        def _build(self):
            super()._build()
            try:
                self.tabs.addTab(self._build_simple_tab("Info vehículos · Marcas", F_CATALOG_INFO_MARCAS, DEFAULT_MARCAS), "Info vehículos")
            except Exception:
                pass
        def refresh_all(self):
            try:
                super().refresh_all()
            except Exception:
                pass
            try:
                if hasattr(self, "simple_tabs") and "Info vehículos · Marcas" in self.simple_tabs:
                    self.refresh_simple("Info vehículos · Marcas")
            except Exception:
                pass

    _ORIG_COTIZACIONES_PAGE_V50 = CotizacionesPage
    class CotizacionesPage(_ORIG_COTIZACIONES_PAGE_V50):
        def _build(self):
            lay = QVBoxLayout(self)
            lay.addWidget(make_title("Cotizaciones", "Vehículos disponibles, clientes interesados y cotizaciones generales sin borrar historial."))
            self.tabs = QTabWidget(); lay.addWidget(self.tabs)
            wv = QWidget(); vl = QVBoxLayout(wv)
            brow = QHBoxLayout(); bnew = QPushButton("Nueva cotización"); bnew.setObjectName("orange"); bnew.clicked.connect(self.new_quote); brow.addWidget(bnew); brow.addStretch(1); vl.addLayout(brow)
            self.vehicle_table = QTableWidget(0, 8)
            self.vehicle_table.setHorizontalHeaderLabels(["Código", "Vehículo", "Precio", "Costo", "Ganancia", "Fecha disponible", "Cotizaciones", "Días disponible"])
            self.vehicle_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); _v46_apply_table_hardening(self.vehicle_table); self.vehicle_table.cellDoubleClicked.connect(self.open_vehicle_quotes); vl.addWidget(self.vehicle_table)
            self.tabs.addTab(wv, "Vehículos disponibles")
            wg = QWidget(); gl = QVBoxLayout(wg)
            self.quote_table = QTableWidget(0, 10)
            self.quote_table.setHorizontalHeaderLabels(["Cliente", "Teléfono", "Medio", "Vehículo cotizado", "Fecha disponible", "Fecha cotización", "Última gestión", "Días", "Estado", "Color"])
            self.quote_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); _v46_apply_table_hardening(self.quote_table); self.quote_table.cellDoubleClicked.connect(self.open_quote); gl.addWidget(self.quote_table)
            self.tabs.addTab(wg, "Cotizaciones general")
            try:
                # Mantener la base de datos de clientes si el sistema ya la traia en versiones anteriores.
                if hasattr(super(), "_build_clients_tab"):
                    self.tabs.addTab(super()._build_clients_tab(), "Base de datos clientes")
            except Exception:
                pass
        def refresh(self):
            vehicles = [v for v in load_vehicles() if ensure_vehicle_runtime_fields(v) and v.get("estado_actual") == STAGE_DISPONIBLE and v.get("estado_comercial") == COMM_DISPONIBLE]
            vehicles.sort(key=lambda v: _parse_date(_v50_available_date(v)) or date.min, reverse=True)
            self._vehicle_ids = []
            self.vehicle_table.setRowCount(len(vehicles))
            for r, v in enumerate(vehicles):
                self._vehicle_ids.append(v.get("id"))
                fd = _v50_available_date(v)
                vals = [v.get("codigo"), _vehicle_display_name(v), _fmt_usd(v.get("precio_venta_usd")), _fmt_usd(vehicle_total_cost(v)), _fmt_usd(vehicle_expected_profit(v)), _fmt_date(fd), str(len(quotes_for_vehicle(v.get("id")))), str(current_stage_days(v))]
                for c, val in enumerate(vals):
                    it = QTableWidgetItem(str(val)); it.setFlags(it.flags() & ~Qt.ItemFlag.ItemIsEditable); self.vehicle_table.setItem(r, c, it)
            vehicle_by_id = {v.get("id"): v for v in load_vehicles()}
            qs = load_quotes()
            qs.sort(key=lambda q: _parse_date(q.get("fecha_cotizacion")) or date.min, reverse=True)
            self._quote_ids = []
            self.quote_table.setRowCount(len(qs))
            for r, q in enumerate(qs):
                self._quote_ids.append(q.get("id")); cl = q.get("cliente", {}) or {}; snap = q.get("vehicle_snapshot", {}) or {}; lvl = quote_alert_level(q)
                vv = vehicle_by_id.get(q.get("vehicle_id")) or find_vehicle(q.get("vehicle_code", "")) or {}
                fd = _v50_available_date(vv) if vv else str(q.get("fecha_disponible") or "")
                vals = [cl.get("nombre"), cl.get("telefono"), cl.get("medio_contacto"), f"{_vehicle_display_name(snap)} · {q.get('vehicle_code')}", _fmt_date(fd), _fmt_date(q.get("fecha_cotizacion")), _fmt_date(q.get("ultima_gestion")), str(quote_days_without_purchase(q)), QUOTE_STATUS_LABELS.get(q.get("estado"), q.get("estado")), lvl]
                for c, val in enumerate(vals):
                    item = QTableWidgetItem(str(val)); item.setBackground(QColor(_quote_alert_color(lvl))); item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable); self.quote_table.setItem(r, c, item)

    _ORIG_QUOTE_EDITOR_DIALOG_V50 = QuoteEditorDialog
    class QuoteEditorDialog(_ORIG_QUOTE_EDITOR_DIALOG_V50):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            try:
                self.vehicle_combo.setMaxVisibleItems(10)
                self.vehicle_combo.setEditable(True)
                self.vehicle_combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
                from PySide6.QtWidgets import QCompleter
                items = [self.vehicle_combo.itemText(i) for i in range(self.vehicle_combo.count())]
                comp = QCompleter(items, self.vehicle_combo)
                comp.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
                comp.setFilterMode(Qt.MatchFlag.MatchContains)
                self.vehicle_combo.setCompleter(comp)
            except Exception:
                pass
        def save(self):
            # Guardar fecha disponible en la cotizacion para que quede snapshot historico.
            try:
                idx = self.vehicle_combo.currentIndex()
                vid = self._vehicle_ids[idx] if 0 <= idx < len(self._vehicle_ids) else ""
                v = find_vehicle(vid) if vid else None
                self._fecha_disponible_snapshot = _v50_available_date(v) if v else ""
            except Exception:
                self._fecha_disponible_snapshot = ""
            before_id = getattr(self, "quote_id", "")
            super().save()
            try:
                # Si el save fue exitoso, el dialog puede quedar aceptado; se refuerza el snapshot en la cotizacion.
                qid = getattr(self, "quote_id", "") or before_id
                if qid:
                    q = find_quote(qid)
                    if q:
                        q["fecha_disponible"] = getattr(self, "_fecha_disponible_snapshot", "") or q.get("fecha_disponible", "")
                        upsert_quote(q)
            except Exception:
                pass

    _ORIG_CREATE_QUOTE_V50 = create_quote
    def create_quote(data: dict, user: dict, device: DeviceInfo, quote_id: str = "") -> tuple[bool, str, str]:
        ok, msg, qid = _ORIG_CREATE_QUOTE_V50(data, user, device, quote_id)
        if ok and qid:
            q = find_quote(qid)
            if q:
                v = find_vehicle(q.get("vehicle_id", "")) or find_vehicle(q.get("vehicle_code", ""))
                if v:
                    q["fecha_disponible"] = _v50_available_date(v)
                if quote_id:
                    q.setdefault("seguimientos", []).append({"fecha": _now_iso(), "usuario": user.get("usuario", ""), "accion": "COTIZACION_RECALCULADA", "comentario": "Se editaron/recalcularon los meses o valores de la cotización. Genera una nueva propuesta para conservar el PDF/Word histórico."})
                upsert_quote(q)
        return ok, msg, qid

    _ORIG_QUOTE_DETAIL_DIALOG_V50 = QuoteDetailDialog
    class QuoteDetailDialog(_ORIG_QUOTE_DETAIL_DIALOG_V50):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            try:
                # Evitar duplicar botones si ya existen.
                existing = [b.text().strip() for b in self.findChildren(QPushButton)]
                row = QHBoxLayout()
                if "Ver propuestas guardadas" not in existing:
                    bver = QPushButton("Ver propuestas guardadas"); bver.setObjectName("ghost"); bver.clicked.connect(self.view_saved_proposals); row.addWidget(bver)
                if "Cliente sin interés" not in existing:
                    bsin = QPushButton("Cliente sin interés"); bsin.setObjectName("danger"); bsin.clicked.connect(self.mark_no_interest); row.addWidget(bsin)
                if "Reactivar cliente" not in existing:
                    bact = QPushButton("Reactivar cliente"); bact.setObjectName("ghost"); bact.clicked.connect(self.reactivate_client); row.addWidget(bact)
                row.addStretch(1)
                self.layout().insertLayout(2, row)
            except Exception:
                pass
        def refresh(self):
            try:
                super().refresh()
            except Exception:
                pass
            q = find_quote(self.quote_id)
            if not q:
                return
            try:
                lines = []
                for h in q.get('seguimientos', []) or []:
                    lines.append(f"[{h.get('fecha')}] {h.get('usuario')} · {h.get('accion')} · {h.get('comentario')}")
                props = q.get("propuestas", []) or []
                if props:
                    lines.append("\n--- PROPUESTAS GENERADAS ---")
                    for i, p in enumerate(props, 1):
                        path = Path(str(p.get("archivo") or p.get("path") or ""))
                        lines.append(f"{i}. [{_fmt_date(p.get('fecha'))}] {p.get('tipo','')} · {path.name}")
                self.hist.setPlainText("\n".join(lines))
            except Exception:
                pass
        def edit_quote(self):
            dlg = QuoteEditorDialog(self, self.user, self.device, quote_id=self.quote_id)
            if dlg.exec() == QDialog.DialogCode.Accepted:
                self.refresh()
                if QMessageBox.question(self, "Nueva propuesta", "La cotización fue recalculada. ¿Deseas generar una nueva propuesta PDF/Word y guardarla en el historial?") == QMessageBox.StandardButton.Yes:
                    self.generate_proposal()
        def view_saved_proposals(self):
            q = find_quote(self.quote_id) or {}
            props = q.get("propuestas", []) or []
            if not props:
                QMessageBox.information(self, "Propuestas", "Este cliente aún no tiene propuestas guardadas."); return
            dlg = QDialog(self); dlg.setWindowTitle("Propuestas guardadas"); dlg.setMinimumSize(860, 440); lay = QVBoxLayout(dlg)
            tbl = QTableWidget(0, 4); tbl.setHorizontalHeaderLabels(["Fecha", "Tipo", "Archivo", "Ruta"]); tbl.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch); _v46_apply_table_hardening(tbl); lay.addWidget(tbl)
            paths = []
            tbl.setRowCount(len(props))
            for r, p in enumerate(props):
                path = Path(str(p.get("archivo") or p.get("path") or "")); paths.append(path)
                vals = [_fmt_date(p.get("fecha")), p.get("tipo", ""), p.get("nombre") or path.name, str(path)]
                for c, val in enumerate(vals):
                    tbl.setItem(r, c, QTableWidgetItem(str(val)))
            row = QHBoxLayout(); lay.addLayout(row)
            bopen = QPushButton("Abrir seleccionado"); bopen.setObjectName("orange"); bclose = QPushButton("Cerrar"); bclose.setObjectName("ghost"); row.addWidget(bopen); row.addWidget(bclose); row.addStretch(1)
            def open_sel():
                rr = tbl.currentRow()
                if rr < 0 or rr >= len(paths):
                    QMessageBox.information(dlg, "Abrir", "Selecciona una propuesta."); return
                p = paths[rr]
                if not p.exists():
                    QMessageBox.warning(dlg, "Abrir", f"No existe el archivo:\n{p}"); return
                QDesktopServices.openUrl(QUrl.fromLocalFile(str(p)))
            bopen.clicked.connect(open_sel); bclose.clicked.connect(dlg.reject); tbl.cellDoubleClicked.connect(lambda r, c: open_sel())
            dlg.exec()
        def mark_no_interest(self):
            q = find_quote(self.quote_id)
            if not q:
                return
            if QMessageBox.question(self, "Cliente sin interés", "¿Marcar este cliente como sin interés y detener seguimiento?") != QMessageBox.StandardButton.Yes:
                return
            q["estado"] = QUOTE_PERDIDA
            q["sin_interes"] = True
            q["ultima_gestion"] = date.today().isoformat()
            q.setdefault("seguimientos", []).append({"fecha": _now_iso(), "usuario": self.user.get("usuario", ""), "accion": "CLIENTE_SIN_INTERES", "comentario": "Cliente marcado sin interés. No requiere seguimiento hasta reactivarlo."})
            upsert_quote(q); self.refresh(); QMessageBox.information(self, "Cliente sin interés", "Cliente marcado como sin interés.")
        def reactivate_client(self):
            q = find_quote(self.quote_id)
            if not q:
                return
            q["sin_interes"] = False
            q["estado"] = QUOTE_SEGUIMIENTO if 'QUOTE_SEGUIMIENTO' in globals() else q.get("estado", "SEGUIMIENTO")
            q["ultima_gestion"] = date.today().isoformat()
            q.setdefault("seguimientos", []).append({"fecha": _now_iso(), "usuario": self.user.get("usuario", ""), "accion": "CLIENTE_REACTIVADO", "comentario": "Cliente reactivado para seguimiento comercial."})
            upsert_quote(q); self.refresh(); QMessageBox.information(self, "Reactivar", "Cliente reactivado.")


# Refuerzo V5.0 fuera de la GUI: snapshot de fecha disponible y seguimiento al recalcular cotización.
_ORIG_CREATE_QUOTE_V50_GLOBAL = create_quote

def create_quote(data: dict, user: dict, device: DeviceInfo, quote_id: str = "") -> tuple[bool, str, str]:
    ok, msg, qid = _ORIG_CREATE_QUOTE_V50_GLOBAL(data, user, device, quote_id)
    if ok and qid:
        q = find_quote(qid)
        if q:
            v = find_vehicle(q.get("vehicle_id", "")) or find_vehicle(q.get("vehicle_code", ""))
            if v:
                q["fecha_disponible"] = _v50_available_date(v)
            if quote_id:
                q.setdefault("seguimientos", []).append({"fecha": _now_iso(), "usuario": user.get("usuario", ""), "accion": "COTIZACION_RECALCULADA", "comentario": "Se editaron/recalcularon meses o valores. Para conservar evidencia, genera una nueva propuesta PDF/Word desde esta cotización."})
            upsert_quote(q)
    return ok, msg, qid

# =============================================================================
# SELF TEST / MAIN
# =============================================================================

def run_self_test() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="lym_auto_control_test_"))
    set_active_folder(tmp)
    bootstrap_system()
    device = collect_device_info()
    ok, msg = create_user("admin", "admin123", ROLE_ADMIN, device)
    if not ok:
        print("ERROR create_user:", msg); return 1
    ok, msg, user = authenticate("admin", "admin123", device)
    if not ok or not user:
        print("ERROR authenticate:", msg); return 1
    proof = tmp / "comprobante_demo.pdf"
    proof.write_bytes(b"COMPROBANTE DEMO LYM")
    data = {
        "marca": "TOYOTA", "modelo": "COROLLA", "anio": 2021, "millaje": 45200,
        "estado_usa": "FLORIDA", "subasta": "COPART", "lote": "ABC12345",
        "precio_ganado_usd": 8500, "fecha_compra": date.today().isoformat(), "observaciones": "Prueba automática",
    }
    ok, msg, vid = create_vehicle_purchase(data, proof, user, device)
    if not ok:
        print("ERROR create_vehicle_purchase:", msg); return 1
    veh = find_vehicle(vid)
    assert veh and re.match(r"LYM-CV-\d{4}-\d{4}", veh["codigo"]), veh
    ok, msg = update_vehicle_stage(vid, STAGE_TRASLADO_USA, {"fecha_inicio": date.today().isoformat(), "fecha_fin": None, "costo_usd": 400, "proveedor": "GRUA DEMO", "comentario": "Traslado interno"}, None, user, device)
    if not ok:
        print("ERROR update stage 1:", msg); return 1
    ok, msg = update_vehicle_stage(vid, STAGE_TRANSITO, {"fecha_inicio": date.today().isoformat(), "fecha_fin": None, "costo_usd": 1200, "proveedor": "NAVIERA DEMO", "comentario": "En tránsito"}, None, user, device)
    if not ok:
        print("ERROR update stage 2:", msg); return 1
    out = generate_html_report(load_vehicles(), user)
    if not out or not out.exists():
        print("ERROR report"); return 1
    okx, msgx, pathx = generate_kpi_excel_report(load_vehicles(), user)
    if not okx or not pathx or not pathx.exists():
        print("ERROR kpi excel:", msgx); return 1
    okb, msgb, pathb = create_backup_snapshot(user.get("usuario",""), "Self test")
    if not okb or not pathb or not pathb.exists():
        print("ERROR backup:", msgb); return 1
    print("SELF TEST OK")
    print("Carpeta prueba:", tmp)
    print("Reporte:", out)
    return 0


def main() -> int:
    if "--self-test" in sys.argv:
        return run_self_test()
    if not PYSIDE_OK:
        print("ERROR: falta PySide6. Instala con: pip install PySide6 cryptography")
        return 2
    app = QApplication(sys.argv)
    app.setStyleSheet(LYM_STYLESHEET)
    login = LoginDialog()
    if login.exec() != QDialog.DialogCode.Accepted or not login.user:
        return 0
    win = MainWindow(login.user, login.device)
    win.show()
    return app.exec()

if __name__ == "__main__":
    raise SystemExit(main())
